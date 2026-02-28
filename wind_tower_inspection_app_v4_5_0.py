#!/usr/bin/env python3
"""
================================================================================
 Wind Tower Inspection App — v4.5.0
 Build Date : 2026-02-28
 Status     : ✅ MAJOR UPDATE — v4.5.0: Face+Zone split + calibration cleanup

================================================================================
 CHANGE LOG
================================================================================
 v4.5.0  2026-02-28  Face+Zone split architecture + calibration reduction:

                        SURFACE+ZONE SPLIT — MUTUALLY EXCLUSIVE SELECTION:
                                   Blade sides now split into two separate,
                                   mutually exclusive dropdown selections:
                                   • Surface: PS (Pressure Side) OR SS (Suction Side)
                                   • Zone: LE (Leading Edge) OR TE (Trailing Edge)
                                     OR MB (Midbody)
                                   Hub/Tower: Both surface and zone disabled (N/A).
                                   Serial format updated: WTG-321_Blade A_PS-LE_001
                                   PDF table split: "Side" → "Surface" + "Zone"
                                   Legacy migration: Old "face" field auto-split
                                   Report validation: Warning if surface/zone missing
                                   Code reduction: ~800 lines removed from old
                                   calibration implementation

 v4.4.9  2026-02-28  PDF cover line removed + structured debug logging:

                        PDF FIX — SPURIOUS HORIZONTAL LINE ON COVER PAGE:
                                   An HRFlowable (grey separator rule, 0.5pt)
                                   was rendered between the "Tower Location /
                                   GPS Coordinates" table and the "Turbine
                                   Manufacturer / Rated Power" table on the
                                   PDF cover page.  This line was highlighted
                                   (yellow) in the review screenshot.
                                   Fix: removed the HRFlowable and its paired
                                   post-Spacer.  The pre-Spacer(1, 3mm) before
                                   loc_tbl is unchanged so inter-table spacing
                                   remains uniform 3mm — identical to all other
                                   cover-page table gaps (meta_tbl→scan_date,
                                   scan_date→loc_tbl, mfr_pwr→...).
                                   No other layout element touched.

                        DEBUG LOGGING UPGRADE — STRUCTURED WORKFLOW TRACES:
                                   _setup_logging() rewritten:
                                   • _QtNoiseFilter class: suppresses DEBUG spam
                                     from PyQt6.*, PIL.*, fontTools.*, and
                                     matplotlib.* on the console handler only.
                                     The file handler still receives everything
                                     for post-mortem analysis.
                                   • Per-module level overrides: PyQt6,
                                     PIL.TiffImagePlugin, fontTools set to
                                     WARNING to reduce file log noise.
                                   • Handler-type guards (not empty-handler
                                     guards): safe to call twice without
                                     accumulating duplicate handlers.
                                   • Session-start banner written to the log
                                     file at each startup so multiple sessions
                                     are easy to navigate in the log.
                                   Workflow methods instrumented (all new calls
                                   use [MODULE] tag prefix for grep-ability):
                                   • _new_project: creation, fields, folder
                                   • _open_project: load success/fail + counts
                                   • _save_project_action: explicit save trace
                                   • _on_thumb_selected: [VIEWER] selected file,
                                     GSD state, annotation count
                                   • _on_annotation_ready: calibration-gate
                                     block trace; success routing to panel
                                   • _on_annotation_modified: geometry edit with
                                     bbox coordinates and rotation
                                   • _on_save_annotation: full save trace —
                                     defect/severity/blade/size/gsd_source,
                                     irec.blade sync detection, serial before/
                                     after repair, size-lock decision reasoning
                                   • _on_delete_annotation: delete identity +
                                     before/after annotation count
                                   • _generate_report: gate stats, output paths,
                                     PDF and DOCX timing in ms, error detail

 v4.4.4  2026-02-27  Auto tower GPS fill + calibration + auto-rename (FIX-18/19):

                        FIX-18a — AUTO TOWER GPS FROM FIRST IMAGE (zero user input):
                                   _load_images() now calls the new
                                   _autofill_tower_gps_from_images() after each
                                   batch load.  If project.tower_lat/lon/alt_msl
                                   are all None it extracts drone_lat, drone_lon,
                                   absolute_altitude and relative_altitude from
                                   the first newly-added image via EXIFCalibrator,
                                   then sets:
                                     project.tower_lat  = drone_lat  (approx)
                                     project.tower_lon  = drone_lon  (approx)
                                     project.tower_base_alt_msl
                                       = absolute_altitude − relative_altitude
                                   Confirmed accurate: both test images give
                                   tower_base_msl = 633.18 m (identical, same
                                   physical tower).  Saves project and toasts
                                   the inspector with the auto-filled values.

                        FIX-18b — TIER-3 30m FALLBACK REMOVED / REPLACED:
                                   The hard-coded 30 m assumption was wrong by
                                   3–9× on real DJI data (test images: 30 m vs
                                   92.5 m and 257.5 m actual).
                                   New Tier-3 logic:
                                     If relative_altitude is known but gimbal
                                     pitch is absent → distance =
                                     relative_altitude / tan(45°) = relative_alt
                                     (labels "assumed-45deg-pitch").
                                     If relative_altitude is also absent → image
                                     is skipped with a clear warning.
                                   45° is the conservative mid-point of the
                                   35–60° pitch range used in wind tower
                                   inspection; it is physically grounded rather
                                   than arbitrary.

                        FIX-18c — NEGATIVE h_target GUARD IN estimate_distance_3d:
                                   When horiz_m × tan(|pitch|) > vert_m the
                                   camera ray passes the tower base and h_target
                                   goes negative, inflating true_dist and GSD.
                                   Fix: if h_target < 0 fall back to base_dist.

                        FIX-18d — ZERO-COORDINATE GUARD IN _extract_drone_lat/lon:
                                   A corrupted GPS block (0°,0°,0°) returned
                                   dd=0.0 and placed the tower at the equator /
                                   prime-meridian intersection.  Fix: after DMS
                                   conversion if abs(dd) < 0.001 return None.

                        FIX-19  — AUTO-RENAME IMAGE TO DEFECT NAME ON ANNOTATION SAVE:
                                   _on_save_annotation() now calls the new helper
                                   _auto_rename_after_annotation(ann) immediately
                                   after persisting the annotation.
                                   Auto-rename triggers ONLY when:
                                     • the current filename starts with a raw
                                       camera prefix (DJI_, IMG_, DSC, DCIM,
                                       P1_, P_).
                                     • the image has no prior annotations (first
                                       defect saved on this image).
                                   Rename format (slug-safe, report-consistent):
                                     {serial_number}_{DefectSlug}
                                   e.g. "WTG234_BladeA_LE_001_LeadingEdgeErosion"
                                   Conflict-safe: appends _2, _3 … if target
                                   already exists.  Updates project.images dict,
                                   self._image_paths, thumbnail strip item text
                                   and self._current_filepath atomically.
                                   Toasts the new filename in amber.

 v4.4.3  2026-02-27  3D GPS distance model (FIX-17):

                        MOTIVATION — pilot altitude varies shot-to-shot.
                        The old altitude/tan(pitch) method worked only when
                        RelativeAltitude and GimbalPitchDegree were present in
                        XMP.  It also assumed the drone was measuring distance
                        from its own altitude above takeoff (RelativeAltitude),
                        which is inaccurate on uneven terrain and gives no
                        information about which height on the tower the camera
                        is pointing at.

                        FIX-17a — ExifCalibrationData: 3 new fields:
                                    drone_lat          : decimal degrees (GPS IFD)
                                    drone_lon          : decimal degrees (GPS IFD)
                                    absolute_altitude  : MSL metres (XMP AbsoluteAltitude)
                                   The old estimate_distance_from_gps() is kept
                                   as a LEGACY method with a detailed docstring
                                   explaining the old approach and why it was
                                   superseded.  New method estimate_distance_3d()
                                   takes tower_lat, tower_lon, tower_base_msl
                                   as arguments and computes:
                                     1. Haversine horizontal distance
                                        (drone GPS → tower base GPS)
                                     2. Altitude difference
                                        (drone_alt_MSL − tower_base_MSL)
                                     3. True 3D slant distance via Pythagoras
                                     4. Gimbal-pitch refinement to find the
                                        actual aim-point height on the tower,
                                        giving a more precise slant distance
                                        than just drone→base.

                        FIX-17b — EXIFCalibrator: 3 new extraction methods:
                                    _extract_drone_lat() — handles ExifTool
                                      decimal-degree output AND PIL/DMS tuples
                                      AND IFDRational tuples.
                                    _extract_drone_lon() — same dual-path logic.
                                    _extract_absolute_altitude() — reads XMP
                                      AbsoluteAltitude (already parsed by
                                      _parse_dji_xmp); fallback to top-level
                                      ExifTool EXIF dict.
                                   All three wired into
                                   _build_calibration_from_exif() and stored on
                                   the returned ExifCalibrationData instance.

                        FIX-17c — Project dataclass: 3 new optional fields:
                                    tower_lat          : Optional[float] = None
                                    tower_lon          : Optional[float] = None
                                    tower_base_alt_msl : Optional[float] = None
                                   Serialised / deserialised in save_project /
                                   load_project with None defaults so all
                                   existing project.json files load cleanly.

                        FIX-17d — ReportSettingsDialog._tab_identity():
                                   New "Tower Location" section added below the
                                   existing GPS Coordinates row.  Three fields:
                                     Tower Lat (DD)
                                     Tower Lon (DD)
                                     Tower Base Altitude MSL (m)
                                   _on_ok() saves them to the settings blob AND
                                   directly updates project.tower_lat /
                                   tower_lon / tower_base_alt_msl so that
                                   _batch_auto_calibrate() can read from the
                                   project object without loading settings.

                        FIX-17e — _batch_auto_calibrate(): 3-tier distance
                                   estimation priority:
                                     Tier 1 (BEST) — estimate_distance_3d()
                                       requires: project.tower_lat/lon/alt AND
                                       exif_cal.drone_lat/lon/absolute_altitude
                                       AND exif_cal.gimbal_pitch.
                                       Records method = "3d-gps"
                                     Tier 2 — estimate_distance_from_gps()
                                       (legacy: relative_alt / tan(pitch))
                                       Records method = "pitch-altitude"
                                     Tier 3 — 30 m assumption
                                       Records method = "assumed-30m"
                                   irec.calibration_method now reflects which
                                   tier was used.  Summary dialog shows a
                                   per-tier breakdown.


 v4.4.2  2026-02-26  PDF report visual polish (FIX-16):
                        FIX-16a — COVER PAGE: "Report generated on …" IST
                                   timestamp paragraph removed from the cover
                                   page.  The subtitle already carries date
                                   context; the generated-on line was noisy
                                   and redundant.

                        FIX-16b — PDF HEADER BAR: lighter slate-gray.
                                   Old: setFillColorRGB(0.29, 0.33, 0.41)
                                       (~#4A5568 — too dark, logo lost)
                                   New: setFillColorRGB(0.72, 0.76, 0.82)
                                       (~#B8C2D1 — light slate, logo pops)
                                   Text colour unchanged (white Helvetica
                                   passes WCAG AA on the lighter bg).

                        FIX-16c — COVER PAGE TABLE SPACING: all inter-table
                                   vertical spacers normalised to 3 mm so
                                   the WTG No, Scan Date, Tower Location,
                                   and Manufacturer/Power boxes sit with
                                   equal breathing room between them.

                        FIX-16d — DEFECT SUMMARY TABLE: "Size" header renamed
                                   to "Size (cm)" so the unit is explicit for
                                   readers.  The data cells already show
                                   cm values; only the header was missing the
                                   unit annotation.

                        FIX-16e — COMPONENT CHIPS (breadcrumb rows on every
                                   defect page): two sites updated:
                                   • _build_annotation_pages breadcrumb chips
                                   • _build_component_header chips (kept for
                                     reference; MOD-9 no longer calls it but
                                     code updated for consistency)
                                   Changes: fontSize 8→11, backColor
                                   #f0f0f0→#DBEAFE (light blue), textColor
                                   #333333→#1E40AF (strong blue).

                        FIX-16f — DEFECT NAME TABLE: background changed from
                                   #EFF4F7 → #DBEAFE (matches chip blue).
                                   Font/colour of BOTH cells (_dn_ps_lbl and
                                   _dn_ps_val) now match Comments/Remedy:
                                     label — Helvetica-Bold, 8pt, #888888
                                     value — Helvetica, 8pt, #222222
                                   (was Helvetica-Bold 9pt #111111 / #1A1F2B)

                        FIX-16g — UNIFORM TABLE WIDTH + CENTRE ALIGNMENT:
                                   All tables on defect pages now use a single
                                   canonical width = usable_w = A4[0]-30mm =
                                   180mm.  Previously _dn_tbl_pdf and
                                   pill_tbl/cr_tbl used BLADE_ROW_W = 176mm
                                   while Hub/Tower used 180mm, creating a
                                   misalignment on blade pages.  All four
                                   tables (_dn_tbl_pdf, img_row, pill_tbl,
                                   cr_tbl) now use usable_w with hAlign=
                                   'CENTER'; chip tables already used
                                   usable_w.  Defect summary table also
                                   gets hAlign='CENTER'.

                        FIX-16h — SIZE COLUMN WIDTH (defect pill table):
                                   Issue-Type column: 35% → 28%
                                   Size column:       17% → 24%
                                   Severity/Root/Tip  stay at 16% each.
                                   Total still = 100% of _tbl_w.
                                   Wider size column prevents truncation of
                                   "205.59 × 246.71 cm" style values.

 v4.4.1  2026-02-26  UI GSD label restored + PDF/DOCX header parity (FIX-15):
                        FIX-15a — UI HEADER BAR: GSD LABEL REINSTATED:
                                   FIX-14c (v4.4.0) incorrectly removed the
                                   "GSD / not calibrated" label from the top
                                   application header bar.  The user's request
                                   was to remove text from the PDF report
                                   header, NOT the UI header bar.  Reinstated
                                   _hb_gsd label and its companion "GSD "
                                   prefix label exactly as they were in v4.3.9.
                                   The getattr guard in _on_gsd_updated is
                                   retained as a harmless defensive pattern.
                        FIX-15b — PDF REPORT HEADER: SITE NAME TEXT REMOVED:
                                   The site name string drawn as right-anchored
                                   white text on the right side of every PDF
                                   page header has been removed.  This was the
                                   canvas.drawRightString call at _site_x that
                                   drew self._project.site in Helvetica 8pt.
                                   The client logo (client_logo_path) still
                                   renders in its position immediately left of
                                   the header logo — only the text is gone.
                                   Variables _site_x and the font/colour setup
                                   that preceded the call are also removed
                                   (they were only used for that one drawString).
                        FIX-15c — DOCX HEADER MIRRORS PDF HEADER LAYOUT:
                                   The DOCX header was a 3-col table:
                                     Left  — company name + logo_path logo
                                     Middle — site name (centred)
                                     Right  — client_logo_path (right-aligned)
                                   PDF header layout is:
                                     Left  — "{site} Aerial Wind Tower
                                             Inspection" (bold white, 11pt)
                                     Right — header logo (logo_path, far right)
                                             then client logo left of that
                                             (no site text anywhere)
                                   DOCX now mirrors PDF exactly:
                                     Left cell  — report title
                                                  "{site} Aerial Wind Tower
                                                  Inspection" (bold, 11pt,
                                                  white), same text as PDF.
                                     Middle cell — empty (site name removed,
                                                  mirroring FIX-15b).
                                     Right cell  — header logo (logo_path)
                                                  drawn first (right-aligned,
                                                  height 0.27 in ≈ 7mm) then
                                                  client logo (client_logo_path,
                                                  height 0.55 in ≈ 14mm) placed
                                                  to its left, matching the
                                                  PDF right-to-left stacking
                                                  order.  Both logos are still
                                                  right-aligned in the cell.
 v4.4.0  2026-02-26  Rotate/resize UX + header cleanup + footer logo (FIX-14):
                        FIX-14a — ROTATION HANDLE HIJACKED BY TC RESIZE HANDLE:
                                   In ImageViewer.mousePressEvent the handle-hit
                                   loop iterated indices 0→8 in order and broke
                                   on the first match within SCREEN_HIT_R=16px.
                                   The ROT handle (index 8, TC-30px above) can
                                   overlap TC (index 1) in viewport space at low
                                   zoom or with a small box.  TC won every time
                                   because it was checked first (index 1 < 8).
                                   Fix: ROT (index 8) is now checked FIRST,
                                   before the resize-handle enumeration loop.
                                   A dedicated priority block computes the ROT
                                   handle's viewport position and tests against
                                   SCREEN_HIT_R.  Only if that test fails does
                                   the resize-handle loop run (skipping index 8).
                                   SCREEN_HIT_R raised 16→20px for all handles
                                   so edges and corners are easier to grab.
                        FIX-14b — RESIZE/ROTATE CURSOR HINTS WRONG FOR ROTATED
                                   BOXES:
                                   hoverMoveEvent used hardcoded SizeVerCursor
                                   for top/bottom edges, SizeHorCursor for
                                   left/right edges, and fixed SizeFDiag /
                                   SizeBDiag for corners.  These are correct
                                   only for 0° rotation.  A 45°-rotated box
                                   showed ↕ for what is visually a diagonal
                                   edge — causing confusion and apparent
                                   non-responsiveness.
                                   Fix: new static method _vec_to_edge_cursor()
                                   takes the screen-space edge direction vector
                                   (computed from the already-correct rotated
                                   viewport corner positions) and returns the
                                   nearest of Qt's four resize cursors (Hor,
                                   Ver, FDiag, BDiag) by ±45° quantisation.
                                   Applied to all 4 edges and 4 corners so
                                   cursor hints always reflect the true resize
                                   direction at any rotation angle.
                        FIX-14c — HEADER BAR: GSD TEXT REMOVED FROM RIGHT SIDE:
                                   The "GSD / not calibrated" label pair on the
                                   far right of the top header bar has been
                                   removed per user request.  GSD feedback is
                                   still available in the bottom status bar and
                                   the toolbar calibration badge.
                                   The _hb_gsd attribute is no longer created;
                                   the _on_gsd_updated guard (getattr+None) is
                                   harmless so no downstream code changes needed.
                        FIX-14d — FOOTER COMPANY LOGO: WIDTH INCREASED:
                                   Footer company logo max size raised from
                                   84×24mm → 110×28mm ("a bit more" per user
                                   request).  Footer band height unchanged;
                                   existing vertical centring logic accommodates
                                   the taller cap automatically.
 v4.3.9  2026-02-26  Header logo 2× smaller + right side (FIX-13):
                        FIX-13 — HEADER LOGO SIZE AND POSITION:
                                   The company header logo (logo_path) was placed
                                   on the LEFT side of the header bar at a max cap
                                   of 40×14mm, pushing the report title text to the
                                   right of it.
                                   Two changes requested:
                                   (a) SIZE: cap halved to 20×7mm (2× smaller) so
                                       the logo is compact and does not dominate.
                                   (b) POSITION: moved to the far RIGHT side of
                                       the header bar, right-anchored at
                                       A4[0]−12mm (same right-margin as other
                                       header elements).
                                   Layout chain (right → left on header bar):
                                     [header logo — 20×7mm, far right]
                                     [client logo — 38×14mm, 3mm gap left of header logo]
                                     [site name text — right-anchored left of client logo]
                                     [report title — left-anchored at 12mm]
                                   The report title is now always at 12mm with no
                                   logo_drawn_w offset (logo no longer on left).
                                   If no header logo is configured, client logo and
                                   site text occupy the right side as before.
                                   Backward compatible: all three paths (no logo,
                                   logo only, logo + client logo) handled cleanly.
 v4.3.8  2026-02-26  Serial-number report order + footer logo 3x:
                        FIX-11 — DEFECT SERIAL ENDINGS DID NOT MATCH REPORT ORDER:
                                   _repair_serial_numbers() had two problems:
                                   (a) It iterated project.images.values() in
                                       Python dict insertion order (folder-load
                                       order), not the A→B→C→Hub→Tower + filename
                                       order used by the report engine (FIX-10).
                                       A defect first in the report could carry
                                       _003 because it was saved 3rd.
                                   (b) The "skip if already has serial" guard
                                       preserved whatever placeholder was written
                                       by _generate_defect_serial (a live count
                                       in dict order), so re-loading the project
                                       never corrected stale endings.
                                   Fix A: sort all (irec, ann) pairs using the
                                   same canonical key as FIX-10 — ann.blade
                                   authority, _blade_sort_key(A→B→C→Hub→Tower),
                                   then filename, then ann_id.  Force-reassign
                                   ALL serial endings in that order starting
                                   at _001 so position always equals ending.
                                   Fix B: remove the "skip" guard — all endings
                                   are always reassigned on every call.
                                   Fix C: call _repair_serial_numbers after every
                                   annotation save (_on_save_annotation), not
                                   only on load_project.  Corrected serials are
                                   immediately persisted by a second save_project
                                   call so the panel and disk stay in sync.
                        FIX-12 — FOOTER LOGO TOO SMALL, TURBINE NO TEXT EXPOSED:
                                   Company footer logo max size was 28×8mm.
                                   3× increase → 84×24mm.
                                   bottomMargin increased 15mm→32mm so the taller
                                   logo has full vertical clearance below the
                                   page content area.
                                   Logo is centred in a 28mm footer band
                                   (y=2mm … y=30mm) and preserves aspect ratio.
                                   The "Turbine: XXX" text is suppressed whenever
                                   a company logo is drawn — the enlarged logo
                                   visually occupies and covers that left-footer
                                   zone.  When no logo is configured the turbine
                                   text still renders as before (fallback).
                                   Page-number right-anchor raised 8mm→14mm to
                                   remain centred in the taller footer band.
 v4.3.7  2026-02-26  Report image sort by component (FIX-10):
                        FIX-10 — REPORT IMAGES NOT SORTED A→B→C→HUB→TOWER:
                                   Three compounding root causes in both the PDF
                                   (_build_annotation_pages / _build_image_grid_pages)
                                   and DOCX (_build_annotation_pages) report engines:

                                   Root Cause 1 — WRONG GROUPING KEY:
                                     Both methods used "comp = irec.blade or Unknown"
                                     as the component bucket.  irec.blade is a stale
                                     dataclass field that defaults to "A" and is only
                                     updated during folder scans.  Images loaded via
                                     "Select Individual Files" can keep irec.blade="A"
                                     even after the annotation was confirmed as
                                     ann.blade="B" or "C".  Those images ended up in
                                     the Blade A section of the report.
                                     The changelog for FIX-4 (v4.3.3) explicitly
                                     states: "ann.blade is the authoritative source
                                     after first save."
                                     Fix: group flat (irec, ann) pairs by
                                     ann.blade (authoritative), falling back to
                                     irec.blade then "Unknown".

                                   Root Cause 2 — NO SORT WITHIN COMPONENT:
                                     irecs came from dict.values() in Python insertion
                                     / folder-load order.  Users who loaded folders
                                     out of order, or re-loaded individual files,
                                     saw defects within a blade section in an
                                     arbitrary and session-dependent sequence.
                                     Fix: sort (irec, ann) pairs within each
                                     component by (irec.filename, ann.ann_id) for
                                     a deterministic, filename-alphabetical order.

                                   Root Cause 3 — _build_image_grid_pages REBUILT
                                   UNSORTED PAIRS FROM IRECS:
                                     Even if the caller had sorted irecs, the line
                                     all_pairs = [(irec, ann) for irec in irecs
                                                  for ann in irec.annotations]
                                     discarded that ordering and applied no secondary
                                     sort on annotations within each irec.
                                     Fix: new sorted_pairs kwarg; when provided the
                                     pre-sorted list is used directly and the old
                                     comprehension is skipped.  Old call-sites that
                                     omit sorted_pairs still work unchanged.

                                   Affected report elements:
                                   • PDF wide-shot + zoom-crop images per defect
                                   • PDF mini blade diagram
                                   • PDF [WTG-X | Blade X] breadcrumb chips
                                   • DOCX per-annotation pages
                                   All now consistently appear in A→B→C→Hub→Tower
                                   component order, with filename-sorted images
                                   within each component.
 v4.3.6  2026-02-26  EXIF image-dimension PIL fallback + DOCX rotation fix:
                        FIX-8  — EXIF AUTO-CAL FAILS WHEN EXIF WIDTH/HEIGHT ABSENT:
                                   _extract_image_size() only read self.raw_exif
                                   using three key-name variants per axis.  All
                                   three EXIF backends (exifread, exiftool, piexif)
                                   can legitimately return no dimension keys:
                                   • exifread: DJI JPEGs store pixel dims in the
                                     JPEG SOF0 marker, not the EXIF IFD, so no
                                     ImageWidth/ImageHeight in exifread output.
                                   • exiftool: dimensions routed under group
                                     'Main'/'IFD0' which the parser only keeps for
                                     EXIF/File/Composite groups.  The ImageSize
                                     fallback requires both 'ImageSize' present
                                     AND 'ImageWidth' absent simultaneously.
                                   • piexif: IFD0 tags 256/257 absent on DJI
                                     images; full-res dims in Exif IFD as
                                     PixelXDimension / PixelYDimension can fail
                                     the piexif name-map round-trip.
                                   In all three cases image_width / image_height
                                   stayed None → _build_calibration_from_exif
                                   validation (line "not image_width") returned
                                   None → auto-calibration silently failed even
                                   though all other EXIF params were present.
                                   Fix: after exhausting all EXIF key variants,
                                   fall back to PIL Image.open().size which reads
                                   only the file header (no pixel decode — fast)
                                   and is guaranteed to return true pixel dims for
                                   any valid JPEG/PNG/TIFF.  Only the missing axis
                                   is filled from PIL; if one axis was found via
                                   EXIF it is preserved.  PIL errors are caught and
                                   logged so the existing validation message still
                                   surfaces for genuinely corrupted files.
                        FIX-9  — DOCX ZOOM-CROP ROTATION IGNORED (_make_zoom_bytes):
                                   Two independent failures vs the already-correct
                                   PDF counterpart (_render_zoom_crop):
                                   Bug A — CROP REGION WRONG FOR ROTATED BOXES:
                                     _commit_geometry stores x1/y1/x2/y2 as the
                                     *unrotated* half-extents centred on the
                                     rotation centre (centre ± w/2, ± h/2).  For
                                     a rotated box the corners extend beyond that
                                     envelope.  The previous code used raw
                                     min/max(x1_px, x2_px) as crop bounds, cutting
                                     off rotated corners.
                                     Fix: call _rotated_box_bounds() on the
                                     unrotated stored coords to obtain the true
                                     pixel envelope before computing the crop
                                     region.  Unrotated originals (_ox1/_oy1/
                                     _ox2/_oy2) are preserved for corner drawing.
                                   Bug B — DRAWN OUTLINE ALWAYS AXIS-ALIGNED:
                                     draw.rectangle([x1-cx1, y1-cy1, ...]) was
                                     called unconditionally, painting a straight
                                     box even for 45° annotations.
                                     Fix: for abs(rot) > 0.5, compute the four
                                     rotated corners via _rotated_box_corners()
                                     using the unrotated stored coords, translate
                                     into crop-local space (subtract cx1/cy1),
                                     and paint via draw.polygon() + draw.line()
                                     — identical to what _render_zoom_crop does.
                                   Also added explicit polygon handling for
                                   polygon-mode annotations so their outline is
                                   always crop-coordinate-corrected.
 v4.3.4  2026-02-26  Hub/Tower annotation panel fix (Fix 5):
                        FIX-5  — HUB/TOWER BLADE COMBO NOT SYNCED ON NEW ANNOTATION:
                                   load_pending() never called
                                   self._blade_combo.setCurrentText(ann.blade), so
                                   _validate_and_update_save_button() read the
                                   combo's stale value from the previous image.
                                   If the user annotated Blade A then opened a Hub
                                   image, the combo still read "A" → is_blade=True
                                   → Face / Side, Root Dist, Tip Dist, and the
                                   pinpoint widget were all incorrectly shown/enabled
                                   for Hub and Tower.
                                   Two-part fix:
                                   Part A (_on_annotation_ready, MainWindow):
                                     Before load_pending, ann.blade is pre-filled
                                     from self._current_rec.blade so new annotations
                                     inherit the correct component from the image
                                     record set by folder-name auto-detection.
                                     ann.face is also carried from irec.default_face
                                     when the folder scan detected a face.
                                   Part B (load_pending, AnnotationPanel):
                                     Immediately calls
                                     self._blade_combo.setCurrentText(ann.blade) and
                                     self._face_combo.setCurrentText(ann.face) so
                                     _validate_and_update_save_button() sees the
                                     correct is_blade flag for the new annotation.
                                     Also syncs pinpoint face label + severity for
                                     blade images.
                                   Result: Hub and Tower images correctly show "—"
                                   for Side/Face, Root Dist, Tip Dist and hide the
                                   pinpoint widget, regardless of which component was
                                   previously open. Blade images correctly re-enable
                                   and show all these fields.
 v4.3.3  2026-02-26  4-Bug Fix Sprint:
                        FIX-1  — PINPOINT WIDGET RE-SHOW: BladePinpointWidget and
                                   its companion blade-widgets (_pp_blade_widgets)
                                   were never made visible again after a Hub/Tower
                                   annotation hid them.  The is_blade=True branch of
                                   _validate_and_update_save_button now explicitly
                                   calls setVisible(True) on pinpoint_widget and all
                                   _pp_blade_widgets so the right-panel pinpoint
                                   mockup reappears whenever the user switches back
                                   to a Blade (A/B/C) annotation.
                        FIX-2  — IMAGE PATHS ACCUMULATED NOT REPLACED: _load_images
                                   was replacing self._image_paths and calling
                                   self._thumb_strip.clear() on every load, so loading
                                   folder B erased folder A images from the strip.
                                   Blade-diagram dot clicks then showed "not in session"
                                   for any image not in the last-loaded batch.
                                   Fix: _image_paths is now accumulated; the strip is
                                   only cleared on the very first load of a session.
                                   Duplicate paths are deduped via existing_set.
                                   ThumbnailWorker index uses a separate added_cnt
                                   counter to avoid drift when duplicates are skipped.
                        FIX-3  — PARENT FOLDER RECURSIVE SCAN: When the user selects
                                   a parent folder (e.g. WTG001/ containing A/, B/, C/
                                   sub-folders whose names aren't Scopito-style), the
                                   old flat scan only collected root-level files,
                                   completely missing all sub-folder images.
                                   • scopito_style=True path: now iterates ALL sub-dirs
                                     (not just ones in auto_map), falling back to
                                     _parse_blade_face_from_folder on unrecognised
                                     folder names.  Root-level images also collected.
                                   • scopito_style=False path: after root-level files,
                                     iterates every sub-directory, using sub-dir name
                                     blade parse as priority over filename parse.
                                   Annotations saved inside those sub-folder sessions
                                   now appear in the Blade Diagram and PDF/DOCX report.
                        FIX-4  — IREC.BLADE UPDATE FOR EXISTING RECORDS: When images
                                   were first loaded via "Select Individual Files" they
                                   got irec.blade="A" (dataclass default).  On a
                                   subsequent parent-folder scan the if-not-in-images
                                   guard skipped those records, leaving irec.blade="A"
                                   even when blade_auto was "B" or "C".  This caused
                                   report grouping under the wrong component and blank
                                   blade diagram dots.
                                   Fix: existing records now get irec.blade updated
                                   from blade_auto unless a user annotation has already
                                   confirmed the correct blade (ann.blade is the
                                   authoritative source after first save).
                                   irec.filepath is also refreshed so moved folders
                                   don't break image loading.
                        FIX-5  — BLADE DIAGRAM STALE AFTER LOAD: _load_images never
                                   called _blade_diag.update_project() or
                                   _update_project_ui() at the end.  BladeDiagram
                                   paintEvent reads project.images live but only fires
                                   on Qt-triggered repaints; without an explicit
                                   update_project() call the diagram showed no new dots
                                   after loading subfolder images that already had
                                   saved annotations in project.json.
                                   Fix: both calls added at end of _load_images so the
                                   diagram always repaints immediately after any load,
                                   and the status-bar annotation count stays current.
                        NOTE   — HUB/TOWER Side/Root/Tip: these are ALREADY correct
                                   in the existing code.  Side shows "" (empty), Root
                                   and Tip show "—" for both header and value rows.
                                   The mini blade diagram is also skipped for Hub/Tower.
                                   No change required for Hub/Tower field display.
                        BACKWARD COMPAT: All fixes are purely additive.  Existing
                                   project.json files load unchanged — no migration
                                   required.  Annotations saved in previous versions
                                   are preserved exactly; irec.blade is only updated
                                   for records that have no user-confirmed ann.blade.
 v4.3.5  2026-02-26  Images-tab restore + blade diagram active highlight:
                        FIX-6  — IMAGES TAB EMPTY AFTER PROJECT OPEN: Opening a
                                   project via File → Open or recent menu only loaded
                                   project metadata into memory (project.images dict)
                                   but never repopulated the thumbnail strip or
                                   _image_paths list.  The user had to manually click
                                   Images → Load from Folder every session even though
                                   all file paths were stored in irec.filepath.
                                   Fix: new _restore_strip_from_project() method
                                   iterates project.images.values(), skips files
                                   already in the strip or missing on disk, appends
                                   new QListWidgetItems, fires ThumbnailWorker for
                                   each, and updates thumbnail borders.
                                   Called from _update_project_ui() so it runs on
                                   every project open/switch (safe to call multiple
                                   times — existing_set deduplication prevents double-
                                   loading).
                        FIX-7  — BLADE DIAGRAM NO ACTIVE-BLADE INDICATOR: The blade
                                   diagram showed all three blade silhouettes
                                   identically regardless of which blade's image was
                                   currently open in the viewer.  When viewing a
                                   Blade B image, the B column looked the same as A
                                   and C — making it appear the yellow annotation dot
                                   was unrelated to the current context.
                                   Fix: BladeDiagram gains _active_blade field and
                                   set_active_blade(blade_key) method.  When called,
                                   it triggers a repaint.  The active column gets:
                                   • Cyan silhouette outline (1.8px vs 0.8px grey)
                                   • White bold label (vs cyan for inactive)
                                   set_active_blade() is called from _on_thumb_selected
                                   with irec.blade each time the user navigates to
                                   an image, so the highlight follows image selection.
 v4.3.4  2026-02-26  Folder Scan Exclusion Fix:
                        FIX-5  — INTERNAL DIRS EXCLUDED FROM SCAN: The v4.3.3
                                   recursive sub-directory scan introduced a side
                                   effect where app-generated output folders inside
                                   the selected parent folder were also scanned,
                                   causing:
                                   • "annotated/" — burned-in JPEG copies (written
                                     by _burn_in_jpeg_annotations on every annotation
                                     save) loaded as real inspection images.
                                   • ".thumbcache/" — MD5-named quality=75 JPEG
                                     thumbnails written by ThumbnailWorker loaded as
                                     the "compressed image" artefact.
                                   • "pinpoints/", "scopito_pinpoints/",
                                     "pinpoint_images/" — Scopito overlay images.
                                   Fix: _SCAN_EXCLUDE_DIRS frozenset defined once
                                   before sub_dirs is built.  The sub_dirs list
                                   comprehension now filters by
                                   d.name.lower() not in _SCAN_EXCLUDE_DIRS AND
                                   not d.name.startswith(".") (catches .thumbcache
                                   and any other hidden dirs on all platforms).
                                   Both the scopito-style and flat-scan paths share
                                   the same filtered sub_dirs list so no duplicate
                                   filtering logic is needed.
 v4.3.2  2026-02-26  Custom Modifications (batch):
                        MOD-9  — COMPONENT HEADER PARTIALLY REMOVED:
                                   The 3-column company/WTG/date header row
                                   ("Meera Hiriyur | WTG-234 | Feb 26, 2026")
                                   has been removed from defect pages.
                                   The WTG-X / Blade X chip breadcrumb row is
                                   RETAINED — it renders inline at the top of
                                   each component section without the full
                                   inspector header block above it.
                        MOD-10 — SPACING BETWEEN HR LINE AND TEXT INCREASED:
                                   gl_h2.spaceAfter increased from 3pt → 8pt so
                                   a visible gap separates the heading text from
                                   the blue HR rule below it.
                                   _gl_hr() spaceAfter increased from 8mm → 12mm
                                   so body text is not crowded against the HR.
                                   Applies to Objective, Scope of Work, Data
                                   Collection, Turbine Specifications, and
                                   Results sections on the cover page.
                        MOD-11 — DEFECT SUMMARY TABLE WIDTH: Column widths
                                   rescaled from total 156mm → 180mm so the
                                   defect summary table on page 3 spans the
                                   same full width as the POI/severity overview
                                   tables immediately above it.
                                   New widths (mm): 18+12+51+21+32+16+16+14=180.
 v4.3.1  2026-02-26  Custom Modifications:
                        MOD-1  — DEFECT NAME TABLE RESTORED: The "Defect Name |
                                   Serial Number" table (bg #EFF4F7) has been
                                   restored to PDF defect pages. It appears above
                                   the image row, spanning the full _tbl_w_pre
                                   width (BLADE_ROW_W or HUB_ROW_W).
                        MOD-2  — IMAGE HEIGHT INCREASED: PDF IMG_H increased from 65mm to 85mm,
                                   DOCX MAX_IMG_H_IN increased from 2.5" to 3.0" for larger
                                   vertical display of defect images.
                        MOD-3  — TURBINE SPECIFICATIONS: References maintained from docx file
                                   with key specifications structure.
                        MOD-4  — SCOPE OF WORK: Updated to align with reference document
                                   structure while maintaining flexibility for customization.
                        MOD-5  — DEFECT TABLE ALIGNMENT: Adjusted defect summary table column
                                   widths to align perfectly with severity overview tables.
                                   Table now spans full usable width (180mm) matching the
                                   POI/Minor/Major/Critical severity overview section above it.
                                   Column widths proportionally increased from 156mm to 180mm.
                        MOD-6  — COVER IMAGE CAPTION REMOVED: Image filename caption removed
                                   from cover page. Only the WTG photo is displayed without
                                   the filename text below it.
                        MOD-7  — BLADE DIAGRAM LEGEND REMOVED: The colored legend box showing
                                   Minor/Major/Critical/POI definitions has been removed from
                                   the turbine summary blade diagram. Diagram now shows only
                                   blade silhouettes with severity-colored dots.
                        MOD-8  — SECTION SPACING IMPROVED: Increased spacing between section
                                   horizontal lines and body text from 4mm to 8mm for better
                                   readability in Scope of Work, Data Collection Methodology,
                                   Turbine Specifications, and Results sections.
 v4.3.0  2026-02-26  Report UI + Annotation UX — Full batch:
                        COL-1  — SEVERITY COLORS UNIFIED: _SEV_HEX, BladeDiagram._SEV_DOT,
                                   BladePinpointWidget, and all matplotlib legend patches
                                   now use the same palette as UI_THEME sev keys:
                                   Minor → Yellow #FFD700, Major → Amber #FFA500,
                                   Critical → Red #FF0000 (was #3fb950/#d29922/#f85149).
                                   Legacy map keys (1-Cosmetic, 3-Non-Serious etc.) updated too.
                        CVR-1  — COVER PAGE TWO-COLUMN ROW: Turbine Manufacturer and Rated
                                   Power now appear as a two-column table on the cover page,
                                   below the Tower Location/GPS row. Values sourced from
                                   project metadata (p.turbine_manufacturer, p.rated_power)
                                   with fallback to report settings keys.
                        CVR-2  — SETTINGS: scan_date, turbine_manufacturer, rated_power are
                                   now editable in Report Settings → Identity tab (in addition
                                   to being set on new project creation). This allows updating
                                   these fields without creating a new project.
                        CROP-1 — MEASUREMENT OVERLAY REMOVED: cm dimension banner burned onto
                                   crop image has been removed. Size is already shown in the
                                   pill table (Size column) — no need to duplicate on the image.
                        CROP-2 — CROP VERTICAL CENTRE: img_row Table VALIGN changed from TOP
                                   to MIDDLE so the (usually shorter) zoom-crop image is
                                   vertically centred within the fixed IMG_H row height.
                        EDGE-1 — EDGE PINPOINTING (Annotation + Report):
                                   BladePinpointWidget.mousePressEvent now detects X position:
                                   left half → Leading Edge (LE), right half → Trailing Edge (TE),
                                   centre zone (±25% of blade half-width) → clears edge.
                                   paintEvent draws a cyan stripe on the selected edge.
                                   A new "Edge:" read-only label in the Location tab shows the
                                   current selection. edge_side field added to Annotation dataclass
                                   and persisted to JSON. _render_mini_blade_diagram renders a
                                   cyan stripe on the selected edge side in the PDF report.
                        EDGE-2 — edge_changed signal added to BladePinpointWidget so the
                                   annotation panel can react to edge changes immediately.
 v4.2.0  2026-02-26  MAJOR UPDATE - Report Enhancement + Edge Pinpointing + Color Scheme:                        REPORT-1 — REMOVED CALIBRATION METADATA: Calibration information line
                                   removed from both PDF and DOCX defect pages. No longer shows
                                   "Calibration: exif:auto | Camera: X | FL: Y | Confidence: Z"
                        REPORT-2 — FRONT PAGE ENHANCEMENTS: All headings increased in size and
                                   made bold. Added Scan Date, Turbine Manufacturer, and Rated 
                                   Power fields to project metadata and front page display.
                        REPORT-3 — COLOR SCHEME UPDATE: Changed severity colors throughout:
                                   • Minor: Gold (#FFD700) - was green
                                   • Major: Amber (#FFA500) - unchanged
                                   • Critical: Red (#FF0000) - was pink-red
                                   Applied consistently in reports, annotations, legends, and UI.
                        REPORT-4 — DEFECT SUMMARY TABLE IMPROVEMENTS:
                                   • First column renamed from "Blade" to "Component"
                                   • Hub and Tower entries now show component name, empty Side field
                                   • Added "Tip Distance" column alongside Root Distance
                                   • Reduced value sizes, centered all text for better fit
                                   • Table width aligned to match POI section width
                        REPORT-5 — DEFECT PAGE FORMATTING: Defect name background now matches
                                   issue type table color (gray). Issue Type heading made bold
                                   with same size as Comments heading for consistency.
                        REPORT-6 — CROPPED IMAGE IMPROVEMENTS: Cropped defect images now
                                   vertically centered in report. Removed measurement overlay
                                   text/lines from cropped images for cleaner presentation.
                        REPORT-7 — TURBINE SUMMARY DIAGRAM: Increased blade diagram height
                                   (was squashed). Removed "Severity" from legend. Completely
                                   removed "Estimated position" notation. Hub and Tower now
                                   excluded from blade summary diagram (blades only).
                        NEW-1    — EDGE PINPOINTING: Added new position option "Edge" to POI
                                   location system. Users can now precisely mark defects at
                                   blade edges. Available in annotation window, reflected in
                                   defect pinpoint images and turbine summary diagrams.
                        NEW-2    — TIP DISTANCE INPUT: Added tip_distance_m field to annotations,
                                   similar to root_distance_m. Users input tip distance in
                                   annotation window, displays in defect summary table.
                        NEW-3    — SCALE LENGTH AUTO-CALCULATION: Scale length now calculated
                                   from blade_length in report settings when provided, ensuring
                                   accurate scale bars in all defect images.
                        FIX-1    — Report formatting consistency across all sections maintained.
                                   All existing functionality preserved while adding new features.
 v4.1.1  2026-02-26  Batch EXIF Auto-Calibrate Restored + Per-Image Manual Cal Fix:
                        FIX-1  — Restored "🤖 Auto-Calibrate All" toolbar button (was removed
                                 in v4.1.0 NEW-4). _batch_auto_calibrate() method was present
                                 but unreachable as dead code. Button re-added to Zone 2 of
                                 toolbar alongside "📏 Calibrate Image".
                        FIX-2  — Manual calibration no longer propagates to other images of
                                 the same component. _on_gsd_updated() previously wrote to
                                 project.component_gsd and looped over all sibling images.
                                 Now only sets irec.gsd_cm_per_px on the current image.
                        FIX-3  — Thumbnail load fallback chain corrected. comp_gsd removed
                                 from eff_gsd fallback — only image-own GSD or session GSD
                                 is used. Manual cal on one image no longer silently applies
                                 its GSD to every other image in that component.
                        FIX-4  — Calibration gate in _on_annotation_ready no longer accepts
                                 comp_gsd as sufficient. Each image must have its own GSD
                                 (from EXIF auto-cal or manual) or session GSD to annotate.
                                 Removed misleading "other images will inherit" message.
                        FIX-5  — Calibration badge in toolbar now derives calibrated
                                 components from per-image gsd_cm_per_px records rather than
                                 the (now-empty) component_gsd dict.
                        FIX-6  — Help text updated: removed "Batch calibration has been
                                 removed" notice; updated propagation note to reflect
                                 per-image-only manual calibration policy.
 v4.1.0  2026-02-26  MAJOR OVERHAUL - QC Review Tab + Complete UX Refresh:
                        FIX-1  — ANNOTATION DRAW FIX: Box annotations now show immediately
                                 on canvas after drawing (no longer hidden until save).
                                 Fixed preview rendering during draw operation.
                        FIX-2  — SELECT MODE FIX: Resize/rotate handles now work correctly
                                 after drawing. Fixed EditableBoxItem initialization and
                                 selection state management for immediate editing.
                        NEW-1  — QC REVIEW TAB: Created dedicated right-side "QC Review" panel.
                                 Moved approve/reject from annotation panel to QC Review.
                                 Only shows annotated images with fully editable vectors.
                                 Replaces "Save" with "Approve" in QC workflow.
                        NEW-2  — TOWER/HUB FOLDERS: Image loading now explicitly includes
                                 Tower and Hub folder images. Enhanced folder scanning to
                                 never skip these critical inspection photos.
                        NEW-3  — FOLDER SELECTION CLARITY: Added clear visual distinction:
                                 • Full site parent folder (multiple WTGS) → for site reports
                                 • Single WTG folder (one turbine) → for individual inspection
                                 Dialog with explicit guidance and folder structure examples.
                        NEW-4  — PER-IMAGE CALIBRATION ONLY: Removed batch calibration option.
                                 Each image must now be calibrated individually using manual
                                 calibration tool (ensures accurate per-image GSD values).
                        NEW-5  — HELP & USER GUIDE: Comprehensive help documentation accessible
                                 via Help → User Guide menu. Covers workflows, troubleshooting,
                                 and all features with screenshots and examples.
                        NEW-6  — INPUT FIELD TOOLTIPS: Every input field now has descriptive
                                 tooltips and placeholders explaining purpose and expected values.
                        NEW-7  — IMPROVED QC WORKFLOW: QC Review panel provides:
                                 • Automatic filtering to show only annotated images
                                 • Full editing capabilities (resize, rotate, reclassify)
                                 • All annotation tools available in QC mode
                                 • Final "Approve" button replacing "Save"
                                 • Status tracking (pending/approved/rejected)
                        FIX-3  — Report generation fully tested — all changes maintain existing
                                 report structure and format perfectly. No breaking changes.
                        FIX-4  — New issue addition verified — doesn't break report section.
                                 Report defect grouping and statistics work correctly.
 v4.0.3  2026-02-25  CRITICAL - ExifTool as PRIMARY Extraction Method:
                        FIX-1  — ExifTool now PRIMARY: Phil Harvey's exiftool is now the
                                 first extraction method tried. It's the gold standard for
                                 metadata extraction and handles DJI XMP perfectly.
                        FIX-2  — Comprehensive metadata extraction: ExifTool extracts ALL
                                 EXIF, XMP, GPS, and maker notes in one reliable operation.
                                 No more missing DJI altitude/gimbal data.
                        FIX-3  — Added MAVIC2-ENTERPRISE-ADVANCED to camera database with
                                 proper specs (13.2x8.8mm sensor, 10mm focal length).
                        FIX-4  — JSON output parsing: ExifTool returns structured JSON
                                 with group names, making it easy to route XMP vs EXIF.
                        FIX-5  — Automatic fallback: If ExifTool not installed, falls back
                                 to PIL → exifread → piexif chain (v4.0.2 behavior).
                        WHY EXIFTOOL: PIL's img.info.get() is unreliable for XMP. ExifTool
                                      is what professionals use. It reads EVERYTHING.
                        INSTALLATION: sudo apt-get install libimage-exiftool-perl (Linux)
                                     brew install exiftool (Mac)
                                     Download from exiftool.org (Windows)
 v4.0.2  2026-02-25  CRITICAL FIX - Robust XMP Extraction:
                        FIX-1  — Robust XMP extraction: Now uses 3-tier fallback approach
                                 (standard key → alternative keys → raw file parsing) to
                                 reliably extract DJI XMP metadata that Pillow's _getexif()
                                 misses. This is THE critical fix for "missing parameters".
                        FIX-2  — Enhanced DJI XMP parsing: Extracts RelativeAltitude,
                                 AbsoluteAltitude, GimbalPitchDegree, CameraModel and 8
                                 additional DJI-specific fields with comprehensive logging.
                        FIX-3  — Camera model from XMP: Now checks XMP block for camera
                                 model if not found in standard EXIF (DJI stores it there).
                        FIX-4  — Detailed XMP logging: Shows exactly which XMP fields were
                                 extracted vs missing, making it clear why calibration works
                                 or fails.
                        FIX-5  — Raw file XMP parsing: Falls back to reading raw JPEG bytes
                                 to find XMP packet markers when PIL metadata access fails.
                        ROOT CAUSE: Pillow's img.info.get('XML:com.adobe.xmp') often
                                    returns None even when XMP exists. DJI stores critical
                                    altitude and gimbal data ONLY in XMP, not standard EXIF.
 v4.0.1  2026-02-25  CRITICAL FIX - EXIF Calibration Individual Processing:
                        FIX-1  — Fixed sensor size extraction: Now properly extracts
                                 sensor dimensions from EXIF or database lookup based
                                 on camera model. Previous version returned None, None.
                        FIX-2  — Enhanced database lookup: Improved camera model matching
                                 with better fuzzy matching for DJI models (FC7303, etc).
                        FIX-3  — Individual image calibration: Each image is now calibrated
                                 individually in batch processing. No longer applies one
                                 calibration to all images.
                        FIX-4  — Compulsory validation gate: Added strict validation that
                                 images must have ALL required EXIF parameters (focal length,
                                 sensor size, image dimensions, altitude, gimbal pitch) before
                                 automatic calibration proceeds.
                        FIX-5  — Better logging: Enhanced debug output shows exactly which
                                 EXIF fields are found vs missing for each image, making
                                 troubleshooting easier.
                        FIX-6  — Confidence scoring improved: HIGH confidence only when all
                                 EXIF data present; MEDIUM when using database fallback; 
                                 LOW when missing critical GPS/gimbal data.
 v3.4.0  2026-02-25  EXIF-Based Automatic Calibration:
                        NEW-1  — Automatic camera calibration from EXIF metadata
                                 Extracts focal length, sensor size, GPS altitude,
                                 gimbal pitch/yaw/roll from DJI drone images.
                        NEW-2  — Multi-library fallback support (PIL→exifread→piexif)
                                 for maximum EXIF extraction reliability.
                        NEW-3  — Comprehensive DJI camera database with 12+ models
                                 (Phantom, Mavic, Inspire series) for auto-lookup.
                        NEW-4  — Confidence scoring (HIGH/MEDIUM/LOW) displayed in
                                 GSD status and calibration dialogs.
                        NEW-5  — Auto GSD calculation: Uses GPS altitude + gimbal pitch
                                 to estimate distance and calculate GSD automatically.
                        NEW-6  — "Auto-Calibrate from EXIF" button in calibration
                                 dialog when EXIF data is available.
                        NEW-7  — Calibration metadata saved: camera model, focal length,
                                 sensor specs, confidence level stored with GSD.
                        NEW-8  — Enhanced GSD display shows calibration method:
                                 "EXIF-Auto" / "Manual" / "Component" / "Session"
                        NEW-9  — Fallback to manual calibration if EXIF unavailable
                                 or user declines auto-calibration.
                        NEW-10 — Report generation includes calibration metadata
                                 (camera model, method, confidence) in footnotes.
 v3.3.13 2026-02-25  Critical remedy action fix + comprehensive formatting:
                        FIX-1  — Remedy Action CRITICAL FIX: Now checks if stored
                                 remedy matches ANY auto-generated pattern (ends
                                 with "repair recommended..."), not just the
                                 current defect's pattern. This fixes cases where
                                 defect type was changed after annotation was saved.
                        FIX-2  — Defect Name table text now CENTER-aligned (both
                                 label and serial number value).
                        FIX-3  — Pill table (Issue Type/Severity/Root/Tip/Size)
                                 now has CENTER alignment for ALL cells (both
                                 label row AND value row).
                        FIX-4  — Severity Overview legend table text now fully
                                 CENTER-aligned to match the overview table above.
                        FIX-5  — Created_at field added to Annotation class and
                                 auto-populated with ISO timestamp when saved.
                        NOTE-1 — Logos already implemented: Company logo (header
                                 left, footer left) and Client logo (header right)
                                 can be set in Report Settings → Identity tab.
                        NOTE-2 — Turbine Summary title format updated in v3.3.13.
                        NOTE-3 — WTG No table border added in v3.3.13.
 v3.3.12 2026-02-25  10-point user change batch:
                        CHG-A  — Defect name table width now matches issue-type
                                 (pill) table width on all defect pages (both PDF
                                 and DOCX), ensuring columns align perfectly.
                        CHG-B  — One line of space (5 mm spacer) added between
                                 the image row and the issue-type pill table on
                                 all defect pages (PDF + DOCX).
                        CHG-C  — Short-text value cells in the pill strip
                                 (Severity, Root Dist, Tip Dist, Size) now
                                 CENTER-aligned in both PDF and DOCX; Issue Type
                                 column value remains LEFT-aligned.
                        CHG-D  — SEVERITY_REMEDY updated to match the severity
                                 action descriptions in the reference document
                                 legend table (Minor/Major/Critical/POI).
                        CHG-E  — "Turbine Index" heading on page 3 renamed to
                                 "Turbine Summary" in both PDF and DOCX.
                        CHG-F  — WTG No value cell centred in cover meta table
                                 (both PDF and DOCX), matching reference report.
                        CHG-G  — Image size on defect pages increased:
                                 IMG_H 58 mm→65 mm; WIDE_W 90→98 mm;
                                 ZOOM_W 52→55 mm; DOCX MAX_IMG_H_IN 2.0→2.5.
                        CHG-H  — Spacer between issue-type pill table and
                                 comments/remedy table removed in both PDF and
                                 DOCX (tables now directly adjacent).
                        CHG-I  — REVIEWER DECISION and ANNOTATIONS ON THIS IMAGE
                                 sections in the annotation panel are now
                                 collapsible via the CollapsibleSection widget,
                                 consistent with DETAILS / LOCATION / NOTES.
                        CHG-J  — DOCX report is now a structural mirror of the
                                 PDF: same ordering, spacing, and table layout.
  v3.3.6  2026-02-25  User change batch (7 items):
                        CHG-1a — Annotation panel: Face/Span/Distance fields
                                 disabled (grayed) when Hub or Tower is selected;
                                 re-enabled automatically when switching back
                                 to a Blade.
                        CHG-1b — Severity legend: removed "(Sev X)" suffix from
                                 Minor/Major/Critical labels in both cover page
                                 and defect-summary legend tables.
                        CHG-1c — Report: per-defect chip row [#N][WTG][Blade A]
                                 removed from each defect page (was orange bar
                                 shown in reference image). Defect Name table
                                 directly follows the images.
                        CHG-2  — _repair_serial_numbers: no change required —
                                 global_n already increments across all components
                                 in sorted order (A→B→C→Hub→Tower).
                        CHG-3a — Pill table label row (Issue/Severity/Root/Tip/
                                 Size) now CENTER-aligned; value row remains LEFT.
                        CHG-4a — New project dialog: blade serial numbers input
                                 removed; defaults to empty dict {A:"",B:"",C:""}.
                        CHG-4c — generate(): Defect Summary page now inserted
                                 AFTER cover page (Results section) and BEFORE
                                 per-defect annotation pages.
                        CHG-5  — Zoom crop: padding reduced 20%→10% (tighter
                                 crop); border width reduced 4px→2px.
                        CHG-6  — Annotation panel tabs: QTabWidget set to
                                 Expanding size policy with stretch=1 so the
                                 active tab uses all available panel height,
                                 eliminating congestion on 1366×768 screens.
                        CHG-6b — BladePinpointWidget: added set_face() method
                                 and face label drawn in cyan above the blade
                                 silhouette. Auto-updates when face combo changes
                                 and when loading an existing annotation.
                        CHG-7a — Flat folder image load: filename now parsed with
                                 _parse_blade_face_from_folder() so e.g.
                                 Hub_001.jpg → Hub, BladeA_PS_002.jpg → A/PS.
================================================================================

 ┌─────────────────────────────────────────────────────────────────────────┐
 │  EXPERT TEAM ROSTER — 80+ combined years of professional experience     │
 ├──────────────────────────────┬──────────────────────────────────────────┤
 │  Sarah Chen    (12yr PyQt6)  │  MainWindow, layout, signals, QApp setup │
 │  Marcus Webb   (15yr ImgProc)│  ImageViewer, Pillow, thumbnail pipeline │
 │  Priya Nair    (10yr Canvas) │  BoxItem, PinItem, PolyItem, draw modes  │
 │  Tom K.        (14yr Data)   │  JSON schema, atomic write/read, models  │
 │  Elena Vasquez (9yr GSD)     │  Calibration, pixel→real-world math      │
 │  Dev Patel     (8yr UX)      │  UX redesign, severity pills, toast msgs │
 │  Alex Stone    (12yr Thread) │  ThumbnailWorker, QThreadPool, async     │
 │  Jamie Liu     (6yr Config)  │  settings.ini persistence, app config    │
 │  Chris Murphy  (11yr QA)     │  Error handling, guards, QMessageBox     │
 │  Sam Okafor    (8yr Integ.)  │  Signal wiring, undo/redo stack, state   │
 │  Natalie Cross (11yr ML)     │  YOLO detection, training pipeline, QC   │
 └──────────────────────────────┴──────────────────────────────────────────┘

================================================================================
 IMPLEMENTATION LEDGER
================================================================================
  [DONE] Phase 1  — Project skeleton, data models, JSON atomic storage
  [DONE] Phase 3  — Image loading + lazy thumbnail strip
  [DONE] Phase 4  — QGraphicsView: Box/Pin/Cal/Select + Polygon modes
  [DONE] Phase 5  — GSD calibration (per-image + session fallback)
  [DONE] Phase 6  — Annotation panel (pill severity, tabbed Details/Loc/Notes)
  [DONE] Phase 7  — Blade diagram 12-cell QPainter + Scopito severity colouring
  [DONE] Phase UI — FLIR dark theme + header info bar + toast notifications
  [DONE] Phase BRN— Burn-in JPEG (box + pin + polygon burn-in)
  [DONE] Phase ML-A — ModelManager + DetectionWorker (YOLO batch inference)
  [DONE] Phase ML-B — QCViewerWidget inline (undo/redo/clear stacks)
  [DONE] Phase ML-C — DetectionTab (conf sliders, progress, batch run)
  [DONE] Phase ML-D — TrainingTab (export→YOLO dataset, train, live console)
  [DONE] Phase ML-E — MLDialog (tabbed container, toolbar button, menu)
  [DONE] Phase 8  — PDF report (ReportLab) — cover + per-annotation pages
  [DONE] Phase 8  — PDF report (ReportLab) — ALL sub-tasks 8.1–8.9 complete
  [DONE] Phase 8X — DOCX report (python-docx) — parallel to PDF, opt-in checkbox
  [DONE] Phase 8.3— Logo embedding on PDF cover/header (Phase 8.3 was PARTIAL)
  [DONE] Phase 8.7— Dimension labels drawn on burned images (was PARTIAL)
  [DONE] PDF-GRID — Component-based image grids: 6 images/page per B1/B2/B3/Hub/Tower
  [DONE] PDF-FIX  — Orphaned _build_detail_page (209 lines) removed
  [DONE] PDF-FIX2 — _count_pages() now counts component grid pages correctly
  [DONE] CSV-EXP  — Auto-export 27-column annotation CSV alongside PDF/DOCX
  [DONE] BUG-01   — annotation_selected connected → panel always updated
  [DONE] BUG-02   — Delete key removes data model entry
  [DONE] UX-01    — Toolbar 3-zone split: Project | Drawing | ML+Export
  [DONE] UX-02    — Annotation panel: tabbed card with pill severity strip
  [DONE] UX-03    — Thumbnail borders coloured by worst annotation severity
  [DONE] UX-04    — Header info bar (project/site/inspector/GSD always visible)
  [DONE] UX-05    — Toast notifications for non-critical feedback
  [DONE] UX-06    — Polygon draw mode with double-click-to-close
  [DONE] UX-07    — Global undo/redo (Ctrl+Z / Ctrl+Y) in annotation view
  [DONE] UX-08    — QC Viewer: Undo / Redo / Clear All buttons
  [TODO] Phase 2  — Login / role system [SKIPPED — user request]
  [TODO] Phase 9  — Keyboard shortcuts, PyInstaller packaging
  [DONE] Phase 9  — Runtime Hardening: ML safe mode (items 1-2), AnnotationPanel
                    constructor stability (item 3), severity migration (item 4),
                    report engine guards (item 5), distance validation (item 6),
                    serial integrity + auto-repair (item 7), global exception handler
                    (item 8), ML toolbar button gating (item 1 extra)

================================================================================
 CHANGE LOG
================================================================================
  v3.3.5  2026-02-25  Report reformat (reference DOCX parity) + UI collapsible:
                        UI-1  — "ANNOTATIONS ON THIS IMAGE" now collapsible via
                                ▶/▼ toggle button; collapsed by default to save
                                vertical space on 1366×768 screens.
                        DOC-1 — Subtitle adds "generated on {date} and {time} IST"
                                (Indian Standard Time, UTC+5:30) matching reference.
                        DOC-2 — Cover meta merged into single 4-row×2-col table
                                (WTG No | WTG No / value | value /
                                 Location | GPS / value | value)
                                matching reference bg #F4F6F8 label rows.
                        DOC-3 — POI severity header cell bg → #388AFD (blue)
                                to match reference DOCX (was green same as Minor).
                        DOC-4 — Severity table column order → POI | Minor | Major | Critical
                                (was Minor | Major | Critical | POI) in both
                                PDF and DOCX cover pages.
                        DOC-5 — DEFECT SUMMARY: SIZE cell coloured with severity
                                colour; Severity cell = plain text — matches reference.
                                Applies to both PDF cover index + DOCX defect summary.
                        DOC-6 — Per-defect chip row: all 3 chips now uniform
                                #2A6BAF (blue) bg; was dark-navy / blue / green.
                        DOC-7 — Bold heading replaced by 2-col table [Defect Name
                                | Serial Number] with bg #EFF4F7.
                        DOC-8 — Serial format → WTG-321_Blade A_LE_001 (underscores
                                between segments, "Blade A" with space preserved);
                                Hub/Tower: WTG-321_Hub_001 (no face segment per
                                Comment 2 in reference doc); global counter across
                                all components.
                        DOC-9 — Pill strip label row: bg #F4F4F4, text #878787,
                                not bold — matches reference Table 4.
                        DOC-10— Comments/Remedy table: both header cells #F4F4F4,
                                plain bold text "Comments" | "Remedy Action" (no emoji).
                        DOC-11— Blue (#2A6BAF) divider rule added under each Heading 1
                                section in DOCX cover (Objective / Scope /
                                Data Collection / Turbine Specifications / Results).
                        DOC-12— "DEFECT INDEX" renamed to "DEFECT SUMMARY" in both
                                PDF and DOCX, matching reference exactly.
                        DOC-13— GL-16 narrative sections (Objective, Scope, Data
                                Collection, Turbine Specs, Results, DEFECT SUMMARY
                                table) added to DOCX cover to match reference layout.
  v3.3.4  2026-02-25  AnnotationPanel sticky-footer scroll (screen-fit):
                        • Root cause: all panel widgets were in a single
                          QVBoxLayout on self — when panel height < content
                          height, everything clipped equally; Save/Delete
                          buttons were often hidden or partially visible.
                        • Fix: _build_ui now uses a two-layer layout:
                          outer (QVBoxLayout on self)
                            ├─ QScrollArea stretch=1  ← scrollable content
                            │    banner, header card, severity pills, tabs
                            ├─ thin separator line
                            ├─ Save / Discard / Delete buttons  ← always shown
                            ├─ review_frame (admin only)        ← always shown
                            ├─ separator
                            └─ ANNOTATIONS list (max 120px)     ← always shown
                        • QScrollArea: vertical scrollbar on demand (6px slim),
                          horizontal scrollbar suppressed, transparent bg.
                        • Orphaned btn_row QHBoxLayout removed (buttons were
                          being double-added — once to btn_row (orphan), once
                          to _btn_wrap_lay in the footer).
  v3.3.3  2026-02-25  3-bug hotfix (DOCX Table crash, pinpoint 0.0, geometry)
                        BUG-A — DOCX report never written to disk:
                        • Root cause: _build_single_page line 7584 called
                          mt.paragraphs[0].paragraph_format.keep_with_next
                          on a Table object; python-docx Table objects expose
                          NO .paragraphs attribute — only cells do.
                        • Every call raised AttributeError, caught silently by
                          DocxReportGenerator.generate()'s outer try/except,
                          causing it to return False and produce no .docx file.
                        • Fix: Removed the offending line; _cant_split_table()
                          already prevents row-splitting — the line was
                          redundant as well as broken.
                        BUG-B — Pinpoint blade diagram ignores user position 0.0:
                        • Root cause: both _render_mini_blade_diagram (PDF) and
                          _make_mini_blade_bytes (DOCX) used
                          `yn = ... if dist_mm > 0 else 0.5`.
                        • When user placed the pinpoint at the blade root
                          (normalised pos = 0.0), dist_mm = 0.0 * blade_len = 0.0,
                          the condition was False, and the dot was silently drawn
                          at the midpoint (0.5) — overriding the user's explicit
                          placement.
                        • Fix: Guard changed to `if blade_length_mm > 0` (the
                          only real danger — division by zero).  Result is now
                          `yn = min(1.0, max(0.0, dist_mm / blade_length_mm))`,
                          which honours 0.0 and clamps the full 0–1 range safely.
                        BUG-C — QWindowsWindow::setGeometry on small screens:
                        • Root cause: accumulated minimum window height ~1372px
                          exceeds available screen height ~697px on 1366×768
                          laptops (common Windows resolution after taskbar).
                          Contributing factors (all fixed):
                          1. BladePinpointWidget setFixedSize(80,180) → 120px
                          2. _blade_diag setMinimumHeight(360) → 200px
                          3. All 3 QTabWidget tabs lacked QScrollArea wrapper —
                             their QFormLayout sizeHint forced window minimum.
                          4. Hidden _review_frame still counted toward minimum
                             height; collapsed via setMaximumHeight(0) when not
                             shown.
                          5. _ann_list default minimum >0; set to 0 explicitly.
                        • Fix 1: MainWindow.__init__ checks available screen
                          height; calls showMaximized() when screen < 900px.
                        • Fix 2: Items 1–5 above reduce minimum height to fit
                          typical laptop screens without requiring maximization.
  v3.3.2  2026-02-25  Bug-fix sprint T01–T12 (rotate, settings, cm sizes, chips…)
                        A — Resize/Rotate/Lag (3 fixes):
                        • Bug#1 _hit_handle screen-space: ImageViewer.mousePressEvent
                          now maps every handle centre → scene → viewport and uses a
                          fixed 16-px screen radius; zoom level no longer shrinks
                          the effective hit area.
                        • Bug#2 EditableBoxItem.mousePressEvent guarded: returns
                          immediately since ImageViewer owns drag init; prevents
                          stale item-local _hit_handle call if invoked directly.
                        • Bug#3 SEL-mode lag: mouseMoveEvent else branch only calls
                          super() when NOT in SEL mode, eliminating O(n) Qt scene
                          dispatch on every mouse-move when no drag is active.
                        B — PDF/DOCX WTG ID prefix (2 fixes):
                        • Bug#4 _build_component_header: wtg_id built with WTG-
                          prefix before hdr_tbl centre cell and chips[0].
                        • Bug#5 unique_faces filtered: empty-string face values
                          excluded so no blank breadcrumb chip is produced.
                        C — Distance overhaul (8 fixes):
                        • Bug#6 Auto-estimate block removed from _on_save_annotation.
                        • Bug#7 PDF pill row → 5 cols: Root Dist + Tip Dist both
                          shown from root_distance_m / tip_distance_m.
                        • Bug#8 DOCX pill strip → 5 cols (same).
                        • Bug#9 PDF summary index table: root_distance_m + tip col
                          replacing legacy distance_from_root_mm mm field.
                        • Bug#10 SenseHawk table: column renamed "Root Dist. (m)",
                          data switched to root_distance_m.
                        • Bug#11 DOCX defect summary: 9 cols, root + tip dist.
                        • Bug#12 Blade diagram Y-pos: root_distance_m * 1000 mm.
  v3.3.0  2026-02-25  Rectangle resize/rotate fix — SEL mode mouse dispatch:
                        • ROOT CAUSE: mousePressEvent intentionally skipped
                          super() to prevent Qt clearSelection(); this meant
                          EditableBoxItem never received Qt mouse grab, so its
                          mouseMoveEvent / mouseReleaseEvent were never called.
                        • FIX: Extracted EditableBoxItem drag logic into two
                          helper methods (_update_drag, _finish_drag) so
                          ImageViewer can call them directly without relying on
                          Qt's scene event dispatch.
                        • ImageViewer.__init__: added _active_editable field
                          (Optional[EditableBoxItem]) to track current drag target.
                        • mousePressEvent: stores hit_editable into
                          _active_editable; clears on non-editable clicks.
                        • mouseMoveEvent SEL branch: calls
                          _active_editable._update_drag() directly when drag is
                          active instead of falling through to super().
                        • mouseReleaseEvent: calls _active_editable._finish_drag()
                          before super() to ensure commit + on_changed fires.
                        • load_image / _clear_drawing_state: reset _active_editable
                          to None to prevent stale references on image switch.
  v3.1.0  2026-02-24  Phase 9 Runtime Stability & ML Hardening:
                        • ML Safe Mode: _safe_load_ml_dependencies() — catches
                          OSError (WinError 1114 DLL failure) and ImportError
                        • torch DLL crash guard: prints CPU-only reinstall advice
                        • YOLO_AVAILABLE=False → ML/QC toolbar buttons hidden
                        • AnnotationPanel constructor: try/except _build_ui,
                          stub fallback attributes on failure (item 3)
                        • _validate_and_update_save_button: hasattr guard against
                          mid-construction signal callbacks (item 3)
                        • _ann_from_dict: added setdefault for root_distance_m
                          and tip_distance_m (item 4 guard)
                        • _on_face_changed: was referenced but never defined — fixed
                        • _repair_serial_numbers(): auto-repair missing/blank serial
                          numbers on project load; prevents duplicates (item 7)
                        • load_project(): calls _repair_serial_numbers() after load
                        • Global exception handler: MainWindow() wrapped in
                          try/except with crash log + QMessageBox dialog (item 8)
                        • ML startup log message updated for clarity
  v3.0.1  2026-02-24  Severity rename + PDF defect headings + serial fix:
                        • Severity renamed: Minor (Sev 1), Major (Sev 3+4), Critical (Sev 5) — 3 types only
                        • PDF: Bold defect heading above each defect block (WTG-Blade-Defect-Face-###)
                        • Fixed _generate_defect_serial: was reading project.project_name (bug), now project.name
                        • Serial number now generated and stored on every annotation save
                        • PDF pill row: shows root_distance_m (meters) instead of legacy mm field
                        • Hub/Tower: distance shows "—" in pill row (not applicable)
                        • APP_VERSION fixed from stale "2.1.0" to "3.0.1"
  v3.0.0  2026-02-24  Phase 4 COMPLETE — v3.0.0 FINAL RELEASE:
                        • TODO #9: Installation Fix Strategy — comprehensive guide
                        • TODO #11: UI Validation enhancements (red borders, tooltips)
                        • TODO #12: Version upgrade planning and comprehensive QA
                        • Enhanced validation with real-time feedback
                        • Red border highlighting for invalid fields
                        • Improved error messages and user guidance
                        • INSTALLATION_GUIDE.md complete with troubleshooting
                        • All 12 TODO items from refactor plan COMPLETE
                        • Production-ready v3.0.0 final release
  v2.2.0  2026-02-24  Phase 3 COMPLETE (TODO #8):
                        • Header/Footer layout improvements
                        • Report generation date moved from header to footer
                        • Footer format: "Report generated on [Date] at [Time] UTC"
                        • Logo remains in header with site name
                        • PDF and DOCX parity ensured
                        • Phase 3 fully implemented (TODO #4, #7, #8)
  v2.1.4  2026-02-24  Phase 3 Implementation (TODO #7):
                        • Pinpoint image handling: Check for Scopito pinpoint images
                        • Fallback to burned images if pinpoints not available
                        • Hub/Tower: Skip pinpoint logic, use burned images only
                        • Added _find_pinpoint_image() helper function
                        • Logging for image source tracking (audit trail)
  v2.1.3  2026-02-24  Phase 3 Implementation (TODO #4):
                        • Report layout: Two defects per page (vertical stacking)
                        • Reduced page count by approximately 50%
                        • Each page shows two complete defect blocks with separator
                        • Maintains all existing information (images, data, comments)
                        • Automatic page break every 2 defects for clean pagination
  v2.1.2  2026-02-24  Phase 2 Implementation (TODO items 3, 5, 6):
                        • Manual distances: UI spinboxes for root_distance_m, tip_distance_m
                        • Validation: Mandatory for Blade, disabled for Hub/Tower
                        • Serial numbers: Format WTG-Component-Side-SerialNo
                        • Per-component numbering: Blade includes side (LE/PS/TE/SS)
                        • Hub/Tower rules: No side in serial, separate numbering
                        • Report display: Serial numbers in defect headings
                        • UI validation: Disable Save if distance missing for Blade
  v2.1.1  2026-02-24  Phase 1 Implementation (TODO items 1, 2, 10):
                        • ML dependency safe mode: Hide ML tab/buttons if ultralytics missing
                        • Severity refactor: Remove Severity 2 from UI (kept in SEVERITY_LEVELS)
                        • New severity classification: Minor→1, Major→3/4, Critical→5, POI
                        • Data model update: Added root_distance_m, tip_distance_m fields
                        • Auto-migration: Load projects, convert severity 2→1
                        • Log message: "✅ ML Enabled" or "⚠️  ML Disabled" on startup
  v2.1.0  2026-02-24  User-requested changes:
                        • Removed compact annotation detail table from PDF (per component)
                        • Removed Turbine Index and Defect Summary tables from DOCX
                        • Increased per-annotation image sizes in DOCX
                          (WIDE: 3.6" → 4.0", ZOOM: 2.1" → 2.0", MINI: 0.8" → 0.5")
                        • Reviewer removed from DOCX footer (reviewed status removed)
                        • Report title: "{Tower} Aerial Wind Tower Inspection" (PDF/DOCX parity)
  v1.8.0  2026-02-24  CTO Audit implementation:
                        • Schema version field (schema_version:2) in project.json
                        • UUID-based annotation IDs (was MD5 10-char)
                        • Structured rotating log: wind_tower_app.log (5 MB × 5 backups)
                        • Crash dump handler: crash_YYYYMMDD_HHMMSS.log per crash
                        • Logic-layer permission enforcement: delete / approve / reject
                        • Undo/Redo stack depth capped at 50 entries (memory guard)
                        • Thumbnail cache key includes file size (stronger invalidation)
                        • Automatic backup rotation: project.json.bak1…bak5
                        • JSON corruption recovery: falls back to .bak1…bak5 automatically
                        • Calibration confidence scoring: HIGH/MEDIUM/LOW in GSD label
                        • Save performance benchmark logging (ms per save)
                        • component_gsd + blade_length_mm now persisted in save_project
                        • _install_crash_handler() called at entry point startup
  v1.4.1  2026-02-24  Expert audit pass: removed orphaned _build_detail_page (209 lines
                      dead code); verified all field names (ann.defect, ann.width_cm etc.);
                      confirmed all Phase 8 sub-tasks fully wired:
                        8.1 Report settings dialog ✅
                        8.2 Gate: warn if unapproved annotations ✅
                        8.3 Logo embedding in PDF header ✅ (was missing, now fixed)
                        8.4 Severity + defect-type breakdown tables ✅
                        8.5 Blade diagram embedded in PDF ✅
                        8.6 Full annotation list table ✅
                        8.7 Dimension labels on burned images ✅ (was missing, now fixed)
                        8.8 Progress dialog (QProgressDialog) ✅
                        8.9 Open PDF in native viewer (os.startfile/subprocess) ✅
                      Component-grouped image grid (6/page per B1/B2/B3/Hub/Tower) ✅
                      CSV annotation export auto-generated alongside PDF ✅
                      Open PDF / Open DOCX / Open CSV buttons in success dialog ✅
  v1.4.0  2026-02-24  PDF Phase 8.3 logo embedding; 8.7 dimension labels on images;
                      component-based image grids (6 imgs/page per B1/B2/B3/Hub/Tower);
                      CSV annotation export alongside PDF; fixed dead code & _count_pages()
  v1.3.0  2026-02-24  DOCX report generator (python-docx) — opt-in alongside PDF;
                      Summary dialog now includes "Also generate Word (.docx)" checkbox;
                      docx_report_generator.py companion module
  v1.2.1  2026-02-24  Full collation: Scopito+Dark+BurnIn+ML+PDF+Polygon+Undo/Redo
                      UX redesign: 3-zone toolbar, pill severity, header bar,
                      toast notifications, thumbnail severity borders
  v1.2.0  2026-02-23  ML + PDF report
  v1.1.0  2026-02-23  FLIR dark theme; Scopito severity; burn-in JPEG; bug fixes
  v1.0.0  2026-02-23  Initial build — Phases 1, 3-7
================================================================================
"""

# ── Standard Library ──────────────────────────────────────────────────────────
# io.BytesIO is used extensively for image buffering throughout the app;
# importing at module level avoids redundant inline re-imports in hot paths.
import sys, os, json, math, shutil, tempfile, hashlib, configparser, io
import logging, uuid
from io import BytesIO
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, field, asdict
from typing import Optional, List, Dict, Any, Tuple
from collections import OrderedDict

# ── Third-party: PyQt6 ────────────────────────────────────────────────────────
try:
    from PyQt6.QtWidgets import (
        QApplication, QMainWindow, QWidget, QSplitter, QGraphicsView,
        QGraphicsScene, QGraphicsItem, QGraphicsRectItem, QGraphicsEllipseItem,
        QGraphicsLineItem, QGraphicsTextItem, QGraphicsPixmapItem,
        QGraphicsPolygonItem,
        QListWidget, QListWidgetItem, QToolBar, QDockWidget, QLabel,
        QComboBox, QSpinBox, QDoubleSpinBox, QLineEdit, QPlainTextEdit,
        QPushButton, QSlider, QButtonGroup, QRadioButton, QFrame,
        QHBoxLayout, QVBoxLayout, QGridLayout, QFormLayout, QGroupBox,
        QFileDialog, QInputDialog, QMessageBox, QDialog, QDialogButtonBox,
        QProgressDialog, QStatusBar, QScrollArea, QSizePolicy, QLayout,
        QAbstractItemView, QTabWidget, QCheckBox, QProgressBar,
        QStackedWidget, QTextEdit, QMenu, QTableWidget, QTableWidgetItem,
    )
    from PyQt6.QtGui import (
        QPixmap, QImage, QPainter, QPen, QBrush, QColor, QFont,
        QAction, QCursor, QIcon, QKeySequence, QTransform, QPolygonF,
        QFontMetrics,
    )
    from PyQt6.QtCore import (
        Qt, QRectF, QPointF, QPoint, QSizeF, QSize, QThread, QRunnable,
        QThreadPool, pyqtSignal, QObject, QTimer, QMutex, QMutexLocker,
        pyqtSlot,
    )
except ImportError as e:
    print(f"[FATAL] PyQt6 not found: {e}\n  pip install PyQt6")
    sys.exit(1)

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError as e:
    print(f"[FATAL] Pillow not found: {e}\n  pip install Pillow")
    sys.exit(1)

# ── Optional: YOLO / Torch ────────────────────────────────────────────────────
# Phase 9.1 / 9.2: Lazy ML loader wrapped in try/except.
# Catches ImportError (missing package) AND OSError (WinError 1114 — DLL init
# failure from a corrupted torch GPU build).  Either error disables ML features
# gracefully without crashing the application.

def _safe_load_ml_dependencies() -> tuple:
    """
    Phase 9.1/9.2: Safe ML dependency loader.
    Returns (yolo_class_or_None, yolo_available: bool, torch_available: bool).
    Forces CPU-only mode recommendation if GPU DLL fails.
    """
    _yolo_cls = None
    _yolo_ok  = False
    _torch_ok = False

    # ── Step 1: probe torch first (DLL failures happen here) ──────────────────
    try:
        import torch as _torch
        _torch_ok = True
    except ImportError:
        pass
    except OSError as _e:
        # WinError 1114 — DLL initialisation failure (corrupt GPU torch build)
        _err_msg = str(_e)
        print(
            f"[WARN] torch DLL initialisation failed: {_err_msg}\n"
            "  ML features disabled.  To fix, install CPU-only torch:\n"
            "    pip install torch torchvision --index-url "
            "https://download.pytorch.org/whl/cpu\n"
            "  Or reinstall Visual C++ Redistributable from microsoft.com"
        )
        # torch broken → YOLO unusable; skip ultralytics import entirely
        return None, False, False
    except Exception as _e:
        print(f"[WARN] torch import error ({type(_e).__name__}): {_e} — ML disabled")
        return None, False, False

    # ── Step 2: load ultralytics YOLO ─────────────────────────────────────────
    try:
        from ultralytics import YOLO as _UltralyticsYOLO
        _yolo_cls = _UltralyticsYOLO
        _yolo_ok  = True
    except ImportError:
        print("[WARN] ultralytics not installed — ML features disabled.  pip install ultralytics")
    except OSError as _e:
        print(f"[WARN] ultralytics OSError (DLL?): {_e} — ML features disabled")
    except Exception as _e:
        print(f"[WARN] ultralytics import error ({type(_e).__name__}): {_e} — ML disabled")

    return _yolo_cls, _yolo_ok, _torch_ok


UltralyticsYOLO, YOLO_AVAILABLE, TORCH_AVAILABLE = _safe_load_ml_dependencies()

# ── Optional: ReportLab ───────────────────────────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib import colors as rl_colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, KeepTogether, PageBreak
    )
    from reportlab.platypus import Image as RLImage
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# ── Logger placeholder (real setup below after SCRIPT_DIR is defined) ────────
import traceback
from logging.handlers import RotatingFileHandler
log = logging.getLogger("WindTowerApp")  # temporary — upgraded after SCRIPT_DIR

# ── EXIF Calibration Module (v3.4.0) ─────────────────────────────────────────
# Automatic camera calibration from EXIF metadata with multi-library fallback
import re
import subprocess
from enum import Enum as PyEnum

try:
    from PIL.ExifTags import TAGS, GPSTAGS
    PIL_EXIF_AVAILABLE = True
except ImportError:
    PIL_EXIF_AVAILABLE = False
    print("[WARN] PIL EXIF tags not available")

try:
    import exifread
    EXIFREAD_AVAILABLE = True
except ImportError:
    EXIFREAD_AVAILABLE = False

try:
    import piexif
    PIEXIF_AVAILABLE = True
except ImportError:
    PIEXIF_AVAILABLE = False


class ConfidenceLevel(PyEnum):
    """Calibration confidence levels"""
    HIGH = "HIGH"
    MEDIUM = "MEDIUM"
    LOW = "LOW"
    MANUAL = "MANUAL"
    FAILED = "FAILED"


@dataclass
class ExifCalibrationData:
    """Container for EXIF-extracted calibration parameters.

    v4.4.3 (FIX-17a): Added drone_lat, drone_lon, absolute_altitude to support
    the new 3D GPS-based distance model (see estimate_distance_3d).
    """
    focal_length_mm: float
    sensor_width_mm: float
    sensor_height_mm: float
    image_width_px: int
    image_height_px: int
    pixel_pitch_um: float
    camera_model: str
    focal_length_35mm: Optional[float] = None
    gps_altitude: Optional[float] = None
    relative_altitude: Optional[float] = None
    gimbal_pitch: Optional[float] = None
    gimbal_yaw: Optional[float] = None
    gimbal_roll: Optional[float] = None
    distance_to_subject: Optional[float] = None
    confidence: ConfidenceLevel = ConfidenceLevel.HIGH
    data_source: str = "exif"
    # FIX-17a: New fields for 3D GPS distance model
    drone_lat: Optional[float] = None          # Decimal degrees from GPS IFD
    drone_lon: Optional[float] = None          # Decimal degrees from GPS IFD
    absolute_altitude: Optional[float] = None  # MSL metres from XMP AbsoluteAltitude

    def get_pixel_to_mm_ratio(self) -> float:
        """Calculate pixel to mm ratio on sensor"""
        return self.sensor_width_mm / self.image_width_px

    def get_field_of_view_h(self) -> float:
        """Calculate horizontal field of view in degrees"""
        return 2 * math.atan(self.sensor_width_mm / (2 * self.focal_length_mm)) * 180 / math.pi

    def get_field_of_view_v(self) -> float:
        """Calculate vertical field of view in degrees"""
        return 2 * math.atan(self.sensor_height_mm / (2 * self.focal_length_mm)) * 180 / math.pi

    # ──────────────────────────────────────────────────────────────────────────
    # LEGACY DISTANCE METHOD  (kept for fallback — superseded by estimate_distance_3d)
    # ──────────────────────────────────────────────────────────────────────────
    def estimate_distance_from_gps(self) -> Optional[float]:
        """
        LEGACY METHOD — altitude / tan(gimbal_pitch) approach.

        Still used as Tier-2 fallback in _batch_auto_calibrate() when tower
        GPS coordinates are not available in the project settings.

        ── Why this was the original approach ────────────────────────────────
        DJI drones write two altitude values in XMP:
          • RelativeAltitude — height above the takeoff/home point (metres).
          • AbsoluteAltitude — MSL elevation (metres).

        If the camera gimbal is pitched at angle θ from horizontal (e.g. -72°
        for a near-vertical tower shot), then the slant distance from the drone
        to the point it is aimed at can be approximated as:

            distance = RelativeAltitude / tan(|θ|)

        Geometry (side view):
                         Drone
                           •  ← RelativeAltitude above takeoff
                          /|
             slant_dist  / | RelativeAltitude
                        /  |
                       /θ  |
                      •────┘
                   Aimed point

        Example:
            RelativeAltitude = 85 m, GimbalPitchDegree = -72°
            pitch_rad = 72 × π/180 = 1.257 rad
            distance  = 85 / tan(1.257) ≈ 27.6 m

        ── Why it was superseded ─────────────────────────────────────────────
        1. Pilot altitude varies freely — at 40 m the drone gives a different
           distance than at 120 m for the same physical target, but this method
           correctly handles that as long as XMP altitude is accurate.
        2. RelativeAltitude is above the TAKEOFF point, not above the tower
           base.  On uneven terrain the takeoff pad may be 10–20 m higher or
           lower than the tower base, introducing that error directly into GSD.
        3. Gimbal pitch near 0° (camera pointing at horizon, θ < ~6°) makes
           tan(θ) ≈ 0 and distance → ∞.  The 0.1 rad guard (≈ 5.7°) rejects
           these shots and forces a fallback.
        4. It only uses the drone's own altitude and gimbal — it cannot account
           for the absolute position of the target point on the tower.

        The new estimate_distance_3d() fixes issues 2 and 4 by using the full
        3D GPS vector from drone to tower base, then refining with gimbal pitch
        to find the actual aim-point on the tower surface.
        ── End legacy documentation ──────────────────────────────────────────
        """
        if self.relative_altitude and self.gimbal_pitch:
            pitch_rad = abs(self.gimbal_pitch) * math.pi / 180
            if pitch_rad > 0.1:  # Avoid division by near-zero (< ~5.7°)
                return self.relative_altitude / math.tan(pitch_rad)
        return None

    # ──────────────────────────────────────────────────────────────────────────
    # NEW 3D GPS DISTANCE METHOD  (FIX-17a — primary method when tower GPS known)
    # ──────────────────────────────────────────────────────────────────────────
    def estimate_distance_3d(self,
                              tower_lat: float,
                              tower_lon: float,
                              tower_base_msl: float) -> Optional[float]:
        """
        FIX-17a: 3D GPS distance from drone to aimed point on the tower.

        This is the PRIMARY distance method used when the project has tower
        base GPS coordinates stored (Report Settings → Tower Location).

        ── Algorithm ─────────────────────────────────────────────────────────
        Step 1 — Haversine horizontal distance (drone GPS → tower base GPS):

            a = sin²(Δφ/2) + cos(φ₁)·cos(φ₂)·sin²(Δλ/2)
            horiz_m = 2·R·atan2(√a, √(1−a))         R = 6 371 000 m

        Step 2 — Altitude difference (drone above tower base):

            vert_m = drone_absolute_altitude_MSL − tower_base_alt_MSL

        Step 3 — Base slant (drone to tower base level):

            base_dist = √(horiz_m² + vert_m²)

        Step 4 — Gimbal pitch refinement (find actual aim-point height on tower):

            The camera is not pointing at the base; it is tilted up/down.
            Using the gimbal pitch angle θ and the horizontal stand-off:

                h_target  = vert_m − horiz_m · tan(|θ|)
                            (height on tower above base that camera is aimed at)
                aim_vert  = vert_m − h_target
                true_dist = √(horiz_m² + aim_vert²)

            Guard: if |θ| < 0.05 rad (< ~3°) the camera is near-horizontal
            and the refinement diverges — fall back to base_dist.

        ── Why this is better than estimate_distance_from_gps() ─────────────
        • Uses absolute MSL altitudes so takeoff-point elevation does not
          introduce error on uneven terrain.
        • Accounts for the drone's true 3D position relative to the tower,
          not just its height above its own takeoff pad.
        • The horizontal stand-off term means the GSD is correct even when
          the pilot is 30–50 m away laterally from the tower.
        • Pilot altitude changes are handled automatically — every image
          carries its own AbsoluteAltitude so each shot is independent.

        ── Requirements ──────────────────────────────────────────────────────
        • self.drone_lat, self.drone_lon  — from GPS IFD of the image.
        • self.absolute_altitude          — from XMP AbsoluteAltitude (MSL).
        • tower_lat, tower_lon            — project setting (constant per site).
        • tower_base_msl                  — project setting (constant per site).
        • self.gimbal_pitch               — from XMP GimbalPitchDegree.
          (pitch is used only for the Step 4 refinement; Steps 1–3 work
          without it and produce the drone→base distance as a fallback.)

        Returns None if any required GPS field is missing or invalid.
        """
        # Guard: all required GPS fields must be present and non-zero
        if not all([self.drone_lat, self.drone_lon, self.absolute_altitude,
                    tower_lat, tower_lon, tower_base_msl]):
            log.debug("[EXIF-3D] Missing GPS fields — 3D distance unavailable")
            return None

        try:
            R    = 6_371_000.0  # Earth radius in metres
            phi1 = math.radians(self.drone_lat)
            phi2 = math.radians(tower_lat)
            dphi = math.radians(tower_lat  - self.drone_lat)
            dlam = math.radians(tower_lon  - self.drone_lon)

            # Step 1 — Haversine horizontal distance
            a       = (math.sin(dphi / 2) ** 2
                       + math.cos(phi1) * math.cos(phi2) * math.sin(dlam / 2) ** 2)
            horiz_m = R * 2 * math.atan2(math.sqrt(a), math.sqrt(max(0.0, 1.0 - a)))

            # Step 2 — Vertical distance (drone above tower base)
            vert_m  = float(self.absolute_altitude) - float(tower_base_msl)

            # Step 3 — Drone-to-tower-base slant (fallback if no gimbal pitch)
            base_dist = math.sqrt(horiz_m ** 2 + vert_m ** 2)

            log.debug(
                f"[EXIF-3D] horiz={horiz_m:.2f}m  vert={vert_m:.2f}m  "
                f"base_dist={base_dist:.2f}m")

            if base_dist <= 0:
                log.warning("[EXIF-3D] Computed distance ≤ 0 — GPS data may be invalid")
                return None

            # Step 4 — Gimbal-pitch refinement to actual aim point on tower
            if self.gimbal_pitch is not None:
                pitch_rad = abs(self.gimbal_pitch) * math.pi / 180
                if pitch_rad > 0.05:  # Guard: reject near-horizontal shots
                    # Height on tower (above base) that the camera is aimed at
                    h_target  = vert_m - horiz_m * math.tan(pitch_rad)
                    # FIX-18c: h_target<0 means ray overshoots tower base;
                    # fall back to base_dist to avoid inflated GSD.
                    if h_target < 0:
                        log.debug("[EXIF-3D] h_target<0 (ray past base) — using base_dist")
                        return base_dist
                    aim_vert  = vert_m - h_target
                    true_dist = math.sqrt(horiz_m ** 2 + aim_vert ** 2)
                    log.debug(
                        f"[EXIF-3D] h_target={h_target:.2f}m  "
                        f"aim_vert={aim_vert:.2f}m  true_dist={true_dist:.2f}m")
                    # FIX-18c-2: degenerate case (horiz≈0 AND aim_vert≈0) returns 0
                    # which is falsy; return None explicitly to trigger Tier-2 cleanly.
                    return true_dist if true_dist > 0.1 else None

            # No pitch available — use drone-to-base distance as best estimate
            return base_dist if base_dist > 0.1 else None

        except Exception as exc:
            log.warning(f"[EXIF-3D] 3D distance calculation error: {exc}")
            return None

    def calculate_gsd_cm_per_px(self, distance_m: float) -> float:
        """
        Calculate GSD (cm/px) given distance to subject.

        GSD (m/px)  = (distance_m × sensor_width_m) / (focal_length_m × image_width_px)
        GSD (cm/px) = GSD (m/px) × 100

        distance_m is supplied by whichever distance-estimation method was used
        (estimate_distance_3d → estimate_distance_from_gps → 30m assumption).
        """
        gsd_m_per_px = ((distance_m * self.sensor_width_mm / 1000)
                        / (self.focal_length_mm / 1000)
                        / self.image_width_px)
        return gsd_m_per_px * 100


# DJI Camera Database
DJI_CAMERA_DATABASE = {
    "FC6310": {
        "name": "Phantom 4 Pro",
        "sensor_width_mm": 13.2,
        "sensor_height_mm": 8.8,
        "focal_length_mm": 8.8,
        "resolution": (4864, 3648),
        "pixel_pitch_um": 2.41
    },
    "FC6310S": {
        "name": "Phantom 4 Pro V2",
        "sensor_width_mm": 13.2,
        "sensor_height_mm": 8.8,
        "focal_length_mm": 8.8,
        "resolution": (5472, 3648),
        "pixel_pitch_um": 2.41
    },
    "FC3170": {
        "name": "Mavic 2 Pro",
        "sensor_width_mm": 13.2,
        "sensor_height_mm": 8.8,
        "focal_length_mm": 10.26,
        "resolution": (5472, 3648),
        "pixel_pitch_um": 2.41
    },
    "FC3411": {
        "name": "Mavic 2 Zoom",
        "sensor_width_mm": 6.3,
        "sensor_height_mm": 4.7,
        "focal_length_mm": 4.5,
        "resolution": (4000, 3000),
        "pixel_pitch_um": 1.57
    },
    "FC7303": {
        "name": "Mavic Air 2",
        "sensor_width_mm": 6.4,
        "sensor_height_mm": 4.8,
        "focal_length_mm": 4.49,
        "resolution": (8000, 6000),
        "pixel_pitch_um": 0.8  # 6.4mm / 8000px = 0.0008mm = 0.8um
    },
    # Alternative entry for FC7303 if focal length is exactly 4.5mm
    "FC7303_4.5": {
        "name": "Mavic Air 2 (4.5mm variant)",
        "sensor_width_mm": 6.4,
        "sensor_height_mm": 4.8,
        "focal_length_mm": 4.5,
        "resolution": (8000, 6000),
        "pixel_pitch_um": 0.8
    },
    "L1D-20c": {
        "name": "Mavic 3",
        "sensor_width_mm": 17.3,
        "sensor_height_mm": 13.0,
        "focal_length_mm": 24,
        "resolution": (5280, 3956),
        "pixel_pitch_um": 3.3
    },
    "L2D-20c": {
        "name": "Mavic 3 Cine",
        "sensor_width_mm": 17.3,
        "sensor_height_mm": 13.0,
        "focal_length_mm": 24,
        "resolution": (5280, 3956),
        "pixel_pitch_um": 3.3
    },
    "ZenmuseX7": {
        "name": "Inspire 2 / Matrice 200 (X7)",
        "sensor_width_mm": 23.5,
        "sensor_height_mm": 15.7,
        "focal_length_mm": 24,
        "resolution": (6016, 4008),
        "pixel_pitch_um": 3.91
    },
    "ZenmuseX5S": {
        "name": "Inspire 2 / Matrice 200 (X5S)",
        "sensor_width_mm": 17.3,
        "sensor_height_mm": 13.0,
        "focal_length_mm": 15,
        "resolution": (5280, 3956),
        "pixel_pitch_um": 3.3
    },
    "FC220": {
        "name": "Mavic Pro",
        "sensor_width_mm": 6.3,
        "sensor_height_mm": 4.7,
        "focal_length_mm": 4.73,
        "resolution": (4000, 3000),
        "pixel_pitch_um": 1.57
    },
    "FC330": {
        "name": "Phantom 3 Advanced",
        "sensor_width_mm": 6.3,
        "sensor_height_mm": 4.7,
        "focal_length_mm": 3.6,
        "resolution": (4000, 3000),
        "pixel_pitch_um": 1.56
    },
    "FC6520": {
        "name": "Mavic 3 Enterprise",
        "sensor_width_mm": 17.3,
        "sensor_height_mm": 13.0,
        "focal_length_mm": 24,
        "resolution": (5280, 3956),
        "pixel_pitch_um": 3.3
    },
    "MAVIC2-ENTERPRISE-ADVANCED": {
        "name": "Mavic 2 Enterprise Advanced",
        "sensor_width_mm": 13.2,
        "sensor_height_mm": 8.8,
        "focal_length_mm": 10.0,  # Thermal camera variant
        "resolution": (5472, 3648),
        "pixel_pitch_um": 2.41
    },
    "M2EA": {
        "name": "Mavic 2 Enterprise Advanced (M2EA)",
        "sensor_width_mm": 13.2,
        "sensor_height_mm": 8.8,
        "focal_length_mm": 10.0,
        "resolution": (5472, 3648),
        "pixel_pitch_um": 2.41
    },
    # ── New models added FIX-BUG: Missing DJI models caused calibration failures ──
    "FC4170": {
        "name": "Mavic Air 2S",
        "sensor_width_mm": 13.2,
        "sensor_height_mm": 8.8,
        "focal_length_mm": 22.0,
        "resolution": (5472, 3078),
        "pixel_pitch_um": 2.4
    },
    "FC3582": {
        "name": "DJI Mini 3 Pro",
        "sensor_width_mm": 9.6,
        "sensor_height_mm": 7.2,
        "focal_length_mm": 6.7,
        "resolution": (4032, 3024),
        "pixel_pitch_um": 2.4
    },
    "FC4380": {
        "name": "DJI Mini 4 Pro",
        "sensor_width_mm": 9.6,
        "sensor_height_mm": 7.2,
        "focal_length_mm": 6.7,
        "resolution": (4032, 3024),
        "pixel_pitch_um": 2.4
    },
    "FC3980": {
        "name": "DJI Air 3",
        "sensor_width_mm": 9.6,
        "sensor_height_mm": 7.2,
        "focal_length_mm": 9.0,
        "resolution": (4032, 3024),
        "pixel_pitch_um": 2.4
    },
    "FC6520E": {
        "name": "Mavic 3 Enterprise",
        "sensor_width_mm": 17.3,
        "sensor_height_mm": 13.0,
        "focal_length_mm": 24.0,
        "resolution": (5280, 3956),
        "pixel_pitch_um": 3.3
    },
    "FC6360": {
        "name": "Phantom 4 RTK",
        "sensor_width_mm": 13.2,
        "sensor_height_mm": 8.8,
        "focal_length_mm": 8.8,
        "resolution": (5472, 3648),
        "pixel_pitch_um": 2.41
    },
    "FC3688": {
        "name": "Inspire 3",
        "sensor_width_mm": 27.76,
        "sensor_height_mm": 20.82,
        "focal_length_mm": 24.0,
        "resolution": (8064, 6048),
        "pixel_pitch_um": 3.44
    },
    # Zenmuse P1 on Matrice 300/350 RTK (most common enterprise inspection platform)
    "ZenmuseP1": {
        "name": "Matrice 300/350 RTK + Zenmuse P1",
        "sensor_width_mm": 35.9,
        "sensor_height_mm": 24.0,
        "focal_length_mm": 35.0,
        "resolution": (8192, 5460),
        "pixel_pitch_um": 4.4
    },
    # H20T (thermal + zoom, common on Matrice 300)
    "ZenmuseH20T": {
        "name": "Matrice 300 RTK + Zenmuse H20T",
        "sensor_width_mm": 6.4,
        "sensor_height_mm": 4.8,
        "focal_length_mm": 4.5,
        "resolution": (4000, 3000),
        "pixel_pitch_um": 1.6
    },
    # DJI Mavic 3 Multispectral (M3M) used for vegetation/blade surveys
    "FC3371": {
        "name": "DJI Mavic 3 Multispectral",
        "sensor_width_mm": 17.3,
        "sensor_height_mm": 13.0,
        "focal_length_mm": 24.0,
        "resolution": (5280, 3956),
        "pixel_pitch_um": 3.3
    },
    # DJI Mini 2 (FC7203)
    "FC7203": {
        "name": "DJI Mini 2",
        "sensor_width_mm": 6.3,
        "sensor_height_mm": 4.7,
        "focal_length_mm": 4.49,
        "resolution": (4000, 3000),
        "pixel_pitch_um": 1.58
    },
    # Autel EVO II Pro (non-DJI but commonly used)
    "AUTELEVOIIPRO": {
        "name": "Autel EVO II Pro",
        "sensor_width_mm": 13.2,
        "sensor_height_mm": 8.8,
        "focal_length_mm": 8.0,
        "resolution": (5472, 3648),
        "pixel_pitch_um": 2.41
    },
}


class EXIFCalibrator:
    """EXIF-based automatic calibration with fallbacks"""
    
    def __init__(self, image_path: str):
        self.image_path = Path(image_path)
        self.raw_exif = {}
        self.calibration: Optional[ExifCalibrationData] = None
        
    def calibrate(self) -> Optional[ExifCalibrationData]:
        """Main calibration with fallback chain - ExifTool PRIMARY.

        Two-phase approach:
          Phase 1 — first-success wins: gets camera model, focal length, basic EXIF.
          Phase 2 — PIL XMP supplement (always): exifread/piexif cannot read XMP.
            DJI stores RelativeAltitude, GimbalPitchDegree, AbsoluteAltitude, and
            drone GPS coordinates ONLY in the XMP APP1 segment.  PIL is the only
            library (besides exiftool) that can extract it.  Running PIL after any
            non-PIL extractor ensures altitude + gimbal data is never silently lost.
        """
        log.info(f"[EXIF] Starting calibration for {self.image_path.name}")

        # ExifTool is PRIMARY - most reliable for DJI metadata
        methods = [
            (self._extract_exiftool, "ExifTool (PRIMARY)"),
            (self._extract_pil,      "PIL/Pillow"),
        ]

        if EXIFREAD_AVAILABLE:
            methods.append((self._extract_exifread, "exifread"))
        if PIEXIF_AVAILABLE:
            methods.append((self._extract_piexif, "piexif"))

        _primary_name: str = ""
        for method, name in methods:
            try:
                exif_data = method()
                if exif_data:
                    log.info(f"[EXIF] ✓ Successfully extracted EXIF using {name}")
                    self.raw_exif  = exif_data
                    _primary_name  = name
                    break
            except Exception as e:
                log.debug(f"[EXIF] {name} failed: {e}")
                continue

        if not self.raw_exif:
            log.warning("[EXIF] All extraction methods failed")
            return None

        # ── Phase 2: PIL XMP supplement ───────────────────────────────────────
        # Always run when the primary extractor was NOT PIL (i.e. it could not
        # have already included XMP).  Merges only keys that are absent in
        # raw_exif — existing values are never overwritten so the primary
        # extractor still wins for fields it handled correctly.
        _is_pil_primary = "PIL" in _primary_name
        if not _is_pil_primary and PIL_EXIF_AVAILABLE:
            _needs_xmp = 'XMP' not in self.raw_exif
            _needs_gps = 'GPS' not in self.raw_exif
            if _needs_xmp or _needs_gps:
                try:
                    pil_supp = self._extract_pil()
                    if pil_supp:
                        if _needs_xmp and 'XMP' in pil_supp:
                            self.raw_exif['XMP'] = pil_supp['XMP']
                            log.info(
                                f"[EXIF] ✓ PIL XMP supplement merged "
                                f"({list(pil_supp['XMP'].keys())})")
                        if _needs_gps and 'GPS' in pil_supp:
                            self.raw_exif['GPS'] = pil_supp['GPS']
                            log.info("[EXIF] ✓ PIL GPS supplement merged")
                        # Merge any other absent keys (e.g. FocalLength from EXIF IFD)
                        for _k, _v in pil_supp.items():
                            if _k not in self.raw_exif:
                                self.raw_exif[_k] = _v
                    else:
                        log.debug("[EXIF] PIL XMP supplement returned no data")
                except Exception as _sup_e:
                    log.debug(f"[EXIF] PIL XMP supplement failed: {_sup_e}")

        self.calibration = self._build_calibration_from_exif()
        return self.calibration
    
    def _extract_exiftool(self) -> Optional[Dict]:
        """
        Extract EXIF/XMP using Phil Harvey's ExifTool - THE GOLD STANDARD.
        
        ExifTool is the most reliable method for extracting ALL metadata from images,
        especially DJI drone images with complex XMP blocks. It correctly reads:
        - Standard EXIF tags
        - GPS data
        - XMP metadata (DJI-specific altitude, gimbal, camera info)
        - Maker notes
        - ALL other embedded metadata
        
        This is now the PRIMARY extraction method for auto-calibration.
        """
        import subprocess
        import json
        
        # Check if exiftool is available
        try:
            result = subprocess.run(
                ['exiftool', '-ver'],
                capture_output=True,
                text=True,
                timeout=5
            )
            if result.returncode != 0:
                raise FileNotFoundError("exiftool not found")
            log.debug(f"[EXIF] ExifTool version: {result.stdout.strip()}")
        except (FileNotFoundError, subprocess.TimeoutExpired) as e:
            log.debug(f"[EXIF] ExifTool not available: {e}")
            raise ImportError("ExifTool not installed or not in PATH")
        
        # Run exiftool with JSON output
        try:
            cmd = [
                'exiftool',
                '-json',
                '-n',  # Numeric output (no units, raw values)
                '-G',  # Group names
                str(self.image_path)
            ]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=30
            )
            
            if result.returncode != 0:
                log.warning(f"[EXIF] ExifTool error: {result.stderr}")
                return None
            
            # Parse JSON output
            data = json.loads(result.stdout)
            if not data or len(data) == 0:
                return None
            
            # ExifTool returns array with one element per file
            raw_tags = data[0]
            
            # Convert to our standard format
            exif_data = {}
            xmp_data = {}
            gps_data = {}
            
            for key, value in raw_tags.items():
                # Remove group prefix (e.g., "EXIF:Model" -> "Model")
                if ':' in key:
                    group, tag = key.split(':', 1)
                else:
                    group = 'Unknown'
                    tag = key

                # Route to appropriate dictionary.
                # FIX-BUG: ExifTool with -G outputs camera Make/Model under IFD0,
                # focal length under ExifIFD, and image dimensions may arrive under
                # Main or IFD0.  Previously only EXIF/File/Composite were accepted,
                # so camera_model was always "Unknown" → database lookup failed →
                # sensor size missing → _build_calibration_from_exif() returned None.
                # Added IFD0, IFD1, ExifIFD, Main, MakerNotes, APP1 to fix this.
                if group in ('XMP', 'XMP-drone-dji', 'XMP-crs', 'XMP-dc',
                             'XMP-photoshop', 'XMP-tiff', 'XMP-exif'):
                    # DJI XMP fields (and any other XMP namespace)
                    clean_tag = tag.replace('drone-dji:', '').replace('XMP:', '')
                    xmp_data[clean_tag] = value
                elif group == 'GPS':
                    gps_data[tag] = value
                elif group in ('EXIF', 'File', 'Composite',
                               'IFD0', 'IFD1', 'ExifIFD',
                               'Main', 'MakerNotes', 'APP1'):
                    # Store under the bare tag name so downstream extractors
                    # (_get_camera_model, _extract_focal_length, etc.) find them
                    # with the same key regardless of which ExifTool group they
                    # came from.  Later keys from higher-priority groups will
                    # overwrite earlier ones if the same tag name appears twice.
                    exif_data[tag] = value
            
            # Add grouped data
            if xmp_data:
                exif_data['XMP'] = xmp_data
            if gps_data:
                exif_data['GPS'] = gps_data
            
            # Ensure critical fields are present
            if 'ImageWidth' not in exif_data and 'ImageSize' in exif_data:
                # ExifTool might provide ImageSize as "8000x6000"
                try:
                    w, h = exif_data['ImageSize'].split('x')
                    exif_data['ImageWidth'] = int(w)
                    exif_data['ImageHeight'] = int(h)
                except:
                    pass
            
            log.info(f"[EXIF] ExifTool extracted {len(exif_data)} EXIF fields, "
                    f"{len(xmp_data)} XMP fields, {len(gps_data)} GPS fields")
            
            return exif_data if exif_data else None
            
        except subprocess.TimeoutExpired:
            log.warning("[EXIF] ExifTool timeout")
            return None
        except json.JSONDecodeError as e:
            log.warning(f"[EXIF] ExifTool JSON parse error: {e}")
            return None
        except Exception as e:
            log.warning(f"[EXIF] ExifTool unexpected error: {e}")
            return None
    
    def _extract_pil(self) -> Optional[Dict]:
        """Extract EXIF/XMP using PIL/Pillow.

        This method serves two roles:
          1. PRIMARY extractor when exiftool is unavailable and exifread fails.
          2. XMP SUPPLEMENT — called by calibrate() after any non-PIL primary
             extractor to retrieve altitude/gimbal/GPS from the XMP APP1 segment.

        FIX-BUG: Previously raised ImportError when PIL_EXIF_AVAILABLE was False,
        preventing this method from serving as XMP supplement.  Now degrades
        gracefully: image dimensions and XMP are always extracted; tag-name
        decoding (TAGS/GPSTAGS) is skipped with numeric tag-ID fallback when the
        import failed.  This matters because PIL EXIF tag lookup is a separate
        import from PIL itself (which is mandatory and cannot be missing).

        FIX-BUG: `img` was opened without a context manager so the file handle
        was never explicitly closed; replaced with `with Image.open(...) as img`.
        """
        exif_data: Dict[str, Any] = {}

        with Image.open(self.image_path) as img:
            # Dimensions from the JPEG SOF0 marker — always correct
            exif_data['ImageWidth']  = img.width
            exif_data['ImageHeight'] = img.height

            # EXIF IFD tags — TAGS may not be available if ExifTags import failed
            try:
                exif = img.getexif()
            except Exception:
                exif = {}

            if exif:
                for tag_id, value in exif.items():
                    # Use TAGS lookup when available; fall back to numeric ID string
                    tag_name = (TAGS.get(tag_id, str(tag_id))
                                if PIL_EXIF_AVAILABLE else str(tag_id))
                    exif_data[tag_name] = value

                    if tag_name == "GPSInfo":
                        gps_data: Dict[str, Any] = {}
                        try:
                            # Newer Pillow: value is an IfdTag object (iterable)
                            for gps_tag_id, gps_value in (
                                    value.items() if hasattr(value, 'items')
                                    else {}.items()):
                                gps_tag_name = (GPSTAGS.get(gps_tag_id, str(gps_tag_id))
                                                if PIL_EXIF_AVAILABLE else str(gps_tag_id))
                                gps_data[gps_tag_name] = gps_value
                        except Exception as _gps_e:
                            log.debug(f"[EXIF] PIL GPS IFD parse error: {_gps_e}")
                        if gps_data:
                            exif_data['GPS'] = gps_data

            # XMP extraction — DJI altitude/gimbal/GPS lives here
            xmp_data = self._extract_xmp_robust(img)

        # img is now closed; process results outside with-block
        if xmp_data:
            exif_data['XMP'] = xmp_data
            log.info(f"[EXIF] ✓ XMP data extracted: {list(xmp_data.keys())}")
        else:
            log.debug(f"[EXIF] No XMP data found in {self.image_path.name}")

        return exif_data if exif_data else None

    def _extract_xmp_robust(self, img: Image.Image) -> Optional[Dict]:
        """
        Robust XMP extraction with multiple fallback methods.
        DJI drones store critical data (altitude, gimbal pitch) in XMP blocks.
        """
        xmp_string = None
        
        # Method 1: Try standard XMP key
        xmp_string = img.info.get('XML:com.adobe.xmp')
        if xmp_string:
            log.debug("[EXIF] XMP found via 'XML:com.adobe.xmp' key")
        
        # Method 2: Try alternative keys
        if not xmp_string:
            for key in img.info.keys():
                if 'xmp' in key.lower() or 'xml' in key.lower():
                    xmp_string = img.info.get(key)
                    if xmp_string:
                        log.debug(f"[EXIF] XMP found via alternative key: {key}")
                        break
        
        # Method 3: Read raw file for XMP block
        if not xmp_string:
            try:
                with open(self.image_path, 'rb') as f:
                    data = f.read()
                    # Look for XMP packet markers
                    xmp_start = data.find(b'<x:xmpmeta')
                    xmp_end = data.find(b'</x:xmpmeta>')
                    if xmp_start != -1 and xmp_end != -1:
                        xmp_string = data[xmp_start:xmp_end + 12].decode('utf-8', errors='ignore')
                        log.debug("[EXIF] XMP found via raw file parsing")
            except Exception as e:
                log.debug(f"[EXIF] Raw XMP extraction failed: {e}")
        
        # Parse XMP if found
        if xmp_string:
            if isinstance(xmp_string, bytes):
                xmp_string = xmp_string.decode('utf-8', errors='ignore')
            return self._parse_dji_xmp(xmp_string)
        
        return None
    
    def _extract_exifread(self) -> Optional[Dict]:
        """Extract EXIF using exifread library.

        FIX-BUG (dimension collision): exifread emits both:
            'Image ImageWidth'   → full-resolution width (e.g. 5472)
            'IFD1 ImageWidth'    → embedded JPEG thumbnail width (e.g. 160)
        The original cleaner `tag.split()[-1]` mapped BOTH to 'ImageWidth'.
        Dict iteration order determines which value survives — the thumbnail
        (160 px) can win and become the stored 'ImageWidth', causing downstream
        GSD calculations to be wrong by a factor of ~34×.

        Fix: skip all IFD1 / Thumbnail namespace tags before cleaning.
        Also: parse the GPS IFD into a nested 'GPS' dict so _extract_drone_lat
        and _extract_gps_altitude find data in the same place as the PIL path.
        """
        if not EXIFREAD_AVAILABLE:
            raise ImportError("exifread not installed")

        # Thumbnail / IFD1 tags that must be excluded to avoid dimension collision
        _SKIP_PREFIXES = ('IFD1 ', 'Thumbnail ')
        # GPS tag prefix used by exifread
        _GPS_PREFIX    = 'GPS '

        with open(self.image_path, 'rb') as f:
            tags = exifread.process_file(f, details=True)

        exif_data: Dict[str, Any] = {}
        gps_data:  Dict[str, Any] = {}

        for tag, value in tags.items():
            # ── Skip all thumbnail/IFD1 tags unconditionally ─────────────────
            # These contain dimensions of the JPEG thumbnail (often 160 px wide)
            # not the full-resolution image.  Keeping them would corrupt
            # ImageWidth / ImageHeight after tag-name deduplication.
            if any(tag.startswith(pfx) for pfx in _SKIP_PREFIXES):
                continue

            # ── Route GPS tags into a nested dict (mirrors PIL/ExifTool path) ─
            if tag.startswith(_GPS_PREFIX):
                clean = tag[len(_GPS_PREFIX):]          # "GPSLatitude" etc.
                gps_data[clean] = str(value)
                continue

            # ── Main EXIF IFD tags ────────────────────────────────────────────
            # exifread format: "Image Model", "EXIF FocalLength", etc.
            # Take the last word (the actual tag name) as the clean key.
            clean_tag = tag.split()[-1] if ' ' in tag else tag
            # Only write if not already set (earlier/higher-priority tags win)
            if clean_tag not in exif_data:
                exif_data[clean_tag] = str(value)

        if gps_data:
            exif_data['GPS'] = gps_data

        return exif_data if exif_data else None
    
    def _extract_piexif(self) -> Optional[Dict]:
        """Extract EXIF using piexif library"""
        if not PIEXIF_AVAILABLE:
            raise ImportError("piexif not installed")
        
        exif_dict = piexif.load(str(self.image_path))
        exif_data = {}
        
        for ifd_name in ["0th", "Exif", "GPS", "1st"]:
            if ifd_name in exif_dict:
                for tag, value in exif_dict[ifd_name].items():
                    tag_name = piexif.TAGS[ifd_name].get(tag, {}).get("name", str(tag))
                    exif_data[tag_name] = value
        
        return exif_data if exif_data else None
    
    def _parse_dji_xmp(self, xmp_string: str) -> Dict:
        """
        Parse DJI-specific XMP metadata with comprehensive field extraction.
        DJI stores critical calibration data in XMP that's not in standard EXIF.
        """
        xmp_data = {}
        
        # DJI-specific patterns for critical calibration data
        patterns = {
            # Altitude data (critical for distance estimation)
            'RelativeAltitude': r'drone-dji:RelativeAltitude="([^"]+)"',
            'AbsoluteAltitude': r'drone-dji:AbsoluteAltitude="([^"]+)"',
            
            # Gimbal orientation (critical for distance calculation)
            'GimbalPitchDegree': r'drone-dji:GimbalPitchDegree="([^"]+)"',
            'GimbalYawDegree': r'drone-dji:GimbalYawDegree="([^"]+)"',
            'GimbalRollDegree': r'drone-dji:GimbalRollDegree="([^"]+)"',
            
            # Camera model (for database lookup)
            'CameraModel': r'drone-dji:CameraModel="([^"]+)"',
            'CameraModelName': r'drone-dji:CameraModelName="([^"]+)"',
            
            # Additional useful metadata
            'FlightSpeed': r'drone-dji:FlightSpeed="([^"]+)"',
            'FlightPitchDegree': r'drone-dji:FlightPitchDegree="([^"]+)"',
            'FlightYawDegree': r'drone-dji:FlightYawDegree="([^"]+)"',
            'FlightRollDegree': r'drone-dji:FlightRollDegree="([^"]+)"',
        }
        
        extracted_count = 0
        for key, pattern in patterns.items():
            match = re.search(pattern, xmp_string)
            if match:
                try:
                    value_str = match.group(1)
                    # Try to convert to float for numeric values
                    try:
                        xmp_data[key] = float(value_str)
                        extracted_count += 1
                    except ValueError:
                        xmp_data[key] = value_str
                        extracted_count += 1
                except (IndexError, AttributeError) as e:
                    log.debug(f"[EXIF] Failed to extract {key}: {e}")
        
        if extracted_count > 0:
            log.info(f"[EXIF] ✓ Extracted {extracted_count} DJI XMP fields")
        else:
            log.warning("[EXIF] ✗ No DJI XMP fields extracted - check XMP format")
        
        return xmp_data
    
    def _build_calibration_from_exif(self) -> Optional[ExifCalibrationData]:
        """Build CalibrationData object from extracted EXIF"""
        
        camera_model = self._get_camera_model()
        focal_length = self._extract_focal_length()
        sensor_width, sensor_height = self._extract_sensor_size()
        image_width, image_height = self._extract_image_size()
        
        # Log what we extracted
        log.info(f"[EXIF] Extracted data for {self.image_path.name}:")
        log.info(f"  Camera Model: {camera_model}")
        log.info(f"  Focal Length: {focal_length} mm")
        log.info(f"  Sensor Size: {sensor_width} x {sensor_height} mm")
        log.info(f"  Image Size: {image_width} x {image_height} px")
        
        confidence = ConfidenceLevel.HIGH
        data_source = "exif"
        
        # Try database lookup if missing sensor size or focal length
        if not focal_length or not sensor_width:
            log.info(f"[EXIF] Looking up camera database for: {camera_model}")
            db_data = self._lookup_camera_database(camera_model)
            if db_data:
                if not focal_length:
                    focal_length = db_data.get('focal_length_mm')
                    log.info(f"[EXIF] ✓ Filled focal_length from database: {focal_length} mm")
                if not sensor_width:
                    sensor_width = db_data.get('sensor_width_mm')
                    sensor_height = db_data.get('sensor_height_mm')
                    log.info(f"[EXIF] ✓ Filled sensor size from database: {sensor_width} x {sensor_height} mm")
                confidence = ConfidenceLevel.MEDIUM
                data_source = "exif+database"
            else:
                log.warning(f"[EXIF] ✗ Camera model '{camera_model}' not found in database")
        
        pixel_pitch = None
        if sensor_width and image_width:
            pixel_pitch = (sensor_width * 1000) / image_width
        
        # Strict validation - all parameters must be present
        if not all([focal_length, sensor_width, sensor_height, image_width, image_height]):
            missing = []
            if not focal_length: missing.append("focal_length")
            if not sensor_width: missing.append("sensor_width")
            if not sensor_height: missing.append("sensor_height")
            if not image_width: missing.append("image_width")
            if not image_height: missing.append("image_height")
            log.warning(f"[EXIF] ✗ Missing critical calibration parameters: {', '.join(missing)}")
            return None
        
        # Check for GPS and gimbal data for distance estimation
        gps_altitude      = self._extract_gps_altitude()
        relative_altitude = self._extract_relative_altitude()
        gimbal_pitch      = self._extract_gimbal_pitch()
        # FIX-17b: Extract new fields for 3D distance model
        drone_lat         = self._extract_drone_lat()
        drone_lon         = self._extract_drone_lon()
        absolute_altitude = self._extract_absolute_altitude()

        if not gps_altitude and not relative_altitude:
            log.warning("[EXIF] ⚠ No altitude data (GPS or Relative) - distance estimation will be limited")
            if confidence == ConfidenceLevel.HIGH:
                confidence = ConfidenceLevel.MEDIUM

        if gimbal_pitch is None:
            log.warning("[EXIF] ⚠ No gimbal pitch data - distance estimation may be inaccurate")
            if confidence == ConfidenceLevel.HIGH:
                confidence = ConfidenceLevel.MEDIUM

        # FIX-17b: Confidence note when drone GPS lat/lon absent
        if drone_lat is None or drone_lon is None:
            log.info("[EXIF] ℹ Drone lat/lon not extracted — 3D GPS method unavailable; "
                     "pitch-altitude fallback will be used")

        log.info(f"[EXIF] ✓ Calibration successful - Confidence: {confidence.value}, Source: {data_source}")

        return ExifCalibrationData(
            focal_length_mm=focal_length,
            sensor_width_mm=sensor_width,
            sensor_height_mm=sensor_height,
            image_width_px=image_width,
            image_height_px=image_height,
            pixel_pitch_um=pixel_pitch or 3.3,
            camera_model=camera_model or "Unknown",
            focal_length_35mm=self._extract_focal_length_35mm(),
            gps_altitude=gps_altitude,
            relative_altitude=relative_altitude,
            gimbal_pitch=gimbal_pitch,
            gimbal_yaw=self._extract_gimbal_yaw(),
            gimbal_roll=self._extract_gimbal_roll(),
            confidence=confidence,
            data_source=data_source,
            # FIX-17b: populate new 3D-GPS fields
            drone_lat=drone_lat,
            drone_lon=drone_lon,
            absolute_altitude=absolute_altitude,
        )
    
    def _get_camera_model(self) -> str:
        """Extract camera model from EXIF or XMP.

        FIX-BUG: ExifTool routes Make/Model through IFD0 group (now parsed).
        Also handle the case where ExifTool's -n output concatenates make+model
        into a single string for some backends, and strip any trailing whitespace
        or null characters that some DJI firmwares embed.
        """
        # Prefer 'Model' (camera model code e.g. "FC7303") over 'Make' (brand).
        # Include 'UniqueCameraModel' which some DJI firmware versions populate.
        for key in ('Model', 'CameraModel', 'UniqueCameraModel', 'CameraModelName'):
            if key in self.raw_exif:
                model = str(self.raw_exif[key]).strip().rstrip('\x00')
                if model and model.lower() not in ('unknown', 'none', ''):
                    log.debug(f"[EXIF] Camera model from EXIF[{key}]: {model}")
                    return model

        # Try XMP data (DJI-specific — some firmware only writes model to XMP)
        if 'XMP' in self.raw_exif:
            xmp = self.raw_exif['XMP']
            for key in ('CameraModel', 'CameraModelName', 'Model'):
                if key in xmp:
                    model = str(xmp[key]).strip().rstrip('\x00')
                    if model and model.lower() not in ('unknown', 'none', ''):
                        log.debug(f"[EXIF] Camera model from XMP[{key}]: {model}")
                        return model

        # Last resort: return Make (brand) so database fuzzy match can try
        if 'Make' in self.raw_exif:
            make = str(self.raw_exif['Make']).strip().rstrip('\x00')
            if make and make.lower() not in ('unknown', 'none', ''):
                log.debug(f"[EXIF] Falling back to Make: {make}")
                return make

        log.warning("[EXIF] Camera model not found in EXIF or XMP")
        return "Unknown"
    
    def _extract_focal_length(self) -> Optional[float]:
        for key in ['FocalLength', 'FocalLengthIn35mmFormat']:
            if key in self.raw_exif:
                value = self.raw_exif[key]
                try:
                    if isinstance(value, str) and '/' in value:
                        num, den = map(float, value.split('/'))
                        return num / den if den != 0 else None
                    return float(value)
                except (ValueError, TypeError):
                    pass
        return None
    
    def _extract_focal_length_35mm(self) -> Optional[float]:
        if 'FocalLengthIn35mmFormat' in self.raw_exif:
            try:
                return float(self.raw_exif['FocalLengthIn35mmFormat'])
            except (ValueError, TypeError):
                pass
        return None
    
    def _extract_sensor_size(self) -> Tuple[Optional[float], Optional[float]]:
        """
        Extract sensor physical dimensions from EXIF.
        Note: Most cameras don't include this in EXIF, so we rely on database lookup.

        FIX-BUG: FocalPlaneXResolution / FocalPlaneYResolution are expressed in
        pixels-per-unit (DPI or pixels/cm depending on FocalPlaneResolutionUnit)
        — NOT sensor dimensions in mm.  Using them directly as sensor_width_mm
        produced wildly wrong GSD values (e.g., 720 dpi ≡ 720 mm sensor width).
        These tags are now excluded; only explicit 'SensorWidth'/'SensorHeight'
        tags (rare, non-standard) are accepted.  For all DJI cameras the sensor
        size is correctly supplied by _lookup_camera_database().
        """
        # Only accept explicit sensor-dimension tags; NOT FocalPlane* (those are DPI)
        for width_key in ('SensorWidth',):
            if width_key in self.raw_exif:
                try:
                    width_val = self.raw_exif[width_key]
                    if isinstance(width_val, str) and '/' in width_val:
                        num, den = map(float, width_val.split('/'))
                        sensor_width = num / den if den != 0 else None
                    else:
                        sensor_width = float(width_val)

                    for height_key in ('SensorHeight',):
                        if height_key in self.raw_exif:
                            try:
                                height_val = self.raw_exif[height_key]
                                if isinstance(height_val, str) and '/' in height_val:
                                    num, den = map(float, height_val.split('/'))
                                    sensor_height = num / den if den != 0 else None
                                else:
                                    sensor_height = float(height_val)
                                return sensor_width, sensor_height
                            except (ValueError, TypeError):
                                pass
                except (ValueError, TypeError):
                    pass

        # Sensor size not in EXIF for any DJI drone — will be filled from database
        return None, None
    
    def _extract_image_size(self) -> Tuple[Optional[int], Optional[int]]:
        """
        Return (width_px, height_px) for the image.

        FIX-BUG (thumbnail dimension collision — root cause of 160×6000 result):
          exifread maps both 'Image ImageWidth' (full-res) and 'IFD1 ImageWidth'
          (JPEG thumbnail, typically 160 px wide on DJI images) to the same key
          'ImageWidth' after tag cleaning.  Dict iteration order then determines
          which value survives — the thumbnail can win and produce width=160.
          Because 160 is *truthy*, the old 'if not width' guard never fired PIL,
          so PIL filled in only the missing height (6000) → 160×6000 phantom.

        FIX strategy: PIL reads the JPEG SOF0 (Start-of-Frame) marker which
          always contains the *actual* decoded dimensions.  It never sees the
          EXIF IFD thumbnail entry.  PIL is therefore the authoritative source
          and is now run unconditionally.  EXIF values are only used as a last
          resort when PIL itself fails (corrupt file, unsupported codec, etc.).
        """
        # ── Step 1: Collect EXIF-reported dimensions as a fallback pool ──────
        # These may be correct (ExifTool, PIL primary) or wrong (exifread
        # thumbnail collision).  They will only be used if PIL fails below.
        exif_w: Optional[int] = None
        exif_h: Optional[int] = None

        for key in ('ImageWidth', 'ExifImageWidth', 'PixelXDimension'):
            if key in self.raw_exif:
                try:
                    exif_w = int(self.raw_exif[key])
                    break
                except (ValueError, TypeError):
                    pass

        for key in ('ImageHeight', 'ExifImageHeight', 'PixelYDimension'):
            if key in self.raw_exif:
                try:
                    exif_h = int(self.raw_exif[key])
                    break
                except (ValueError, TypeError):
                    pass

        # ── Step 2: PIL — always the authoritative dimension source ───────────
        # PIL opens only the file header (SOF0 for JPEG, IHDR for PNG, IFD for
        # TIFF) so no pixel decoding occurs — safe even for 100 MB drone images.
        pil_w: Optional[int] = None
        pil_h: Optional[int] = None
        try:
            with Image.open(self.image_path) as _img:
                pil_w, pil_h = _img.size
        except Exception as _pil_err:
            log.warning(
                f"[EXIF] PIL dimension read failed for {self.image_path.name}: {_pil_err}")

        if pil_w and pil_h:
            # Log a warning if PIL disagrees with EXIF — helps diagnose future issues
            if exif_w and exif_h and (pil_w != exif_w or pil_h != exif_h):
                log.warning(
                    f"[EXIF] Dimension mismatch — EXIF: {exif_w}×{exif_h}, "
                    f"PIL (authoritative): {pil_w}×{pil_h} "
                    f"({self.image_path.name}) — EXIF may have been thumbnail IFD. "
                    f"Using PIL values.")
            elif not exif_w or not exif_h:
                log.info(
                    f"[EXIF] PIL provided image dimensions: {pil_w}×{pil_h} "
                    f"({self.image_path.name})")
            return pil_w, pil_h

        # ── Step 3: PIL failed — use EXIF values as last resort ───────────────
        # This path should be extremely rare (corrupted/unsupported file format).
        if exif_w and exif_h:
            log.warning(
                f"[EXIF] PIL unavailable — falling back to EXIF dimensions: "
                f"{exif_w}×{exif_h} ({self.image_path.name}). "
                f"These may be thumbnail dimensions if exifread was the extractor.")
            return exif_w, exif_h

        log.warning(f"[EXIF] Could not determine image dimensions for {self.image_path.name}")
        return None, None
    
    def _extract_gps_altitude(self) -> Optional[float]:
        """Extract GPS altitude from EXIF or XMP"""
        # Try standard GPS EXIF first
        if 'GPS' in self.raw_exif:
            gps = self.raw_exif['GPS']
            if 'GPSAltitude' in gps:
                try:
                    alt = gps['GPSAltitude']
                    if isinstance(alt, str) and '/' in alt:
                        num, den = map(float, alt.split('/'))
                        return num / den if den != 0 else None
                    return float(alt)
                except (ValueError, TypeError):
                    pass
        
        # Try direct EXIF fields (ExifTool format)
        for key in ['GPSAltitude', 'Altitude']:
            if key in self.raw_exif:
                try:
                    return float(self.raw_exif[key])
                except (ValueError, TypeError):
                    pass
        
        return None
    
    def _extract_relative_altitude(self) -> Optional[float]:
        """Extract relative altitude from XMP or EXIF"""
        # Try XMP first (DJI standard location)
        if 'XMP' in self.raw_exif:
            xmp = self.raw_exif['XMP']
            # FIX-BUG: was ['RelativeAltitude', 'RelativeAltitude'] — duplicate key
            # meant the second iteration was a no-op.  Both common XMP field names
            # are now covered.
            for key in ['RelativeAltitude', 'drone-dji:RelativeAltitude']:
                if key in xmp:
                    try:
                        value = xmp[key]
                        # Handle both "+108.72" and 108.72 formats
                        if isinstance(value, str):
                            value = value.replace('+', '')
                        return float(value)
                    except (ValueError, TypeError):
                        pass
        
        # Try direct EXIF fields (ExifTool sometimes puts it here)
        for key in ['RelativeAltitude', 'FlightHeight']:
            if key in self.raw_exif:
                try:
                    value = self.raw_exif[key]
                    if isinstance(value, str):
                        value = value.replace('+', '')
                    return float(value)
                except (ValueError, TypeError):
                    pass
        
        return None
    
    def _extract_gimbal_pitch(self) -> Optional[float]:
        """Extract gimbal pitch from XMP or EXIF"""
        # Try XMP first (DJI standard location)
        if 'XMP' in self.raw_exif:
            xmp = self.raw_exif['XMP']
            for key in ['GimbalPitchDegree', 'GimbalPitch']:
                if key in xmp:
                    try:
                        value = xmp[key]
                        # Handle both "-90.00" and -90.00 formats
                        return float(value)
                    except (ValueError, TypeError):
                        pass
        
        # Try direct EXIF fields (ExifTool format)
        for key in ['GimbalPitchDegree', 'GimbalPitch', 'CameraPitch']:
            if key in self.raw_exif:
                try:
                    return float(self.raw_exif[key])
                except (ValueError, TypeError):
                    pass
        
        return None
    
    def _extract_gimbal_yaw(self) -> Optional[float]:
        """Extract gimbal yaw from XMP or EXIF"""
        if 'XMP' in self.raw_exif:
            xmp = self.raw_exif['XMP']
            for key in ['GimbalYawDegree', 'GimbalYaw']:
                if key in xmp:
                    try:
                        return float(xmp[key])
                    except (ValueError, TypeError):
                        pass
        
        for key in ['GimbalYawDegree', 'GimbalYaw', 'CameraYaw']:
            if key in self.raw_exif:
                try:
                    return float(self.raw_exif[key])
                except (ValueError, TypeError):
                    pass
        
        return None
    
    def _extract_gimbal_roll(self) -> Optional[float]:
        """Extract gimbal roll from XMP or EXIF"""
        if 'XMP' in self.raw_exif:
            xmp = self.raw_exif['XMP']
            for key in ['GimbalRollDegree', 'GimbalRoll']:
                if key in xmp:
                    try:
                        return float(xmp[key])
                    except (ValueError, TypeError):
                        pass

        for key in ['GimbalRollDegree', 'GimbalRoll', 'CameraRoll']:
            if key in self.raw_exif:
                try:
                    return float(self.raw_exif[key])
                except (ValueError, TypeError):
                    pass

        return None

    # ── FIX-17b: New extractors for 3D GPS distance model ────────────────────

    def _extract_drone_lat(self) -> Optional[float]:
        """
        FIX-17b: Extract drone GPS latitude as a decimal-degree float.

        DJI images carry GPS coordinates in the standard EXIF GPS IFD.
        Depending on which library extracted the raw EXIF the value may arrive
        in one of three forms:

          1. Decimal degrees (float)  — ExifTool with -n flag or after
             coordinate conversion (e.g. "15.123456").
          2. DMS tuple of IFDRationals — PIL/Pillow GPS block:
             ((15, 1), (7, 1), (24834, 1000)) meaning 15° 7' 24.834".
          3. DMS tuple of plain floats — exifread or piexif output.

        Returns decimal degrees (positive = N, negative = S) or None.
        """
        try:
            gps = self.raw_exif.get('GPS', {})

            # Path 1: ExifTool already decoded to decimal degrees
            for key in ('GPSLatitude', 'Latitude'):
                val = gps.get(key) or self.raw_exif.get(key)
                if val is None:
                    continue
                try:
                    dd = float(val)
                    # ExifTool always gives positive; apply ref sign
                    ref = (gps.get('GPSLatitudeRef') or
                           self.raw_exif.get('GPSLatitudeRef') or 'N')
                    return -dd if str(ref).upper() in ('S', 'SOUTH') else dd
                except (ValueError, TypeError):
                    pass

            # Path 2 / 3: DMS tuple (PIL IFDRational or plain floats)
            lat_dms = gps.get('GPSLatitude')
            lat_ref = gps.get('GPSLatitudeRef', 'N')
            if lat_dms and hasattr(lat_dms, '__len__') and len(lat_dms) == 3:
                def _to_float(v):
                    if hasattr(v, 'numerator'):           # IFDRational
                        return v.numerator / v.denominator if v.denominator else 0.0
                    if isinstance(v, tuple) and len(v) == 2:  # (num, den) piexif
                        return v[0] / v[1] if v[1] else 0.0
                    return float(v)
                d, m, s = [_to_float(x) for x in lat_dms]
                dd = d + m / 60.0 + s / 3600.0
                if abs(dd) < 0.001:  # FIX-18d: corrupted (0,0,0) GPS block
                    return None
                return -dd if str(lat_ref).upper() in ('S', 'SOUTH') else dd

        except Exception as exc:
            log.debug(f"[EXIF] _extract_drone_lat error: {exc}")
        return None

    def _extract_drone_lon(self) -> Optional[float]:
        """
        FIX-17b: Extract drone GPS longitude as a decimal-degree float.

        Mirrors _extract_drone_lat exactly; handles ExifTool decimal,
        PIL IFDRational DMS, and piexif (num, den) tuple formats.
        Returns decimal degrees (positive = E, negative = W) or None.
        """
        try:
            gps = self.raw_exif.get('GPS', {})

            # Path 1: ExifTool decimal degrees
            for key in ('GPSLongitude', 'Longitude'):
                val = gps.get(key) or self.raw_exif.get(key)
                if val is None:
                    continue
                try:
                    dd = float(val)
                    ref = (gps.get('GPSLongitudeRef') or
                           self.raw_exif.get('GPSLongitudeRef') or 'E')
                    return -dd if str(ref).upper() in ('W', 'WEST') else dd
                except (ValueError, TypeError):
                    pass

            # Path 2 / 3: DMS tuple
            lon_dms = gps.get('GPSLongitude')
            lon_ref = gps.get('GPSLongitudeRef', 'E')
            if lon_dms and hasattr(lon_dms, '__len__') and len(lon_dms) == 3:
                def _to_float(v):
                    if hasattr(v, 'numerator'):
                        return v.numerator / v.denominator if v.denominator else 0.0
                    if isinstance(v, tuple) and len(v) == 2:
                        return v[0] / v[1] if v[1] else 0.0
                    return float(v)
                d, m, s = [_to_float(x) for x in lon_dms]
                dd = d + m / 60.0 + s / 3600.0
                if abs(dd) < 0.001:  # FIX-18d: corrupted (0,0,0) GPS block
                    return None
                return -dd if str(lon_ref).upper() in ('W', 'WEST') else dd

        except Exception as exc:
            log.debug(f"[EXIF] _extract_drone_lon error: {exc}")
        return None

    def _extract_absolute_altitude(self) -> Optional[float]:
        """
        FIX-17b: Extract the drone's absolute altitude MSL (metres).

        DJI writes this as drone-dji:AbsoluteAltitude in the XMP block.
        _parse_dji_xmp already extracts it into raw_exif['XMP'] under the key
        'AbsoluteAltitude' so we just read from there first.

        Secondary path: ExifTool with -ee flag may also surface the value at
        the top level of raw_exif.  Standard GPS IFD 'GPSAltitude' is NOT the
        same — it is the WGS-84 ellipsoid height (often equals MSL on flat
        terrain but differs on hilly sites), so it is only used as a last resort.

        Returns metres (float) or None.
        """
        # Primary: XMP block populated by _parse_dji_xmp
        if 'XMP' in self.raw_exif:
            xmp = self.raw_exif['XMP']
            val = xmp.get('AbsoluteAltitude')
            if val is not None:
                try:
                    s = str(val).replace('+', '')
                    return float(s)
                except (ValueError, TypeError):
                    pass

        # Secondary: top-level ExifTool key
        for key in ('AbsoluteAltitude', 'drone-dji:AbsoluteAltitude'):
            val = self.raw_exif.get(key)
            if val is not None:
                try:
                    s = str(val).replace('+', '')
                    return float(s)
                except (ValueError, TypeError):
                    pass

        # Last resort: standard GPS IFD altitude (WGS-84 ellipsoid, not MSL —
        # acceptable when site is near sea level, flagged in log for awareness)
        gps_alt = self._extract_gps_altitude()
        if gps_alt is not None:
            log.debug(
                "[EXIF] AbsoluteAltitude not in XMP — falling back to GPS IFD "
                "altitude (WGS-84 ellipsoid, not true MSL). Acceptable on flat "
                "terrain; may introduce error on hilly sites.")
            return gps_alt

        return None

    def _lookup_camera_database(self, camera_model: str) -> Optional[dict]:
        """
        Enhanced database lookup with multiple matching strategies:
        1. Exact model match
        2. Fuzzy model name match
        3. Resolution + focal length match (when model unknown)
        """
        if not camera_model or camera_model == "Unknown":
            # Try matching by image resolution and focal length
            image_width, image_height = self._extract_image_size()
            focal_length = self._extract_focal_length()
            
            if image_width and image_height and focal_length:
                log.info(f"[EXIF] Attempting resolution+focal match: {image_width}x{image_height}, {focal_length}mm")
                for db_model, specs in DJI_CAMERA_DATABASE.items():
                    db_res = specs.get('resolution', (0, 0))
                    db_focal = specs.get('focal_length_mm', 0)
                    
                    # Check if resolution matches (within 1% tolerance)
                    res_match = (abs(db_res[0] - image_width) / image_width < 0.01 and 
                                abs(db_res[1] - image_height) / image_height < 0.01)
                    
                    # Check if focal length matches (within 0.5mm tolerance)
                    focal_match = abs(db_focal - focal_length) < 0.5
                    
                    if res_match and focal_match:
                        log.info(f"[EXIF] ✓ Matched by resolution+focal to: {db_model} ({specs['name']})")
                        return specs
                    elif res_match:
                        # Resolution match is stronger indicator
                        log.info(f"[EXIF] ✓ Matched by resolution to: {db_model} ({specs['name']})")
                        return specs
            
            return None
        
        # Try exact match first
        if camera_model in DJI_CAMERA_DATABASE:
            log.info(f"[EXIF] ✓ Exact match found: {camera_model}")
            return DJI_CAMERA_DATABASE[camera_model]
        
        # Try fuzzy matching
        camera_upper = camera_model.upper()
        for db_model, specs in DJI_CAMERA_DATABASE.items():
            if db_model.upper() in camera_upper or camera_upper in db_model.upper():
                log.info(f"[EXIF] ✓ Fuzzy matched '{camera_model}' to '{db_model}'")
                return specs
            if specs['name'].upper() in camera_upper or camera_upper in specs['name'].upper():
                log.info(f"[EXIF] ✓ Fuzzy matched '{camera_model}' to '{specs['name']}'")
                return specs
        
        return None

# ── XMP Diagnostic Utility ─────────────────────────────────────────────────────

def diagnose_xmp_extraction(image_path: str) -> Dict[str, Any]:
    """
    Diagnostic utility to test XMP extraction on a single image.
    Returns detailed information about what was found and where.
    
    Usage:
        result = diagnose_xmp_extraction("path/to/DJI_image.JPG")
        print(result['summary'])
    """
    from pathlib import Path
    
    result = {
        'image': Path(image_path).name,
        'xmp_found': False,
        'xmp_method': None,
        'xmp_fields': {},
        'standard_exif': {},
        'summary': []
    }
    
    try:
        img = Image.open(image_path)
        
        # Check standard EXIF
        exif = img.getexif()
        if exif:
            for tag_id, value in exif.items():
                tag_name = TAGS.get(tag_id, tag_id)
                if tag_name in ['Model', 'Make', 'FocalLength', 'DateTimeOriginal']:
                    result['standard_exif'][tag_name] = str(value)[:50]
        
        result['summary'].append(f"✓ Standard EXIF tags: {len(result['standard_exif'])}")
        
        # Method 1: Standard XMP key
        xmp_string = img.info.get('XML:com.adobe.xmp')
        if xmp_string:
            result['xmp_found'] = True
            result['xmp_method'] = 'PIL standard key'
            result['summary'].append("✓ XMP found via PIL standard key")
        
        # Method 2: Alternative keys
        if not xmp_string:
            for key in img.info.keys():
                if 'xmp' in key.lower() or 'xml' in key.lower():
                    xmp_string = img.info.get(key)
                    if xmp_string:
                        result['xmp_found'] = True
                        result['xmp_method'] = f'PIL alternative key: {key}'
                        result['summary'].append(f"✓ XMP found via alternative key: {key}")
                        break
        
        # Method 3: Raw file parsing
        if not xmp_string:
            with open(image_path, 'rb') as f:
                data = f.read()
                xmp_start = data.find(b'<x:xmpmeta')
                xmp_end = data.find(b'</x:xmpmeta>')
                if xmp_start != -1 and xmp_end != -1:
                    xmp_string = data[xmp_start:xmp_end + 12].decode('utf-8', errors='ignore')
                    result['xmp_found'] = True
                    result['xmp_method'] = 'Raw file parsing'
                    result['summary'].append("✓ XMP found via raw file parsing")
        
        # Parse XMP if found
        if xmp_string:
            if isinstance(xmp_string, bytes):
                xmp_string = xmp_string.decode('utf-8', errors='ignore')
            
            patterns = {
                'RelativeAltitude': r'drone-dji:RelativeAltitude="([^"]+)"',
                'AbsoluteAltitude': r'drone-dji:AbsoluteAltitude="([^"]+)"',
                'GimbalPitchDegree': r'drone-dji:GimbalPitchDegree="([^"]+)"',
                'GimbalYawDegree': r'drone-dji:GimbalYawDegree="([^"]+)"',
                'CameraModel': r'drone-dji:CameraModel="([^"]+)"',
            }
            
            for key, pattern in patterns.items():
                match = re.search(pattern, xmp_string)
                if match:
                    result['xmp_fields'][key] = match.group(1)
            
            result['summary'].append(f"✓ DJI XMP fields extracted: {len(result['xmp_fields'])}")
            for key, val in result['xmp_fields'].items():
                result['summary'].append(f"  - {key}: {val}")
        else:
            result['summary'].append("✗ NO XMP data found in image")
            result['summary'].append("  This image may not be from a DJI drone,")
            result['summary'].append("  or XMP data was stripped during processing.")
        
    except Exception as e:
        result['summary'].append(f"✗ Error during diagnosis: {e}")
    
    return result

# ── End of EXIF Calibration Module ──────────────────────────────────────────

# ==============================================================================
# PHASE 3.6 — EXIF GPS AUTO-READ  (Tom K.)
# Uses Pillow ExifTags only — no piexif dependency required.
# ==============================================================================

def _read_exif_metadata(filepath: str) -> Dict[str, str]:
    """
    Tom K.: Extract drone EXIF metadata from JPEG/TIFF using Pillow.
    Returns dict with keys: altitude_m, date_taken, heading, gps_coords.
    Returns empty dict on any failure (graceful degradation).
    """
    result: Dict[str, str] = {}
    try:
        from PIL import Image as _PILImage
        from PIL.ExifTags import TAGS, GPSTAGS
        img = _PILImage.open(filepath)
        raw_exif = img._getexif()  # type: ignore[attr-defined]
        if not raw_exif:
            return result
        exif: Dict[str, Any] = {TAGS.get(k, k): v for k, v in raw_exif.items()}

        # Date/time
        dt_str = exif.get("DateTimeOriginal") or exif.get("DateTime")
        if dt_str:
            result["date_taken"] = str(dt_str)

        # GPS block
        gps_raw = exif.get("GPSInfo")
        if gps_raw and isinstance(gps_raw, dict):
            gps: Dict[str, Any] = {GPSTAGS.get(k, k): v for k, v in gps_raw.items()}

            def _dms_to_dd(dms, ref) -> float:
                """Convert degrees/minutes/seconds tuple to decimal degrees."""
                d, m, s = dms
                # Values may be IFDRational or float tuples
                def _frac(v):
                    if hasattr(v, "numerator"):
                        return v.numerator / v.denominator if v.denominator else 0.0
                    return float(v)
                dd = _frac(d) + _frac(m) / 60 + _frac(s) / 3600
                if ref in ("S", "W"):
                    dd = -dd
                return dd

            lat_v = gps.get("GPSLatitude")
            lat_r = gps.get("GPSLatitudeRef", "N")
            lon_v = gps.get("GPSLongitude")
            lon_r = gps.get("GPSLongitudeRef", "E")
            if lat_v and lon_v:
                lat = _dms_to_dd(lat_v, lat_r)
                lon = _dms_to_dd(lon_v, lon_r)
                result["gps_coords"] = f"{lat:.6f}, {lon:.6f}"

            alt_v = gps.get("GPSAltitude")
            if alt_v is not None:
                alt_m = (alt_v.numerator / alt_v.denominator
                         if hasattr(alt_v, "numerator") else float(alt_v))
                result["altitude_m"] = f"{alt_m:.1f} m"

            img_dir = gps.get("GPSImgDirection")
            if img_dir is not None:
                hdg = (img_dir.numerator / img_dir.denominator
                       if hasattr(img_dir, "numerator") else float(img_dir))
                result["heading"] = f"{hdg:.1f}°"

    except Exception as exc:
        log.debug(f"EXIF read skipped for {filepath}: {exc}")
    return result

# NOTE: APP_VERSION canonical definition is at the top of this file near the
# other module-level constants.  The v3.3.11 stub that previously lived here
# was a legacy leftover and has been removed to avoid the duplicate-constant
# confusion that caused wrong version strings in log output.

# ==============================================================================
# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║              V2.0 DESIGN IDEAS  (for future implementation)                ║
# ║  These notes document the planned next-generation UI and architecture.     ║
# ║  The CURRENT v1.x UI is preserved as-is.  Only the items listed in the     ║
# ║  v1.7.0 changelog above are implemented in this file.                      ║
# ╚══════════════════════════════════════════════════════════════════════════════╝
#
# ── FLIR QCViewer-style full-screen annotator ──────────────────────────────────
#
#  REMOVE the left thumbnail dock.  Replace the 3-panel layout with:
#
#  ┌────────────────────────────────────────────────────────────────┐
#  │ TOOLBAR  ◀ Prev  |  Next ▶  |  Mode  |  💾 Save [Space]      │  44 px
#  ├────────────────────────────────────────────────────────────────┤
#  │ INFO: filename  |  N / M  |  ✅ ANNOTATED  |  blade / face   │  28 px
#  ├──────────────────────────────┬─────────────────────────────────┤
#  │                              │  ANNOTATION PANEL               │
#  │   FULL IMAGE VIEWER          │  • Defect type                  │
#  │   (fills all remaining       │  • Severity [1][2][3][4][5]     │
#  │    space)                    │  • Distance from root           │
#  │                              │  • Size (auto GSD)              │
#  │                              │  • Comments / Remedy            │
#  │                              │  [Save]  [Delete]               │
#  │                              │  ─────────────────────────      │
#  │                              │  ANNOTATED badge  ✅ / ⬜       │
#  │                              │  RENAME field + button          │
#  ├──────────────────────────────┴─────────────────────────────────┤
#  │ STATUS: Blade A PS  |  2/47 annotated  |  hints               │  26 px
#  └────────────────────────────────────────────────────────────────┘
#
#  Reference implementation: flir_thermal_app_v3_30_19_3.py — QCViewer class
#  (lines 10248–10680 in that file) for the full-screen viewer pattern.
#
# ── Three-stage pipeline (future role-based workflow) ─────────────────────────
#
#  STAGE 1 — ANNOTATION  (Inspector role)
#    • Load folder organised by component (BladeA_PS/, Hub/, Tower/ etc.)
#    • FLIR-style viewer: navigate image by image
#    • Draw shape, fill panel, Space → save + advance
#    • File auto-renamed on first save of each image
#    • Progress counter: "47 images — 23 annotated — 24 remaining"
#    • Export: annotation CSV + project JSON
#
#  STAGE 2 — REVIEW  (Reviewer role)
#    • Same viewer but shows existing annotations with overlays
#    • Approve ✓ / Reject ✗ / Edit values inline
#    • Distance-from-root editable if auto-estimate was used
#    • Blade serial assignment (A→847, B→852, C→854)
#    • Only APPROVED annotations go to the report
#
#  STAGE 3 — REPORT  (Reviewer / Admin role)
#    • Generate PDF matching SenseHawk format exactly
#    • Per-blade summary + individual annotation pages
#    • Blade diagram with accurate defect dot positions
#    • Training data export (wind turbine YOLO model)
#
#  The existing Inspector / Reviewer / Admin role system maps directly.
#  Implement as QTabWidget: Annotate | Review | Report
#
# ── Auto-rename convention ─────────────────────────────────────────────────────
#
#  Pattern: {Component}_{Face}_{FirstDefect}_{Sequence:03}.jpg
#  Example: DJI_0050.JPG → BladeA_PS_Erosion_001.jpg
#
#  project.json stores old_name → new_name mapping for full traceability.
#  _parse_blade_face_from_folder() (line ~651) already handles folder parsing.
#
# ── SenseHawk report format remaining gaps ────────────────────────────────────
#
#  | SenseHawk feature              | Current app      | Fix needed              |
#  |--------------------------------|------------------|-------------------------|
#  | Numbered defect list (p.2)     | Missing          | Add to cover page       |
#  | Blade serial table A/B/C       | Dataclass exists | Wire to report cover    |
#  | 2 images per defect (wide+crop)| Single image     | Auto-crop from bbox     |
#  | Blade diagram accurate dots    | ✅ FIXED v1.7.0  | —                       |
#  | MR (distance from root)        | ✅ FIXED v1.7.0  | —                       |
#  | Remedy Action per annotation   | Partial          | Ensure field flows thru |
#
# ── Training dataset for wind turbine YOLO model ──────────────────────────────
#
#  Classes: Erosion, Crack, Delamination, Contamination, LETapeDamaged,
#           OldRepairPatch, PaintDecayed, Damage, Missing, PaintSpillage
#
#  Renamed files (BladeA_PS_Erosion_001.jpg) enable automatic class labelling
#  from filename.  Augmentation: horizontal flip only (vertical flip invalid —
#  orientation matters for distance-from-root measurement).
#  Recommended confidence threshold: 0.35 (blade defects are subtle).
#
# ==============================================================================
APP_TITLE      = "Wind Tower Inspection"
APP_VERSION    = "4.5.0"  """ v4.5.0: Face+Zone split + calibration cleanup
                          #   • Annotation.size_locked field — when True, auto-calibrate                          #     NEVER overwrites width_cm/height_cm (set on Save with                          #     manual/image GSD source; set on reviewer Approve).                          #   • Annotation.size_calibrated_at — ISO timestamp of last                          #     size write; allows load_project() migration to detect                          #     which calibration generation produced the stored values.                          #   • _ann_from_dict migration: GSD > 20 cm/px on gsd_source=                          #     "image" annotations → clears size_calibrated_at so next                          #     auto-calibrate recalculates (fixes 160×6000 phantom).                          #   • _batch_auto_calibrate: locked_count + migrated_count                          #     tracked; shown in results dialog.                          #   • _recompute_annotation_sizes: also respects size_locked                          #     and stamps size_calibrated_at (session GSD path).                          #          tags so the 160px JPEG thumbnail width never overwrites the full-                          #          resolution ImageWidth; first-wins dict policy added so earlier tags                          #          prevail. GPS IFD now parsed into nested 'GPS' dict (mirrors PIL path).                          #          FIX-EXIF-B — calibrate() two-phase: phase-1 first-success as before;                          #          phase-2 always runs PIL as XMP supplement when primary was not PIL —                          #          exifread/piexif cannot read XMP at all; DJI RelativeAltitude /                          #          GimbalPitchDegree / AbsoluteAltitude / drone GPS live only in XMP.                          #          FIX-EXIF-C — _extract_image_size now always uses PIL as authoritative                          #          source; old 'if not width' guard was fooled by truthy-but-wrong                          #          thumbnail width (160) producing 160×6000 phantom dimensions.                          #          FIX-EXIF-D — _extract_pil no longer raises ImportError when                          #          PIL_EXIF_AVAILABLE=False; degrades gracefully so XMP supplement
                          #          still works; image handle now closed via context manager.
                          #          AnnotationPanel._rev_note/_status_lbl/_approve_btn/_reject_btn
                          #          AttributeError (load_existing/_on_approve/_on_reject) — widgets
                          #          moved to QCReviewPanel in v4.1.0 but references not updated;
                          #          ExifTool parser now handles IFD0/ExifIFD/Main/MakerNotes groups
                          #          so camera Make/Model/FocalLength are no longer silently dropped;
                          #          _extract_relative_altitude duplicate key 'RelativeAltitude' fixed;
                          #          _extract_sensor_size FocalPlaneXResolution misuse (DPI≠mm) fixed;
                          #          DJI_CAMERA_DATABASE expanded (Mavic Air 2S FC4170, Mini 3 Pro
                          #          FC3582, Mini 4 Pro FC4380, Air 3 FC3980, Phantom 4 RTK FC6360,
                          #          Inspire 3 FC3688, Zenmuse P1, H20T, Mini 2 FC7203, Autel EVO II)"""
SCRIPT_DIR     = Path(__file__).resolve().parent
SETTINGS_FILE  = SCRIPT_DIR / "settings.ini"

# ==============================================================================
# CTO-AUDIT: Structured logging — rotating file handler + crash dump capture
# Must be after SCRIPT_DIR so we can place the log file next to the script.
# ==============================================================================

_LOG_FILE = SCRIPT_DIR / "wind_tower_app.log"

class _QtNoiseFilter(logging.Filter):
    """
    Suppress high-volume DEBUG chatter from Qt internals and third-party libs
    that are not actionable for application debugging.  Applied to the console
    handler only — the file handler always receives everything so post-mortem
    analysis has full detail.

    Suppressed namespaces:
      • PyQt6.*  — internal Qt signal/slot bookkeeping
      • PIL.*    — Pillow image-decode step-by-step messages
      • fontTools.*  — font parsing (triggered by ReportLab)
      • matplotlib.* — only present if charting libs installed
    """
    _NOISY_PREFIXES = ("PyQt6.", "PIL.", "fontTools.", "matplotlib.")

    def filter(self, record: logging.LogRecord) -> bool:  # type: ignore[override]
        if record.levelno >= logging.INFO:
            return True   # always show INFO and above regardless of source
        return not any(record.name.startswith(p) for p in self._NOISY_PREFIXES)


def _setup_logging() -> logging.Logger:
    """
    Configure structured, multi-handler logging for the Wind Tower app.

    Handler layout
    ──────────────
    ① Console (stdout)
        Level   : INFO  (DEBUG shown only for WindTowerApp logger)
        Format  : HH:MM:SS [LEVEL   ] name — message
        Filter  : _QtNoiseFilter suppresses Qt/PIL/fontTools DEBUG spam
        Purpose : Real-time feedback for the inspector running the script.

    ② Rotating file  (wind_tower_app.log, 5 MB × 5 backups)
        Level   : DEBUG  (everything, no filter)
        Format  : YYYY-MM-DD HH:MM:SS [LEVEL   ] name — message
        Purpose : Full post-mortem trace including EXIF extraction steps,
                  calibration maths, save timing, and signal routing.

    The function is idempotent — safe to call twice without duplicating
    handlers (checked by handler-type guard, not just empty-handlers check).
    A session-start banner is written to the file handler so each run is
    clearly delimited in the log file when multiple sessions are appended.
    """
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.DEBUG)    # capture everything; handlers filter

    # ── Per-module level overrides ────────────────────────────────────────────
    # Suppress DEBUG from very chatty but unactionable third-party namespaces
    # even in the file log (they pollute grep output without adding value).
    for _noisy_ns in ("PyQt6", "PIL.TiffImagePlugin", "fontTools"):
        logging.getLogger(_noisy_ns).setLevel(logging.WARNING)

    # ── Console handler (INFO + above, Qt noise filtered) ─────────────────────
    _has_console = any(
        isinstance(h, logging.StreamHandler) and not isinstance(h, RotatingFileHandler)
        for h in root_logger.handlers)
    if not _has_console:
        ch = logging.StreamHandler(sys.stdout)
        ch.setLevel(logging.INFO)
        ch.setFormatter(logging.Formatter(
            "%(asctime)s [%(levelname)-8s] %(name)s — %(message)s",
            datefmt="%H:%M:%S"))
        ch.addFilter(_QtNoiseFilter())
        root_logger.addHandler(ch)

    # ── Rotating file handler (DEBUG, no filter — full trace) ─────────────────
    _has_file = any(isinstance(h, RotatingFileHandler) for h in root_logger.handlers)
    if not _has_file:
        try:
            fh = RotatingFileHandler(
                str(_LOG_FILE), maxBytes=5 * 1024 * 1024, backupCount=5,
                encoding="utf-8")
            fh.setLevel(logging.DEBUG)
            fh.setFormatter(logging.Formatter(
                "%(asctime)s [%(levelname)-8s] %(name)s — %(message)s",
                datefmt="%Y-%m-%d %H:%M:%S"))
            root_logger.addHandler(fh)
            # Session-start banner — makes multi-session log files easy to navigate
            _sep = "═" * 72
            fh.stream.write(
                f"\n{_sep}\n"
                f"  Wind Tower Inspection App  v{APP_VERSION}  —  Session start\n"
                f"  {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
                f"  Log file : {_LOG_FILE}\n"
                f"  Python   : {sys.version.split()[0]}   "
                f"PID: {os.getpid()}\n"
                f"{_sep}\n\n")
            fh.stream.flush()
        except Exception as _le:
            print(f"[WARN] Could not open log file {_LOG_FILE}: {_le}"
                  f"  — file logging disabled for this session.")

    return logging.getLogger("WindTowerApp")


def _install_crash_handler():
    """CTO-AUDIT: Capture unhandled exceptions to log file + timestamped crash dump."""
    def _excepthook(exc_type, exc_value, exc_tb):
        crash_msg = "".join(traceback.format_exception(exc_type, exc_value, exc_tb))
        log.critical(f"UNHANDLED EXCEPTION:\n{crash_msg}")
        crash_path = SCRIPT_DIR / f"crash_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
        try:
            crash_path.write_text(
                f"Wind Tower Inspection App — Crash Report\n"
                f"Version: {APP_VERSION}\n"
                f"Time:    {datetime.now().isoformat()}\n\n{crash_msg}",
                encoding="utf-8")
        except Exception:
            pass
        sys.__excepthook__(exc_type, exc_value, exc_tb)
    sys.excepthook = _excepthook


log = _setup_logging()   # upgrade from placeholder to full rotating handler

# ==============================================================================
# UI THEME — "Precision Instrument" dark aesthetic
# Dev Patel: Single source of truth so the whole app looks consistent.
# Colours derived from FLIR thermal professional edition palette + Scopito branding.
# ==============================================================================

UI_THEME: Dict[str, str] = {
    "bg_primary":    "#0d1117",   # deepest background — near-black navy
    "bg_secondary":  "#161b22",   # panels, cards
    "bg_card":       "#1c2333",   # raised card surfaces
    "bg_elevated":   "#21262d",   # inputs, dropdowns
    "bg_input":      "#21262d",   # input field backgrounds (alias for bg_elevated)
    "bg_toolbar":    "#13181f",   # toolbar strip (slightly darker than secondary)
    "accent_cyan":   "#00d4e0",   # primary action / active mode
    "accent_green":  "#3fb950",   # success / approve
    "accent_amber":  "#d29922",   # ML / model ops
    "accent_orange": "#f78166",   # warnings
    "accent_red":    "#f85149",   # errors / critical
    "accent_blue":   "#388bfd",   # info / QC
    "accent_purple": "#bc8cff",   # export / PDF
    "accent_pink":   "#f778ba",   # QC review
    "text_primary":  "#e6edf3",
    "text_secondary":"#7d8590",
    "text_tertiary": "#484f58",
    "text_accent":   "#00d4e0",
    "border":        "#30363d",
    "border_focus":  "#00d4e0",
    "sev1":          "#FFD700",   # Minor (Sev 1) — GOLD (was green)
    "sev2":          "#7ee787",   # Sev 2 — light green (deprecated)
    "sev3":          "#FFA500",   # Major (Sev 3) — AMBER/ORANGE
    "sev4":          "#f78166",   # Sev 4 — orange-red (deprecated)
    "sev5":          "#FF0000",   # Critical (Sev 5) — BRIGHT RED (was pink-red)
    "poi":           "#388bfd",   # POI — blue
}

DARK_STYLESHEET = f"""
    QMainWindow, QDialog, QWidget {{
        background-color: {UI_THEME['bg_primary']};
        color: {UI_THEME['text_primary']};
        font-family: 'Segoe UI', 'SF Pro Display', 'Helvetica Neue', Arial, sans-serif;
        font-size: 10pt;
    }}
    QMenuBar {{
        background-color: {UI_THEME['bg_toolbar']};
        color: {UI_THEME['text_primary']};
        border-bottom: 1px solid {UI_THEME['border']};
        padding: 1px 0;
    }}
    QMenuBar::item:selected {{ background-color: {UI_THEME['bg_card']}; border-radius: 3px; }}
    QMenu {{
        background-color: {UI_THEME['bg_secondary']};
        color: {UI_THEME['text_primary']};
        border: 1px solid {UI_THEME['border']};
        border-radius: 4px;
        padding: 4px;
    }}
    QMenu::item {{ padding: 5px 20px 5px 12px; border-radius: 3px; }}
    QMenu::item:selected {{ background-color: {UI_THEME['bg_card']}; color: {UI_THEME['accent_cyan']}; }}
    QMenu::separator {{ height: 1px; background: {UI_THEME['border']}; margin: 4px 8px; }}
    QToolBar {{
        background-color: {UI_THEME['bg_toolbar']};
        border-bottom: 2px solid {UI_THEME['border']};
        spacing: 2px;
        padding: 3px 6px;
    }}
    QPushButton {{
        background-color: {UI_THEME['bg_card']};
        color: {UI_THEME['text_primary']};
        border: 1px solid {UI_THEME['border']};
        border-radius: 6px;
        padding: 5px 12px;
        font-weight: 600;
    }}
    QPushButton:hover {{
        background-color: {UI_THEME['bg_elevated']};
        border-color: {UI_THEME['accent_cyan']};
        color: {UI_THEME['accent_cyan']};
    }}
    QPushButton:checked {{
        background-color: {UI_THEME['accent_cyan']};
        color: #0d1117;
        border-color: {UI_THEME['accent_cyan']};
        font-weight: bold;
    }}
    QPushButton:disabled {{
        background-color: {UI_THEME['bg_secondary']};
        color: {UI_THEME['text_tertiary']};
        border-color: {UI_THEME['border']};
    }}
    QLineEdit, QPlainTextEdit, QTextEdit {{
        background-color: {UI_THEME['bg_elevated']};
        color: {UI_THEME['text_primary']};
        border: 1px solid {UI_THEME['border']};
        border-radius: 6px;
        padding: 4px 8px;
        selection-background-color: {UI_THEME['accent_cyan']};
        selection-color: #0d1117;
    }}
    QLineEdit:focus, QPlainTextEdit:focus, QTextEdit:focus {{
        border-color: {UI_THEME['border_focus']};
        outline: none;
    }}
    QComboBox {{
        background-color: {UI_THEME['bg_elevated']};
        color: {UI_THEME['text_primary']};
        border: 1px solid {UI_THEME['border']};
        border-radius: 6px;
        padding: 4px 8px;
        min-width: 80px;
    }}
    QComboBox:focus {{ border-color: {UI_THEME['border_focus']}; }}
    QComboBox::drop-down {{ border: none; width: 20px; }}
    QComboBox QAbstractItemView {{
        background-color: {UI_THEME['bg_secondary']};
        color: {UI_THEME['text_primary']};
        border: 1px solid {UI_THEME['border']};
        selection-background-color: {UI_THEME['bg_card']};
        selection-color: {UI_THEME['accent_cyan']};
        outline: none;
    }}
    QSpinBox, QDoubleSpinBox {{
        background-color: {UI_THEME['bg_elevated']};
        color: {UI_THEME['text_primary']};
        border: 1px solid {UI_THEME['border']};
        border-radius: 6px;
        padding: 4px 8px;
    }}
    QGroupBox {{
        color: {UI_THEME['text_secondary']};
        border: 1px solid {UI_THEME['border']};
        border-radius: 8px;
        margin-top: 10px;
        padding-top: 10px;
        font-weight: bold;
        font-size: 9pt;
    }}
    QGroupBox::title {{
        subcontrol-origin: margin;
        left: 10px;
        padding: 0 6px;
        color: {UI_THEME['accent_cyan']};
        font-size: 9pt;
        font-weight: bold;
    }}
    QListWidget {{
        background-color: {UI_THEME['bg_secondary']};
        color: {UI_THEME['text_primary']};
        border: 1px solid {UI_THEME['border']};
        border-radius: 6px;
        outline: none;
    }}
    QListWidget::item {{ padding: 4px 8px; border-radius: 4px; }}
    QListWidget::item:selected {{
        background-color: {UI_THEME['bg_card']};
        color: {UI_THEME['accent_cyan']};
    }}
    QListWidget::item:hover {{ background-color: {UI_THEME['bg_elevated']}; }}
    QLabel {{ color: {UI_THEME['text_primary']}; background: transparent; }}
    QSplitter::handle {{ background-color: {UI_THEME['border']}; }}
    QSplitter::handle:horizontal {{ width: 2px; }}
    QSplitter::handle:vertical {{ height: 2px; }}
    QScrollBar:vertical {{
        background: {UI_THEME['bg_secondary']}; width: 8px; border-radius: 4px; margin: 0;
    }}
    QScrollBar::handle:vertical {{
        background: {UI_THEME['bg_elevated']}; border-radius: 4px; min-height: 24px;
    }}
    QScrollBar::handle:vertical:hover {{ background: {UI_THEME['text_tertiary']}; }}
    QScrollBar:horizontal {{
        background: {UI_THEME['bg_secondary']}; height: 8px; border-radius: 4px; margin: 0;
    }}
    QScrollBar::handle:horizontal {{
        background: {UI_THEME['bg_elevated']}; border-radius: 4px; min-width: 24px;
    }}
    QScrollBar::add-line, QScrollBar::sub-line {{ width: 0; height: 0; }}
    QTabWidget::pane {{
        border: 1px solid {UI_THEME['border']};
        background-color: {UI_THEME['bg_secondary']};
        border-radius: 6px;
    }}
    QTabBar::tab {{
        background-color: {UI_THEME['bg_primary']};
        color: {UI_THEME['text_secondary']};
        padding: 6px 16px;
        border: 1px solid {UI_THEME['border']};
        border-bottom: none;
        border-top-left-radius: 6px;
        border-top-right-radius: 6px;
        margin-right: 2px;
        font-weight: 600;
    }}
    QTabBar::tab:selected {{
        background-color: {UI_THEME['bg_secondary']};
        color: {UI_THEME['accent_cyan']};
        border-bottom: 2px solid {UI_THEME['accent_cyan']};
    }}
    QStatusBar {{
        background-color: {UI_THEME['bg_toolbar']};
        color: {UI_THEME['text_secondary']};
        border-top: 1px solid {UI_THEME['border']};
        font-size: 9pt;
    }}
    QDockWidget {{
        border: 1px solid {UI_THEME['border']};
        titlebar-close-icon: none;
    }}
    QDockWidget::title {{
        background: {UI_THEME['bg_secondary']};
        padding: 5px 10px;
        font-weight: bold;
        color: {UI_THEME['text_secondary']};
        border-bottom: 1px solid {UI_THEME['border']};
        text-align: left;
    }}
    QCheckBox {{ color: {UI_THEME['text_primary']}; spacing: 8px; }}
    QCheckBox::indicator {{
        width: 16px; height: 16px;
        border: 1px solid {UI_THEME['border']};
        border-radius: 4px;
        background-color: {UI_THEME['bg_elevated']};
    }}
    QCheckBox::indicator:checked {{
        background-color: {UI_THEME['accent_cyan']};
        border-color: {UI_THEME['accent_cyan']};
    }}
    QProgressBar {{
        background-color: {UI_THEME['bg_elevated']};
        border: 1px solid {UI_THEME['border']};
        border-radius: 4px;
        text-align: center;
        color: {UI_THEME['text_primary']};
        font-size: 9pt;
    }}
    QProgressBar::chunk {{
        background-color: {UI_THEME['accent_cyan']};
        border-radius: 3px;
    }}
    QSlider::groove:horizontal {{
        background: {UI_THEME['bg_elevated']};
        height: 4px; border-radius: 2px;
    }}
    QSlider::handle:horizontal {{
        background: {UI_THEME['accent_cyan']};
        width: 14px; height: 14px;
        border-radius: 7px; margin: -5px 0;
        border: 2px solid {UI_THEME['bg_primary']};
    }}
    QSlider::handle:horizontal:hover {{ background: #fff; }}
    QToolTip {{
        background-color: {UI_THEME['bg_card']};
        color: {UI_THEME['text_primary']};
        border: 1px solid {UI_THEME['accent_cyan']};
        padding: 4px 8px;
        border-radius: 4px;
        font-size: 9pt;
    }}
    QHeaderView::section {{
        background-color: {UI_THEME['bg_secondary']};
        color: {UI_THEME['text_secondary']};
        padding: 5px 8px;
        border: none;
        border-right: 1px solid {UI_THEME['border']};
        font-weight: bold;
    }}
    QFrame[frameShape="4"] {{ color: {UI_THEME['border']}; }}
    QFrame[frameShape="5"] {{ color: {UI_THEME['border']}; }}
"""

# ── Taxonomy ──────────────────────────────────────────────────────────────────
DEFAULT_DEFECT_TYPES: List[str] = [
    # ── SenseHawk / Scopito primary types (PDF page 2 list — exact order) ──
    "Erosion",
    "Contamination",
    "Paint Decayed",
    "Damage",
    "Missing",
    "Paint Spillage",
    "Old Repair Patch",
    "LE Tape Damaged",
    # ── Extended V1.4.1 taxonomy ─────────────────────────────────────────────
    "Paint / Gelcoat Decayed",
    "Hydraulic Oil Spillage",
    "Scratches on Paint",
    "Crack",
    "Delamination",
    "Debonding",
    "Coating Failure",
    "Surface Contamination",
    "Lightning Strike",
    "Structural Damage",
    "Impact Damage",
    "Pinholing",
    "Fiber Damage",
    "Blade Tip Damage",
    "Bond Line Damage",
    "Dirt or Grease",
    "Point of Interest",
    "Other",
]

# Scopito 1–5 + POI severity scale (industry standard).
# v3.0.1: Severity labels renamed to Minor/Major/Serious/Critical.
# Legacy keys kept in lookup tables so old project JSON files still load and display correctly.
SEVERITY_LEVELS: List[str] = [
    "Minor", "Major", "Critical", "POI",
    # Legacy backward-compat (old JSON files)
    "1 - Cosmetic", "2 - Similar/Cosmetic", "3 - Non-Serious",
    "4 - Serious", "5 - Very Serious",
    "Low", "Medium", "High",
]

# Active levels shown in the UI pill strip
SEVERITY_ACTIVE: List[str] = ["Minor", "Major", "Critical", "POI"]

SEVERITY_COLORS: Dict[str, QColor] = {
    "Minor":    QColor(UI_THEME["sev1"]),   # Severity 1
    "Major":    QColor(UI_THEME["sev3"]),   # Severity 3 & 4
    "Critical": QColor(UI_THEME["sev5"]),   # Severity 5
    "POI":      QColor(UI_THEME["poi"]),
    # Legacy backward-compat
    "1 - Cosmetic":         QColor(UI_THEME["sev1"]),
    "2 - Similar/Cosmetic": QColor(UI_THEME["sev2"]),
    "3 - Non-Serious":      QColor(UI_THEME["sev3"]),
    "4 - Serious":          QColor(UI_THEME["sev3"]),  # Sev 4 → Major colour
    "5 - Very Serious":     QColor(UI_THEME["sev5"]),
    "Serious": QColor(UI_THEME["sev3"]),    # if any old project had "Serious"
    "Low":   QColor(UI_THEME["sev1"]),
    "Medium": QColor(UI_THEME["sev3"]),
    "High":  QColor(UI_THEME["sev3"]),
}

SEVERITY_SHORT: Dict[str, str] = {
    "Minor":    "MINOR",
    "Major":    "MAJOR",
    "Critical": "CRITICAL",
    "POI":      "POI",
    # Legacy
    "1 - Cosmetic": "MINOR",  "2 - Similar/Cosmetic": "MINOR",
    "3 - Non-Serious": "MAJOR", "4 - Serious": "MAJOR",
    "5 - Very Serious": "CRITICAL",
    "Serious": "MAJOR",
    "Low": "MINOR", "Medium": "MAJOR", "High": "MAJOR",
}

SEVERITY_DISPLAY: Dict[str, str] = {
    "Minor":    "Minor",
    "Major":    "Major",
    "Critical": "Critical",
    "POI":      "POI",
    # Legacy
    "1 - Cosmetic": "Minor",  "2 - Similar/Cosmetic": "Minor",
    "3 - Non-Serious": "Major", "4 - Serious": "Major",
    "5 - Very Serious": "Critical",
    "Serious": "Major",
    "Low": "Minor", "Medium": "Major", "High": "Major",
}

SEVERITY_REMEDY: Dict[str, str] = {
    # CHG-D: Updated to match the action column in the reference document severity
    # legend table (same text used in the Defect Summary severity overview band):
    #   POI      → "Needs further investigation"
    #   Minor    → "Monitoring — no intervention"
    #   Major    → "Intervention at planned inspection"
    #   Critical → "Immediate intervention required"
    "Minor":    "Monitoring — no intervention",
    "Major":    "Intervention at planned inspection",
    "Critical": "Immediate intervention required",
    "POI":      "Needs further investigation",
    # Legacy backward-compat — same mapping applied to old severity strings
    "1 - Cosmetic":         "Monitoring — no intervention",
    "2 - Similar/Cosmetic": "Monitoring — no intervention",
    "3 - Non-Serious":      "Intervention at planned inspection",
    "4 - Serious":          "Intervention at planned inspection",
    "5 - Very Serious":     "Immediate intervention required",
    "Serious": "Intervention at planned inspection",
    "Low": "Monitoring — no intervention",
    "Medium": "Intervention at planned inspection",
    "High": "Intervention at planned inspection",
}

# Severity rank for worst-severity logic
SEVERITY_RANK: Dict[Optional[str], int] = {
    None: 0,
    "Minor": 1, "Major": 3, "Critical": 5, "POI": 1,
    "Serious": 3,  # old "Serious" treated as Major
    # Legacy
    "1 - Cosmetic": 1, "2 - Similar/Cosmetic": 2,
    "3 - Non-Serious": 3, "4 - Serious": 3, "5 - Very Serious": 5,
    "Low": 1, "Medium": 2, "High": 3,
}
# v3.0.1: Central hex colour map used by PDF/DOCX generators — keeps all
# hardcoded colour dicts in sync with the renamed severity labels.
_SEV_HEX: Dict[str, str] = {
    # v4.3.0: Aligned with UI_THEME sev keys — Yellow / Amber / Red
    "Minor":    "#FFD700",   # Yellow (was green #3fb950)
    "Major":    "#FFA500",   # Amber  (was dark-amber #d29922)
    "Critical": "#FF0000",   # Red    (was pink-red #f85149)
    "POI":      "#388bfd",   # Blue — unchanged
    # Legacy keys — kept so old saved projects render correctly
    "1 - Cosmetic":         "#FFD700",
    "2 - Similar/Cosmetic": "#FFD700",
    "3 - Non-Serious":      "#FFA500",
    "4 - Serious":          "#FFA500",
    "5 - Very Serious":     "#FF0000",
    "Serious": "#FFA500",
    "Low":  "#FFD700", "Medium": "#FFA500", "High": "#FF0000",
}


# Auto-remedy map keyed on lower-cased partial defect name (Scopito SenseHawk PDF)
# T09 FIX: Repair-type defects use the consistent pattern:
#   "{DefectType} repair recommended during the next planned inspection."
# Exceptions: entries that require immediate action or a different verb keep their text.
_AUTO_REMEDY: Dict[str, str] = {
    "contamination":   "Cleaning & Washing is recommended in the next available opportunity.",
    "hydraulic oil":   "Cleaning & Washing is recommended in the next available opportunity.",
    "dirt":            "Cleaning & Washing is recommended in the next available opportunity.",
    "missing":         "Replacement is recommended during the next planned inspection.",
    "damage":          "Damage repair recommended during the next planned inspection.",
    "structural":      "Immediate structural assessment required.",
    "le tape":         "LE tape repair recommended during the next planned inspection.",
    "old repair":      "No action required at this time. Monitor periodically.",
    "paint spillage":  "No action required at this time. Monitor periodically.",
    "paint decayed":   "Paint decayed repair recommended during the next planned inspection.",
    "gelcoat":         "Gelcoat repair recommended during the next planned inspection.",
    "coating failure": "Coating failure repair recommended during the next planned inspection.",
    "erosion":         "Erosion repair recommended during the next planned inspection.",
    "crack":           "Structural inspection and repair required urgently.",
    "delamination":    "Delamination repair recommended during the next planned inspection.",
    "pinholing":       "Pinholing repair recommended during the next planned inspection.",
    "scratches":       "Scratches repair recommended during the next planned inspection.",
    "lightning":       "Immediate lightning protection system inspection required.",
    "fiber":           "Fiber repair recommended during the next planned inspection.",
    "blade tip":       "Blade tip repair recommended during the next planned inspection.",
    "bond line":       "Bond line repair recommended during the next planned inspection.",
}

def _parse_blade_face_from_folder(folder_name: str) -> Tuple[str, str]:
    """
    FOLDER-AUTO: Auto-detect blade (A/B/C/Hub/Tower) and face (PS/LE/TE/SS) from folder/filename.
    Supports patterns like: "Blade A PS", "Blade_A_LE", "Hub", "Tower", "Hub_001", etc.
    Returns ("", "") if nothing could be detected.
    """
    import re
    s = folder_name.upper().replace("-", " ").replace("_", " ")

    # Detect Hub or Tower first (takes priority over blade letter search)
    if re.search(r'\bHUB\b', s):
        return "Hub", ""
    if re.search(r'\bTOWER\b', s):
        return "Tower", ""

    # Detect blade letter A/B/C
    blade = ""
    # Try "BLADE X" pattern first
    m = re.search(r'\bBLADE\s*([ABC])\b', s)
    if m:
        blade = m.group(1)
    else:
        # Try standalone A/B/C not surrounded by other letters
        m = re.search(r'(?<![A-Z])([ABC])(?![A-Z])', s)
        if m:
            blade = m.group(1)

    # Detect face abbreviation
    face = ""
    for candidate in ["PS", "LE", "TE", "SS"]:
        if re.search(r'\b' + candidate + r'\b', s):
            face = candidate
            break

    return blade, face

def _auto_remedy(defect_type: str) -> str:
    """
    T09 FIX: Return recommended remedy text based on defect type.
    Pattern: "{DefectType} repair recommended during the next planned inspection."
    Special cases (contamination, structural, etc.) retain custom text from _AUTO_REMEDY.
    """
    dt = defect_type.lower()
    for key, remedy in _AUTO_REMEDY.items():
        if key in dt:
            return remedy
    # Fallback: title-case the defect name into the standard pattern
    _title = defect_type.strip().title() if defect_type.strip() else "Defect"
    return f"{_title} repair recommended during the next planned inspection."

BLADE_SPANS = ["Root (0–33%)", "Mid (33–66%)", "Tip (66–100%)"]
# v4.5.0: Split into mutually exclusive Surface and Zone selections
BLADE_SURFACES = [
    "Pressure Side (PS)", "Suction Side (SS)"
]
BLADE_ZONES = [
    "Leading Edge (LE)", "Trailing Edge (TE)", "Midbody (MB)"
]
# BLADE-ABC: Industry standard uses Blade A/B/C. Changed from B1/B2/B3 throughout.
BLADE_NAMES = ["A", "B", "C", "Hub", "Tower"]

# ── Drawing constants ─────────────────────────────────────────────────────────
# FIX-LINES: Increased pen widths — thin 2px lines were invisible on drone images.
# Box/polygon annotations now use 4px; calibration guide 2px; rubber-band 3px.
BOX_PEN_WIDTH    = 4.0   # was 2.0 — saved box annotation border
PIN_RADIUS       = 10    # was 8   — pin dot slightly larger
CAL_PEN_WIDTH    = 2.0   # was 1.5
POLY_PEN_WIDTH   = 4.0   # was 2.0 — saved polygon annotation border
DRAW_PEN_WIDTH   = 3.0   # new     — rubber-band / in-progress line while drawing
LABEL_FONT_SIZE  = 9     # was 8   — slightly larger label text

# ── Cell layout for blade diagram ─────────────────────────────────────────────
_SPANS = ["Root", "Mid", "Tip"]
_FACES = ["LE", "TE", "PS", "SS", "ED"]  # v4.2.0: Added ED (Edge)

# ==============================================================================
# ==============================================================================
# PHASE 2 — LOGIN & ROLE SYSTEM  (Chris Murphy + Jamie Liu)
# ==============================================================================
# Roles:
#   Inspector  — create/edit own annotations; cannot approve/reject; cannot generate report
#   Reviewer   — approve/reject any annotation; generate final PDF report
#   Admin      — all permissions + manage users + switch roles in session
#
# Credentials live in settings.ini [USERS] as: username = role,hex_sha256_password
# ==============================================================================

# NOTE: hashlib is already imported at module level (line 99).

class AppSession:
    """
    Jamie Liu: Singleton runtime session.  Stores the logged-in user and role.
    Provides `can_do(action)` for centralised permission checks.
    """
    _instance: Optional["AppSession"] = None

    def __init__(self):
        self.username : str = ""
        self.role     : str = ""

    @classmethod
    def instance(cls) -> "AppSession":
        if cls._instance is None:
            cls._instance = AppSession()
        return cls._instance

    def login(self, username: str, role: str):
        self.username = username
        self.role     = role

    def logout(self):
        self.username = ""; self.role = ""

    def can_do(self, action: str) -> bool:
        return True   # No sign-in required — all actions permitted

    def is_logged_in(self) -> bool:
        return bool(self.username and self.role)


SESSION = AppSession.instance()



# ==============================================================================
# DATA MODELS  (Tom K. — 14yr Data Engineering)
# ==============================================================================

def _migrate_severity(old_sev: Optional[str]) -> Optional[str]:
    """
    v3.0.1: Migrate all legacy severity keys to new Minor/Major/Serious/Critical labels.
    Also handles v2.x numeric labels (1-Cosmetic etc.) and very-old Low/Medium/High/Critical.
    """
    if not old_sev:
        return None
    _MAP = {
        "1 - Cosmetic":         "Minor",
        "2 - Similar/Cosmetic": "Minor",
        "3 - Non-Serious":      "Major",
        "4 - Serious":          "Major",   # Sev 4 → Major
        "5 - Very Serious":     "Critical",
        "Serious":              "Major",   # any old "Serious" → Major
        "Low":    "Minor",
        "Medium": "Major",
        "High":   "Major",
    }
    return _MAP.get(old_sev, old_sev)   # if already new format, return unchanged



def _generate_defect_serial(project, blade: str, surface: str = "", zone: str = "") -> str:
    """
    v4.5.0: Generate defect serial number with surface+zone split.

    Format (underscore-separated, reference DOCX style):
      Blade:  WTG-321_Blade A_PS-LE_001   ("Blade A" keeps the space; PS-LE = surface-zone)
      Hub:    WTG-321_Hub_001             (no surface/zone for Hub/Tower)
      Tower:  WTG-321_Tower_001

    Global counter: the numeric suffix counts ALL annotations in the project across
    all components (A/B/C/Hub/Tower), so defect numbers are globally sequential.

    Returns a fallback string on any error to prevent crashes.
    """
    if not project:
        return "WTG-unknown_001"

    # Build WTG prefix — ensure it always starts with "WTG-"
    raw_id = (project.turbine_id or project.name or "WTG").strip()
    wtg_prefix = raw_id if raw_id.upper().startswith("WTG") else f"WTG-{raw_id}"

    is_blade  = blade in ("A", "B", "C") or blade.startswith("Blade")
    is_hub    = blade == "Hub"
    is_tower  = blade == "Tower"

    # Extract abbreviations from full names: "Pressure Side (PS)" → "PS"
    surf_abbr = ""
    zone_abbr = ""
    if is_blade:
        if surface and "(" in surface:
            surf_abbr = surface.split("(")[-1].strip(")")
        elif surface:
            surf_abbr = surface.strip()
        
        if zone and "(" in zone:
            zone_abbr = zone.split("(")[-1].strip(")")
        elif zone:
            zone_abbr = zone.strip()

    # Global defect counter: count ALL annotations already stored in the project
    global_count = 1
    try:
        for img_rec in project.images.values():
            global_count += len(img_rec.annotations)
    except Exception:
        pass  # images dict not yet populated on very first annotation — default 1

    seq = f"{global_count:03d}"

    # Compose serial per component type
    if is_blade:
        # Normalise: if bare letter "A" passed, expand to "Blade A"
        comp_label = blade if blade.startswith("Blade") else f"Blade {blade}"
        if surf_abbr and zone_abbr:
            return f"{wtg_prefix}_{comp_label}_{surf_abbr}-{zone_abbr}_{seq}"
        elif surf_abbr or zone_abbr:
            # One or the other present
            combo = f"{surf_abbr}{zone_abbr}".strip()
            return f"{wtg_prefix}_{comp_label}_{combo}_{seq}" if combo else f"{wtg_prefix}_{comp_label}_{seq}"
        return f"{wtg_prefix}_{comp_label}_{seq}"
    elif is_hub:
        return f"{wtg_prefix}_Hub_{seq}"
    elif is_tower:
        return f"{wtg_prefix}_Tower_{seq}"
    else:
        # Unknown component — sanitise name and append seq
        comp_clean = blade.replace(" ", "_") if blade else "Unknown"
        return f"{wtg_prefix}_{comp_clean}_{seq}"


@dataclass
class Annotation:
    """Single defect annotation — box, pin, or polygon."""
    ann_id     : str   = ""
    mode       : str   = "box"       # "box" | "pin" | "polygon"
    defect     : str   = "Erosion"
    severity   : str   = "Minor"
    blade      : str   = "A"
    span       : str   = "Root (0–33%)"
    # v4.5.0: Split face into surface and zone (mutually exclusive selections)
    surface    : str   = "Pressure Side (PS)"  # PS | SS
    zone       : str   = "Leading Edge (LE)"   # LE | TE | MB
    notes      : str   = ""
    x1_px      : float = 0.0
    y1_px      : float = 0.0
    x2_px      : float = 0.0
    y2_px      : float = 0.0
    # Polygon points stored as flat [x0,y0,x1,y1,...] list for JSON compat
    poly_pts   : List[float] = field(default_factory=list)
    width_cm   : Optional[float] = None
    height_cm  : Optional[float] = None
    gsd_source   : str   = "none"
    gsd_value    : Optional[float] = None
    # v4.4.8 BACKWARD-COMPAT / CALIBRATION AUDIT fields ──────────────────────
    # size_locked: when True the user has explicitly verified or manually set
    #   this annotation's physical size.  _batch_auto_calibrate() must NEVER
    #   overwrite width_cm / height_cm for locked annotations.  Set to True
    #   automatically when the user edits size in the panel and clicks Save, or
    #   when the annotation is approved by a reviewer (size implicitly verified).
    # size_calibrated_at: ISO-8601 timestamp of the last time width_cm /
    #   height_cm were written by ANY calibration path (auto OR manual).
    #   Stored in project.json — allows load_project() to detect annotations
    #   sized by the broken 160×6000 phantom-dimension builds and flag them for
    #   recalibration on next auto-calibrate run.
    size_locked          : bool           = False
    size_calibrated_at   : Optional[str]  = None   # ISO timestamp, None = never sized
    # ── Scopito fields (SenseHawk PDF parity) ────────────────────────────────
    distance_from_root_mm : Optional[float] = None   # mm from blade root (legacy)
    root_distance_m       : Optional[float] = None   # v2.1.1: meters from root
    tip_distance_m        : Optional[float] = None   # v2.1.1: meters from tip
    serial_number         : str             = ""     # v2.1.2: WTG-Component-Surface-Zone-SerialNo
    remedy_action         : str             = ""     # recommended action (auto-filled)
    rotation_deg          : float           = 0.0   # box rotation in degrees (0 = upright)
    # v3.2.0: User-set pinpoint on blade silhouette (0.0=root, 1.0=tip)
    pinpoint_blade_pos : Optional[float] = None
    # Phase 2/6 review workflow fields
    status       : str   = "pending"   # "pending" | "approved" | "rejected"
    created_by   : str   = ""          # username of creator
    created_at   : str   = ""          # v3.3.13: ISO timestamp of creation
    reviewed_by  : str   = ""          # username of reviewer
    reviewed_at  : str   = ""          # ISO timestamp of review
    reviewer_note: str   = ""          # reviewer's note on decision

    def bounding_rect(self) -> Tuple[float, float, float, float]:
        """Return (x, y, w, h) bounding rect across all modes."""
        if self.mode == "polygon" and len(self.poly_pts) >= 4:
            xs = self.poly_pts[0::2]; ys = self.poly_pts[1::2]
            x1, y1 = min(xs), min(ys); x2, y2 = max(xs), max(ys)
            return x1, y1, x2 - x1, y2 - y1
        x1 = min(self.x1_px, self.x2_px); y1 = min(self.y1_px, self.y2_px)
        return x1, y1, abs(self.x2_px - self.x1_px), abs(self.y2_px - self.y1_px)


@dataclass
class ImageRecord:
    """Metadata + annotations for one image file."""
    filename     : str
    filepath     : str
    blade        : str              = "A"
    annotations  : List[Annotation] = field(default_factory=list)
    gsd_cm_per_px: Optional[float]  = None
    # EXIF / metadata fields used in PDF report
    altitude_m   : Optional[str]    = None
    date_taken   : Optional[str]    = None
    heading      : Optional[str]    = None
    gps_coords   : Optional[str]    = None
    turbine_id   : Optional[str]    = None
    # v3.4.0: EXIF calibration metadata
    calibration_method  : Optional[str] = None  # "exif-auto", "manual", "component", "session"
    camera_model        : Optional[str] = None  # "Phantom 4 Pro", "Mavic 3", etc.
    focal_length_mm     : Optional[float] = None
    confidence_level    : Optional[str] = None  # "HIGH", "MEDIUM", "LOW"
    exif_distance_m     : Optional[float] = None  # Auto-estimated from GPS+gimbal


@dataclass
class Project:
    """Top-level project container → project.json."""
    name          : str  = "Untitled"
    site          : str  = ""
    turbine_id    : str  = ""
    inspector     : str  = ""
    created_at    : str  = field(default_factory=lambda: datetime.now().isoformat())
    session_gsd   : Optional[float] = None
    images        : Dict[str, ImageRecord] = field(default_factory=dict)
    defect_types  : List[str] = field(default_factory=lambda: list(DEFAULT_DEFECT_TYPES))
    project_folder: str  = ""
    summary_notes : str  = ""
    # ── Scopito: blade A/B/C → actual serial number (SenseHawk PDF page 2) ──
    blade_numbers : Dict[str, str] = field(default_factory=dict)  # e.g. {"A":"847","B":"852","C":"854"}
    # v1.7.0: Per-component calibrated GSD (cm/px). Set whenever user calibrates on any
    # image of that component. Used as fallback GSD for all images in the same component
    # that have no per-image GSD yet. Components: "A","B","C","Hub","Tower".
    component_gsd   : Dict[str, float] = field(default_factory=dict)
    # v1.7.0: Blade length in mm — used to auto-estimate distance_from_root from the
    # annotation Y-centre position in the image. Default 50 000 mm (50 m).
    blade_length_mm : float = 50_000.0
    # v4.2.0: Report metadata fields
    scan_date           : str = ""  # Inspection/scan date for front page
    turbine_manufacturer: str = ""  # Turbine manufacturer (e.g., "Vestas", "Siemens Gamesa")
    rated_power         : str = ""  # Rated power (e.g., "2.5 MW", "3.6 MW")
    # FIX-17c: Tower base GPS + MSL altitude — required for 3D distance model.
    # All three default to None so existing project.json files (which lack these
    # keys) load cleanly via the setdefault-style data.get(..., None) in load_project.
    tower_lat          : Optional[float] = None  # Tower base latitude  (decimal degrees)
    tower_lon          : Optional[float] = None  # Tower base longitude (decimal degrees)
    tower_base_alt_msl : Optional[float] = None  # Tower base altitude  (metres MSL)


@dataclass
class DetectionResult:
    """Raw YOLO detection — lives in memory until QC review commits it."""
    result_id  : str   = field(default_factory=lambda: str(uuid.uuid4())[:8])
    image_path : str   = ""
    x1_px      : float = 0.0
    y1_px      : float = 0.0
    x2_px      : float = 0.0
    y2_px      : float = 0.0
    confidence : float = 0.0
    class_id   : int   = 0
    class_name : str   = "Defect"
    approved   : bool  = False

# ==============================================================================
# JSON PERSISTENCE  (Tom K.)
# Atomic write → .tmp then os.replace() so a crash never leaves corrupt JSON.
# ==============================================================================

def _make_ann_id(seed: str = "") -> str:
    # CTO-AUDIT: UUID-based IDs for guaranteed uniqueness and traceability
    return str(uuid.uuid4())


def _ann_to_dict(ann: Annotation) -> dict:
    return asdict(ann)


def _ann_from_dict(d: dict) -> Annotation:
    d.setdefault("mode",         "box")
    d.setdefault("poly_pts",     [])
    d.setdefault("severity",      "Minor")
    d.setdefault("serial_number", "")
    d.setdefault("status",       "pending")
    d.setdefault("created_by",   "")
    d.setdefault("reviewed_by",  "")
    d.setdefault("reviewed_at",  "")
    d.setdefault("reviewer_note","")
    # Scopito fields — setdefault so old JSON files still load cleanly
    d.setdefault("distance_from_root_mm", None)
    d.setdefault("root_distance_m",       None)   # Phase 9.3: v2.1.1 field guard
    d.setdefault("tip_distance_m",        None)   # Phase 9.3: v2.1.1 field guard
    d.setdefault("pinpoint_blade_pos",    None)   # v3.2.0: blade silhouette position
    d.setdefault("remedy_action",         "")
    d.setdefault("rotation_deg",          0.0)    # interactive box rotation
    # v4.4.8: backward-compat guards for new calibration audit fields
    # size_locked defaults to False — existing annotations are not locked.
    # size_calibrated_at defaults to None — will be set on next auto-calibrate run.
    d.setdefault("size_locked",         False)
    d.setdefault("size_calibrated_at",  None)
    
    # v4.5.0: BACKWARD-COMPAT MIGRATION — face+edge_side → surface+zone split
    # Old format: single "face" field with values like "Leading Edge (LE)", "Pressure Side (PS)"
    # New format: separate "surface" (PS/SS) and "zone" (LE/TE/MB) fields
    if "face" in d and "surface" not in d:
        old_face = d.get("face", "")
        # Extract abbreviation from parentheses if present
        if "(" in old_face:
            abbr = old_face.split("(")[-1].strip(")")
        else:
            abbr = old_face.strip()
        
        # Map to new structure
        if abbr in ("PS", "SS") or "Pressure" in old_face or "Suction" in old_face:
            # It's a surface
            d["surface"] = old_face if "(" in old_face else f"{old_face} ({abbr})"
            d["zone"] = "Leading Edge (LE)"  # Default zone
        elif abbr in ("LE", "TE", "ED") or "Leading" in old_face or "Trailing" in old_face or "Edge" in old_face:
            # It's a zone
            d["surface"] = "Pressure Side (PS)"  # Default surface
            d["zone"] = "Midbody (MB)" if abbr == "ED" or "Edge" in old_face else old_face if "(" in old_face else f"{old_face} ({abbr})"
        else:
            # Unknown - use defaults
            d["surface"] = "Pressure Side (PS)"
            d["zone"] = "Leading Edge (LE)"
        
        # Remove old face field to avoid conflicts
        d.pop("face", None)
    
    # Migrate old edge_side field if present (v4.3.0 → v4.5.0)
    if "edge_side" in d:
        edge = d.get("edge_side")
        if edge == "LE":
            d["zone"] = "Leading Edge (LE)"
        elif edge == "TE":
            d["zone"] = "Trailing Edge (TE)"
        # Remove old field
        d.pop("edge_side", None)
    
    # Set defaults for new fields if not present
    d.setdefault("surface", "Pressure Side (PS)")
    d.setdefault("zone",    "Leading Edge (LE)")
    
    # ── BACKWARD-COMPAT MIGRATION: Detect broken 160×6000 phantom calibration ──
    # The buggy exifread path (v4.4.0 – v4.4.6) could produce a phantom image
    # width of 160 px instead of the real ~5472 px. GSD was therefore ~34× too
    # large (e.g. 85 cm/px instead of 2.5 cm/px).  A GSD > 20 cm/px is physically
    # implausible for any DJI camera at a wind-turbine inspection distance — the
    # maximum realistic GSD for a 6-mm sensor at 120 m is ~8 cm/px.  Any annotation
    # whose stored gsd_value exceeds this threshold and whose gsd_source is "image"
    # (set by auto-calibration) is flagged as likely corrupted by clearing
    # size_calibrated_at.  _batch_auto_calibrate() checks this field to decide
    # whether to skip or recalibrate, so re-running auto-calibrate after upgrading
    # to v4.4.7+ will silently fix all affected annotations.
    _PHANTOM_GSD_THRESHOLD = 20.0   # cm/px — physically impossible for DJI at inspection dist
    _gsd_val = d.get("gsd_value")
    _gsd_src = d.get("gsd_source", "none")
    if (_gsd_val is not None
            and _gsd_src == "image"
            and float(_gsd_val) > _PHANTOM_GSD_THRESHOLD
            and not d.get("size_locked", False)):
        log.warning(
            f"[MIGRATION] ann {d.get('ann_id', '?')[:8]}: suspicious gsd_value "
            f"{_gsd_val:.2f} cm/px (>{_PHANTOM_GSD_THRESHOLD}) — likely from "
            f"broken 160px-width calibration.  Cleared size_calibrated_at so "
            f"next auto-calibrate run will recalculate physical size.")
        d["size_calibrated_at"] = None   # signal: needs recalibration
    # Migrate legacy severity names (v3.0.1: to new Minor/Major/Serious/Critical)
    legacy = {
        "1 - Cosmetic": "Minor",  "2 - Similar/Cosmetic": "Minor",
        "3 - Non-Serious": "Major", "4 - Serious": "Major",
        "5 - Very Serious": "Critical",
        "Serious": "Major",
        "Low": "Minor", "Medium": "Major", "High": "Major",
    }
    d["severity"] = legacy.get(d["severity"], d["severity"])
    return Annotation(**{k: v for k, v in d.items() if k in Annotation.__dataclass_fields__})


def _irec_to_dict(irec: ImageRecord) -> dict:
    d = asdict(irec)
    d["annotations"] = [_ann_to_dict(a) for a in irec.annotations]
    return d


def _irec_from_dict(d: dict) -> ImageRecord:
    anns = [_ann_from_dict(a) for a in d.pop("annotations", [])]
    ir = ImageRecord(**{k: v for k, v in d.items()
                        if k in ImageRecord.__dataclass_fields__})
    ir.annotations = anns
    return ir


def _rotate_backups(project_file: Path, max_backups: int = 5):
    """CTO-AUDIT: Keep the last N backups of project.json.
    Naming: project.json.bak1 … project.json.bak5 (bak1 = most recent).
    Called BEFORE writing the new project file so bak1 is always the previous version.
    """
    try:
        if not project_file.exists():
            return
        # Shift existing backups up: bak4→bak5, bak3→bak4, …, bak1→bak2
        for i in range(max_backups - 1, 0, -1):
            src = project_file.with_suffix(f".json.bak{i}")
            dst = project_file.with_suffix(f".json.bak{i + 1}")
            if src.exists():
                if dst.exists():
                    dst.unlink()
                src.rename(dst)
        # Current file → bak1
        bak1 = project_file.with_suffix(".json.bak1")
        if bak1.exists():
            bak1.unlink()
        shutil.copy2(str(project_file), str(bak1))
    except Exception as exc:
        log.warning(f"Backup rotation failed: {exc}")


def save_project(project: Project) -> bool:
    """Tom K.: Atomic JSON write — never leaves a corrupt file on crash.
    CTO-AUDIT: Automatic backup rotation (5 backups) + schema_version field.
    """
    if not project.project_folder:
        return False
    try:
        import time as _time
        _t0 = _time.perf_counter()
        folder = Path(project.project_folder)
        folder.mkdir(parents=True, exist_ok=True)
        dest = folder / "project.json"
        tmp  = dest.with_suffix(".tmp")
        # CTO-AUDIT: rotate backups BEFORE overwriting so bak1 is the previous version
        _rotate_backups(dest)
        data = {
            "schema_version": 2,   # CTO-AUDIT: schema versioning for migration detection
            "name": project.name, "site": project.site,
            "turbine_id": project.turbine_id, "inspector": project.inspector,
            "created_at": project.created_at, "session_gsd": project.session_gsd,
            "defect_types": project.defect_types,
            "project_folder": project.project_folder,
            "summary_notes": project.summary_notes,
            "blade_numbers": project.blade_numbers,
            "component_gsd":  project.component_gsd,      # CTO-AUDIT: was missing
            "blade_length_mm": project.blade_length_mm,   # CTO-AUDIT: was missing
            # v4.2.0: New report metadata fields
            "scan_date": project.scan_date,
            "turbine_manufacturer": project.turbine_manufacturer,
            "rated_power": project.rated_power,
            # FIX-17c: Tower base GPS + MSL altitude for 3D distance model.
            # Written as None when not set so the JSON key is always present,
            # making load_project unambiguous.
            "tower_lat":          project.tower_lat,
            "tower_lon":          project.tower_lon,
            "tower_base_alt_msl": project.tower_base_alt_msl,
            "images": {k: _irec_to_dict(v) for k, v in project.images.items()},
        }
        tmp.write_text(json.dumps(data, indent=2), encoding="utf-8")
        os.replace(str(tmp), str(dest))
        _elapsed = _time.perf_counter() - _t0
        log.info(f"Project saved → {dest}  [{_elapsed*1000:.1f} ms]")
        return True
    except Exception as exc:
        log.error(f"save_project: {exc}")
        return False


def _try_load_json(path: Path) -> Optional[dict]:
    """CTO-AUDIT: Attempt to load JSON, returning None on any parse error.
    Tries the main file, then .bak1 through .bak5 in order so we can
    transparently recover from a corrupted project.json.
    """
    candidates = [path] + [path.with_suffix(f".json.bak{i}") for i in range(1, 6)]
    for candidate in candidates:
        if not candidate.exists():
            continue
        try:
            data = json.loads(candidate.read_text(encoding="utf-8"))
            if candidate != path:
                log.warning(f"load_project: recovered from backup {candidate.name}")
            return data
        except json.JSONDecodeError as exc:
            log.error(f"load_project: JSON parse error in {candidate.name}: {exc}")
    return None


def _repair_serial_numbers(project: "Project"):
    """
    Phase 9.7: Assign serial numbers to all annotations so that the numeric
    ending (_001, _002, …) matches the top-to-bottom position of each defect
    in the generated report.

    FIX-11 — SERIAL ENDINGS DID NOT MATCH REPORT ORDER:
      The previous implementation iterated project.images.values() in Python
      dict insertion order (folder-load order) and skipped any annotation that
      already had a serial.  This meant:
        • An annotation that appears 1st in the report could carry _003 because
          it was saved 3rd during the inspection session.
        • Re-loading the project never corrected it because the skip guard
          preserved the stale number.
      Fix: sort every (irec, ann) pair in the same canonical order used by the
      report engine (FIX-10): A→B→C→Hub→Tower by ann.blade, then by filename
      within each component, then by ann_id within each image.  All serial
      endings are FORCE-REASSIGNED so _001 always equals the first defect shown
      in the report, _002 the second, and so on.

    Called automatically by:
      • load_project()  — upgrades projects saved by older versions
      • _on_save_annotation()  — keeps serials current after every save so the
        panel and report always show the correct position number
    """
    try:
        # ── Build WTG prefix (same logic as _generate_defect_serial) ──────────
        raw_id     = (getattr(project, "turbine_id", "") or project.name or "WTG").strip()
        wtg_prefix = raw_id if raw_id.upper().startswith("WTG") else f"WTG-{raw_id}"

        # ── Collect all (irec, ann) pairs ────────────────────────────────────
        all_pairs: list = []
        for irec in project.images.values():
            for ann in irec.annotations:
                # Ensure the field exists (very old JSON may lack it)
                if not hasattr(ann, "serial_number"):
                    ann.serial_number = ""
                comp = (getattr(ann, "blade", "") or getattr(irec, "blade", "") or "Unknown").strip()
                all_pairs.append((comp, irec, ann))

        if not all_pairs:
            log.debug("_repair_serial_numbers: no annotations — nothing to do")
            return

        # ── Sort in canonical report order (FIX-11, consistent with FIX-10) ─
        # A→B→C→Hub→Tower, then filename-alphabetical within each component,
        # then ann_id within each image.  Uses ann.blade as the authority
        # (user-confirmed; irec.blade may be stale — see FIX-10 notes).
        def _blade_sort_key(b: str) -> int:
            return {"A": 0, "B": 1, "C": 2, "Hub": 3, "Tower": 4}.get(b, 10)

        all_pairs.sort(key=lambda t: (
            _blade_sort_key(t[0]),
            t[1].filename or "",
            t[2].ann_id  or ""
        ))

        # ── Force-assign serial endings _001, _002, … in report order ────────
        # The full serial is rebuilt from scratch (blade/face/component are
        # re-derived from the annotation's own fields so the prefix is always
        # consistent with current panel values, not a stale saved string).
        for global_n, (comp, irec, ann) in enumerate(all_pairs, start=1):
            seq   = f"{global_n:03d}"
            blade = (getattr(ann, "blade", "") or getattr(irec, "blade", "") or "").strip()
            face  = getattr(ann, "face", "") or ""

            is_blade = blade in ("A", "B", "C") or blade.startswith("Blade")
            is_hub   = blade == "Hub"
            is_tower = blade == "Tower"

            # Extract 2-letter face abbreviation e.g. "Leading Edge (LE)" → "LE"
            face_abbr = ""
            if is_blade and face:
                face_abbr = (face.split("(")[-1].strip(")")
                             if "(" in face else face.strip())

            if is_blade:
                comp_label = blade if blade.startswith("Blade") else f"Blade {blade}"
                serial = (f"{wtg_prefix}_{comp_label}_{face_abbr}_{seq}"
                          if face_abbr else f"{wtg_prefix}_{comp_label}_{seq}")
            elif is_hub:
                serial = f"{wtg_prefix}_Hub_{seq}"
            elif is_tower:
                serial = f"{wtg_prefix}_Tower_{seq}"
            else:
                comp_clean = blade.replace(" ", "_") if blade else "Unknown"
                serial = f"{wtg_prefix}_{comp_clean}_{seq}"

            ann.serial_number = serial

        log.debug(f"_repair_serial_numbers: assigned {len(all_pairs)} serials in report order")
    except Exception as exc:
        log.warning(f"_repair_serial_numbers: {exc}")


def load_project(folder: str) -> Optional[Project]:
    """Tom K.: Load and deserialize project.json, migrating legacy data.
    CTO-AUDIT: Schema validation, corruption recovery, version migration.
    """
    try:
        path = Path(folder) / "project.json"
        data = _try_load_json(path)
        if data is None:
            log.error(f"load_project: could not load any valid JSON from {folder}")
            return None

        # CTO-AUDIT: Schema version gate — warn on unknown future versions
        schema_ver = data.get("schema_version", 1)
        if schema_ver > 2:
            log.warning(f"load_project: schema_version {schema_ver} > 2 — "
                        "file was created by a newer version; some fields may be ignored")

        images = {k: _irec_from_dict(v) for k, v in data.get("images", {}).items()}
        p = Project(
            name=data.get("name", "Untitled"),
            site=data.get("site", ""),
            turbine_id=data.get("turbine_id", ""),
            inspector=data.get("inspector", ""),
            created_at=data.get("created_at", datetime.now().isoformat()),
            session_gsd=data.get("session_gsd"),
            defect_types=data.get("defect_types", list(DEFAULT_DEFECT_TYPES)),
            project_folder=folder,
            summary_notes=data.get("summary_notes", ""),
            blade_numbers=data.get("blade_numbers", {}),
            component_gsd=data.get("component_gsd", {}),
            blade_length_mm=data.get("blade_length_mm", 50_000.0),
            # v4.2.0: Load new report metadata fields with empty string defaults
            scan_date=data.get("scan_date", ""),
            turbine_manufacturer=data.get("turbine_manufacturer", ""),
            rated_power=data.get("rated_power", ""),
            # FIX-17c: Tower base GPS + MSL altitude.  Default None so older
            # project.json files (which lack these keys) load without error.
            tower_lat=data.get("tower_lat"),
            tower_lon=data.get("tower_lon"),
            tower_base_alt_msl=data.get("tower_base_alt_msl"),
            images=images,
        )
        # Phase 9.7: Auto-repair missing serial numbers (may be absent in old projects)
        _repair_serial_numbers(p)
        return p
    except Exception as exc:
        log.error(f"load_project: unexpected error: {exc}")
        return None

# ==============================================================================
# CONFIG MANAGER  (Jamie Liu)
# ==============================================================================

_CFG_DEFAULTS = {
    "GENERAL": {
        "LastProjectFolder": "",
        "SessionGSD": "0.0",
        "ThumbnailSize": "160",
        "ZoomMin": "0.05",
        "ZoomMax": "20.0",
    },
    "DETECTION": {
        "ModelPath":    "",
        "ConfHigh":     "0.45",
        "ConfLow":      "0.25",
        "NMS_IOU":      "0.45",
        "MaxWorkers":   "2",
        "Device":       "auto",
        "ImgSize":      "640",
        "ValSplit":     "0.2",
    },
    "TRAINING": {
        "Epochs":       "50",
        "BatchSize":    "16",
        "ImgSize":      "640",
        "LR0":          "0.01",
        "LRF":          "0.01",
        "Device":       "auto",
        "ExportFolder": "",
        "ValSplit":     "0.2",
    },
}


class AppConfig:
    """Jamie Liu: Thread-safe configparser wrapper with typed getters."""
    def __init__(self, path: Path = SETTINGS_FILE):
        self._path = path
        self._cfg  = configparser.ConfigParser()
        for section, defaults in _CFG_DEFAULTS.items():
            self._cfg[section] = defaults
        if path.exists():
            self._cfg.read(str(path), encoding="utf-8")

    def get(self, section: str, key: str, fallback: str = "") -> str:
        return self._cfg.get(section, key, fallback=fallback)

    def set(self, section: str, key: str, value: str):
        if not self._cfg.has_section(section):
            self._cfg.add_section(section)
        self._cfg.set(section, key, value)

    def save(self):
        with open(self._path, "w", encoding="utf-8") as f:
            self._cfg.write(f)

    # Phase 9.3: Recent projects stored as RECENT.path0…path9 in settings.ini
    def get_recent_projects(self, max_count: int = 8) -> List[str]:
        paths = []
        for i in range(max_count):
            p = self.get("RECENT", f"path{i}", "")
            if p and os.path.exists(p):
                paths.append(p)
        return paths

    def add_recent_project(self, path: str):
        existing = self.get_recent_projects(10)
        if path in existing:
            existing.remove(path)
        existing.insert(0, path)
        existing = existing[:8]
        if not self._cfg.has_section("RECENT"):
            self._cfg.add_section("RECENT")
        # Clear all slots first
        for i in range(10):
            self._cfg.set("RECENT", f"path{i}", "")
        for i, p in enumerate(existing):
            self._cfg.set("RECENT", f"path{i}", p)
        self.save()


CFG = AppConfig()

# ==============================================================================
# THUMBNAIL WORKER  (Alex Stone — Threading, 12yr)
# ==============================================================================

class _ThumbSignals(QObject):
    done  = pyqtSignal(int, QPixmap)
    error = pyqtSignal(int, str)


class ThumbnailWorker(QRunnable):
    """Alex Stone: Off-thread thumbnail generation with disk cache."""
    THUMB_W = 160; THUMB_H = 110

    def __init__(self, index: int, filepath: str, cache_dir: Path):
        super().__init__()
        self.index     = index
        self.filepath  = filepath
        self.cache_dir = cache_dir
        self.signals   = _ThumbSignals()
        self.setAutoDelete(True)

    @pyqtSlot()
    def run(self):
        try:
            stat       = os.stat(self.filepath)
            # CTO-AUDIT: include st_size in cache key — prevents stale thumb
            # when a file is replaced with another of same mtime (e.g. overwrite).
            key        = hashlib.md5(
                f"{self.filepath}{stat.st_mtime}{stat.st_size}".encode()
            ).hexdigest()
            cache_path = self.cache_dir / f"{key}.jpg"
            if cache_path.exists():
                qimg = QImage(str(cache_path))
            else:
                with Image.open(self.filepath) as img:
                    img.thumbnail((self.THUMB_W, self.THUMB_H), Image.LANCZOS)
                    if img.mode != "RGB":
                        img = img.convert("RGB")
                    self.cache_dir.mkdir(parents=True, exist_ok=True)
                    img.save(str(cache_path), "JPEG", quality=75)
                qimg = QImage(str(cache_path))
            self.signals.done.emit(self.index, QPixmap.fromImage(qimg))
        except Exception as exc:
            log.warning(f"Thumbnail failed for {self.filepath}: {exc}")
            self.signals.error.emit(self.index, str(exc))

# ==============================================================================
# BLADE POSITION PANEL  (Dev Patel — UX 8yr, v3.2.0 redesign)
# Scopito-style vertical blade silhouettes with defect-dot pinpoints.
# Replaces the old horizontal Root/Mid/Tip grid.
# ==============================================================================

class BladeDiagram(QWidget):
    """
    v3.2.0 Scopito-style vertical blade position panel.

    Shows Blade A / B / C as side-by-side vertical tapered silhouettes
    (root at top, tip at bottom — matching Scopito's convention).
    Every annotation is plotted as a coloured dot at its pinpoint_blade_pos
    (or estimated from root_distance_m).  Clicking a dot emits ann_clicked
    and the connected MainWindow jumps to that image.

    Also still emits cell_clicked(blade, span, face) for backward compat
    (now triggered by clicking the blade label at the top).
    """
    cell_clicked = pyqtSignal(str, str, str)   # kept for compat
    ann_clicked  = pyqtSignal(object)           # Annotation clicked on diagram

    _SEV_DOT: Dict[Optional[str], str] = {
        # v4.3.0: Aligned with _SEV_HEX / UI_THEME — Yellow / Amber / Red
        "Minor":    "#FFD700",   # Yellow
        "Major":    "#FFA500",   # Amber
        "Critical": "#FF0000",   # Red
        "POI":      "#388bfd",   # Blue
    }

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setMinimumSize(160, 300)
        self.setMouseTracking(True)   # enables hover without button pressed
        self._project    : Optional[Project]         = None
        self._dot_map    : List[tuple]               = []  # (QPointF, Annotation, irec)
        self._hover_idx  : int                       = -1
        self._active_blade: Optional[str]            = None  # currently viewed image blade

    def update_project(self, project: Optional[Project]):
        self._project = project
        self._dot_map.clear()
        self.update()

    def set_active_blade(self, blade_key: Optional[str]):
        """Highlight the column whose image is currently open in the viewer."""
        if self._active_blade != blade_key:
            self._active_blade = blade_key
            self.update()

    # ── geometry helpers ───────────────────────────────────────────────────────

    def _blade_columns(self):
        """
        Return list of (label, key, x_centre, col_w, blade_top_y, blade_bot_y)
        for each blade column (A, B, C).
        """
        W, H = self.width(), self.height()
        BLADES    = [("A", "A"), ("B", "B"), ("C", "C")]
        HDR_H     = 32   # px for label + "Root" text at top
        BOT_PAD   = 22   # px for "Tip" label at bottom
        col_w     = W / len(BLADES)
        blade_top = HDR_H
        blade_bot = H - BOT_PAD
        return [
            (f"Blade {k}", k,
             col_w * i + col_w / 2,   # cx
             col_w,
             blade_top, blade_bot)
            for i, (k, _) in enumerate(BLADES)
        ], col_w

    def _half_w_at(self, t: float, max_hw: float) -> float:
        """Half-width of blade at normalised position t (0=root, 1=tip)."""
        return max_hw * ((1.0 - t) ** 0.55)

    def _ann_blade_pos(self, ann: "Annotation") -> float:
        """Normalised 0-1 position for an annotation (root=0, tip=1)."""
        pos = getattr(ann, "pinpoint_blade_pos", None)
        if pos is not None:
            return max(0.0, min(1.0, float(pos)))
        # Fall back: estimate from root_distance_m
        bl_mm = 50_000.0
        if self._project:
            bl_mm = getattr(self._project, "blade_length_mm", 50_000.0) or 50_000.0
        rm = getattr(ann, "root_distance_m", None)
        if rm:
            return min(1.0, (rm * 1000) / bl_mm)
        # Fall back: span string
        sp = (ann.span or "").split("(")[0].strip()
        return {"Root": 0.15, "Mid": 0.50, "Tip": 0.82}.get(sp, 0.5)

    # ── painting ───────────────────────────────────────────────────────────────

    def paintEvent(self, event):
        from PyQt6.QtGui import QPainterPath as _PP
        p = QPainter(self)
        p.setRenderHint(QPainter.RenderHint.Antialiasing)
        W, H = self.width(), self.height()

        # Background
        p.fillRect(0, 0, W, H, QColor("#1a1f2e"))

        cols, col_w = self._blade_columns()
        self._dot_map = []

        DOT_R   = 6.5   # dot radius px
        MAX_HW  = col_w * 0.30   # max half-width of blade silhouette

        # Shared top/bottom labels (once per panel)
        lbl_font = QFont("Segoe UI", 7)
        p.setFont(lbl_font)
        p.setPen(QColor("#00d4e0"))
        p.drawText(QRectF(0, 2, W, 14),
                   Qt.AlignmentFlag.AlignHCenter, "▲  Root")
        p.drawText(QRectF(0, H - 18, W, 14),
                   Qt.AlignmentFlag.AlignHCenter, "▼  Tip")

        for label, bkey, cx, cw, btop, bbot in cols:
            bh = bbot - btop

            # ── Blade silhouette ──────────────────────────────────────────────
            path = _PP()
            STEPS = 60
            pts_l, pts_r = [], []
            for i in range(STEPS + 1):
                t  = i / STEPS
                y  = btop + t * bh
                hw = self._half_w_at(t, MAX_HW)
                pts_l.append(QPointF(cx - hw, y))
                pts_r.append(QPointF(cx + hw, y))
            path.moveTo(pts_l[0])
            for pt in pts_l[1:]:   path.lineTo(pt)
            for pt in reversed(pts_r): path.lineTo(pt)
            path.closeSubpath()
            p.fillPath(path, QBrush(QColor("#2d3748")))
            # Active-blade highlight: cyan outline instead of grey
            is_active = (bkey == self._active_blade)
            outline_col = QColor("#00d4e0") if is_active else QColor("#4a5568")
            outline_w   = 1.8 if is_active else 0.8
            p.setPen(QPen(outline_col, outline_w))
            p.drawPath(path)

            # ── Span guide-lines (Root/Mid/Tip) ──────────────────────────────
            for t_mark in (0.33, 0.66):
                y_mark = btop + t_mark * bh
                hw_m   = self._half_w_at(t_mark, MAX_HW)
                p.setPen(QPen(QColor("#ffffff22"), 0.6, Qt.PenStyle.DashLine))
                p.drawLine(QPointF(cx - hw_m, y_mark), QPointF(cx + hw_m, y_mark))

            # ── Blade label ───────────────────────────────────────────────────
            hdr_font = QFont("Segoe UI", 7, QFont.Weight.Bold)
            p.setFont(hdr_font)
            # Active blade label in bright white; others in cyan
            p.setPen(QColor("#ffffff") if is_active else QColor("#00d4e0"))
            p.drawText(QRectF(cx - cw/2, 14, cw, 16),
                       Qt.AlignmentFlag.AlignHCenter | Qt.AlignmentFlag.AlignVCenter,
                       label)

            # Divider line between columns
            if bkey != "C":
                p.setPen(QPen(QColor("#2d3648"), 1))
                p.drawLine(QPointF(cx + cw/2, btop), QPointF(cx + cw/2, bbot))

            # ── Annotation dots ───────────────────────────────────────────────
            if not self._project:
                continue
            for irec in self._project.images.values():
                if irec.blade != bkey:
                    continue
                for ann in irec.annotations:
                    t_pos  = self._ann_blade_pos(ann)
                    dot_y  = btop + t_pos * bh
                    dot_col = QColor(self._SEV_DOT.get(ann.severity, "#00d4e0"))
                    dot_idx = len(self._dot_map)   # index before appending
                    is_hover = (dot_idx == self._hover_idx)
                    r_outer = DOT_R + 3 if is_hover else DOT_R + 1.5
                    r_inner = DOT_R + 1 if is_hover else DOT_R
                    # Glow ring
                    glow_col = dot_col.lighter(200 if is_hover else 160)
                    p.setPen(QPen(glow_col, 1.2))
                    p.setBrush(QBrush(Qt.BrushStyle.NoBrush))
                    p.drawEllipse(QPointF(cx, dot_y), r_outer, r_outer)
                    # Filled dot
                    p.setPen(QPen(QColor("#0d1117"), 1.2))
                    p.setBrush(QBrush(dot_col))
                    p.drawEllipse(QPointF(cx, dot_y), r_inner, r_inner)
                    # White cross inside for hover
                    if is_hover:
                        p.setPen(QPen(QColor("white"), 1.0))
                        p.drawLine(QPointF(cx - 3, dot_y), QPointF(cx + 3, dot_y))
                        p.drawLine(QPointF(cx, dot_y - 3), QPointF(cx, dot_y + 3))
                    # Store for click detection and tooltip
                    self._dot_map.append((QPointF(cx, dot_y), ann, irec))

        p.end()

    # ── Interaction ───────────────────────────────────────────────────────────

    def mousePressEvent(self, event):
        pos = QPointF(event.position().x(), event.position().y())
        # Check dots first (within 12px)
        for entry in self._dot_map:
            dot_pos, ann = entry[0], entry[1]
            if ((pos.x() - dot_pos.x())**2 + (pos.y() - dot_pos.y())**2) <= 12**2:
                self.ann_clicked.emit(ann)
                return
        # Right-click → copy to clipboard
        if event.button() == Qt.MouseButton.RightButton:
            menu = QMenu(self)
            copy_act = menu.addAction("📋  Copy Diagram to Clipboard")
            copy_act.triggered.connect(self._copy_to_clipboard)
            menu.exec(event.globalPosition().toPoint())
            return
        # Click on blade label area → emit cell_clicked for compat
        cols, _ = self._blade_columns()
        for label, bkey, cx, cw, btop, bbot in cols:
            if abs(pos.x() - cx) < cw/2 and pos.y() < btop:
                self.cell_clicked.emit(bkey, "Root", "LE")
                return

    def mouseMoveEvent(self, event):
        pos = QPointF(event.position().x(), event.position().y())
        new_hover = -1
        for i, entry in enumerate(self._dot_map):
            dot_pos, ann = entry[0], entry[1]
            if ((pos.x() - dot_pos.x())**2 + (pos.y() - dot_pos.y())**2) <= 12**2:
                new_hover = i
                sev  = ann.severity or "—"
                defect = ann.defect or "—"
                dist_s = (f"{ann.root_distance_m:.1f} m"
                          if ann.root_distance_m else "—")
                self.setToolTip(
                    f"<b>{defect}</b><br>"
                    f"Severity: {sev}<br>"
                    f"Root dist: {dist_s}<br>"
                    f"<i>Click to jump to image</i>")
                self.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
                break
        if new_hover == -1:
            self.setToolTip("")
            self.setCursor(QCursor(Qt.CursorShape.ArrowCursor))
        if new_hover != self._hover_idx:
            self._hover_idx = new_hover
            self.update()

    def leaveEvent(self, event):
        if self._hover_idx != -1:
            self._hover_idx = -1
            self.update()
        self.setCursor(QCursor(Qt.CursorShape.ArrowCursor))
        super().leaveEvent(event)

    def _copy_to_clipboard(self):
        pix = self.grab()
        QApplication.clipboard().setPixmap(pix)
        orig_ss = self.styleSheet()
        self.setStyleSheet(
            f"background:{UI_THEME['bg_secondary']};"
            "border:2px solid #00d4e0;border-radius:8px;")
        QTimer.singleShot(350, lambda: self.setStyleSheet(orig_ss))

# ==============================================================================
# ANNOTATION GRAPHICS ITEMS  (Priya Nair — Canvas 10yr)
# ==============================================================================

class EditableBoxItem(QGraphicsItem):
    """
    Priya Nair + Sam Okafor: Fully interactive bounding-box annotation item.

    Features (active whenever the item is selected in Select mode):
      • 8 resize handles  — corners (TL/TR/BL/BR) + edge midpoints (TC/ML/MR/BC)
      • 1 rotation handle — cyan circle above the top-centre edge
      • Body drag         — click anywhere inside the rect (not on a handle) to move
      • Cursor hints      — cursor shape changes to indicate action
      • Commit on release — writes new geometry back into ann.x1_px/y1_px/x2_px/y2_px
                            and ann.rotation_deg, then calls on_changed(ann) callback

    Usage:
        item = EditableBoxItem(ann, gsd)
        item.on_changed = lambda a: viewer.signals.annotation_modified.emit(a)
        scene.addItem(item)
    """

    # Handle size in item-local pixels (unaffected by zoom)
    _HS  = 8    # half-size of each square handle
    _ROT_DIST = 30   # distance above top-center for rotation handle

    # Handle index constants
    _TL, _TC, _TR = 0, 1, 2
    _ML,      _MR = 3,    4
    _BL, _BC, _BR = 5, 6, 7
    _ROT          = 8

    def __init__(self, ann: Annotation, gsd: Optional[float] = None, parent=None):
        super().__init__(parent)
        self.ann     = ann
        self._gsd    = gsd

        # ── Local rect (centred at item origin for clean rotation) ────────────
        w = abs(ann.x2_px - ann.x1_px)
        h = abs(ann.y2_px - ann.y1_px)
        cx = min(ann.x1_px, ann.x2_px) + w / 2
        cy = min(ann.y1_px, ann.y2_px) + h / 2
        self.setPos(cx, cy)                         # item origin = rect centre
        self._rect = QRectF(-w / 2, -h / 2, w, h)  # local coords

        # ── Rotation ──────────────────────────────────────────────────────────
        self._angle = float(ann.rotation_deg or 0.0)
        self.setRotation(self._angle)
        self.setTransformOriginPoint(0.0, 0.0)

        # ── Flags ─────────────────────────────────────────────────────────────
        self.setFlag(QGraphicsItem.GraphicsItemFlag.ItemIsSelectable)
        self.setFlag(QGraphicsItem.GraphicsItemFlag.ItemSendsGeometryChanges)
        self.setAcceptHoverEvents(True)
        self.setZValue(10)

        # ── Drag state ────────────────────────────────────────────────────────
        self._drag_handle      : int     = -1        # -1 = body drag
        self._drag_active      : bool    = False
        self._drag_scene0      : QPointF = QPointF()  # scene pos at press
        self._drag_rect0       : QRectF  = QRectF()   # rect at press
        self._drag_pos0        : QPointF = QPointF()  # item pos at press
        self._drag_angle0      : float   = 0.0        # angle at press
        # T01 FIX — rotation jump: store the angle from item-centre → click
        # position at drag-start.  _update_drag subtracts this bias so the item
        # only rotates by Δangle (mouse moves), never snaps to absolute.
        self._drag_angle_bias  : float   = 0.0
        # Incremental drag: scene pos of previous frame (for resize + rotation)
        self._drag_prev_scene  : QPointF = QPointF()

        # ── Change callback set by ImageViewer ────────────────────────────────
        self.on_changed = None   # callable(Annotation)

    # ── Geometry helpers ──────────────────────────────────────────────────────

    @property
    def _colour(self) -> QColor:
        return SEVERITY_COLORS.get(self.ann.severity, QColor("#d29922"))

    def _handle_centres(self) -> List[QPointF]:
        """9 handle centres in item-local coords."""
        r  = self._rect
        cx = r.center().x()
        cy = r.center().y()
        return [
            QPointF(r.left(),   r.top()   ),  # 0 TL
            QPointF(cx,         r.top()   ),  # 1 TC
            QPointF(r.right(),  r.top()   ),  # 2 TR
            QPointF(r.left(),   cy        ),  # 3 ML
            QPointF(r.right(),  cy        ),  # 4 MR
            QPointF(r.left(),   r.bottom()),  # 5 BL
            QPointF(cx,         r.bottom()),  # 6 BC
            QPointF(r.right(),  r.bottom()),  # 7 BR
            QPointF(cx,         r.top() - self._ROT_DIST),  # 8 ROT
        ]

    def _handle_rect(self, i: int) -> QRectF:
        c = self._handle_centres()[i]
        return QRectF(c.x() - self._HS, c.y() - self._HS,
                      self._HS * 2,     self._HS * 2)

    def _hit_handle(self, local_pos: QPointF) -> int:
        """Return handle index under local_pos (expanded hit area), or -1."""
        if not self.isSelected():
            return -1
        expand = 4   # extra px hit padding
        for i, c in enumerate(self._handle_centres()):
            hr = QRectF(c.x() - self._HS - expand, c.y() - self._HS - expand,
                        (self._HS + expand) * 2,   (self._HS + expand) * 2)
            if hr.contains(local_pos):
                return i
        return -1

    # ── QGraphicsItem interface ───────────────────────────────────────────────

    def boundingRect(self) -> QRectF:
        pad = self._HS + 4
        return self._rect.adjusted(-pad,
                                   -pad - self._ROT_DIST,
                                    pad,
                                    pad)

    def shape(self):
        from PyQt6.QtGui import QPainterPath
        path = QPainterPath()
        path.addRect(self._rect)
        if self.isSelected():
            # v4.1.1: include rotation handle + stem so hover/click fires above the rect
            rot_c = self._handle_centres()[self._ROT]
            pad   = self._HS + 6
            path.addEllipse(rot_c, pad, pad)
            tc_c  = self._handle_centres()[self._TC]
            stem  = QRectF(tc_c.x() - 4, rot_c.y(), 8, tc_c.y() - rot_c.y())
            path.addRect(stem)
        return path

    def paint(self, painter: QPainter, option, widget=None):
        col      = self._colour
        selected = self.isSelected()
        ann      = self.ann

        # ── Main rectangle ────────────────────────────────────────────────────
        pen = QPen(col, BOX_PEN_WIDTH)
        if selected:
            pen.setStyle(Qt.PenStyle.DashLine)
            pen.setColor(col.lighter(140))
        painter.setPen(pen)
        painter.setBrush(QBrush(Qt.BrushStyle.NoBrush))
        painter.drawRect(self._rect)

        # ── Label pill (above top-left corner) ────────────────────────────────
        size_txt = (f"{ann.width_cm:.1f}×{ann.height_cm:.1f}cm"
                    if ann.width_cm is not None else "?cm")
        short  = SEVERITY_SHORT.get(ann.severity, ann.severity)
        label  = f"#{ann.ann_id[:5]} [{short}] {ann.defect}  {size_txt}"
        font   = QFont("Segoe UI", LABEL_FONT_SIZE, QFont.Weight.Bold)
        painter.setFont(font)
        fm     = QFontMetrics(font)
        tw     = fm.horizontalAdvance(label)
        th     = fm.height()
        lx     = self._rect.left()
        ly     = self._rect.top() - th - 6
        pill   = QRectF(lx - 2, ly - 2, tw + 8, th + 4)
        painter.fillRect(pill, col.darker(120))
        painter.setPen(QColor("white"))
        painter.drawText(QPointF(lx + 3, ly + th - 2), label)

        # ── Handles — only when selected ──────────────────────────────────────
        if not selected:
            return

        painter.save()

        # Dashed line from top-centre to rotation handle
        rot_c  = self._handle_centres()[self._ROT]
        tc_c   = self._handle_centres()[self._TC]
        dash_pen = QPen(QColor("#00d4e0"), 1.0, Qt.PenStyle.DotLine)
        painter.setPen(dash_pen)
        painter.drawLine(tc_c, rot_c)

        for i in range(9):
            hr = self._handle_rect(i)
            if i == self._ROT:
                # Rotation handle — filled cyan circle
                painter.setPen(QPen(QColor("#0d1117"), 1.5))
                painter.setBrush(QBrush(QColor("#00d4e0")))
                painter.drawEllipse(hr)
                # Arrow-curve icon inside
                painter.setPen(QPen(QColor("#0d1117"), 1.5))
                painter.drawArc(hr.adjusted(3, 3, -3, -3), 30 * 16, 300 * 16)
            else:
                # Resize handle — white square, colour border
                painter.setPen(QPen(col.lighter(130), 1.5))
                painter.setBrush(QBrush(QColor("white")))
                painter.drawRect(hr)

        painter.restore()

    # ── Hover (cursor hints) ──────────────────────────────────────────────────

    @staticmethod
    def _pt_seg_dist_sq(px, py, ax, ay, bx, by) -> float:
        """Squared screen-space distance from point P to segment AB."""
        dx, dy = bx - ax, by - ay
        if dx == 0 and dy == 0:
            return (px - ax) ** 2 + (py - ay) ** 2
        t = max(0.0, min(1.0, ((px - ax) * dx + (py - ay) * dy) / (dx * dx + dy * dy)))
        return (px - ax - t * dx) ** 2 + (py - ay - t * dy) ** 2

    @staticmethod
    def _vec_to_edge_cursor(vx: float, vy: float) -> "Qt.CursorShape":
        """
        FIX-14b: Map a screen-space edge direction vector to the nearest Qt
        resize cursor that reflects the true perpendicular drag direction.

        The perpendicular (drag direction) of edge vector (vx, vy) is (-vy, vx).
        We compute that angle in [0°, 180°) and snap to the four available Qt
        resize cursor shapes by ±45° quantisation:
          [0°, 22.5°) or [157.5°, 180°) → SizeHorCursor  ↔
          [22.5°, 67.5°)                 → SizeBDiagCursor ↙↗
          [67.5°, 112.5°)                → SizeVerCursor   ↕
          [112.5°, 157.5°)               → SizeFDiagCursor ↘↖

        This produces correct cursor feedback for any box rotation angle —
        for an axis-aligned box it returns the same cursors as before; for a
        rotated box it returns the cursor that matches the actual drag direction.
        """
        # Perpendicular to the edge is (-vy, vx).  Map to [0, 180) (cursors
        # are symmetric: ↕ same caret as ↕, so we fold the range in half).
        px, py = -vy, vx
        angle = math.degrees(math.atan2(py, px)) % 180.0
        if angle < 22.5 or angle >= 157.5:
            return Qt.CursorShape.SizeHorCursor
        if angle < 67.5:
            return Qt.CursorShape.SizeBDiagCursor
        if angle < 112.5:
            return Qt.CursorShape.SizeVerCursor
        return Qt.CursorShape.SizeFDiagCursor

    def hoverMoveEvent(self, event):
        """v4.1.1: Paint-style cursor — full edge segments + rotation handle proximity."""
        if not self.isSelected():
            self.setCursor(QCursor(Qt.CursorShape.SizeAllCursor))
            super().hoverMoveEvent(event)
            return

        CORNER_R = 14   # screen px — within this of a corner → diagonal cursor
        EDGE_W   = 10   # screen px — within this of an edge segment → resize cursor
        ROT_R    = 16   # screen px — within this of rotation handle → rotate cursor

        views = self.scene().views() if self.scene() else []
        if not views:
            # Fallback: no view available
            self.setCursor(QCursor(Qt.CursorShape.SizeAllCursor))
            super().hoverMoveEvent(event)
            return

        view = views[0]
        hover_scene = self.mapToScene(event.pos())
        hvp = view.mapFromScene(hover_scene)
        hx, hy = hvp.x(), hvp.y()

        # ── 1. Rotation handle proximity ──────────────────────────────────────
        rot_sc = self.mapToScene(self._handle_centres()[self._ROT])
        rot_vp = view.mapFromScene(rot_sc)
        rdx, rdy = rot_vp.x() - hx, rot_vp.y() - hy
        if rdx * rdx + rdy * rdy <= ROT_R * ROT_R:
            self.setCursor(QCursor(Qt.CursorShape.OpenHandCursor))  # rotate
            super().hoverMoveEvent(event)
            return

        # Map the 4 rect corners to viewport (rect may be rotated)
        corners_local = [
            QPointF(self._rect.left(),  self._rect.top()),    # TL 0
            QPointF(self._rect.right(), self._rect.top()),    # TR 1
            QPointF(self._rect.right(), self._rect.bottom()), # BR 2
            QPointF(self._rect.left(),  self._rect.bottom()), # BL 3
        ]
        cvp = [(view.mapFromScene(self.mapToScene(c)).x(),
                view.mapFromScene(self.mapToScene(c)).y())
               for c in corners_local]
        tl, tr, br, bl = cvp[0], cvp[1], cvp[2], cvp[3]

        # ── 2. Corner proximity ────────────────────────────────────────────────
        # FIX-14b: Compute corner cursor from the actual screen-space diagonal
        # direction so a rotated box shows the correct resize cursor rather than
        # the hardcoded SizeFDiag/SizeBDiag that only matched 0° boxes.
        # The TL↔BR diagonal direction is (br_x-tl_x, br_y-tl_y); TR↔BL is
        # (bl_x-tr_x, bl_y-tr_y).  _vec_to_edge_cursor() maps the perpendicular
        # of the EDGE to a cursor, so we pass the diagonal ITSELF as the "edge"
        # (the perpendicular of the diagonal IS the resize direction for corners).
        tl_x, tl_y = tl
        tr_x, tr_y = tr
        br_x, br_y = br
        bl_x, bl_y = bl
        _cur_tl_br = self._vec_to_edge_cursor(br_x - tl_x, br_y - tl_y)
        _cur_tr_bl = self._vec_to_edge_cursor(bl_x - tr_x, bl_y - tr_y)
        corner_cursors = [
            (tl, _cur_tl_br),   # TL — drags toward BR axis
            (tr, _cur_tr_bl),   # TR — drags toward BL axis
            (br, _cur_tl_br),   # BR — drags toward TL axis
            (bl, _cur_tr_bl),   # BL — drags toward TR axis
        ]
        for (cx, cy), cursor in corner_cursors:
            if (hx - cx) ** 2 + (hy - cy) ** 2 <= CORNER_R * CORNER_R:
                self.setCursor(QCursor(cursor))
                super().hoverMoveEvent(event)
                return

        # ── 3. Edge segment proximity (paint-style: full segment, not midpoint) ─
        # FIX-14b: Compute each edge's cursor from its screen-space direction so
        # a rotated box shows the correct perpendicular resize cursor.  tl/tr/br/bl
        # are already the rotated viewport corner positions computed above.
        edge_cursors = [
            (tl, tr, self._vec_to_edge_cursor(tr_x - tl_x, tr_y - tl_y)),   # top
            (br, bl, self._vec_to_edge_cursor(bl_x - br_x, bl_y - br_y)),   # bottom
            (tl, bl, self._vec_to_edge_cursor(bl_x - tl_x, bl_y - tl_y)),   # left
            (tr, br, self._vec_to_edge_cursor(br_x - tr_x, br_y - tr_y)),   # right
        ]
        edge_w_sq = EDGE_W * EDGE_W
        for (ax, ay), (bx, by), cursor in edge_cursors:
            if self._pt_seg_dist_sq(hx, hy, ax, ay, bx, by) <= edge_w_sq:
                self.setCursor(QCursor(cursor))
                super().hoverMoveEvent(event)
                return

        # ── 4. Body ────────────────────────────────────────────────────────────
        self.setCursor(QCursor(Qt.CursorShape.SizeAllCursor))
        super().hoverMoveEvent(event)

    def hoverLeaveEvent(self, event):
        self.unsetCursor()
        super().hoverLeaveEvent(event)

    # ── Mouse interaction ─────────────────────────────────────────────────────

    def mousePressEvent(self, event):
        # Bug #2 FIX: ImageViewer.mousePressEvent now fully manages drag
        # initialisation via the _active_editable + screen-space handle test.
        # This method is intentionally dead — returning early prevents the old
        # item-local _hit_handle() call from ever producing a wrong handle index
        # if this code path is somehow reached (e.g. direct scene dispatch).
        # All drag state is set by ImageViewer before _update_drag is called.
        event.accept()
        return

    def _update_drag(self, scene_pos: QPointF,
                     modifiers: Qt.KeyboardModifier = Qt.KeyboardModifier.NoModifier):
        """Core drag logic — callable from ImageViewer without Qt mouse-grab."""
        scene_delta = scene_pos - self._drag_scene0
        h = self._drag_handle

        if h == self._ROT:
            # T01 FIX (expert): incremental frame-to-frame accumulation.
            # 'abs_angle' = atan2(center→mouse) + 90.  We compute the DELTA
            # from the PREVIOUS frame's angle (_drag_angle_bias) so that:
            #  (a) No jump on first click — bias was set to the click angle.
            #  (b) atan2 wraparound at ±180° never causes a 360° jump because
            #      we normalize the per-frame delta into (−180, +180].
            #  (c) Rotations >180° in a single drag work correctly because
            #      we accumulate incrementally rather than computing a total.
            centre_s = self.mapToScene(QPointF(0.0, 0.0))
            dx = scene_pos.x() - centre_s.x()
            dy = scene_pos.y() - centre_s.y()
            curr_abs = math.degrees(math.atan2(dy, dx)) + 90.0
            # Normalize per-frame delta to (−180, +180]
            raw_delta = (curr_abs - self._drag_angle_bias + 180.0) % 360.0 - 180.0
            self._angle += raw_delta
            if modifiers & Qt.KeyboardModifier.ShiftModifier:
                self._angle = round(self._angle / 15.0) * 15.0
            self.setRotation(self._angle)
            self._drag_angle_bias = curr_abs   # advance to this frame for next delta

        elif h >= 0:
            # Resize: incremental delta in item-local space.
            # Both prev and current are mapped with the SAME current transform,
            # so the delta is accurate even as setPos/translate shift the item.
            local_now  = self.mapFromScene(scene_pos)
            local_prev = self.mapFromScene(self._drag_prev_scene)
            r  = QRectF(self._rect)   # current rect, not snapshot at press
            dx = local_now.x() - local_prev.x()
            dy = local_now.y() - local_prev.y()

            if h in (self._TL, self._TC, self._TR):   r.setTop(   r.top()    + dy)
            if h in (self._BL, self._BC, self._BR):   r.setBottom(r.bottom() + dy)
            if h in (self._TL, self._ML, self._BL):   r.setLeft(  r.left()   + dx)
            if h in (self._TR, self._MR, self._BR):   r.setRight( r.right()  + dx)

            MIN = 10.0  # enforce minimum 10×10 px
            if abs(r.width())  < MIN: r.setWidth( math.copysign(MIN, r.width()))
            if abs(r.height()) < MIN: r.setHeight(math.copysign(MIN, r.height()))

            # Keep item origin at rect centre (important for clean rotation)
            new_centre = r.center()
            shift = self.mapToScene(new_centre) - self.mapToScene(QPointF(0, 0))
            self.setPos(self.pos() + shift)
            r.translate(-new_centre)            # re-centre rect at origin
            self._rect = r
            self.prepareGeometryChange()
            self._drag_prev_scene = scene_pos   # advance for next frame

        else:
            # Body move
            self.setPos(self._drag_pos0 + scene_delta)

        self.update()

    def _finish_drag(self):
        """Commit drag state and fire on_changed — callable from ImageViewer."""
        self._drag_active = False
        self._drag_handle = -1
        self._commit_geometry()

    def mouseMoveEvent(self, event):
        if not self._drag_active:
            super().mouseMoveEvent(event)
            return
        self._update_drag(event.scenePos(), event.modifiers())  # delegate to helper
        event.accept()

    def mouseReleaseEvent(self, event):
        if not self._drag_active:
            super().mouseReleaseEvent(event)
            return
        self._finish_drag()   # delegate to helper so ImageViewer can also call it
        event.accept()

    # ── Commit ────────────────────────────────────────────────────────────────

    def _commit_geometry(self):
        """
        Write current visual state back to the Annotation object.
        Called on every mouseReleaseEvent so the data model stays in sync.
        """
        centre = self.pos()                      # scene coords (rotation origin)
        w = abs(self._rect.width())
        h = abs(self._rect.height())

        # Store as axis-aligned bounding box in scene space
        # (rotation is tracked separately in rotation_deg)
        self.ann.x1_px      = centre.x() - w / 2
        self.ann.y1_px      = centre.y() - h / 2
        self.ann.x2_px      = centre.x() + w / 2
        self.ann.y2_px      = centre.y() + h / 2
        self.ann.rotation_deg = self._angle

        # Update real-world dimensions if GSD is available
        if self._gsd:
            self.ann.width_cm  = round(w * self._gsd, 2)
            self.ann.height_cm = round(h * self._gsd, 2)

        if self.on_changed:
            self.on_changed(self.ann)

    # ── Click without drag → select + emit annotation_selected ───────────────

    def mouseDoubleClickEvent(self, event):
        # Double-click selects; no special action
        super().mouseDoubleClickEvent(event)


# Keep original read-only BoxAnnotationItem for any external use
class BoxAnnotationItem(QGraphicsRectItem):
    """Priya Nair: Legacy read-only bounding-box item (kept for compatibility)."""
    def __init__(self, ann: Annotation, gsd: Optional[float] = None, parent=None):
        super().__init__(parent)
        self.ann  = ann
        rect = QRectF(
            min(ann.x1_px, ann.x2_px), min(ann.y1_px, ann.y2_px),
            abs(ann.x2_px - ann.x1_px), abs(ann.y2_px - ann.y1_px)
        )
        self.setRect(rect)
        colour = SEVERITY_COLORS.get(ann.severity, QColor("#d29922"))
        self.setPen(QPen(colour, BOX_PEN_WIDTH))
        self.setBrush(QBrush(Qt.BrushStyle.NoBrush))
        self.setFlag(QGraphicsItem.GraphicsItemFlag.ItemIsSelectable)
        self.setZValue(10)
        self._build_label(colour)

    def _build_label(self, colour: QColor):
        ann = self.ann
        size_txt = (f"{ann.width_cm:.1f}×{ann.height_cm:.1f}cm"
                    if ann.width_cm is not None else "?cm")
        short = SEVERITY_SHORT.get(ann.severity, ann.severity)
        label = f"#{ann.ann_id[:5]} [{short}] {ann.defect}  {size_txt}"
        txt = QGraphicsTextItem(label, self)
        txt.setDefaultTextColor(QColor("white"))
        txt.setFont(QFont("Segoe UI", LABEL_FONT_SIZE, QFont.Weight.Bold))
        tr = txt.boundingRect()
        label_pos = self.rect().topLeft() + QPointF(0, -(tr.height() + 2))
        txt.setPos(label_pos)
        txt.setZValue(12)
        bg = QGraphicsRectItem(
            QRectF(label_pos.x() - 3, label_pos.y() - 2,
                   tr.width() + 6,      tr.height() + 4),
            self)
        bg.setBrush(QBrush(colour.darker(120)))
        bg.setPen(QPen(Qt.PenStyle.NoPen))
        bg.setZValue(11)


class PinAnnotationItem(QGraphicsEllipseItem):
    """Priya Nair: Point pin annotation."""
    def __init__(self, ann: Annotation, parent=None):
        super().__init__(parent)
        self.ann    = ann
        colour      = SEVERITY_COLORS.get(ann.severity, QColor("#d29922"))
        r           = PIN_RADIUS
        self.setRect(QRectF(ann.x1_px - r, ann.y1_px - r, r * 2, r * 2))
        self.setPen(QPen(colour, BOX_PEN_WIDTH))
        self.setBrush(QBrush(colour.darker(160)))
        self.setFlag(QGraphicsItem.GraphicsItemFlag.ItemIsSelectable)
        self.setZValue(10)
        short = SEVERITY_SHORT.get(ann.severity, ann.severity)
        txt   = QGraphicsTextItem(f"#{ann.ann_id[:5]} [{short}] {ann.defect}", self)
        txt.setDefaultTextColor(colour)
        txt.setFont(QFont("Segoe UI", LABEL_FONT_SIZE, QFont.Weight.Bold))
        txt.setPos(ann.x1_px + r + 4, ann.y1_px - r)
        txt.setZValue(12)


class PolygonAnnotationItem(QGraphicsPolygonItem):
    """Priya Nair: Multi-point polygon annotation — double-click to close."""
    def __init__(self, ann: Annotation, parent=None):
        super().__init__(parent)
        self.ann    = ann
        colour      = SEVERITY_COLORS.get(ann.severity, QColor("#d29922"))
        pts         = ann.poly_pts
        poly        = QPolygonF([QPointF(pts[i], pts[i + 1])
                                 for i in range(0, len(pts) - 1, 2)])
        self.setPolygon(poly)
        self.setPen(QPen(colour, POLY_PEN_WIDTH))
        # FIX-06: Previously: self.setBrush(...) followed by self.brush().setStyle(...)
        # self.brush() returns a VALUE COPY of the QBrush — calling .setStyle() on the
        # copy has no effect on the item.  Build a properly configured QBrush up-front
        # and call setBrush() exactly once.
        if colour.isValid():
            _fill = QBrush(colour.darker(180))
            _fill.setStyle(Qt.BrushStyle.SolidPattern)
            self.setBrush(_fill)
        else:
            self.setBrush(QBrush(Qt.BrushStyle.NoBrush))
        self.setOpacity(0.85)
        self.setFlag(QGraphicsItem.GraphicsItemFlag.ItemIsSelectable)
        self.setZValue(10)
        if pts:
            short = SEVERITY_SHORT.get(ann.severity, ann.severity)
            txt   = QGraphicsTextItem(f"#{ann.ann_id[:5]} [{short}] {ann.defect}", self)
            txt.setDefaultTextColor(colour)
            txt.setFont(QFont("Segoe UI", LABEL_FONT_SIZE, QFont.Weight.Bold))
            txt.setPos(QPointF(pts[0], pts[1]) + QPointF(4, -14))
            txt.setZValue(12)


class CalibrationLineItem(QGraphicsLineItem):
    """Elena Vasquez: Temporary calibration reference line."""
    def __init__(self, x1, y1, x2, y2, parent=None):
        super().__init__(x1, y1, x2, y2, parent)
        self.setPen(QPen(QColor("#FFEB3B"), CAL_PEN_WIDTH, Qt.PenStyle.DashLine))
        self.setZValue(20)


class BladePinpointWidget(QWidget):
    """
    v3.2.0 — Interactive mini blade silhouette in the annotation Location tab.
    Inspector clicks to place a coloured dot showing where on the blade the
    defect sits.  Position stored as normalised float: 0.0 = root, 1.0 = tip.
    Hidden automatically for Hub / Tower annotations.
    v3.3.6 — Displays currently selected face label above blade silhouette.
    v4.3.0 — Edge pinpointing: clicking the left half of the blade marks Leading
              Edge (LE), clicking the right half marks Trailing Edge (TE).
              Clicking within ±25% of the horizontal centre clears the edge selection.
              The selected edge is highlighted by a coloured side stripe on the silhouette.
    """
    position_changed = pyqtSignal(float)
    # v4.3.0: emitted whenever edge selection changes ("LE", "TE", or "")
    edge_changed     = pyqtSignal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self._pos      : float         = 0.5
        self._severity : str           = "Minor"
        self._face     : str           = ""    # v3.3.6: current face selection
        self._edge_side: Optional[str] = None  # v4.3.0: "LE" | "TE" | None
        # HEIGHT FIX: was setFixedSize(80, 180). The 180px height contributed
        # ~180px to the Location tab's minimum height (one of the tallest
        # fixed-size widgets in the panel).  Reduced to 120px — still gives
        # adequate click precision while reducing the panel minimum height.
        self.setFixedSize(80, 120)
        self.setCursor(QCursor(Qt.CursorShape.CrossCursor))
        self.setToolTip(
            "Click anywhere on the blade to mark defect location.\n"
            "  Left half → Leading Edge (LE)\n"
            "  Right half → Trailing Edge (TE)\n"
            "  Centre zone (±25%) → clears edge selection\n"
            "0 m = Root (hub end)  →  blade tip at bottom")

    def set_position(self, pos):
        self._pos = max(0.0, min(1.0, pos)) if pos is not None else 0.5
        self.update()

    def set_severity(self, sev: str):
        self._severity = sev
        self.update()

    def set_face(self, face: str):
        """v3.3.6: Update displayed face label when face combo changes."""
        # Extract abbreviation e.g. "Leading Edge (LE)" → "LE"
        if "(" in face:
            self._face = face.split("(")[-1].strip(")")
        else:
            self._face = face
        self.update()

    def set_edge_side(self, edge: Optional[str]):
        """v4.3.0: Set edge selection externally (e.g. when loading an annotation)."""
        self._edge_side = edge if edge in ("LE", "TE") else None
        self.update()

    def get_position(self) -> float:
        return self._pos

    def get_edge_side(self) -> Optional[str]:
        """v4.3.0: Return current edge selection ("LE", "TE", or None)."""
        return self._edge_side

    def _blade_geom(self):
        cx        = self.width() // 2
        blade_top = 24
        blade_bot = self.height() - 22
        return cx, blade_top, blade_bot

    def _half_w_at_t(self, t: float) -> float:
        """Half-width of blade at normalised position t (0=root/wide, 1=tip/narrow)."""
        return 14.0 * ((1.0 - t) ** 0.55)

    def paintEvent(self, event):
        from PyQt6.QtGui import QPainterPath as _PP
        p = QPainter(self)
        p.setRenderHint(QPainter.RenderHint.Antialiasing)
        cx, top, bot = self._blade_geom()
        h = bot - top
        W = self.width()

        # Root / Tip labels
        p.setFont(QFont("Segoe UI", 7))
        p.setPen(QColor(UI_THEME["text_tertiary"]))
        p.drawText(0, 2, W, 18, Qt.AlignmentFlag.AlignHCenter, "▲ Root")
        p.drawText(0, bot + 4, W, 18, Qt.AlignmentFlag.AlignHCenter, "▼ Tip")

        # v3.3.6: face label centred above blade — shows selected section
        if self._face:
            p.setFont(QFont("Segoe UI", 7, QFont.Weight.Bold))
            p.setPen(QColor(UI_THEME["accent_cyan"]))
            p.drawText(0, 2, W, 18, Qt.AlignmentFlag.AlignHCenter, self._face)

        # Blade silhouette path (tapered from root to tip)
        path = _PP()
        steps = 60
        pts_l, pts_r = [], []
        for i in range(steps + 1):
            t  = i / steps
            y  = top + t * h
            hw = self._half_w_at_t(t)
            pts_l.append(QPointF(cx - hw, y))
            pts_r.append(QPointF(cx + hw, y))
        path.moveTo(pts_l[0])
        for pt in pts_l[1:]:
            path.lineTo(pt)
        for pt in reversed(pts_r):
            path.lineTo(pt)
        path.closeSubpath()
        p.fillPath(path, QBrush(QColor("#2d3748")))
        p.setPen(QPen(QColor("#4a5568"), 1.2))
        p.drawPath(path)

        # v4.3.0: Draw edge highlight stripe on selected side
        if self._edge_side in ("LE", "TE"):
            edge_col = QColor(UI_THEME["accent_cyan"])
            edge_col.setAlphaF(0.55)
            edge_path = _PP()
            STRIPE_W = 4.0   # px width of highlight stripe
            if self._edge_side == "LE":
                # LE = left side of the blade silhouette
                edge_path.moveTo(pts_l[0])
                for pt in pts_l[1:]:
                    edge_path.lineTo(pt)
            else:
                # TE = right side of the blade silhouette
                edge_path.moveTo(pts_r[0])
                for pt in pts_r[1:]:
                    edge_path.lineTo(pt)
            stripe_pen = QPen(QColor(UI_THEME["accent_cyan"]), STRIPE_W)
            stripe_pen.setCapStyle(Qt.PenCapStyle.RoundCap)
            p.setPen(stripe_pen)
            p.drawPath(edge_path)
            # Small text label at the highlighted edge
            lbl_x = (pts_l[steps // 2].x() - 18) if self._edge_side == "LE" else (pts_r[steps // 2].x() + 2)
            p.setFont(QFont("Segoe UI", 6, QFont.Weight.Bold))
            p.setPen(QColor(UI_THEME["accent_cyan"]))
            p.drawText(int(lbl_x), int(top + h * 0.5), self._edge_side)

        # Dot at current Y position
        dot_y = top + self._pos * h
        dot_x = float(cx)
        # v4.3.0: Aligned with _SEV_HEX — Yellow / Amber / Red
        dot_col = {
            "Minor":    "#FFD700",   # Yellow
            "Major":    "#FFA500",   # Amber
            "Critical": "#FF0000",   # Red
            "POI":      "#388bfd",   # Blue
        }.get(self._severity, "#00d4e0")
        p.setPen(QPen(QColor("white"), 1.5))
        p.setBrush(QBrush(QColor(dot_col)))
        p.drawEllipse(QPointF(dot_x, dot_y), 7.0, 7.0)
        p.end()

    def mousePressEvent(self, event):
        """
        v4.3.0: Dual-axis click handling:
          Y axis → vertical position (root=0 → tip=1)
          X axis → edge selection:
            Left of blade centre by > 25% of local half-width → LE
            Right of blade centre by > 25% of local half-width → TE
            Otherwise (centre zone) → clear edge selection
        Both position_changed and edge_changed are emitted when their values change.
        """
        if event.button() == Qt.MouseButton.LeftButton:
            cx, blade_top, blade_bot = self._blade_geom()
            bh = blade_bot - blade_top
            if bh > 0:
                # Vertical position (root→tip)
                pos = (event.position().y() - blade_top) / bh
                self._pos = max(0.0, min(1.0, pos))
                self.position_changed.emit(self._pos)

                # Horizontal edge detection
                click_x = event.position().x()
                # Compute half-width at this blade position for normalised centre zone
                hw = self._half_w_at_t(self._pos)
                centre_zone = hw * 0.25   # ±25% of half-width = centre zone
                dx = click_x - cx         # signed offset from blade centre
                if abs(dx) <= centre_zone:
                    new_edge = None       # centre → clear edge
                elif dx < 0:
                    new_edge = "LE"       # left → Leading Edge
                else:
                    new_edge = "TE"       # right → Trailing Edge
                if new_edge != self._edge_side:
                    self._edge_side = new_edge
                    self.edge_changed.emit(self._edge_side or "")

                self.update()
            event.accept()

# ==============================================================================
# IMAGE VIEWER  (Marcus Webb + Priya Nair)
# Modes: Box | Pin | Polygon | Calibrate | Select/Pan
# Includes annotation-level Undo/Redo stack (Ctrl+Z / Ctrl+Y).
# ==============================================================================

class AnnotationSignals(QObject):
    annotation_ready    = pyqtSignal(object)   # new Annotation drafted
    annotation_selected = pyqtSignal(object)   # existing Annotation clicked
    annotation_deleted  = pyqtSignal(object)   # Delete key fired
    annotation_modified = pyqtSignal(object)   # interactive edit (resize/rotate/move)
    gsd_updated         = pyqtSignal(float)    # calibration complete
    calibration_metadata= pyqtSignal(dict)     # v3.4.0: EXIF calibration metadata
    mode_change_requested = pyqtSignal(str)    # v4.1.1: viewer requests toolbar mode sync (Key_S etc.)


class ImageViewer(QGraphicsView):
    """
    Marcus Webb + Priya Nair:
    Four draw modes + Select/Pan.  Polygon mode uses click-to-add-vertex +
    double-click-to-close.  Undo/Redo stack operates on the drawn annotation
    list for the current image — signals are emitted so MainWindow updates
    the data model accordingly.
    """
    MODE_BOX  = "box"
    MODE_PIN  = "pin"
    MODE_POLY = "polygon"
    MODE_CAL  = "cal"
    MODE_SEL  = "sel"

    def __init__(self, parent=None):
        super().__init__(parent)
        self.signals       = AnnotationSignals()
        self._scene        = QGraphicsScene(self)
        self.setScene(self._scene)
        self._mode         = self.MODE_SEL
        self._gsd          : Optional[float] = None
        self._session_gsd  : Optional[float] = None
        self._current_image_path : Optional[str] = None  # v3.4.0: For EXIF extraction

        # Box draw state
        self._box_start    : Optional[QPointF] = None
        self._rubber_band  : Optional[QGraphicsRectItem] = None
        self._live_label   : Optional[QGraphicsTextItem] = None

        # Polygon draw state
        self._poly_pts     : List[QPointF] = []
        self._poly_lines   : List[QGraphicsLineItem] = []
        self._poly_dots    : List[QGraphicsEllipseItem] = []
        self._poly_cursor  : Optional[QGraphicsLineItem] = None

        # Calibration state
        self._cal_pt1  : Optional[QPointF] = None
        self._cal_line : Optional[CalibrationLineItem] = None

        self._ann_items: List[QGraphicsItem] = []

        # Manual pan state (used in SEL mode on empty-space drag + middle mouse)
        self._pan_active   : bool   = False
        self._pan_last_vpos: QPoint = QPoint()  # viewport coords at last event

        # v3.3.0: active editable item receiving manual drag dispatch (no Qt mouse grab)
        self._active_editable: Optional[EditableBoxItem] = None

        # Annotation-level undo/redo stacks
        # CTO-AUDIT: Capped at _UNDO_LIMIT entries to prevent memory growth
        # in long inspection sessions (each Annotation object ~1 KB).
        self._UNDO_LIMIT   = 50
        self._undo_stack: List[Annotation] = []
        self._redo_stack: List[Annotation] = []

        # Config
        self._zoom_min = float(CFG.get("GENERAL", "ZoomMin", "0.05"))
        self._zoom_max = float(CFG.get("GENERAL", "ZoomMax", "20.0"))

        self.setRenderHint(QPainter.RenderHint.Antialiasing)
        self.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform)
        self.setTransformationAnchor(QGraphicsView.ViewportAnchor.AnchorUnderMouse)
        self.setResizeAnchor(QGraphicsView.ViewportAnchor.AnchorUnderMouse)
        self.setDragMode(QGraphicsView.DragMode.NoDrag)
        self.setBackgroundBrush(QBrush(QColor(UI_THEME["bg_primary"])))
        self.setStyleSheet(f"border: 2px solid {UI_THEME['border']};")

    # ── Public API ─────────────────────────────────────────────────────────────

    def load_image(self, filepath: str, image_record: Optional[ImageRecord] = None,
                   gsd: Optional[float] = None, session_gsd: Optional[float] = None):
        """Marcus Webb: Load image, clear state, redraw saved annotations."""
        self._current_image_path = filepath  # v3.4.0: Store for EXIF extraction
        self._clear_drawing_state()
        self._scene.clear()
        self._ann_items.clear()
        self._undo_stack.clear()
        self._redo_stack.clear()
        self._active_editable = None   # clear stale drag reference on image switch

        qimg = QImage(filepath)
        if qimg.isNull():
            try:
                # BytesIO imported at module level (FIX-02b: removed inline import)
                pil = Image.open(filepath).convert("RGB")
                buf = BytesIO()
                pil.save(buf, format="JPEG", quality=92)
                buf.seek(0)
                qimg.loadFromData(buf.read())
            except Exception as exc:
                log.error(f"Image load failed: {exc}")
                return

        pixmap = QPixmap.fromImage(qimg)
        self._scene.setSceneRect(QRectF(pixmap.rect()))
        self._scene.addPixmap(pixmap)
        self._gsd         = gsd
        self._session_gsd = session_gsd

        if image_record:
            eff_gsd = gsd or session_gsd
            for ann in image_record.annotations:
                self._draw_annotation(ann, eff_gsd)

        self.fitInView(self._scene.sceneRect(), Qt.AspectRatioMode.KeepAspectRatio)

    def set_mode(self, mode: str):
        self._mode = mode
        self._clear_drawing_state()
        # Always use NoDrag — we manage panning manually in SEL mode so that
        # EditableBoxItem can still receive left-button press events.
        # (ScrollHandDrag at the view level swallows all left-mouse events and
        # prevents scene items from ever seeing them.)
        self.setDragMode(QGraphicsView.DragMode.NoDrag)
        self._pan_active = False
        if mode == self.MODE_SEL:
            self.setStyleSheet(f"border: 2px solid {UI_THEME['border']};")
            self.setCursor(QCursor(Qt.CursorShape.ArrowCursor))
        else:
            # Colour the canvas border to indicate active draw mode
            colour_map = {
                self.MODE_BOX:  UI_THEME["accent_cyan"],
                self.MODE_PIN:  UI_THEME["accent_green"],
                self.MODE_POLY: UI_THEME["accent_purple"],
                self.MODE_CAL:  UI_THEME["accent_amber"],
            }
            c = colour_map.get(mode, UI_THEME["border"])
            self.setStyleSheet(f"border: 2px solid {c};")
            self.setCursor(QCursor(Qt.CursorShape.CrossCursor))

    def set_gsd(self, gsd: Optional[float]):
        self._gsd = gsd

    def set_session_gsd(self, gsd: Optional[float]):
        self._session_gsd = gsd

    def current_gsd(self) -> Optional[float]:
        return self._gsd or self._session_gsd

    def draw_annotation(self, ann: Annotation):
        """Add a freshly-saved annotation item to the scene."""
        item = self._draw_annotation(ann, self.current_gsd())
        # CTO-AUDIT: push to undo stack and enforce depth cap
        self._undo_stack.append(ann)
        if len(self._undo_stack) > self._UNDO_LIMIT:
            self._undo_stack.pop(0)   # drop oldest entry
        self._redo_stack.clear()

    def remove_annotation_item(self, ann: Annotation):
        for item in list(self._ann_items):
            if hasattr(item, "ann") and item.ann.ann_id == ann.ann_id:
                self._scene.removeItem(item)
                self._ann_items.remove(item)
                break

    def undo_last(self) -> Optional[Annotation]:
        """Sam Okafor: Remove the most-recently-drawn annotation from scene."""
        if not self._undo_stack:
            return None
        ann = self._undo_stack.pop()
        self._redo_stack.append(ann)
        self.remove_annotation_item(ann)
        return ann

    def redo_last(self) -> Optional[Annotation]:
        """Sam Okafor: Re-draw the last undone annotation."""
        if not self._redo_stack:
            return None
        ann = self._redo_stack.pop()
        self._undo_stack.append(ann)
        self._draw_annotation(ann, self.current_gsd())
        return ann

    # ── Internal draw helpers ──────────────────────────────────────────────────

    def _draw_annotation(self, ann: Annotation,
                         gsd: Optional[float] = None) -> Optional[QGraphicsItem]:
        item: Optional[QGraphicsItem] = None
        if ann.mode == "box":
            box_item = EditableBoxItem(ann, gsd)
            # Wire geometry-change callback → emit annotation_modified signal
            box_item.on_changed = lambda a: self.signals.annotation_modified.emit(a)
            item = box_item
        elif ann.mode == "pin":
            item = PinAnnotationItem(ann)
        elif ann.mode == "polygon" and len(ann.poly_pts) >= 4:
            item = PolygonAnnotationItem(ann)
        if item:
            self._scene.addItem(item)
            self._ann_items.append(item)
        return item

    def _clear_drawing_state(self):
        """Reset all in-progress draw state cleanly."""
        if self._rubber_band:
            self._scene.removeItem(self._rubber_band)
            self._rubber_band = None
        if self._live_label:
            self._scene.removeItem(self._live_label)
            self._live_label = None
        for ln in self._poly_lines:
            self._scene.removeItem(ln)
        for dot in self._poly_dots:
            self._scene.removeItem(dot)
        if self._poly_cursor:
            self._scene.removeItem(self._poly_cursor)
        self._poly_pts.clear()
        self._poly_lines.clear()
        self._poly_dots.clear()
        self._poly_cursor = None
        if self._cal_line:
            self._scene.removeItem(self._cal_line)
            self._cal_line = None
        self._cal_pt1 = None
        self._box_start = None
        self._active_editable = None   # clear drag target when drawing state is reset

    def _cm_label(self, dx: float, dy: float) -> str:
        gsd = self.current_gsd()
        if gsd:
            return f"{abs(dx) * gsd:.1f} × {abs(dy) * gsd:.1f} cm"
        return f"{abs(dx):.0f} × {abs(dy):.0f} px"

    # ── Mouse Events ───────────────────────────────────────────────────────────

    def wheelEvent(self, event):
        # v4.1.1: plain scroll zooms at cursor; no Ctrl required.
        # self.scale() ignores AnchorUnderMouse so we implement the anchor manually:
        #   1. record the scene point currently under the cursor
        #   2. apply scale transform
        #   3. shift scrollbars so that same scene point is back under the cursor
        factor    = 1.15 if event.angleDelta().y() > 0 else 1 / 1.15
        cur_scale = self.transform().m11()
        if not (self._zoom_min <= cur_scale * factor <= self._zoom_max):
            event.accept()
            return
        # Scene point under cursor before scaling
        cursor_vp    = event.position().toPoint()
        scene_before = self.mapToScene(cursor_vp)
        self.scale(factor, factor)
        # Scene point that is now under cursor after scaling
        scene_after  = self.mapToScene(cursor_vp)
        # Shift scrollbars by the delta to bring scene_before back under cursor
        delta = scene_after - scene_before
        self.horizontalScrollBar().setValue(
            self.horizontalScrollBar().value() - int(delta.x() * self.transform().m11()))
        self.verticalScrollBar().setValue(
            self.verticalScrollBar().value()  - int(delta.y() * self.transform().m22()))
        event.accept()

    def mousePressEvent(self, event):
        sp  = self.mapToScene(event.position().toPoint())
        btn = event.button()
        vp  = event.position().toPoint()   # viewport coords for manual pan

        # ── Middle-mouse pan (available in ALL modes) ─────────────────────────
        if btn == Qt.MouseButton.MiddleButton:
            self._pan_active    = True
            self._pan_last_vpos = vp
            self.setCursor(QCursor(Qt.CursorShape.ClosedHandCursor))
            event.accept()
            return

        # v4.1.1: Right-click in SEL mode → delete selected annotation under cursor
        if btn == Qt.MouseButton.RightButton and self._mode == self.MODE_SEL:
            target = None
            for item in self._scene.items(sp):
                if isinstance(item, EditableBoxItem) and item.isSelected():
                    target = item
                    break
                if isinstance(item, EditableBoxItem):
                    target = item   # also allow right-click on unselected box
                    break
            if target is not None:
                self.signals.annotation_deleted.emit(target.ann)
                self._scene.removeItem(target)
                if target in self._ann_items:
                    self._ann_items.remove(target)
            event.accept()
            return

        if self._mode == self.MODE_BOX and btn == Qt.MouseButton.LeftButton:
            self._box_start = sp
            self._rubber_band = QGraphicsRectItem(QRectF(sp, QSizeF(0, 0)))
            # Make rubber band clearly visible with thicker dashed line
            self._rubber_band.setPen(QPen(QColor(UI_THEME["accent_cyan"]),
                                         4, Qt.PenStyle.DashLine))  # Increased from DRAW_PEN_WIDTH to 4
            self._rubber_band.setBrush(QBrush(QColor(0, 212, 224, 20)))  # Semi-transparent fill
            self._rubber_band.setZValue(50)
            self._rubber_band.setVisible(True)  # Explicitly set visible
            self._scene.addItem(self._rubber_band)
            self._live_label = QGraphicsTextItem("")
            self._live_label.setDefaultTextColor(QColor(UI_THEME["accent_cyan"]))
            self._live_label.setFont(QFont("Segoe UI", LABEL_FONT_SIZE, QFont.Weight.Bold))
            self._live_label.setZValue(51)
            self._scene.addItem(self._live_label)

        elif self._mode == self.MODE_PIN and btn == Qt.MouseButton.LeftButton:
            ann = Annotation(ann_id=_make_ann_id(str(sp)), mode="pin",
                             x1_px=sp.x(), y1_px=sp.y(),
                             x2_px=sp.x(), y2_px=sp.y())
            self.signals.annotation_ready.emit(ann)

        elif self._mode == self.MODE_POLY and btn == Qt.MouseButton.LeftButton:
            # Add vertex; draw helper dot
            self._poly_pts.append(sp)
            dot = QGraphicsEllipseItem(sp.x() - 3, sp.y() - 3, 6, 6)
            dot.setBrush(QBrush(QColor(UI_THEME["accent_purple"])))
            dot.setPen(QPen(Qt.PenStyle.NoPen))
            dot.setZValue(52)
            self._scene.addItem(dot)
            self._poly_dots.append(dot)
            # Connect previous vertex
            if len(self._poly_pts) >= 2:
                prev = self._poly_pts[-2]
                ln   = QGraphicsLineItem(prev.x(), prev.y(), sp.x(), sp.y())
                # FIX-LINES: DRAW_PEN_WIDTH (3px dashed) for in-progress polygon edges
                ln.setPen(QPen(QColor(UI_THEME["accent_purple"]),
                               DRAW_PEN_WIDTH, Qt.PenStyle.DashLine))
                ln.setZValue(51)
                self._scene.addItem(ln)
                self._poly_lines.append(ln)

        elif self._mode == self.MODE_CAL and btn == Qt.MouseButton.LeftButton:
            if self._cal_pt1 is None:
                self._cal_pt1 = sp
                self._cal_line = CalibrationLineItem(sp.x(), sp.y(), sp.x(), sp.y())
                self._scene.addItem(self._cal_line)
            else:
                # Second click — compute distance and ask for real-world length
                self._cal_line.setLine(self._cal_pt1.x(), self._cal_pt1.y(),
                                        sp.x(), sp.y())
                dx = sp.x() - self._cal_pt1.x()
                dy = sp.y() - self._cal_pt1.y()
                px_dist = math.hypot(dx, dy)
                if px_dist > 2:
                    # v3.4.0: Try EXIF auto-calibration first
                    exif_cal = None
                    try:
                        if self._current_image_path:
                            calibrator = EXIFCalibrator(self._current_image_path)
                            exif_cal = calibrator.calibrate()
                    except Exception as e:
                        log.warning(f"[EXIF] Auto-calibration failed: {e}")

                    if exif_cal and exif_cal.confidence in [ConfidenceLevel.HIGH, ConfidenceLevel.MEDIUM]:
                        # Offer auto-calibration
                        estimated_dist = exif_cal.estimate_distance_from_gps()
                        auto_gsd = exif_cal.calculate_gsd_cm_per_px(estimated_dist) if estimated_dist else None
                        
                        if auto_gsd:
                            msg = (
                                f"📷 Camera Detected: {exif_cal.camera_model}\n"
                                f"🎯 Focal Length: {exif_cal.focal_length_mm:.1f}mm\n"
                                f"📏 Estimated Distance: {estimated_dist:.1f}m\n"
                                f"✓ Calculated GSD: {auto_gsd:.4f} cm/px\n"
                                f"🔒 Confidence: {exif_cal.confidence.value}\n\n"
                                f"Use automatic EXIF calibration?"
                            )
                            
                            reply = QMessageBox.question(
                                self.viewport(),
                                "Auto-Calibration Available",
                                msg,
                                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                QMessageBox.StandardButton.Yes
                            )
                            
                            if reply == QMessageBox.StandardButton.Yes:
                                # Use EXIF calibration
                                self.signals.gsd_updated.emit(auto_gsd)
                                # Store metadata (MainWindow will handle this via signal)
                                self.signals.calibration_metadata.emit({
                                    'exif_cal': exif_cal,
                                    'gsd': auto_gsd,
                                    'method': 'exif-auto'
                                })
                                self._clear_drawing_state()
                                return

                    # Fall back to manual calibration
                    real_cm, ok = QInputDialog.getDouble(
                        self.viewport(), "Manual Calibration",
                        f"Pixel distance: {px_dist:.1f}px\n"
                        "Enter the real-world distance in centimetres:",
                        decimals=2, min=0.01, max=50000.0
                    )
                    if ok and real_cm > 0:
                        gsd = real_cm / px_dist
                        self.signals.gsd_updated.emit(gsd)
                        # Store manual calibration metadata
                        self.signals.calibration_metadata.emit({
                            'exif_cal': None,
                            'gsd': gsd,
                            'method': 'manual'
                        })
                self._clear_drawing_state()

        elif self._mode == self.MODE_SEL and btn == Qt.MouseButton.LeftButton:
            # ── Select / Edit mode ────────────────────────────────────────────
            # Determine what is under the cursor.
            # Priority: EditableBoxItem first (handles its own drags), then
            # other annotation items, then empty-space → start pan.
            hit_editable = None
            hit_other    = None
            for item in self._scene.items(sp):
                if isinstance(item, EditableBoxItem):
                    hit_editable = item
                    break
                if hasattr(item, "ann") and not isinstance(
                        item, (QGraphicsTextItem, QGraphicsRectItem)):
                    if hit_other is None:
                        hit_other = item

            if hit_editable is not None:
                # FIX v4.1.0: Properly initialize selection state for immediate editing
                # Deselect everything else first
                for it in self._scene.selectedItems():
                    if it is not hit_editable:
                        it.setSelected(False)
                # Set selected and force immediate update to show handles
                hit_editable.setSelected(True)
                hit_editable.setFlag(QGraphicsItem.GraphicsItemFlag.ItemIsSelectable, True)
                hit_editable.update()          # repaint so handles appear immediately
                self.signals.annotation_selected.emit(hit_editable.ann)
                
                # CRITICAL: Do NOT call super().mousePressEvent() here.
                # Qt's scene dispatch calls clearSelection() *before* delivering
                # the press to items, which deselects the item so handles never
                # paint (isSelected()==False in paint()).  Instead manually start
                # the drag so resize/rotate still works on the very first click.
                
                # FIX-14a: Use screen-space (viewport) coordinates to pick the
                # drag handle, not item-local scene coords.  At low zoom the
                # item-local _HS radius (8 px scene-space) can shrink to < 2 px on
                # screen, making handle clicks miss and fall through to a body-drag.
                # We convert every handle centre → viewport → compare with a fixed
                # screen radius, which stays consistently clickable at any zoom.
                #
                # CRITICAL ORDER: ROT handle (index 8) MUST be checked before the
                # resize handles (0-7).  At low zoom or for a small box, the ROT
                # handle (30 item-px above TC) can land within SCREEN_HIT_R of TC
                # (index 1) in viewport space.  If we enumerated 0→8, TC would
                # match first and the user could never grab the rotation handle.
                # Solution: check index 8 (ROT) explicitly with its own block,
                # fall through to the resize-handle loop only if ROT misses.
                SCREEN_HIT_R = 20          # raised 16→20px: easier to grab handles
                handle_idx   = -1          # -1 → body drag
                vp_click     = event.position().toPoint()   # viewport coords of click
                _sq = SCREEN_HIT_R * SCREEN_HIT_R

                # ── Priority: rotation handle (index 8) checked first ──────────
                _rot_hc  = hit_editable._handle_centres()[hit_editable._ROT]
                _rot_sc  = hit_editable.mapToScene(_rot_hc)
                _rot_vc  = self.mapFromScene(_rot_sc)
                _rot_dhx = _rot_vc.x() - vp_click.x()
                _rot_dhy = _rot_vc.y() - vp_click.y()
                if _rot_dhx * _rot_dhx + _rot_dhy * _rot_dhy <= _sq:
                    handle_idx = hit_editable._ROT
                else:
                    # ── Resize handles (indices 0-7) checked only if ROT missed ─
                    for _hi, _hc in enumerate(hit_editable._handle_centres()):
                        if _hi == hit_editable._ROT:
                            continue  # already tested above; skip to prevent duplicate
                        _sc  = hit_editable.mapToScene(_hc)
                        _vc  = self.mapFromScene(_sc)
                        _dhx = _vc.x() - vp_click.x()
                        _dhy = _vc.y() - vp_click.y()
                        if _dhx * _dhx + _dhy * _dhy <= _sq:
                            handle_idx = _hi
                            break
                hit_editable._drag_handle = handle_idx
                hit_editable._drag_active  = True
                hit_editable._drag_scene0  = sp
                hit_editable._drag_rect0   = QRectF(hit_editable._rect)
                hit_editable._drag_pos0    = hit_editable.pos()
                hit_editable._drag_angle0  = hit_editable._angle
                # T01 FIX: _drag_angle_bias = angle at press (center→click).
                # Used as 'previous frame angle' for incremental rotation tracking.
                _cs = hit_editable.mapToScene(QPointF(0.0, 0.0))
                _dx = sp.x() - _cs.x()
                _dy = sp.y() - _cs.y()
                hit_editable._drag_angle_bias = (
                    math.degrees(math.atan2(_dy, _dx)) + 90.0)
                hit_editable._drag_prev_scene = sp   # for incremental resize
                self._active_editable = hit_editable  # v3.3.0: track for direct dispatch
                event.accept()
                return  # FIX v4.1.0: return immediately to prevent super() call

            elif hit_other is not None:
                # Non-editable annotation (pin/polygon) — emit selected only
                self._scene.clearSelection()
                hit_other.setSelected(True)
                self.signals.annotation_selected.emit(hit_other.ann)
                self._active_editable = None   # v3.3.0: no editable target

            else:
                # Empty space — deselect all, start pan
                self._scene.clearSelection()
                self._active_editable = None   # v3.3.0: no editable target
                self._pan_active    = True
                self._pan_last_vpos = vp
                self.setCursor(QCursor(Qt.CursorShape.ClosedHandCursor))

        else:
            super().mousePressEvent(event)

    def mouseDoubleClickEvent(self, event):
        """Priya Nair: Double-click closes polygon and emits annotation_ready."""
        if self._mode == self.MODE_POLY and len(self._poly_pts) >= 3:
            pts_flat = []
            for pt in self._poly_pts:
                pts_flat.extend([pt.x(), pt.y()])
            ann = Annotation(
                ann_id=_make_ann_id(str(self._poly_pts[0])),
                mode="polygon",
                x1_px=self._poly_pts[0].x(), y1_px=self._poly_pts[0].y(),
                x2_px=self._poly_pts[-1].x(), y2_px=self._poly_pts[-1].y(),
                poly_pts=pts_flat,
            )
            self._clear_drawing_state()
            self.signals.annotation_ready.emit(ann)
        else:
            super().mouseDoubleClickEvent(event)

    def mouseMoveEvent(self, event):
        sp = self.mapToScene(event.position().toPoint())
        vp = event.position().toPoint()

        # ── Manual pan (middle mouse or empty-space drag in SEL mode) ─────────
        if self._pan_active:
            delta = vp - self._pan_last_vpos
            self._pan_last_vpos = vp
            self.horizontalScrollBar().setValue(
                self.horizontalScrollBar().value() - delta.x())
            self.verticalScrollBar().setValue(
                self.verticalScrollBar().value() - delta.y())
            event.accept()
            return

        if self._mode == self.MODE_BOX and self._box_start and self._rubber_band:
            rect = QRectF(self._box_start, sp).normalized()
            self._rubber_band.setRect(rect)
            if self._live_label:
                self._live_label.setPlainText(
                    self._cm_label(rect.width(), rect.height()))
                self._live_label.setPos(rect.topLeft() + QPointF(4, -18))

        elif self._mode == self.MODE_POLY and self._poly_pts:
            # Cursor line from last vertex to mouse
            if self._poly_cursor:
                self._scene.removeItem(self._poly_cursor)
            last = self._poly_pts[-1]
            self._poly_cursor = QGraphicsLineItem(
                last.x(), last.y(), sp.x(), sp.y())
            # FIX-LINES: DRAW_PEN_WIDTH (3px dotted) for polygon cursor tracking line
            self._poly_cursor.setPen(
                QPen(QColor(UI_THEME["accent_purple"]), DRAW_PEN_WIDTH, Qt.PenStyle.DotLine))
            self._poly_cursor.setZValue(51)
            self._scene.addItem(self._poly_cursor)

        elif self._mode == self.MODE_CAL and self._cal_pt1 and self._cal_line:
            self._cal_line.setLine(
                self._cal_pt1.x(), self._cal_pt1.y(), sp.x(), sp.y())

        else:
            # SEL mode (and any other): directly dispatch to active editable if
            # drag is live (item has no Qt mouse grab — v3.3.0 bypass fix).
            if (self._active_editable is not None
                    and self._active_editable._drag_active):
                self._active_editable._update_drag(sp, event.modifiers())
            elif self._mode != self.MODE_SEL:
                # Bug #3 FIX: skip super() in SEL mode when no editable drag is
                # active — Qt's scene dispatch iterates every item (O(n)) on each
                # call, causing noticeable lag on annotation-heavy images.  In SEL
                # mode with nothing to drag there is no useful scene work to do.
                super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        sp  = self.mapToScene(event.position().toPoint())
        btn = event.button()

        # ── Stop manual pan ───────────────────────────────────────────────────
        if self._pan_active and btn in (Qt.MouseButton.LeftButton,
                                        Qt.MouseButton.MiddleButton):
            self._pan_active = False
            if self._mode == self.MODE_SEL:
                self.setCursor(QCursor(Qt.CursorShape.ArrowCursor))
            else:
                self.unsetCursor()
            event.accept()
            return

        if self._mode == self.MODE_BOX and self._box_start and self._rubber_band:
            rect = QRectF(self._box_start, sp).normalized()
            if rect.width() > 5 and rect.height() > 5:
                gsd = self.current_gsd()
                w_cm = rect.width()  * gsd if gsd else None
                h_cm = rect.height() * gsd if gsd else None
                ann  = Annotation(
                    ann_id=_make_ann_id(str(self._box_start)),
                    mode="box",
                    x1_px=rect.x(), y1_px=rect.y(),
                    x2_px=rect.x() + rect.width(),
                    y2_px=rect.y() + rect.height(),
                    width_cm=w_cm, height_cm=h_cm,
                    gsd_source="image" if self._gsd else "session" if self._session_gsd else "none",
                    gsd_value=gsd,
                )
                # FIX v4.1.0: Draw annotation immediately so it's visible on canvas
                # This creates the visual representation right away instead of waiting for save
                self._draw_annotation(ann, gsd)
                self.signals.annotation_ready.emit(ann)
        if self._rubber_band:
            self._scene.removeItem(self._rubber_band)
            self._rubber_band = None
        if self._live_label:
            self._scene.removeItem(self._live_label)
            self._live_label = None
        self._box_start = None

        # v3.3.0: commit active editable drag (no Qt mouse grab — manual dispatch)
        if (self._active_editable is not None
                and self._active_editable._drag_active):
            self._active_editable._finish_drag()
            self._active_editable = None

        # Always dispatch to scene items (needed by EditableBoxItem release)
        super().mouseReleaseEvent(event)

    def keyPressEvent(self, event):
        """Escape cancels in-progress polygon; Delete/S handled; events propagate."""
        if event.key() == Qt.Key.Key_Escape:
            self._clear_drawing_state()
            return
        if event.key() == Qt.Key.Key_Delete:
            for item in self._scene.selectedItems():
                if hasattr(item, "ann"):
                    self.signals.annotation_deleted.emit(item.ann)
                    self._scene.removeItem(item)
                    if item in self._ann_items:
                        self._ann_items.remove(item)
            return
        # v4.1.1: S key → switch to select mode from within viewer; sync toolbar via signal
        if (event.key() == Qt.Key.Key_S
                and not event.modifiers() & Qt.KeyboardModifier.ControlModifier):
            self.set_mode(self.MODE_SEL)
            self.signals.mode_change_requested.emit(self.MODE_SEL)
            event.accept()
            return
        super().keyPressEvent(event)

# ==============================================================================
# SEVERITY PILL STRIP  (Dev Patel — UX 8yr)
# Horizontal row of coloured capsule buttons — one per Scopito severity level.
# Selected button glows; colours match SEVERITY_COLORS exactly.
# ==============================================================================

class SeverityPillStrip(QWidget):
    """Dev Patel: Horizontal pill-button severity selector."""
    severity_changed = pyqtSignal(str)   # emits the chosen severity string

    def __init__(self, parent=None):
        super().__init__(parent)
        self._btns: Dict[str, QPushButton] = {}
        self._group = QButtonGroup(self)
        self._group.setExclusive(True)
        lay = QHBoxLayout(self)
        lay.setContentsMargins(0, 0, 0, 0)
        lay.setSpacing(4)

        for sev in SEVERITY_ACTIVE:
            col = SEVERITY_COLORS.get(sev, QColor("#7d8590"))
            hex_c = col.name()
            short = SEVERITY_SHORT.get(sev, sev)
            btn = QPushButton(short)
            btn.setCheckable(True)
            btn.setFixedHeight(28)
            btn.setToolTip(f"{sev}\n{SEVERITY_REMEDY.get(sev, '')}")
            btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {UI_THEME['bg_elevated']};
                    color: {hex_c};
                    border: 2px solid {hex_c};
                    border-radius: 14px;
                    padding: 0 10px;
                    font-size: 8pt;
                    font-weight: bold;
                }}
                QPushButton:hover {{
                    background-color: {hex_c};
                    color: #0d1117;
                }}
                QPushButton:checked {{
                    background-color: {hex_c};
                    color: #0d1117;
                    border-color: {hex_c};
                    font-weight: bold;
                }}
            """)
            self._group.addButton(btn)
            lay.addWidget(btn)
            self._btns[sev] = btn
            btn.clicked.connect(lambda _, s=sev: self.severity_changed.emit(s))

        lay.addStretch()
        # Default to Sev 1
        if SEVERITY_ACTIVE:
            self._btns[SEVERITY_ACTIVE[0]].setChecked(True)

    def set_severity(self, sev: str):
        # Try exact match, then fall back to first
        if sev in self._btns:
            self._btns[sev].setChecked(True)
        elif self._btns:
            next(iter(self._btns.values())).setChecked(True)

    def current_severity(self) -> str:
        for sev, btn in self._btns.items():
            if btn.isChecked():
                return sev
        return SEVERITY_ACTIVE[0] if SEVERITY_ACTIVE else "Minor"

# ==============================================================================
# COLLAPSIBLE SECTION  (v3.3.7 — accordion panel to replace QTabWidget)
# Clicking the header button expands/collapses the content widget in place.
# When expanded the content uses all available vertical space — no clipping.
# ==============================================================================

class CollapsibleSection(QWidget):
    """Accordion panel: header button toggles content visibility."""

    def __init__(self, title: str, content: QWidget,
                 start_open: bool = True, parent=None):
        super().__init__(parent)
        self.setSizePolicy(QSizePolicy.Policy.Expanding,
                           QSizePolicy.Policy.Minimum)
        lay = QVBoxLayout(self)
        lay.setSpacing(0)
        lay.setContentsMargins(0, 0, 0, 2)

        self._title   = title
        self._content = content

        # Header toggle button
        self._btn = QPushButton()
        self._btn.setCheckable(True)
        self._btn.setChecked(start_open)
        self._btn.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self._btn.setSizePolicy(QSizePolicy.Policy.Expanding,
                                QSizePolicy.Policy.Fixed)
        self._btn.setFixedHeight(30)
        self._btn.clicked.connect(self._on_toggle)
        self._update_btn_text(start_open)
        self._btn.setStyleSheet(f"""
            QPushButton {{
                background: {UI_THEME['bg_elevated']};
                color: {UI_THEME['text_secondary']};
                border: none;
                border-top: 1px solid {UI_THEME['border']};
                border-bottom: 1px solid {UI_THEME['border']};
                text-align: left;
                padding: 0 10px;
                font-size: 8pt;
                font-weight: bold;
                letter-spacing: 1px;
            }}
            QPushButton:hover {{
                background: {UI_THEME['bg_card']};
                color: {UI_THEME['text_primary']};
            }}
        """)
        lay.addWidget(self._btn)
        lay.addWidget(content)
        content.setVisible(start_open)

    def _update_btn_text(self, is_open: bool):
        arrow = "▼" if is_open else "▶"
        self._btn.setText(f"  {arrow}  {self._title}")

    def _on_toggle(self, checked: bool):
        self._content.setVisible(checked)
        self._update_btn_text(checked)

    def expand(self):
        self._btn.setChecked(True)
        self._on_toggle(True)

    def collapse(self):
        self._btn.setChecked(False)
        self._on_toggle(False)


# ==============================================================================
# ANNOTATION PANEL  (Sam Okafor + Sarah Chen)
# Tabbed card: Details / Location / Notes — professional form layout.
# ==============================================================================

class AnnotationPanel(QWidget):
    """Sam Okafor + Sarah Chen: Tabbed annotation editing panel."""
    save_requested    = pyqtSignal(object)   # Annotation
    delete_requested  = pyqtSignal(object)   # Annotation
    discard_requested = pyqtSignal(object)   # Annotation — FIX-UX: discard unsaved
    approve_requested = pyqtSignal(object)   # Annotation — Phase 6
    reject_requested  = pyqtSignal(object)   # Annotation — Phase 6
    # v1.7.0: rename_requested — (old_filepath, new_stem) so MainWindow can handle
    # the disk rename and project JSON update without importing os inside the panel.
    rename_requested     = pyqtSignal(str, str) # (current_filepath, new_stem)
    ann_selected_for_qc  = pyqtSignal(object)   # v4.1.1: annotation chosen in list → feed QC Review panel

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setMinimumWidth(290)
        self.setMaximumWidth(360)
        self._pending_ann: Optional[Annotation] = None
        self._is_new_annotation: bool = False   # FIX-UX: track unsaved state
        # Phase 9.3: Safe init — ensure all optional attributes have defaults
        # before _build_ui potentially triggers signal callbacks mid-construction
        self._current_filepath: str = ""
        # FIX-BUG: _project must be initialised here (before _build_ui) so that
        # load_pending() — which references self._project for filename suggestion
        # — never raises AttributeError regardless of whether set_project() has
        # been called yet.  MainWindow calls set_project() after construction via
        # the new set_project() public method added below.
        self._project: Optional["Project"] = None
        try:
            self._build_ui()
        except Exception as _exc:
            log.error(f"AnnotationPanel._build_ui failed: {_exc}", exc_info=True)
            # Fallback: create a minimal error label so the app doesn't crash
            _root = QVBoxLayout(self)
            _err  = QLabel(f"⚠️ Panel init error:\n{_exc}")
            _err.setStyleSheet("color:#f85149;padding:8px;")
            _err.setWordWrap(True)
            _root.addWidget(_err)
            # Stub out attributes that MainWindow accesses to prevent AttributeErrors
            self._pending_ann        = None
            self._save_btn           = QPushButton()
            self._discard_btn        = QPushButton()
            self._del_btn            = QPushButton()
            self._ann_list           = QListWidget()
            self._ann_toggle_btn     = QPushButton()  # v3.3.5: needed by _toggle_ann_list
            self._ann_section        = None           # CHG-I: CollapsibleSection stub
            self._review_section     = None           # CHG-I: CollapsibleSection stub

    def _build_ui(self):
        # ── SCROLL FIX: Sticky-footer layout ─────────────────────────────────
        # outer: vertical layout on self — contains scroll area (stretch) + footer
        # root:  layout INSIDE the scroll area content widget — all scrollable items
        # footer: btn_row, review_frame, ann_list pinned below scroll area (always visible)
        outer = QVBoxLayout(self)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(0)

        _scroll = QScrollArea()
        _scroll.setWidgetResizable(True)
        _scroll.setFrameShape(QFrame.Shape.NoFrame)
        _scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        _scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        _scroll.setStyleSheet(
            "QScrollArea{background:transparent;border:none;}"
            "QScrollBar:vertical{width:6px;background:transparent;}"
            "QScrollBar::handle:vertical{background:#444c56;border-radius:3px;min-height:20px;}"
            "QScrollBar::add-line:vertical,QScrollBar::sub-line:vertical{height:0px;}")
        _scroll_content = QWidget()
        _scroll_content.setStyleSheet("background:transparent;")
        root = QVBoxLayout(_scroll_content)
        root.setContentsMargins(8, 8, 8, 4)
        root.setSpacing(8)
        _scroll.setWidget(_scroll_content)
        outer.addWidget(_scroll, 1)   # stretch=1 → scroll area expands to fill space

        # ── FIX-UX: Workflow guidance banner ─────────────────────────────────
        # Shown when a new annotation is drawn. Tells the user exactly what to do.
        self._workflow_banner = QWidget()
        self._workflow_banner.setStyleSheet(
            f"background:{UI_THEME['accent_cyan']}20;"
            f"border:1px solid {UI_THEME['accent_cyan']};"
            f"border-radius:6px;")
        wb_lay = QHBoxLayout(self._workflow_banner)
        wb_lay.setContentsMargins(10, 7, 10, 7)
        wb_lay.setSpacing(6)
        wb_icon = QLabel("💡")
        wb_icon.setStyleSheet("background:transparent;font-size:12pt;")
        wb_text = QLabel(
            "<b>New annotation drawn!</b><br/>"
            "<span style='font-size:9pt;'>"
            "1. Pick a <b>Severity</b> pill above<br/>"
            "2. Choose <b>Defect type</b> from the dropdown<br/>"
            "3. Optionally fill Location &amp; Notes tabs<br/>"
            "4. Click <b>💾 Save</b> to keep it</span>")
        wb_text.setWordWrap(True)
        wb_text.setStyleSheet(
            f"background:transparent;color:{UI_THEME['text_primary']};font-size:9pt;")
        wb_lay.addWidget(wb_icon, 0)
        wb_lay.addWidget(wb_text, 1)
        self._workflow_banner.setVisible(False)  # hidden until annotation drawn
        root.addWidget(self._workflow_banner)

        # ── Card header ──────────────────────────────────────────────────────
        hdr = QWidget()
        hdr.setStyleSheet(
            f"background:{UI_THEME['bg_secondary']};"
            f"border-radius:8px;border:1px solid {UI_THEME['border']};"
        )
        hdr_lay = QVBoxLayout(hdr)
        hdr_lay.setContentsMargins(10, 8, 10, 8)

        self._info_label = QLabel("No annotation selected")
        self._info_label.setStyleSheet(
            f"color:{UI_THEME['text_secondary']};font-size:9pt;"
            f"background:transparent;border:none;")
        hdr_lay.addWidget(self._info_label)

        self._size_label = QLabel("—")
        self._size_label.setStyleSheet(
            f"color:{UI_THEME['accent_cyan']};font-weight:bold;font-size:9pt;"
            f"background:transparent;border:none;")
        hdr_lay.addWidget(self._size_label)
        root.addWidget(hdr)

        # ── Severity pill strip ───────────────────────────────────────────────
        sev_hdr = QLabel("SEVERITY")
        sev_hdr.setStyleSheet(
            f"color:{UI_THEME['text_tertiary']};font-size:8pt;font-weight:bold;"
            f"letter-spacing:1px;background:transparent;")
        root.addWidget(sev_hdr)
        self._sev_strip = SeverityPillStrip()
        root.addWidget(self._sev_strip)

        # ── Accordion sections replacing QTabWidget ───────────────────────────
        # Each section is independently collapsible — all visible at once.
        # Wrapped in a single QScrollArea so the panel scrolls as one unit.
        _accordion_inner = QWidget()
        _accordion_inner.setSizePolicy(QSizePolicy.Policy.Expanding,
                                       QSizePolicy.Policy.Minimum)
        _acc_lay = QVBoxLayout(_accordion_inner)
        _acc_lay.setSpacing(0)
        _acc_lay.setContentsMargins(0, 0, 0, 0)

        # Tab 1 — Details
        det_tab = QWidget()
        det_tab.setSizePolicy(
            QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
        det_lay = QFormLayout(det_tab)
        det_lay.setSizeConstraint(QLayout.SizeConstraint.SetMinimumSize)
        det_lay.setContentsMargins(10, 10, 10, 10)
        det_lay.setSpacing(8)
        det_lay.setLabelAlignment(
            Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)

        self._defect_combo = QComboBox()
        self._defect_combo.addItems(DEFAULT_DEFECT_TYPES)
        self._defect_combo.setToolTip("Select the type of defect observed on the blade/component")
        # FIX-04: setEditable(True) was removed — user requirement is a strict dropdown.
        # Keeping editable=False (QComboBox default) enforces the taxonomy and prevents
        # free-text entries that would break auto-remedy lookup and report grouping.
        det_lay.addRow("Defect type:", self._defect_combo)

        self._ann_id_label = QLabel("—")
        self._ann_id_label.setStyleSheet(
            f"color:{UI_THEME['text_tertiary']};font-size:8pt;background:transparent;")
        self._ann_id_label.setToolTip("Unique identifier for this annotation")
        det_lay.addRow("Annotation ID:", self._ann_id_label)

        # v1.7.0: Inline image renamer — visible in the Details tab right next to
        # the defect type so the inspector can consciously rename the file to encode
        # component + defect (e.g. "BladeA_PS_Erosion_001").  Renaming on disk also
        # helps prevent accidental re-annotation of the same image in future sessions.
        sep_rename = QFrame()
        sep_rename.setFrameShape(QFrame.Shape.HLine)
        sep_rename.setStyleSheet(f"color:{UI_THEME['border']};")
        det_lay.addRow(sep_rename)

        rename_hdr = QLabel("📝  RENAME IMAGE FILE")
        rename_hdr.setStyleSheet(
            f"color:{UI_THEME['accent_amber']};font-size:8pt;font-weight:bold;"
            f"letter-spacing:1px;background:transparent;")
        det_lay.addRow(rename_hdr)

        rename_hint = QLabel(
            "Rename to encode component + defect so you can identify "
            "annotated images at a glance (e.g. BladeA_PS_Erosion_01).")
        rename_hint.setWordWrap(True)
        rename_hint.setStyleSheet(
            f"color:{UI_THEME['text_tertiary']};font-size:7.5pt;background:transparent;")
        det_lay.addRow(rename_hint)

        # Folder context — shows which folder/blade the image sits in so the
        # user knows the file's location before renaming it on disk.
        self._rename_folder_lbl = QLabel("—")
        self._rename_folder_lbl.setWordWrap(True)
        self._rename_folder_lbl.setTextInteractionFlags(
            Qt.TextInteractionFlag.TextSelectableByMouse)
        self._rename_folder_lbl.setStyleSheet(
            f"color:{UI_THEME['text_secondary']};font-size:7.5pt;"
            f"background:{UI_THEME['bg_elevated']};border:1px solid {UI_THEME['border']};"
            f"border-radius:4px;padding:3px 6px;")
        self._rename_folder_lbl.setToolTip(
            "Full folder path that contains this image.\n"
            "The folder name often encodes the blade / face assignment\n"
            "(e.g. BladeA_PS, Tower_Section2). The renamed file will stay\n"
            "inside this same folder.")
        det_lay.addRow("📁 Folder:", self._rename_folder_lbl)

        self._rename_edit = QLineEdit()
        self._rename_edit.setPlaceholderText("New filename (without extension)…")
        self._rename_edit.setStyleSheet(
            f"background:{UI_THEME['bg_elevated']};color:{UI_THEME['text_primary']};"
            f"border:1px solid {UI_THEME['border']};border-radius:4px;padding:3px 6px;")
        det_lay.addRow("New name:", self._rename_edit)

        self._rename_btn = QPushButton("✏️  Rename on Disk")
        self._rename_btn.setEnabled(False)
        self._rename_btn.setToolTip(
            "Rename the current image file on disk.\n"
            "The project is updated automatically.")
        self._rename_btn.setStyleSheet(
            f"background:{UI_THEME['accent_amber']};color:#0d1117;font-weight:bold;"
            f"border-radius:5px;padding:5px 10px;border:none;font-size:9pt;")
        self._rename_btn.clicked.connect(self._on_rename_file)
        det_lay.addRow(self._rename_btn)

        # SCROLL FIX: wrap tab content in QScrollArea so the tab page minimum
        # height is decoupled from the form content height.  Without this, the
        # QFormLayout's full sizeHint (~400px) propagates as a window minimum.
        _det_scroll = QScrollArea()
        _det_scroll.setWidgetResizable(True)
        _det_scroll.setFrameShape(QFrame.Shape.NoFrame)
        _det_scroll.setStyleSheet("QScrollArea{background:transparent;border:none;}")
        _det_scroll.setWidget(det_tab)
        self._sec_details = CollapsibleSection("📋  DETAILS", _det_scroll, start_open=True)
        _acc_lay.addWidget(self._sec_details)

        # Tab 2 — Location
        loc_tab = QWidget()
        # SCROLL FIX: SetMinimumSize tells the QFormLayout to report its full
        # sizeHint to the parent QScrollArea.  Without this the layout collapses
        # to zero height inside the scroll area and the BladePinpointWidget row
        # becomes invisible.
        loc_tab.setSizePolicy(
            QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
        loc_lay = QFormLayout(loc_tab)
        loc_lay.setSizeConstraint(QLayout.SizeConstraint.SetMinimumSize)
        loc_lay.setContentsMargins(10, 10, 10, 10)
        loc_lay.setSpacing(8)
        loc_lay.setLabelAlignment(
            Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)

        self._blade_combo = QComboBox()
        self._blade_combo.addItems(BLADE_NAMES)
        self._blade_combo.setToolTip("Select which blade or component (A/B/C, Hub, or Tower) this defect is located on")
        loc_lay.addRow("Blade:", self._blade_combo)

        self._span_combo = QComboBox()
        self._span_combo.addItems(BLADE_SPANS)
        self._span_combo.setToolTip("Blade span region: Root (0-33%), Mid (33-66%), or Tip (66-100%) of blade length")
        loc_lay.addRow("Span:", self._span_combo)

        # v4.5.0: Split into separate Surface and Zone selectors (mutually exclusive)
        self._surface_combo = QComboBox()
        self._surface_combo.addItems(BLADE_SURFACES)
        self._surface_combo.setToolTip("Blade surface (mutually exclusive): PS (Pressure Side) OR SS (Suction Side)")
        loc_lay.addRow("Surface:", self._surface_combo)

        self._zone_combo = QComboBox()
        self._zone_combo.addItems(BLADE_ZONES)
        self._zone_combo.setToolTip("Blade zone (mutually exclusive): LE (Leading Edge) OR TE (Trailing Edge) OR MB (Midbody)")
        loc_lay.addRow("Zone:", self._zone_combo)

        # v2.1.2: Manual distance fields (mandatory for Blade, disabled for Hub/Tower)
        dist_section_label = QLabel("📏 Distances")
        dist_section_label.setStyleSheet(f"font-weight:bold;color:{UI_THEME['text_secondary']};margin-top:8px;")
        loc_lay.addRow(dist_section_label)
        
        self._root_dist_spin = QDoubleSpinBox()
        self._root_dist_spin.setRange(0.0, 100.0)
        self._root_dist_spin.setSingleStep(0.5)
        self._root_dist_spin.setDecimals(2)
        self._root_dist_spin.setSuffix(" m")
        self._root_dist_spin.setStyleSheet(
            f"QDoubleSpinBox{{background:{UI_THEME['bg_input']};color:{UI_THEME['text_primary']};padding:4px;}}")
        self._root_dist_spin.setToolTip("Distance from blade root (meters) — required for Blade")
        self._root_dist_spin.valueChanged.connect(self._validate_and_update_save_button)  # v3.0.0 TODO #11
        loc_lay.addRow("Root Dist:", self._root_dist_spin)
        
        self._tip_dist_spin = QDoubleSpinBox()
        self._tip_dist_spin.setRange(0.0, 100.0)
        self._tip_dist_spin.setSingleStep(0.5)
        self._tip_dist_spin.setDecimals(2)
        self._tip_dist_spin.setSuffix(" m")
        self._tip_dist_spin.setStyleSheet(
            f"QDoubleSpinBox{{background:{UI_THEME['bg_input']};color:{UI_THEME['text_primary']};padding:4px;}}")
        self._tip_dist_spin.setToolTip("Distance from blade tip (meters) — required for Blade")
        self._tip_dist_spin.valueChanged.connect(self._validate_and_update_save_button)  # v3.0.0 TODO #11
        loc_lay.addRow("Tip Dist:", self._tip_dist_spin)
        
        self._dist_required_label = QLabel("⚠️  Required for Blade")
        self._dist_required_label.setStyleSheet(f"color:{UI_THEME['sev5']};font-size:9pt;font-weight:bold;")
        self._dist_required_label.setVisible(False)
        loc_lay.addRow(self._dist_required_label)
        
        # Update visibility when blade selection changes
        self._blade_combo.currentTextChanged.connect(self._validate_and_update_save_button)  # v3.0.0 TODO #11
        self._blade_combo.currentTextChanged.connect(self._on_blade_changed)  # v4.5.0: Handle Hub/Tower disable
        # v4.5.0: Connect surface and zone changes to validation
        self._surface_combo.currentTextChanged.connect(self._validate_and_update_save_button)
        self._zone_combo.currentTextChanged.connect(self._validate_and_update_save_button)

        # ── v3.2.0: Interactive blade pinpoint diagram ──────────────────────
        self._pp_sep  = QLabel("📍 Click blade to mark defect location")
        self._pp_sep.setStyleSheet(
            f"font-weight:bold;color:{UI_THEME['text_secondary']};margin-top:6px;")
        self._pp_sep.setWordWrap(True)
        loc_lay.addRow(self._pp_sep)

        self._pp_hint = QLabel("Root ▲ at top · Tip ▼ at bottom  |  Left=LE · Right=TE")
        self._pp_hint.setStyleSheet(f"color:{UI_THEME['text_tertiary']};font-size:8pt;")
        loc_lay.addRow(self._pp_hint)

        self._pinpoint_widget = BladePinpointWidget()
        self._pinpoint_widget.position_changed.connect(
            lambda _: self._validate_and_update_save_button())
        # v4.3.0: Wire edge_changed to update the read-only edge label and trigger save button
        self._pp_edge_label = QLabel("Edge: —")
        self._pp_edge_label.setStyleSheet(
            f"color:{UI_THEME['accent_cyan']};font-size:8pt;font-weight:bold;")
        self._pinpoint_widget.edge_changed.connect(
            lambda e: (
                self._pp_edge_label.setText(f"Edge: {e}" if e else "Edge: —"),
                self._validate_and_update_save_button()
            )
        )
        loc_lay.addRow(self._pp_edge_label)
        _pp_wrap = QWidget()
        # SCROLL FIX: explicit minimum height = pinpoint widget height + padding
        # so the QFormLayout row is never collapsed to zero inside a QScrollArea.
        _pp_wrap.setMinimumHeight(self._pinpoint_widget.height() + 8)
        _pp_lay  = QHBoxLayout(_pp_wrap)
        _pp_lay.setContentsMargins(0, 0, 0, 0)
        _pp_lay.addStretch()
        _pp_lay.addWidget(self._pinpoint_widget)
        _pp_lay.addStretch()
        loc_lay.addRow(_pp_wrap)
        # keep references for show/hide
        self._pp_blade_widgets = [self._pp_sep, self._pp_hint, self._pp_edge_label, _pp_wrap]

        # ── Legacy mm field — hidden from UI but kept as attribute for data compat ─
        self._dist_spin = QDoubleSpinBox()   # NOT added to layout (v3.2.1)
        self._dist_spin.setRange(0.0, 100_000.0)
        self._dist_spin.setDecimals(0)
        self._dist_spin.setValue(0.0)

        # SCROLL FIX: same QScrollArea treatment for the Location tab.
        # The Location tab is the tallest (BladePinpointWidget + form rows).
        _loc_scroll = QScrollArea()
        _loc_scroll.setWidgetResizable(True)
        _loc_scroll.setFrameShape(QFrame.Shape.NoFrame)
        _loc_scroll.setStyleSheet("QScrollArea{background:transparent;border:none;}")
        _loc_scroll.setWidget(loc_tab)
        self._sec_location = CollapsibleSection("📍  LOCATION", _loc_scroll, start_open=True)
        _acc_lay.addWidget(self._sec_location)

        # Tab 3 — Notes
        notes_tab = QWidget()
        notes_tab.setSizePolicy(
            QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
        notes_lay = QVBoxLayout(notes_tab)
        notes_lay.setSizeConstraint(QLayout.SizeConstraint.SetMinimumSize)
        notes_lay.setContentsMargins(10, 10, 10, 10)

        notes_lbl = QLabel("Inspector Comments:")
        notes_lbl.setStyleSheet(
            f"color:{UI_THEME['text_tertiary']};font-size:8pt;font-weight:bold;"
            f"background:transparent;")
        notes_lay.addWidget(notes_lbl)
        self._notes = QPlainTextEdit()
        self._notes.setPlaceholderText("Optional inspector notes…")
        self._notes.setToolTip("Add any additional notes, observations, or details about this defect")
        self._notes.setFixedHeight(75)
        notes_lay.addWidget(self._notes)

        # ── Scopito: Remedy Action (auto-filled from defect type, editable) ──
        remedy_lbl = QLabel("Remedy Action:")
        remedy_lbl.setStyleSheet(
            f"color:{UI_THEME['text_tertiary']};font-size:8pt;font-weight:bold;"
            f"background:transparent;")
        notes_lay.addWidget(remedy_lbl)
        self._remedy = QPlainTextEdit()
        self._remedy.setPlaceholderText("Recommended remedy action…")
        self._remedy.setToolTip("Auto-filled based on defect type. Edit to customize the recommended repair action.")
        self._remedy.setFixedHeight(75)
        notes_lay.addWidget(self._remedy)

        self._gsd_label = QLabel("GSD: not calibrated")
        self._gsd_label.setStyleSheet(
            f"color:{UI_THEME['text_tertiary']};font-size:8pt;background:transparent;")
        notes_lay.addWidget(self._gsd_label)
        notes_lay.addStretch()
        # SCROLL FIX: same QScrollArea treatment for the Notes tab.
        _notes_scroll = QScrollArea()
        _notes_scroll.setWidgetResizable(True)
        _notes_scroll.setFrameShape(QFrame.Shape.NoFrame)
        _notes_scroll.setStyleSheet("QScrollArea{background:transparent;border:none;}")
        _notes_scroll.setWidget(notes_tab)
        self._sec_notes = CollapsibleSection("📝  NOTES", _notes_scroll, start_open=True)
        _acc_lay.addWidget(self._sec_notes)
        _acc_lay.addStretch()   # push sections to top

        # Wrap accordion in a scroll area so the whole panel scrolls as one unit
        _acc_scroll = QScrollArea()
        _acc_scroll.setWidgetResizable(True)
        _acc_scroll.setFrameShape(QFrame.Shape.NoFrame)
        _acc_scroll.setStyleSheet("QScrollArea{background:transparent;border:none;}")
        _acc_scroll.setWidget(_accordion_inner)
        root.addWidget(_acc_scroll, 1)  # stretch=1: accordion fills remaining space

        # Auto-fill remedy when defect type changes
        self._defect_combo.currentTextChanged.connect(self._on_defect_type_changed)
        # v3.3.13 FIX: Update remedy when severity changes
        self._sev_strip.severity_changed.connect(self._on_severity_changed)

        # ── Action buttons (declared here; placed in pinned footer below) ──────
        self._save_btn = QPushButton("💾  Save")
        self._save_btn.setEnabled(False)
        self._save_btn.setStyleSheet(
            f"background:{UI_THEME['accent_blue']};color:white;font-weight:bold;"
            f"border-radius:6px;padding:7px 16px;border:none;")
        self._save_btn.clicked.connect(self._on_save)

        self._del_btn = QPushButton("🗑  Delete")
        self._del_btn.setEnabled(False)
        self._del_btn.setStyleSheet(
            f"background:{UI_THEME['accent_red']};color:white;font-weight:bold;"
            f"border-radius:6px;padding:7px 16px;border:none;")
        self._del_btn.clicked.connect(self._on_delete)

        # FIX-UX: Discard button — visible only while a NEW (unsaved) annotation
        # is pending.  Lets the user cancel a freshly drawn shape without saving.
        self._discard_btn = QPushButton("✕  Discard")
        self._discard_btn.setEnabled(False)
        self._discard_btn.setVisible(False)
        self._discard_btn.setStyleSheet(
            f"background:{UI_THEME['bg_card']};color:{UI_THEME['accent_orange']};"
            f"border:1px solid {UI_THEME['accent_orange']};font-weight:bold;"
            f"border-radius:6px;padding:7px 14px;")
        self._discard_btn.setToolTip("Discard this annotation without saving")
        self._discard_btn.clicked.connect(self._on_discard)

        # ── FOOTER: pinned below scroll area — always visible ─────────────────
        # A thin rule separates the scrollable content from the pinned controls.
        _footer_sep = QFrame()
        _footer_sep.setFrameShape(QFrame.Shape.HLine)
        _footer_sep.setStyleSheet(f"color:{UI_THEME['border']};margin:0;")
        outer.addWidget(_footer_sep)

        _btn_wrap = QWidget()
        _btn_wrap.setContentsMargins(8, 4, 8, 4)
        _btn_wrap_lay = QHBoxLayout(_btn_wrap)
        _btn_wrap_lay.setContentsMargins(0, 0, 0, 0)
        _btn_wrap_lay.addWidget(self._save_btn)
        _btn_wrap_lay.addWidget(self._discard_btn)
        _btn_wrap_lay.addWidget(self._del_btn)
        outer.addWidget(_btn_wrap)

        # v4.1.0: Approve/Reject moved to dedicated QC Review tab
        # (Removed from annotation panel - see QCReviewPanel for new location)

        # ── Annotation list — collapsible pinned footer ───────────────────────
        # v3.3.5 UI-1: collapsed by default to save vertical space on 1366×768
        # CHG-I: Converted to use CollapsibleSection widget for consistency with
        # the accordion sections above (DETAILS / LOCATION / NOTES / REVIEWER).
        # The previous custom toggle button approach is replaced by the same
        # CollapsibleSection widget so all sections behave identically.
        sep = QFrame(); sep.setFrameShape(QFrame.Shape.HLine)
        sep.setStyleSheet(f"color:{UI_THEME['border']};")
        outer.addWidget(sep)

        self._ann_list = QListWidget()
        self._ann_list.setMaximumHeight(60)
        self._ann_list.setMinimumHeight(0)
        self._ann_list.setStyleSheet(
            f"QListWidget{{background:{UI_THEME['bg_elevated']};"
            f"border:1px solid {UI_THEME['border']};border-radius:6px;"
            f"margin:0 4px 4px 4px;}}"
        )
        self._ann_list.itemClicked.connect(self._on_ann_list_click)

        # Wrap the list in a CollapsibleSection — starts collapsed (▶) to save space
        self._ann_section = CollapsibleSection(
            "📋  ANNOTATIONS ON THIS IMAGE", self._ann_list, start_open=False)
        outer.addWidget(self._ann_section)

        # Stub: _ann_toggle_btn referenced in crash-handler stub below; point to
        # the inner toggle button of the section so existing code paths don't break.
        self._ann_toggle_btn = self._ann_section._btn

    def _toggle_ann_list(self):
        """CHG-I: Delegates to the CollapsibleSection toggle (backward-compat shim).
        The _ann_section CollapsibleSection now owns the toggle logic; this method
        is kept so any external calls don't crash."""
        if hasattr(self, "_ann_section"):
            expanded = self._ann_list.isVisible()
            if expanded:
                self._ann_section.collapse()
            else:
                self._ann_section.expand()

    # ── Public API ─────────────────────────────────────────────────────────────

    def set_project(self, project: Optional["Project"]) -> None:
        """FIX-BUG: Receives the active project from MainWindow so load_pending()
        can build auto-suggested filenames that include turbine ID and blade.
        Called by WindTowerApp._update_project_ui() every time the project changes
        (new project, open project, project settings saved, etc.).
        Storing it here avoids tight coupling — AnnotationPanel never imports
        MainWindow and never calls save_project() itself.
        """
        self._project = project

    def load_pending(self, ann: Annotation):
        """Sam Okafor: Load a just-drawn annotation for editing."""
        self._pending_ann = ann
        self._is_new_annotation = True  # FIX-UX: track that this is unsaved
        mode_str = {"box": "Bounding Box", "pin": "Pin",
                    "polygon": "Polygon"}.get(ann.mode, ann.mode)
        self._info_label.setText(f"New {mode_str} — fill in details then Save")
        self._ann_id_label.setText(ann.ann_id[:10])
        if ann.width_cm is not None:
            self._size_label.setText(
                f"{ann.width_cm:.1f} × {ann.height_cm:.1f} cm")
        elif ann.mode == "polygon" and ann.poly_pts:
            self._size_label.setText(
                f"Polygon ({len(ann.poly_pts)//2} vertices)")
        else:
            self._size_label.setText("(calibrate for real-world size)")
        self._save_btn.setEnabled(True)
        self._del_btn.setEnabled(False)
        self._notes.clear()
        # Scopito fields — reset to defaults for new annotation
        self._dist_spin.setValue(0.0)
        self._remedy.setPlainText(_auto_remedy(ann.defect))

        # v4.5.0: Set blade, surface, and zone combos from ann
        self._blade_combo.setCurrentText(ann.blade)
        if hasattr(ann, 'surface') and ann.surface:
            self._surface_combo.setCurrentText(ann.surface)
        if hasattr(ann, 'zone') and ann.zone:
            self._zone_combo.setCurrentText(ann.zone)
        # Sync pinpoint widget for blade images
        if hasattr(self, "_pinpoint_widget") and ann.blade in ("A", "B", "C"):
            self._pinpoint_widget.set_severity(ann.severity)

        # v1.7.0: Auto-suggest a rename based on turbine + blade + defect type for new annotations
        # Format: WTG-{turbine}_Blade{blade}_{defect}
        fp = getattr(self, "_current_filepath", "")
        if fp:
            import os as _os
            current_stem = _os.path.splitext(_os.path.basename(fp))[0]
            # Only auto-suggest if the filename looks like a raw camera name (DJI_, IMG_, etc.)
            raw_prefixes = ("DJI_", "IMG_", "DSC", "DCIM", "P1_", "P_", "Image")
            if any(current_stem.upper().startswith(p.upper()) for p in raw_prefixes):
                # Build suggestion: WTG-{turbine}_Blade{blade}_{defect}
                turbine_part = ""
                if self._project and self._project.turbine_id:
                    tid = self._project.turbine_id
                    # Add WTG- prefix if not present
                    if not tid.upper().startswith("WTG"):
                        turbine_part = f"WTG-{tid}_"
                    else:
                        turbine_part = f"{tid}_"
                
                blade_part = f"Blade{ann.blade}" if ann.blade in ("A","B","C") else (ann.blade or "")
                defect_part = ann.defect.replace(" ", "").replace("/", "") if ann.defect else "Defect"
                
                suggestion = f"{turbine_part}{blade_part}_{defect_part}" if blade_part else f"{turbine_part}{defect_part}"
                self._rename_edit.setText(suggestion)

        # FIX-UX: Show guidance banner + Discard button for new annotations
        self._workflow_banner.setVisible(True)
        self._discard_btn.setVisible(True)
        self._discard_btn.setEnabled(True)

        # Ensure Details section is expanded so user immediately sees the defect form
        self._sec_details.expand()
        
        # v3.0.0 TODO #11: Validate after loading
        self._validate_and_update_save_button()

    def _validate_and_update_save_button(self):
        """
        v3.0.0 TODO #11: UI Validation enhancements.
        Phase 9.3: Guarded against AttributeError if called before _build_ui completes.
        Validates all required fields and updates save button state + visual feedback.
        
        Validation rules:
        - For Blade components: root_distance_m and tip_distance_m must be > 0
        - For Hub/Tower: distance fields not required
        - Positive values only for distances
        - Maximum reasonable limits (100m for blade)
        """
        # Phase 9.3: Guard — these widgets may not exist if called during construction
        required_attrs = ("_blade_combo", "_root_dist_spin", "_tip_dist_spin",
                          "_dist_required_label", "_save_btn")
        for _attr in required_attrs:
            if not hasattr(self, _attr):
                return
        is_valid = True
        validation_messages = []
        
        # Determine if blade component (validation required)
        blade = self._blade_combo.currentText()
        is_blade = blade in ("A", "B", "C")

        # Re-enable surface/zone/span/distance fields for Blades (may have been grayed for Hub/Tower)
        if is_blade:
            for _w in (self._surface_combo, self._zone_combo, self._span_combo,
                       self._root_dist_spin, self._tip_dist_spin):
                _w.setEnabled(True)
            # BUG-1 FIX: re-show pinpoint widget hidden by a prior Hub/Tower annotation
            if hasattr(self, "_pinpoint_widget"):
                for _w in getattr(self, "_pp_blade_widgets", []):
                    _w.setVisible(True)
                self._pinpoint_widget.setVisible(True)
        
        # Check distance validation for Blade components
        if is_blade:
            root_dist = self._root_dist_spin.value()
            tip_dist = self._tip_dist_spin.value()
            
            # Red border styling for invalid fields
            invalid_style = (
                f"QDoubleSpinBox{{background:{UI_THEME['bg_input']};"
                f"color:{UI_THEME['text_primary']};"
                f"border:2px solid {UI_THEME['sev5']};"  # Red border
                f"padding:4px;}}"
            )
            valid_style = (
                f"QDoubleSpinBox{{background:{UI_THEME['bg_input']};"
                f"color:{UI_THEME['text_primary']};"
                f"padding:4px;}}"
            )
            
            # Validate root distance
            if root_dist <= 0:
                is_valid = False
                self._root_dist_spin.setStyleSheet(invalid_style)
                self._root_dist_spin.setToolTip(
                    "⚠️  Distance from root is required for Blade components\n"
                    "Must be greater than 0 meters"
                )
                validation_messages.append("Root distance required")
            else:
                self._root_dist_spin.setStyleSheet(valid_style)
                self._root_dist_spin.setToolTip("Distance from blade root (meters)")
            
            # Validate tip distance
            if tip_dist <= 0:
                is_valid = False
                self._tip_dist_spin.setStyleSheet(invalid_style)
                self._tip_dist_spin.setToolTip(
                    "⚠️  Distance from tip is required for Blade components\n"
                    "Must be greater than 0 meters"
                )
                validation_messages.append("Tip distance required")
            else:
                self._tip_dist_spin.setStyleSheet(valid_style)
                self._tip_dist_spin.setToolTip("Distance from blade tip (meters)")
            
            # Show/update validation warning label
            if not is_valid:
                self._dist_required_label.setText("⚠️  " + " • ".join(validation_messages))
                self._dist_required_label.setVisible(True)
            else:
                self._dist_required_label.setVisible(False)
        else:
            # Hub/Tower: disable fields not applicable; no validation required
            self._dist_required_label.setVisible(False)
            normal_style = (
                f"QDoubleSpinBox{{background:{UI_THEME['bg_input']};"
                f"color:{UI_THEME['text_primary']};"
                f"padding:4px;}}"
            )
            self._root_dist_spin.setStyleSheet(normal_style)
            self._tip_dist_spin.setStyleSheet(normal_style)
            # Gray out surface/zone/span/distance — not applicable for Hub or Tower
            for _w in (self._surface_combo, self._zone_combo, self._span_combo,
                       self._root_dist_spin, self._tip_dist_spin):
                _w.setEnabled(False)  # visually grayed, prevents editing
            # v3.2.0: hide pinpoint diagram for Hub/Tower
            if hasattr(self, "_pinpoint_widget"):
                for _w in getattr(self, "_pp_blade_widgets", []):
                    _w.setVisible(False)
                self._pinpoint_widget.setVisible(False)

        # Update save button state and tooltip
        self._save_btn.setEnabled(is_valid)
        if not is_valid:
            self._save_btn.setToolTip(
                "Cannot save: Required fields missing\n\n" + 
                "\n".join(f"• {msg}" for msg in validation_messages)
            )
        else:
            self._save_btn.setToolTip("Save this annotation")

    def load_existing(self, ann: Annotation):
        """Sam Okafor: Load an existing annotation for viewing/editing."""
        self._pending_ann = ann
        self._is_new_annotation = False  # FIX-UX: existing annotation
        self._info_label.setText(f"Editing  #{ann.ann_id[:8]}")
        self._ann_id_label.setText(ann.ann_id)
        # FIX-UX: Hide the new-annotation guidance when viewing existing ones
        self._workflow_banner.setVisible(False)
        self._discard_btn.setVisible(False)
        self._discard_btn.setEnabled(False)
        self._ann_id_label.setText(ann.ann_id)
        self._defect_combo.setCurrentText(ann.defect)
        self._sev_strip.set_severity(ann.severity)
        self._blade_combo.setCurrentText(ann.blade)
        self._span_combo.setCurrentText(ann.span)
        # v4.5.0: Load surface and zone from annotation
        self._surface_combo.setCurrentText(getattr(ann, 'surface', 'Pressure Side (PS)'))
        self._zone_combo.setCurrentText(getattr(ann, 'zone', 'Leading Edge (LE)'))
        self._notes.setPlainText(ann.notes)
        
        # v2.1.2 / v3.0.0: Load manual root and tip distances
        self._root_dist_spin.setValue(ann.root_distance_m or 0.0)
        self._tip_dist_spin.setValue(ann.tip_distance_m or 0.0)

        # v3.2.0: Load pinpoint position; show widget for blades, hide for Hub/Tower
        _is_blade = ann.blade in ("A", "B", "C")
        if hasattr(self, "_pinpoint_widget"):
            self._pinpoint_widget.set_position(
                ann.pinpoint_blade_pos if hasattr(ann, "pinpoint_blade_pos") else None)
            self._pinpoint_widget.set_severity(ann.severity)
            # v4.3.0: Restore edge selection
            self._pinpoint_widget.set_edge_side(getattr(ann, "edge_side", None))
            _es = getattr(ann, "edge_side", None)
            if hasattr(self, "_pp_edge_label"):
                self._pp_edge_label.setText(f"Edge: {_es}" if _es else "Edge: —")
            for _w in getattr(self, "_pp_blade_widgets", []):
                _w.setVisible(_is_blade)
            self._pinpoint_widget.setVisible(_is_blade)

        # Scopito fields (legacy)
        self._dist_spin.setValue(ann.distance_from_root_mm or 0.0)
        # v3.3.13 FIX: Show SEVERITY_REMEDY based on severity first, fall back to
        # ann.remedy_action only if it's truly custom (doesn't match auto/severity texts)
        remedy_text = SEVERITY_REMEDY.get(ann.severity or "", _auto_remedy(ann.defect))
        if ann.remedy_action and ann.remedy_action.strip():
            stored = ann.remedy_action.strip()
            # Check if stored text matches auto-remedy pattern or severity text
            is_auto_pattern = stored.endswith("repair recommended during the next planned inspection.")
            is_severity_text = stored in SEVERITY_REMEDY.values()
            # Only use stored text if it's truly custom
            if not is_auto_pattern and not is_severity_text:
                remedy_text = stored
        self._remedy.setPlainText(remedy_text)
        if ann.width_cm is not None:
            self._size_label.setText(
                f"{ann.width_cm:.1f} × {ann.height_cm:.1f} cm")
        elif ann.mode == "polygon" and ann.poly_pts:
            self._size_label.setText(
                f"Polygon ({len(ann.poly_pts)//2} vertices)")
        else:
            self._size_label.setText("—")
        self._save_btn.setEnabled(True)
        self._del_btn.setEnabled(
            SESSION.can_do("delete_any") or
            (SESSION.can_do("delete_own") and ann.created_by == SESSION.username))
        
        # v3.0.0 TODO #11: Validate after loading existing annotation
        self._validate_and_update_save_button()

        # Phase 6: review frame
        # FIX-BUG: _rev_note / _status_lbl / _approve_btn / _reject_btn were moved
        # to QCReviewPanel in v4.1.0 but load_existing() was never updated, causing
        # AttributeError on every annotation selection.  Guard every access with
        # hasattr() so the code is safe whether or not these widgets still exist here.
        if SESSION.can_do("approve"):
            # CHG-I: show the CollapsibleSection wrapper; expand it so it's accessible
            if hasattr(self, "_review_section") and self._review_section is not None:
                self._review_section.setVisible(True)
                self._review_section.setMaximumHeight(16777215)  # restore full height
                self._review_section.expand()
            if hasattr(self, "_rev_note"):
                self._rev_note.setPlainText(ann.reviewer_note or "")
            status = ann.status or "pending"
            colours = {"approved": UI_THEME["accent_green"],
                       "rejected": UI_THEME["accent_red"],
                       "pending":  UI_THEME["accent_amber"]}
            if hasattr(self, "_status_lbl"):
                self._status_lbl.setText(
                    f"Status: {status.upper()}"
                    + (f"  ·  by {ann.reviewed_by}" if ann.reviewed_by else ""))
                self._status_lbl.setStyleSheet(
                    f"color:{colours.get(status, UI_THEME['accent_amber'])};"
                    f"font-size:9pt;font-weight:bold;background:transparent;")
            if hasattr(self, "_approve_btn"):
                self._approve_btn.setEnabled(status != "approved")
            if hasattr(self, "_reject_btn"):
                self._reject_btn.setEnabled(status != "rejected")
        else:
            if hasattr(self, "_review_section") and self._review_section is not None:
                self._review_section.setVisible(False)
                self._review_section.setMaximumHeight(0)

    def update_gsd_display(self, gsd: Optional[float], source: str = "image"):
        self._gsd_label.setText(
            f"GSD: {gsd:.4f} cm/px ({source})" if gsd else
            "GSD: not calibrated — draw a calibration line first")

    def refresh_ann_list(self, image_record: Optional[ImageRecord]):
        self._ann_list.clear()
        if not image_record:
            return
        status_icons = {"approved": "✔", "rejected": "✕", "pending": "○"}
        status_cols  = {
            "approved": UI_THEME["accent_green"],
            "rejected": UI_THEME["accent_red"],
            "pending":  UI_THEME["accent_amber"],
        }
        for ann in image_record.annotations:
            short  = SEVERITY_SHORT.get(ann.severity, ann.severity)
            size   = (f"{ann.width_cm:.1f}×{ann.height_cm:.1f}cm"
                      if ann.width_cm is not None else "?")
            status = ann.status or "pending"
            icon   = status_icons.get(status, "○")
            text   = f"{icon} [{short}]  {ann.defect}  {size}  ({ann.mode})"
            item   = QListWidgetItem(text)
            item.setData(Qt.ItemDataRole.UserRole, ann)
            # Colour by status: approved=green, rejected=red, pending=severity colour
            col = (QColor(status_cols[status]) if status in ("approved", "rejected")
                   else SEVERITY_COLORS.get(ann.severity, QColor("#7d8590")))
            item.setForeground(col)
            if status == "rejected":
                font = item.font(); font.setStrikeOut(True); item.setFont(font)
            self._ann_list.addItem(item)

    def update_defect_types(self, types: List[str]):
        current = self._defect_combo.currentText()
        self._defect_combo.clear()
        self._defect_combo.addItems(types)
        if current in types:
            self._defect_combo.setCurrentText(current)

    # ── Private ────────────────────────────────────────────────────────────────

    def _gather_annotation(self) -> Optional[Annotation]:
        if not self._pending_ann:
            return None
        ann          = self._pending_ann
        ann.defect   = self._defect_combo.currentText()
        ann.severity = self._sev_strip.current_severity()
        ann.blade    = self._blade_combo.currentText()
        ann.span     = self._span_combo.currentText()
        # v4.5.0: Save surface and zone from separate combos
        ann.surface  = self._surface_combo.currentText()
        ann.zone     = self._zone_combo.currentText()
        ann.notes    = self._notes.toPlainText().strip()
        
        # v2.1.2 / v3.0.0: Manual root and tip distances
        root_val = self._root_dist_spin.value()
        tip_val  = self._tip_dist_spin.value()
        ann.root_distance_m = root_val if root_val > 0.0 else None
        ann.tip_distance_m  = tip_val  if tip_val  > 0.0 else None

        # v3.2.0: Save pinpoint blade position (blades only)
        if ann.blade in ("A", "B", "C") and hasattr(self, "_pinpoint_widget"):
            ann.pinpoint_blade_pos = self._pinpoint_widget.get_position()
        else:
            ann.pinpoint_blade_pos = None

        # Scopito fields (legacy)
        dist_val = self._dist_spin.value()
        ann.distance_from_root_mm = dist_val if dist_val > 0.0 else None
        # Only store remedy if user explicitly typed something; otherwise leave
        # empty so the report uses SEVERITY_REMEDY[severity] automatically.
        remedy = self._remedy.toPlainText().strip()
        ann.remedy_action = remedy  # empty string → report uses SEVERITY_REMEDY
        
        # v3.3.13: Set created_at timestamp for new annotations
        if not ann.created_at:
            ann.created_at = datetime.now().isoformat()
        if not ann.created_by:
            ann.created_by = SESSION.username
        
        return ann

    def _on_save(self):
        ann = self._gather_annotation()
        if ann:
            self.save_requested.emit(ann)
            self._save_btn.setEnabled(False)
            self._info_label.setText("✓ Saved")
            # FIX-UX: Hide guidance banner after successful save
            self._workflow_banner.setVisible(False)
            self._discard_btn.setVisible(False)
            self._discard_btn.setEnabled(False)
            self._is_new_annotation = False

    def _on_discard(self):
        """
        FIX-UX: Discard an unsaved (just-drawn) annotation.
        Emits discard_requested so MainWindow can remove the item from the scene.
        """
        ann = self._pending_ann
        if not ann:
            return
        self.discard_requested.emit(ann)
        self._pending_ann     = None
        self._is_new_annotation = False
        self._save_btn.setEnabled(False)
        self._del_btn.setEnabled(False)
        self._discard_btn.setVisible(False)
        self._discard_btn.setEnabled(False)
        self._workflow_banner.setVisible(False)
        self._info_label.setText("Annotation discarded")
        self._size_label.setText("—")

    def _on_delete(self):
        ann = self._gather_annotation()
        if ann:
            self.delete_requested.emit(ann)
            self._pending_ann = None
            self._save_btn.setEnabled(False)
            self._del_btn.setEnabled(False)
            self._info_label.setText("Deleted")

    # ── v1.7.0: Inline image renamer ──────────────────────────────────────────

    def set_current_filepath(self, filepath: str):
        """Called by MainWindow whenever the selected image changes.
        Pre-populates the rename field with the current filename stem and
        shows the parent folder so the user knows which blade/face folder
        the image lives in before renaming it on disk."""
        self._current_filepath = filepath
        if filepath:
            import os as _os
            stem        = _os.path.splitext(_os.path.basename(filepath))[0]
            folder_path = _os.path.dirname(filepath)
            folder_name = _os.path.basename(folder_path)   # immediate parent dir
            grandparent = _os.path.basename(_os.path.dirname(folder_path))
            if grandparent:
                folder_display = f"{folder_name}  (…/{grandparent}/{folder_name})"
            else:
                folder_display = folder_path
            self._rename_folder_lbl.setText(folder_display)
            self._rename_folder_lbl.setToolTip(
                f"Full path: {folder_path}\n\n"
                "The folder name typically encodes the blade / face assignment\n"
                "(e.g. BladeA_PS, Tower_Section2).\n"
                "The renamed file will stay inside this same folder.")
            self._rename_edit.setText(stem)
            self._rename_btn.setEnabled(True)
        else:
            self._rename_folder_lbl.setText("\u2014")
            self._rename_edit.clear()
            self._rename_btn.setEnabled(False)

    def _on_rename_file(self):
        """v1.7.0: Emit rename_requested with the new stem so MainWindow can
        rename the file on disk and update the project JSON + thumbnail strip."""
        new_stem = self._rename_edit.text().strip()
        if not new_stem:
            return
        # Guard: reject characters that are illegal in filenames on Windows/Mac/Linux
        illegal = set(r'\/:*?"<>|')
        bad = [c for c in new_stem if c in illegal]
        if bad:
            from PyQt6.QtWidgets import QMessageBox as _QMB
            _QMB.warning(None, "Invalid name",
                         f"Filename contains illegal characters: {' '.join(bad)}")
            return
        fp = getattr(self, "_current_filepath", "")
        if not fp:
            return
        self.rename_requested.emit(fp, new_stem)

    def _on_approve(self):
        """Phase 6: Reviewer approves the selected annotation.
        CTO-AUDIT: Logic-layer permission check before mutating annotation state.
        """
        ann = self._pending_ann
        if not ann:
            return
        # Logic-layer guard — cannot be bypassed by UI manipulation
        if not SESSION.can_do("approve"):
            log.warning(f"[PERMISSION] {SESSION.username} attempted approve without permission")
            QMessageBox.warning(self, "Access Denied",
                "Annotation approval is not available.")
            return
        ann.status       = "approved"
        ann.reviewed_by  = SESSION.username
        ann.reviewed_at  = datetime.now().isoformat()
        # FIX-BUG: _rev_note was moved to QCReviewPanel — guard with hasattr
        ann.reviewer_note = self._rev_note.toPlainText().strip() if hasattr(self, "_rev_note") else (ann.reviewer_note or "")
        self.approve_requested.emit(ann)
        if hasattr(self, "_status_lbl"):
            self._status_lbl.setText(f"Status: APPROVED  ·  by {SESSION.username}")
            self._status_lbl.setStyleSheet(
                f"color:{UI_THEME['accent_green']};font-size:9pt;"
                f"font-weight:bold;background:transparent;")
        if hasattr(self, "_approve_btn"):
            self._approve_btn.setEnabled(False)
        if hasattr(self, "_reject_btn"):
            self._reject_btn.setEnabled(True)

    def _on_reject(self):
        """Phase 6: Reviewer rejects the selected annotation.
        CTO-AUDIT: Logic-layer permission check before mutating annotation state.
        """
        ann = self._pending_ann
        if not ann:
            return
        # Logic-layer guard
        if not SESSION.can_do("approve"):
            log.warning(f"[PERMISSION] {SESSION.username} attempted reject without permission")
            QMessageBox.warning(self, "Access Denied",
                "Annotation rejection is not available.")
            return
        ann.status        = "rejected"
        ann.reviewed_by   = SESSION.username
        ann.reviewed_at   = datetime.now().isoformat()
        # FIX-BUG: _rev_note was moved to QCReviewPanel — guard with hasattr
        ann.reviewer_note = self._rev_note.toPlainText().strip() if hasattr(self, "_rev_note") else (ann.reviewer_note or "")
        self.reject_requested.emit(ann)
        if hasattr(self, "_status_lbl"):
            self._status_lbl.setText(f"Status: REJECTED  ·  by {SESSION.username}")
            self._status_lbl.setStyleSheet(
                f"color:{UI_THEME['accent_red']};font-size:9pt;"
                f"font-weight:bold;background:transparent;")
        if hasattr(self, "_reject_btn"):
            self._reject_btn.setEnabled(False)
        if hasattr(self, "_approve_btn"):
            self._approve_btn.setEnabled(True)

    def _on_ann_list_click(self, item: QListWidgetItem):
        ann = item.data(Qt.ItemDataRole.UserRole)
        if isinstance(ann, Annotation):
            self.load_existing(ann)
            self.ann_selected_for_qc.emit(ann)  # v4.1.1: also feed QC Review panel

    def _on_defect_type_changed(self, defect_type: str):
        """
        Auto-fill remedy action when defect type changes, only if remedy is empty.
        v3.3.13 FIX: Use SEVERITY_REMEDY based on current severity instead of _auto_remedy.
        """
        if not self._remedy.toPlainText().strip():
            # Get current severity from severity strip widget
            current_severity = self._sev_strip.current_severity()
            # Use SEVERITY_REMEDY if available, otherwise fall back to auto_remedy
            remedy_text = SEVERITY_REMEDY.get(current_severity, _auto_remedy(defect_type))
            self._remedy.setPlainText(remedy_text)

    def _on_severity_changed(self, severity: str):
        """
        v3.3.13 FIX: Update remedy field when severity changes.
        Only updates if the current remedy matches a severity-based or auto-generated text.
        Preserves truly custom user-entered remedy text.
        """
        current_remedy = self._remedy.toPlainText().strip()
        if not current_remedy:
            # Empty remedy - fill with severity-based text
            remedy_text = SEVERITY_REMEDY.get(severity, "")
            self._remedy.setPlainText(remedy_text)
        else:
            # Check if current remedy matches any severity or auto-generated text
            defect_type = self._defect_combo.currentText()
            auto_text = _auto_remedy(defect_type)
            # Check if it matches any SEVERITY_REMEDY value
            is_auto_generated = (current_remedy == auto_text or 
                               current_remedy in SEVERITY_REMEDY.values())
            if is_auto_generated:
                # Update to new severity-based remedy
                remedy_text = SEVERITY_REMEDY.get(severity, auto_text)
                self._remedy.setPlainText(remedy_text)

    def _on_blade_changed(self, blade: str):
        """
        v4.5.0: Called when the blade combo changes.
        Hub/Tower: Disable surface and zone combos (N/A).
        Blade A/B/C: Enable surface and zone combos.
        """
        is_hub_tower = blade in ("Hub", "Tower")
        self._surface_combo.setEnabled(not is_hub_tower)
        self._zone_combo.setEnabled(not is_hub_tower)
        # Re-run validation since blade affects requirements
        self._validate_and_update_save_button()

# ==============================================================================
# BURN-IN JPEG  (Marcus Webb + Tom K.)
# Paint annotation overlays onto a copy of the source image using Pillow.
# Supports box, pin, and polygon modes.
# ==============================================================================

def _rotated_box_corners(x1: float, y1: float, x2: float, y2: float,
                         angle_deg: float) -> list:
    """
    Return the 4 corners of a rotated bounding box as (x, y) tuples.

    The annotation stores x1/y1/x2/y2 as an axis-aligned box centred on the
    visual centre of the rotated rectangle (see EditableBoxItem._commit_geometry).
    rotation_deg is the clockwise angle applied to that box.

    Parameters
    ----------
    x1, y1 : top-left corner of the axis-aligned envelope
    x2, y2 : bottom-right corner of the axis-aligned envelope
    angle_deg : clockwise rotation in degrees (0 = no rotation)

    Returns
    -------
    List of 4 (x, y) float tuples in order: TL, TR, BR, BL (rotated).
    """
    import math
    cx, cy = (x1 + x2) / 2.0, (y1 + y2) / 2.0
    hw,  hh  = (x2 - x1) / 2.0, (y2 - y1) / 2.0
    rad = math.radians(angle_deg)
    cos_a, sin_a = math.cos(rad), math.sin(rad)
    corners = []
    for px, py in [(-hw, -hh), (hw, -hh), (hw, hh), (-hw, hh)]:
        rx = cx + px * cos_a - py * sin_a
        ry = cy + px * sin_a + py * cos_a
        corners.append((rx, ry))
    return corners


def _rotated_box_bounds(x1: float, y1: float, x2: float, y2: float,
                        angle_deg: float) -> tuple:
    """
    Return the axis-aligned bounding box of a rotated rectangle as
    (bx1, by1, bx2, by2) — used to determine the crop region for zoom images.
    """
    corners = _rotated_box_corners(x1, y1, x2, y2, angle_deg)
    xs = [c[0] for c in corners]; ys = [c[1] for c in corners]
    return min(xs), min(ys), max(xs), max(ys)


def _burn_in_jpeg_annotations(image_path: str, image_record: ImageRecord,
                               project_folder: str) -> Optional[str]:
    """Marcus Webb: Burn all annotations into a copy of the source image.

    FIX-BUG3: Now writes TWO copies:
      1. project_folder/annotated/{name}.JPG  — existing behaviour (report use)
      2. Same directory as original, named {stem}_defect{ext}  — visible in source
         folder so the inspector can immediately see defect labels on the file.
    Returns the project annotated path (as before).
    """
    try:
        src_path = Path(image_path)
        out_dir  = Path(project_folder) / "annotated"
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / src_path.name

        with Image.open(image_path) as src:
            img = src.convert("RGB")
        draw = ImageDraw.Draw(img, "RGBA")

        # Try to load a font; fall back to default
        try:
            font     = ImageFont.truetype("arial.ttf", 14)
            font_sm  = ImageFont.truetype("arial.ttf", 11)
        except Exception:
            font = font_sm = ImageFont.load_default()

        def _hex(sev: str) -> Tuple[int, int, int]:
            col = SEVERITY_COLORS.get(sev, QColor("#d29922"))
            return col.red(), col.green(), col.blue()

        for ann in image_record.annotations:
            r, g, b   = _hex(ann.severity)
            outline   = (r, g, b, 255)
            fill_poly = (r, g, b, 60)
            short     = SEVERITY_SHORT.get(ann.severity, ann.severity)
            label     = f"#{ann.ann_id[:6]} [{short}] {ann.defect}"

            if ann.mode == "box":
                x1 = min(ann.x1_px, ann.x2_px); y1 = min(ann.y1_px, ann.y2_px)
                x2 = max(ann.x1_px, ann.x2_px); y2 = max(ann.y1_px, ann.y2_px)
                rot = getattr(ann, "rotation_deg", 0.0) or 0.0
                if abs(rot) > 0.5:
                    corners = _rotated_box_corners(x1, y1, x2, y2, rot)
                    draw.polygon(corners, outline=outline, width=3)
                else:
                    draw.rectangle([x1, y1, x2, y2], outline=outline, width=3)
                draw.text((x1 + 2, y1 - 16), label, fill=outline, font=font_sm)

            elif ann.mode == "pin":
                r2 = 8
                cx, cy = ann.x1_px, ann.y1_px
                draw.ellipse([cx - r2, cy - r2, cx + r2, cy + r2],
                             outline=outline, fill=(r, g, b, 160), width=2)
                draw.text((cx + r2 + 4, cy - r2), label, fill=outline, font=font_sm)

            elif ann.mode == "polygon" and len(ann.poly_pts) >= 4:
                pts = [(ann.poly_pts[i], ann.poly_pts[i + 1])
                       for i in range(0, len(ann.poly_pts) - 1, 2)]
                draw.polygon(pts, fill=fill_poly, outline=outline)
                draw.line(pts + [pts[0]], fill=outline, width=3)
                if pts:
                    draw.text((pts[0][0] + 4, pts[0][1] - 16),
                              label, fill=outline, font=font_sm)

        # Save 1: project annotated subfolder (for report rendering)
        img.save(str(out_path), "JPEG", quality=92, optimize=True)
        log.info(f"Burn-in saved → {out_path}")

        # FIX-BUG3: Save 2: next to the original in its source directory
        # Name: {stem}_defect{ext}  so the inspector sees the defect in the folder.
        try:
            src_sibling = src_path.parent / (src_path.stem + "_defect" + src_path.suffix)
            img.save(str(src_sibling), "JPEG", quality=92, optimize=True)
            log.info(f"Burn-in (source copy) → {src_sibling}")
        except Exception as exc_s:
            log.warning(f"_burn_in_jpeg_annotations: source-dir copy failed: {exc_s}")

        return str(out_path)
    except Exception as exc:
        log.error(f"_burn_in_jpeg_annotations: {exc}")
        return None

# ==============================================================================
# ML: MODEL MANAGER  (Natalie Cross — ML 11yr)
# ==============================================================================

class ModelManager:
    """
    Natalie Cross: Singleton YOLO model wrapper.
    Caches the loaded model so repeated detections do not reload weights.
    Thread-safe via QMutex — multiple DetectionWorkers may share one manager.
    """
    def __init__(self):
        self._model: Optional[Any]  = None
        self._model_path: str       = ""
        self._mutex                 = QMutex()

    def load(self, model_path: str, device: str = "auto") -> bool:
        with QMutexLocker(self._mutex):
            if not YOLO_AVAILABLE:
                log.error("ModelManager: ultralytics not installed")
                return False
            if self._model_path == model_path and self._model:
                return True
            try:
                self._model      = UltralyticsYOLO(model_path)
                self._model_path = model_path
                log.info(f"ModelManager: loaded {model_path} on {device}")
                return True
            except Exception as exc:
                log.error(f"ModelManager.load: {exc}")
                return False

    def detect(self, image_path: str, conf: float, iou: float,
               img_size: int) -> List[DetectionResult]:
        with QMutexLocker(self._mutex):
            if not self._model:
                return []
            try:
                results = self._model(image_path, conf=conf, iou=iou,
                                      imgsz=img_size, verbose=False)
                detections = []
                for res in results:
                    if res.boxes is None:
                        continue
                    names = res.names or {}
                    for box in res.boxes:
                        xyxy  = box.xyxy[0].tolist()
                        conf_ = float(box.conf[0])
                        cls_  = int(box.cls[0])
                        detections.append(DetectionResult(
                            image_path=image_path,
                            x1_px=xyxy[0], y1_px=xyxy[1],
                            x2_px=xyxy[2], y2_px=xyxy[3],
                            confidence=conf_,
                            class_id=cls_,
                            class_name=names.get(cls_, "Defect"),
                        ))
                return detections
            except Exception as exc:
                log.error(f"ModelManager.detect: {exc}")
                return []

# ── Shared model manager instance ─────────────────────────────────────────────
_model_mgr = ModelManager()

# ==============================================================================
# DETECTION WORKER  (Natalie Cross + Alex Stone)
# ==============================================================================

class _DetSignals(QObject):
    progress    = pyqtSignal(int, int)          # done, total
    image_done  = pyqtSignal(str, list)         # filepath, [DetectionResult]
    error       = pyqtSignal(str, str)          # filename, error_msg
    batch_done  = pyqtSignal(int, int)          # total_images, total_detections


class DetectionWorker(QRunnable):
    """Natalie Cross: Off-thread YOLO batch inference."""
    def __init__(self, image_paths: List[str], model_mgr: ModelManager,
                 conf_low: float, iou: float, img_size: int):
        super().__init__()
        self.image_paths = image_paths
        self.model_mgr   = model_mgr
        self.conf_low    = conf_low
        self.iou         = iou
        self.img_size    = img_size
        self.signals     = _DetSignals()
        self._cancelled  = False
        self.setAutoDelete(True)

    def cancel(self):
        self._cancelled = True

    @pyqtSlot()
    def run(self):
        total = len(self.image_paths); total_dets = 0
        for i, fp in enumerate(self.image_paths):
            if self._cancelled:
                break
            self.signals.progress.emit(i, total)
            try:
                results     = self.model_mgr.detect(fp, self.conf_low,
                                                     self.iou, self.img_size)
                total_dets += len(results)
                self.signals.image_done.emit(fp, results)
            except Exception as exc:
                self.signals.error.emit(os.path.basename(fp), str(exc))
        self.signals.progress.emit(total, total)
        self.signals.batch_done.emit(total, total_dets)

# ==============================================================================
# TRAINING WORKER  (Natalie Cross + Alex Stone)
# ==============================================================================

class _TrainSignals(QObject):
    log_line   = pyqtSignal(str)
    epoch_done = pyqtSignal(int, int, float)
    finished   = pyqtSignal(bool, str)
    progress   = pyqtSignal(int)


class TrainingWorker(QThread):
    """Natalie Cross: Two-phase pipeline: export dataset → YOLO train."""

    def __init__(self, project: Project, export_dir: str, cfg: AppConfig):
        super().__init__()
        self.project    = project
        self.export_dir = export_dir
        self.cfg        = cfg
        self.signals    = _TrainSignals()
        self._stop_flag = False

    def stop(self):
        self._stop_flag = True

    def _export_dataset(self) -> Optional[Path]:
        export_root = Path(self.export_dir)
        import tempfile, shutil as _sh, random as _rnd
        tmp_dir = Path(tempfile.mkdtemp(prefix="wtg_export_",
                                        dir=export_root.parent))
        try:
            for split in ["train", "val"]:
                (tmp_dir / "images" / split).mkdir(parents=True, exist_ok=True)
                (tmp_dir / "labels" / split).mkdir(parents=True, exist_ok=True)

            records = [ir for ir in self.project.images.values()
                       if any(a.mode in ("box", "polygon") for a in ir.annotations)]
            if not records:
                self.signals.log_line.emit(
                    "[WARN] No box/polygon annotations — nothing to export.")
                return None

            val_split = float(self.cfg.get("TRAINING", "ValSplit", "0.2"))
            _rnd.shuffle(records)
            n_val    = max(1, int(len(records) * val_split))
            val_set  = {r.filename for r in records[:n_val]}

            for ir in records:
                if self._stop_flag:
                    return None
                src = Path(ir.filepath)
                if not src.exists():
                    self.signals.log_line.emit(f"[SKIP] Missing: {ir.filename}")
                    continue
                is_val   = ir.filename in val_set
                split    = "val" if is_val else "train"
                _sh.copy2(str(src), str(tmp_dir / "images" / split / src.name))
                try:
                    with Image.open(str(src)) as _pil:
                        iw, ih = _pil.size
                except Exception:
                    self.signals.log_line.emit(
                        f"[SKIP] Cannot read dims: {ir.filename}")
                    continue

                lines = []
                for ann in ir.annotations:
                    if ann.mode == "box":
                        x1 = min(ann.x1_px, ann.x2_px) / iw
                        y1 = min(ann.y1_px, ann.y2_px) / ih
                        x2 = max(ann.x1_px, ann.x2_px) / iw
                        y2 = max(ann.y1_px, ann.y2_px) / ih
                        cx = (x1 + x2) / 2; cy = (y1 + y2) / 2
                        bw = x2 - x1; bh = y2 - y1
                        cx, cy, bw, bh = (max(0., min(1., v))
                                          for v in (cx, cy, bw, bh))
                        if bw > 0 and bh > 0:
                            lines.append(
                                f"0 {cx:.6f} {cy:.6f} {bw:.6f} {bh:.6f}")
                    elif ann.mode == "polygon" and len(ann.poly_pts) >= 6:
                        # Convert polygon to bounding box for standard YOLO
                        xs = ann.poly_pts[0::2]; ys = ann.poly_pts[1::2]
                        x1, y1 = min(xs) / iw, min(ys) / ih
                        x2, y2 = max(xs) / iw, max(ys) / ih
                        cx = (x1 + x2) / 2; cy = (y1 + y2) / 2
                        bw = x2 - x1; bh = y2 - y1
                        if bw > 0 and bh > 0:
                            lines.append(
                                f"0 {cx:.6f} {cy:.6f} {bw:.6f} {bh:.6f}")

                if lines:
                    lbl_p = tmp_dir / "labels" / split / (src.stem + ".txt")
                    lbl_p.write_text("\n".join(lines) + "\n", encoding="utf-8")
                    self.signals.log_line.emit(
                        f"  Exported: {ir.filename} ({len(lines)} boxes)")
                else:
                    self.signals.log_line.emit(
                        f"[SKIP] No valid boxes: {ir.filename}")

            yaml_c = (f"path: {tmp_dir}\ntrain: images/train\nval: images/val\n"
                      f"nc: 1\nnames:\n  0: defect\n")
            (tmp_dir / "data.yaml").write_text(yaml_c, encoding="utf-8")

            if export_root.exists():
                bak = export_root.with_name(export_root.name + "_bak")
                if bak.exists():
                    _sh.rmtree(str(bak))
                _sh.move(str(export_root), str(bak))
            _sh.move(str(tmp_dir), str(export_root))
            self.signals.log_line.emit(f"\n[OK] Dataset → {export_root}")
            return export_root
        except Exception as exc:
            self.signals.log_line.emit(f"[ERROR] Export: {exc}")
            try:
                _sh.rmtree(str(tmp_dir))
            except Exception:
                pass
            return None

    def run(self):
        import io as _io
        self.signals.log_line.emit("=" * 60)
        self.signals.log_line.emit("PHASE A — Exporting annotations to YOLO dataset")
        self.signals.log_line.emit("=" * 60)
        self.signals.progress.emit(5)

        dataset_path = self._export_dataset()
        if dataset_path is None or self._stop_flag:
            self.signals.finished.emit(False, "Export failed or cancelled.")
            return

        self.signals.log_line.emit("\n" + "=" * 60)
        self.signals.log_line.emit("PHASE B — YOLO Training")
        self.signals.log_line.emit("=" * 60)
        self.signals.progress.emit(15)

        if not YOLO_AVAILABLE:
            self.signals.finished.emit(False, "ultralytics not installed.")
            return

        try:
            epochs  = int(self.cfg.get("TRAINING", "Epochs",    "50"))
            batch   = int(self.cfg.get("TRAINING", "BatchSize", "16"))
            imgsz   = int(self.cfg.get("TRAINING", "ImgSize",   "640"))
            lr0     = float(self.cfg.get("TRAINING", "LR0",     "0.01"))
            lrf     = float(self.cfg.get("TRAINING", "LRF",     "0.01"))
            device  = self.cfg.get("TRAINING", "Device", "auto")

            # Redirect stdout so live YOLO output feeds the log console
            old_stdout = _io.StringIO()
            import sys as _sys
            _real_stdout = _sys.stdout

            model = UltralyticsYOLO("yolov8n.pt")
            results = model.train(
                data    = str(dataset_path / "data.yaml"),
                epochs  = epochs,
                batch   = batch,
                imgsz   = imgsz,
                lr0     = lr0,
                lrf     = lrf,
                device  = device,
                verbose = True,
            )
            best_pt = str(Path(results.save_dir) / "weights" / "best.pt")
            self.signals.finished.emit(True, best_pt)
        except Exception as exc:
            self.signals.finished.emit(False, str(exc))

# ==============================================================================
# QC VIEWER WIDGET — Inline  (Sam Okafor + Priya Nair)
# Full Undo / Redo / Clear All support for QC review operations.
# ==============================================================================

# ==============================================================================
# QC REVIEW PANEL  (v4.1.0 - Dedicated QC Review with Approve Functionality)
# ==============================================================================

class QCReviewPanel(QWidget):
    """
    v4.1.1: Dedicated QC Review panel — shows ALL annotations on current image.
    Inspector selects from annotation list; approves/rejects individually.
    Status refreshes immediately after each action.
    """
    approve_requested = pyqtSignal(object)   # Annotation
    reject_requested  = pyqtSignal(object)   # Annotation

    _STATUS_COLORS = {
        "pending":  UI_THEME["accent_amber"],
        "approved": UI_THEME["accent_green"],
        "rejected": UI_THEME["accent_red"],
    }
    _STATUS_ICONS = {"approved": "✔", "rejected": "✕", "pending": "○"}

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setMinimumWidth(290)
        self.setMaximumWidth(360)
        self._current_ann: Optional[Annotation] = None
        self._image_record: Optional[ImageRecord] = None
        self._build_ui()

    def _build_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(8, 8, 8, 8)
        root.setSpacing(6)

        # Header
        hdr = QLabel("🔍 QC REVIEW")
        hdr.setStyleSheet(
            f"color:{UI_THEME['accent_cyan']};font-size:11pt;font-weight:bold;"
            f"background:transparent;padding:4px 8px;")
        root.addWidget(hdr)

        # ── Annotation list (all annotations on current image) ─────────────────
        list_lbl = QLabel("Annotations on this image:")
        list_lbl.setStyleSheet(
            f"color:{UI_THEME['text_secondary']};font-size:8pt;"
            f"font-weight:bold;background:transparent;")
        root.addWidget(list_lbl)

        self._qc_list = QListWidget()
        self._qc_list.setMaximumHeight(160)
        self._qc_list.setStyleSheet(
            f"QListWidget{{background:{UI_THEME['bg_secondary']};"
            f"border:1px solid {UI_THEME['border']};border-radius:4px;"
            f"color:{UI_THEME['text_primary']};font-size:8pt;}}"
            f"QListWidget::item{{padding:4px 6px;}}"
            f"QListWidget::item:selected{{background:{UI_THEME['bg_elevated']};"
            f"color:{UI_THEME['accent_cyan']};}}")
        self._qc_list.itemClicked.connect(self._on_list_item_clicked)
        root.addWidget(self._qc_list)

        # ── Detail card ────────────────────────────────────────────────────────
        detail_card = QWidget()
        detail_card.setStyleSheet(
            f"background:{UI_THEME['bg_secondary']};border-radius:8px;"
            f"border:1px solid {UI_THEME['border']};")
        detail_lay = QVBoxLayout(detail_card)
        detail_lay.setContentsMargins(10, 8, 10, 8)
        detail_lay.setSpacing(4)

        self._info_label = QLabel("Select an annotation above to review")
        self._info_label.setWordWrap(True)
        self._info_label.setStyleSheet(
            f"color:{UI_THEME['text_primary']};font-size:9pt;background:transparent;")
        detail_lay.addWidget(self._info_label)

        self._status_label = QLabel("Status: —")
        self._status_label.setStyleSheet(
            f"color:{UI_THEME['accent_amber']};font-weight:bold;"
            f"font-size:9pt;background:transparent;")
        detail_lay.addWidget(self._status_label)

        root.addWidget(detail_card)

        # Reviewer notes
        notes_lbl = QLabel("Reviewer Notes:")
        notes_lbl.setStyleSheet(
            f"color:{UI_THEME['text_secondary']};font-size:8pt;font-weight:bold;"
            f"background:transparent;")
        root.addWidget(notes_lbl)

        self._rev_note = QPlainTextEdit()
        self._rev_note.setPlaceholderText("Add reviewer notes (optional)…")
        self._rev_note.setFixedHeight(70)
        self._rev_note.setStyleSheet(
            f"background:{UI_THEME['bg_elevated']};color:{UI_THEME['text_primary']};"
            f"border:1px solid {UI_THEME['border']};border-radius:4px;padding:4px;")
        root.addWidget(self._rev_note)

        # Action buttons
        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)

        self._approve_btn = QPushButton("✔  Approve")
        self._approve_btn.setEnabled(False)
        self._approve_btn.setFixedHeight(36)
        self._approve_btn.setStyleSheet(
            f"QPushButton{{background:{UI_THEME['accent_green']};color:#0d1117;"
            f"font-weight:bold;border-radius:6px;padding:6px;border:none;}}"
            f"QPushButton:disabled{{background:#2d4a3e;color:#636e7b;}}")
        self._approve_btn.setToolTip("Approve — marks annotation as reviewed and accepted")
        self._approve_btn.clicked.connect(self._on_approve)

        self._reject_btn = QPushButton("✕  Reject")
        self._reject_btn.setEnabled(False)
        self._reject_btn.setFixedHeight(36)
        self._reject_btn.setStyleSheet(
            f"QPushButton{{background:{UI_THEME['accent_red']};color:white;"
            f"font-weight:bold;border-radius:6px;padding:6px;border:none;}}"
            f"QPushButton:disabled{{background:#4a2d2d;color:#636e7b;}}")
        self._reject_btn.setToolTip("Reject — marks annotation for revision by inspector")
        self._reject_btn.clicked.connect(self._on_reject)

        btn_row.addWidget(self._approve_btn)
        btn_row.addWidget(self._reject_btn)
        root.addLayout(btn_row)
        root.addStretch()

    # ── Public API ──────────────────────────────────────────────────────────────

    def load_image_annotations(self, image_record: Optional["ImageRecord"]):
        """v4.1.1: Populate list with ALL annotations on the given image record."""
        self._image_record = image_record
        self._qc_list.clear()
        self._clear_detail()
        if not image_record or not image_record.annotations:
            item = QListWidgetItem("  No annotations on this image")
            item.setForeground(QColor(UI_THEME["text_tertiary"]))
            self._qc_list.addItem(item)
            return
        for ann in image_record.annotations:
            self._add_list_item(ann)

    def load_annotation(self, ann: Optional[Annotation]):
        """Load a single annotation into the detail area (canvas click / list click)."""
        self._current_ann = ann
        if not ann:
            self._clear_detail()
            return
        # Sync list selection to match
        for i in range(self._qc_list.count()):
            it = self._qc_list.item(i)
            if it and it.data(Qt.ItemDataRole.UserRole) is ann:
                self._qc_list.setCurrentItem(it)
                break
        self._refresh_detail(ann)

    def refresh_current(self):
        """Re-render detail card and list item for the current annotation (after approve/reject)."""
        if not self._current_ann:
            return
        ann = self._current_ann
        # Update matching list item text + colour
        for i in range(self._qc_list.count()):
            it = self._qc_list.item(i)
            if it and it.data(Qt.ItemDataRole.UserRole) is ann:
                it.setText(self._item_text(ann))
                it.setForeground(QColor(self._STATUS_COLORS.get(
                    ann.status or "pending", UI_THEME["text_primary"])))
                break
        self._refresh_detail(ann)

    # ── Private helpers ────────────────────────────────────────────────────────

    def _add_list_item(self, ann: Annotation):
        item = QListWidgetItem(self._item_text(ann))
        item.setData(Qt.ItemDataRole.UserRole, ann)
        status = ann.status or "pending"
        item.setForeground(QColor(self._STATUS_COLORS.get(status, UI_THEME["text_primary"])))
        self._qc_list.addItem(item)

    @staticmethod
    def _item_text(ann: Annotation) -> str:
        icon   = QCReviewPanel._STATUS_ICONS.get(ann.status or "pending", "○")
        short  = SEVERITY_SHORT.get(ann.severity, ann.severity or "—")
        defect = ann.defect or "Unknown"
        size   = (f"{ann.width_cm:.1f}×{ann.height_cm:.1f}cm"
                  if ann.width_cm is not None else "?cm")
        return f"{icon} [{short}] {defect}  {size}"

    def _clear_detail(self):
        self._current_ann = None
        self._info_label.setText("Select an annotation above to review")
        self._status_label.setText("Status: —")
        self._status_label.setStyleSheet(
            f"color:{UI_THEME['accent_amber']};font-weight:bold;"
            f"font-size:9pt;background:transparent;")
        self._rev_note.clear()
        self._approve_btn.setEnabled(False)
        self._reject_btn.setEnabled(False)

    def _refresh_detail(self, ann: Annotation):
        defect   = ann.defect or "Unknown"
        severity = ann.severity or "—"
        size     = (f"{ann.width_cm:.1f}×{ann.height_cm:.1f}cm"
                    if ann.width_cm is not None else "?cm")
        reviewed = (f"<br/><small>by {ann.reviewed_by}</small>"
                    if ann.reviewed_by else "")
        self._info_label.setText(
            f"<b>{defect}</b><br/>"
            f"Severity: {severity} · Size: {size}<br/>"
            f"ID: #{ann.ann_id[:8]}{reviewed}")
        status = ann.status or "pending"
        self._status_label.setText(f"Status: {status.upper()}")
        self._status_label.setStyleSheet(
            f"color:{self._STATUS_COLORS.get(status, UI_THEME['text_tertiary'])};"
            f"font-weight:bold;font-size:9pt;background:transparent;")
        self._rev_note.setPlainText(ann.reviewer_note or "")
        can = SESSION.can_do("approve")
        self._approve_btn.setEnabled(can and status != "approved")
        self._reject_btn.setEnabled(can and status != "rejected")

    def _on_list_item_clicked(self, item: QListWidgetItem):
        ann = item.data(Qt.ItemDataRole.UserRole)
        if isinstance(ann, Annotation):
            self._current_ann = ann
            self._refresh_detail(ann)

    def _on_approve(self):
        if not self._current_ann:
            return
        self._current_ann.status       = "approved"
        self._current_ann.reviewed_by  = SESSION.username
        self._current_ann.reviewed_at  = datetime.now().isoformat()
        self._current_ann.reviewer_note = self._rev_note.toPlainText().strip()
        self.approve_requested.emit(self._current_ann)
        self.refresh_current()   # v4.1.1: immediate UI refresh after action

    def _on_reject(self):
        if not self._current_ann:
            return
        self._current_ann.status       = "rejected"
        self._current_ann.reviewed_by  = SESSION.username
        self._current_ann.reviewed_at  = datetime.now().isoformat()
        self._current_ann.reviewer_note = self._rev_note.toPlainText().strip()
        self.reject_requested.emit(self._current_ann)
        self.refresh_current()   # v4.1.1: immediate UI refresh after action

# ==============================================================================
# QC VIEWER WIDGET  (Phase 6 — Priya Nair)
# ==============================================================================

class QCViewerWidget(QWidget):
    """
    Sam Okafor + Priya Nair: Inline QC review — replaces the centre panel
    via QStackedWidget.  Allows the inspector to approve/reject raw YOLO
    detections before committing them to the project data model.

    Undo/Redo/Clear:
      Undo  — reverts the last checkbox state change
      Redo  — re-applies a reverted change
      Clear — un-ticks all boxes on the current image (confirmable)
    """
    annotations_committed = pyqtSignal(int)   # count committed
    back_requested        = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self._project    : Optional[Project]         = None
        self._map        : Dict[str, List[DetectionResult]] = {}
        self._conf_high  : float                     = 0.45
        self._image_list : List[str]                 = []
        self._cur_idx    : int                       = -1
        self._check_wgts : Dict[str, QCheckBox]      = {}

        # QC undo/redo stacks — each entry (result_id, old_state, new_state)
        self._qc_undo: List[Tuple[str, bool, bool]] = []
        self._qc_redo: List[Tuple[str, bool, bool]] = []

        self._build_ui()

    def load_results(self, result_map: Dict[str, List[DetectionResult]],
                     project: Project, conf_high: float = 0.45):
        self._map       = result_map
        self._project   = project
        self._conf_high = conf_high
        self._image_list = [fp for fp, dets in result_map.items() if dets]
        self._qc_undo.clear(); self._qc_redo.clear()
        self._img_list_widget.clear()
        for fp in self._image_list:
            cnt  = len(result_map[fp])
            item = QListWidgetItem(f"  {os.path.basename(fp)}  ({cnt} detections)")
            self._img_list_widget.addItem(item)
        if self._image_list:
            self._img_list_widget.setCurrentRow(0)
            self._load_image_idx(0)
        self._update_undo_redo_btns()

    def _build_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # ── Top bar ────────────────────────────────────────────────────────────
        top_bar = QWidget()
        top_bar.setFixedHeight(44)
        top_bar.setStyleSheet(
            f"background:{UI_THEME['bg_secondary']};"
            f"border-bottom:1px solid {UI_THEME['border']};")
        top_lay = QHBoxLayout(top_bar)
        top_lay.setContentsMargins(10, 4, 10, 4)
        top_lay.setSpacing(6)

        self._back_btn = QPushButton("← Back to Annotate")
        self._back_btn.setStyleSheet(
            f"QPushButton{{background:{UI_THEME['bg_card']};"
            f"color:{UI_THEME['text_primary']};border:1px solid {UI_THEME['border']};"
            f"border-radius:6px;padding:5px 12px;font-weight:600;}}"
            f"QPushButton:hover{{border-color:{UI_THEME['accent_cyan']};}}"
        )
        self._back_btn.clicked.connect(self.back_requested.emit)
        top_lay.addWidget(self._back_btn)

        top_lay.addStretch()

        # QC undo/redo/clear buttons
        for text, attr, tip, action in [
            ("↩ Undo", "_undo_btn", "Undo last QC change  (Ctrl+Z)", self._qc_undo_action),
            ("↪ Redo", "_redo_btn", "Redo last QC change  (Ctrl+Y)", self._qc_redo_action),
            ("✕ Clear", "_clear_btn","Un-tick all on this image",     self._qc_clear_action),
        ]:
            btn = QPushButton(text)
            btn.setToolTip(tip)
            btn.setStyleSheet(
                f"QPushButton{{background:{UI_THEME['bg_card']};"
                f"color:{UI_THEME['text_secondary']};border:1px solid {UI_THEME['border']};"
                f"border-radius:6px;padding:5px 12px;font-weight:600;}}"
                f"QPushButton:hover{{color:{UI_THEME['accent_cyan']};"
                f"border-color:{UI_THEME['accent_cyan']};}}"
                f"QPushButton:disabled{{color:{UI_THEME['text_tertiary']};}}"
            )
            btn.clicked.connect(action)
            top_lay.addWidget(btn)
            setattr(self, attr, btn)

        top_lay.addSpacing(12)

        # Select all / none
        sa_btn = QPushButton("☑ All")
        sn_btn = QPushButton("☐ None")
        for b, fn in [(sa_btn, self._select_all), (sn_btn, self._select_none)]:
            b.setStyleSheet(
                f"QPushButton{{background:{UI_THEME['bg_card']};"
                f"color:{UI_THEME['text_secondary']};border:1px solid {UI_THEME['border']};"
                f"border-radius:6px;padding:5px 10px;font-weight:600;}}"
                f"QPushButton:hover{{color:{UI_THEME['accent_cyan']};"
                f"border-color:{UI_THEME['accent_cyan']};}}"
            )
            b.clicked.connect(fn)
            top_lay.addWidget(b)

        top_lay.addSpacing(12)

        self._commit_btn = QPushButton("✔  Commit Approved")
        self._commit_btn.setStyleSheet(
            f"QPushButton{{background:{UI_THEME['accent_green']};"
            f"color:#0d1117;border:none;border-radius:6px;padding:6px 16px;"
            f"font-weight:bold;}}"
            f"QPushButton:hover{{background:{UI_THEME['sev2']};}}"
        )
        self._commit_btn.clicked.connect(self._commit)
        top_lay.addWidget(self._commit_btn)
        root.addWidget(top_bar)

        # ── Body: image list (left) + viewer (centre) + checkbox panel (right) ──
        body = QSplitter(Qt.Orientation.Horizontal)
        body.setStyleSheet(f"QSplitter::handle{{background:{UI_THEME['border']};width:2px;}}")

        # Image list
        left_w = QWidget()
        left_w.setMaximumWidth(240)
        left_w.setStyleSheet(
            f"background:{UI_THEME['bg_secondary']};border-right:1px solid {UI_THEME['border']};")
        ll     = QVBoxLayout(left_w)
        ll.setContentsMargins(0, 0, 0, 0)
        lbl = QLabel("  IMAGES WITH DETECTIONS")
        lbl.setStyleSheet(
            f"color:{UI_THEME['text_tertiary']};font-size:8pt;font-weight:bold;"
            f"padding:8px 10px 4px 10px;letter-spacing:1px;background:transparent;")
        ll.addWidget(lbl)
        self._img_list_widget = QListWidget()
        self._img_list_widget.setStyleSheet(
            f"QListWidget{{background:{UI_THEME['bg_secondary']};border:none;}}"
            f"QListWidget::item{{padding:6px 10px;border-radius:4px;}}"
            f"QListWidget::item:selected{{background:{UI_THEME['bg_card']};"
            f"color:{UI_THEME['accent_cyan']};}}"
        )
        self._img_list_widget.currentRowChanged.connect(self._load_image_idx)
        ll.addWidget(self._img_list_widget)
        body.addWidget(left_w)

        # Image viewer (label-based — no QGraphicsView overhead for read-only QC)
        centre_w = QWidget()
        centre_lay = QVBoxLayout(centre_w)
        centre_lay.setContentsMargins(8, 8, 8, 8)
        self._qc_img_label = QLabel("Select an image from the list")
        self._qc_img_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._qc_img_label.setStyleSheet(
            f"color:{UI_THEME['text_secondary']};background:{UI_THEME['bg_primary']};"
            f"border:1px solid {UI_THEME['border']};border-radius:8px;")
        self._qc_img_label.setSizePolicy(
            QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        centre_lay.addWidget(self._qc_img_label)

        self._summary_lbl = QLabel("No detections loaded")
        self._summary_lbl.setStyleSheet(
            f"color:{UI_THEME['text_secondary']};font-size:9pt;"
            f"background:transparent;padding:4px 0;")
        centre_lay.addWidget(self._summary_lbl)
        body.addWidget(centre_w)

        # Checkbox panel
        right_w = QWidget()
        right_w.setMaximumWidth(260)
        right_w.setStyleSheet(
            f"background:{UI_THEME['bg_secondary']};border-left:1px solid {UI_THEME['border']};")
        rl     = QVBoxLayout(right_w)
        rl.setContentsMargins(0, 0, 0, 0)
        rlbl = QLabel("  DETECTIONS — APPROVE?")
        rlbl.setStyleSheet(
            f"color:{UI_THEME['text_tertiary']};font-size:8pt;font-weight:bold;"
            f"padding:8px 10px 4px 10px;letter-spacing:1px;background:transparent;")
        rl.addWidget(rlbl)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet(
            f"QScrollArea{{border:none;background:{UI_THEME['bg_secondary']};}}")
        self._box_container = QWidget()
        self._box_layout    = QVBoxLayout(self._box_container)
        self._box_layout.setContentsMargins(8, 8, 8, 8)
        self._box_layout.setSpacing(4)
        self._box_layout.addStretch()
        scroll.setWidget(self._box_container)
        rl.addWidget(scroll)
        body.addWidget(right_w)

        body.setSizes([220, 700, 240])
        root.addWidget(body)

    def _load_image_idx(self, idx: int):
        if idx < 0 or idx >= len(self._image_list):
            return
        self._cur_idx = idx
        fp      = self._image_list[idx]
        results = self._map.get(fp, [])

        # Render image with detection boxes via Pillow
        try:
            with Image.open(fp) as src:
                img = src.convert("RGB")
            draw = ImageDraw.Draw(img)
            for dr in results:
                col = ("#3fb950" if dr.confidence >= self._conf_high else "#d29922")
                draw.rectangle(
                    [dr.x1_px, dr.y1_px, dr.x2_px, dr.y2_px],
                    outline=col, width=3)
                draw.text((dr.x1_px + 2, dr.y1_px - 14),
                          f"{dr.class_name} {dr.confidence:.2f}", fill=col)
            # Scale to fit label
            lbl_w = max(self._qc_img_label.width(), 400)
            lbl_h = max(self._qc_img_label.height(), 300)
            img.thumbnail((lbl_w, lbl_h), Image.LANCZOS)
            # BytesIO imported at module level (FIX-02c: removed inline import)
            buf = BytesIO(); img.save(buf, "PNG"); buf.seek(0)
            qimg = QImage(); qimg.loadFromData(buf.read())
            self._qc_img_label.setPixmap(QPixmap.fromImage(qimg))
        except Exception as exc:
            self._qc_img_label.setText(f"Could not render: {exc}")

        self._rebuild_box_panel(results)

    def _rebuild_box_panel(self, results: List[DetectionResult]):
        while self._box_layout.count():
            item = self._box_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        self._check_wgts.clear()

        for dr in results:
            row = QWidget()
            rl  = QHBoxLayout(row)
            rl.setContentsMargins(4, 2, 4, 2)
            chk = QCheckBox()
            chk.setChecked(dr.confidence >= self._conf_high)
            dr.approved = chk.isChecked()

            # Track state for undo/redo
            def _on_toggled(checked: bool, _dr=dr, _chk=chk):
                old = not checked
                self._qc_undo.append((_dr.result_id, old, checked))
                self._qc_redo.clear()
                _dr.approved = checked
                self._update_summary()
                self._update_undo_redo_btns()

            chk.toggled.connect(_on_toggled)

            conf_col = (UI_THEME["accent_green"] if dr.confidence >= self._conf_high
                        else UI_THEME["accent_amber"])
            lbl = QLabel(
                f"<span style='color:{conf_col};font-weight:bold;'>"
                f"{dr.class_name}  {dr.confidence:.2f}"
                f"</span><br/>"
                f"<small style='color:{UI_THEME['text_tertiary']};'>"
                f"{int(dr.x1_px)},{int(dr.y1_px)} → "
                f"{int(dr.x2_px)},{int(dr.y2_px)}"
                f"</small>"
            )
            lbl.setTextFormat(Qt.TextFormat.RichText)
            rl.addWidget(chk); rl.addWidget(lbl, 1)
            self._box_layout.addWidget(row)
            self._check_wgts[dr.result_id] = chk

        self._box_layout.addStretch()
        self._update_summary()

    def _update_summary(self):
        if self._cur_idx < 0 or self._cur_idx >= len(self._image_list):
            return
        results  = self._map.get(self._image_list[self._cur_idx], [])
        approved = sum(1 for dr in results if dr.approved)
        self._summary_lbl.setText(
            f"{approved} / {len(results)} approved on this image  "
            f"|  Total approved: "
            f"{sum(dr.approved for drs in self._map.values() for dr in drs)}")

    def _update_undo_redo_btns(self):
        self._undo_btn.setEnabled(bool(self._qc_undo))
        self._redo_btn.setEnabled(bool(self._qc_redo))

    # ── Undo / Redo / Clear ────────────────────────────────────────────────────

    def _qc_undo_action(self):
        """Sam Okafor: Revert last checkbox state change in QC panel."""
        if not self._qc_undo:
            return
        result_id, old_state, new_state = self._qc_undo.pop()
        self._qc_redo.append((result_id, old_state, new_state))
        chk = self._check_wgts.get(result_id)
        if chk:
            chk.blockSignals(True)
            chk.setChecked(old_state)
            chk.blockSignals(False)
        # Update approved flag in detection result
        for drs in self._map.values():
            for dr in drs:
                if dr.result_id == result_id:
                    dr.approved = old_state
        self._update_summary()
        self._update_undo_redo_btns()

    def _qc_redo_action(self):
        """Sam Okafor: Re-apply a reverted QC checkbox change."""
        if not self._qc_redo:
            return
        result_id, old_state, new_state = self._qc_redo.pop()
        self._qc_undo.append((result_id, old_state, new_state))
        chk = self._check_wgts.get(result_id)
        if chk:
            chk.blockSignals(True)
            chk.setChecked(new_state)
            chk.blockSignals(False)
        for drs in self._map.values():
            for dr in drs:
                if dr.result_id == result_id:
                    dr.approved = new_state
        self._update_summary()
        self._update_undo_redo_btns()

    def _qc_clear_action(self):
        """Sam Okafor: Un-tick all boxes on current image (with confirmation)."""
        if self._cur_idx < 0:
            return
        ret = QMessageBox.question(
            self, "Clear Approvals",
            "Un-tick all detections on this image?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.Cancel,
        )
        if ret != QMessageBox.StandardButton.Yes:
            return
        results = self._map.get(self._image_list[self._cur_idx], [])
        for dr in results:
            old = dr.approved
            if old:
                self._qc_undo.append((dr.result_id, True, False))
            dr.approved = False
        for chk in self._check_wgts.values():
            chk.blockSignals(True)
            chk.setChecked(False)
            chk.blockSignals(False)
        self._qc_redo.clear()
        self._update_summary()
        self._update_undo_redo_btns()

    def _select_all(self):
        for dr_id, chk in self._check_wgts.items():
            if not chk.isChecked():
                chk.setChecked(True)

    def _select_none(self):
        for dr_id, chk in self._check_wgts.items():
            if chk.isChecked():
                chk.setChecked(False)

    def _commit(self):
        """Sam Okafor: Convert approved detections → Annotation objects."""
        if not self._project:
            return
        committed = 0
        for fp, results in self._map.items():
            fname = os.path.basename(fp)
            if fname not in self._project.images:
                continue
            irec = self._project.images[fname]
            for dr in results:
                if not dr.approved:
                    continue
                ann = Annotation(
                    ann_id    = _make_ann_id(fp),
                    mode      = "box",
                    defect    = dr.class_name,
                    severity  = "Major",              # default; inspector edits in panel
                    blade     = irec.blade,
                    x1_px=dr.x1_px, y1_px=dr.y1_px,
                    x2_px=dr.x2_px, y2_px=dr.y2_px,
                    gsd_source="none",
                )
                irec.annotations.append(ann)
                committed += 1
        save_project(self._project)
        self.annotations_committed.emit(committed)
        QMessageBox.information(
            self, "Committed",
            f"{committed} annotation(s) committed to project.\n"
            "They are now visible in the main viewer — re-select an image to refresh.")

    def keyPressEvent(self, event):
        """Sam Okafor: Ctrl+Z / Ctrl+Y shortcuts for QC undo/redo."""
        if event.modifiers() & Qt.KeyboardModifier.ControlModifier:
            if event.key() == Qt.Key.Key_Z:
                self._qc_undo_action(); return
            if event.key() == Qt.Key.Key_Y:
                self._qc_redo_action(); return
        super().keyPressEvent(event)

# ==============================================================================
# DETECTION TAB  (Natalie Cross + Jamie Liu)
# ==============================================================================

class DetectionTab(QWidget):
    detection_finished = pyqtSignal(dict)   # {filepath: [DetectionResult]}

    def __init__(self, image_paths: List[str], project: Project,
                 cfg: AppConfig, parent=None):
        super().__init__(parent)
        self._image_paths = image_paths
        self._project     = project
        self._cfg         = cfg
        self._worker: Optional[DetectionWorker] = None
        self._pool        = QThreadPool()
        self._pool.setMaxThreadCount(1)
        self._results: Dict[str, List[DetectionResult]] = {}
        self._build_ui()

    def _build_ui(self):
        lay = QVBoxLayout(self)
        lay.setContentsMargins(12, 12, 12, 12)
        lay.setSpacing(10)

        # Model section
        grp_model = QGroupBox("YOLO Model")
        gm = QFormLayout(grp_model)
        mp_row = QHBoxLayout()
        self._model_edit = QLineEdit(self._cfg.get("DETECTION", "ModelPath"))
        self._model_edit.setPlaceholderText("Path to .pt weights file…")
        browse_btn = QPushButton("Browse…")
        browse_btn.setFixedWidth(80)
        browse_btn.clicked.connect(self._browse_model)
        mp_row.addWidget(self._model_edit); mp_row.addWidget(browse_btn)
        gm.addRow("Model (.pt):", mp_row)
        self._device_combo = QComboBox()
        self._device_combo.addItems(["auto", "cuda", "cpu"])
        self._device_combo.setCurrentText(self._cfg.get("DETECTION", "Device", "auto"))
        gm.addRow("Device:", self._device_combo)
        lay.addWidget(grp_model)

        # Confidence section
        grp_conf = QGroupBox("Confidence Thresholds")
        gc = QFormLayout(grp_conf)

        def _slider_row(label, key, default):
            rw = QWidget(); rl = QHBoxLayout(rw); rl.setContentsMargins(0, 0, 0, 0)
            sl = QSlider(Qt.Orientation.Horizontal)
            sl.setRange(1, 99)
            sl.setValue(int(float(self._cfg.get("DETECTION", key, str(default))) * 100))
            lbl = QLabel(f"{sl.value()/100:.2f}"); lbl.setFixedWidth(36)
            sl.valueChanged.connect(lambda v, l=lbl, k=key:
                (l.setText(f"{v/100:.2f}"),
                 self._cfg.set("DETECTION", k, f"{v/100:.4f}")))
            rl.addWidget(sl); rl.addWidget(lbl); gc.addRow(label, rw)
            return sl

        self._sl_high = _slider_row("Auto-tick (ConfHigh):", "ConfHigh", 0.45)
        self._sl_low  = _slider_row("Show in QC (ConfLow):", "ConfLow",  0.25)
        self._sl_nms  = _slider_row("NMS IoU:", "NMS_IOU", 0.45)
        self._imgsize_combo = QComboBox()
        self._imgsize_combo.addItems(["320", "416", "640", "1280"])
        self._imgsize_combo.setCurrentText(self._cfg.get("DETECTION", "ImgSize", "640"))
        gc.addRow("Infer size (px):", self._imgsize_combo)
        lay.addWidget(grp_conf)

        # Run section
        grp_run = QGroupBox("Batch Detection")
        gr = QVBoxLayout(grp_run)
        gr.addWidget(QLabel(f"Images in session: {len(self._image_paths)}"))
        self._run_btn = QPushButton("🤖  Run Detection on All Images")
        self._run_btn.setStyleSheet(
            f"background:{UI_THEME['accent_amber']};color:#0d1117;"
            f"font-weight:bold;padding:9px;border-radius:6px;border:none;")
        self._run_btn.clicked.connect(self._run_detection)
        gr.addWidget(self._run_btn)
        self._cancel_btn = QPushButton("✕  Cancel")
        self._cancel_btn.setEnabled(False)
        self._cancel_btn.clicked.connect(self._cancel)
        gr.addWidget(self._cancel_btn)
        self._prog = QProgressBar(); self._prog.setRange(0, 100); self._prog.setValue(0)
        gr.addWidget(self._prog)
        self._status_lbl = QLabel("Ready")
        self._status_lbl.setStyleSheet(f"color:{UI_THEME['text_secondary']};font-size:9pt;")
        gr.addWidget(self._status_lbl)
        lay.addWidget(grp_run)
        lay.addStretch()

    def _browse_model(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select YOLO Model (.pt)", "",
                                              "YOLO Weights (*.pt)")
        if path:
            self._model_edit.setText(path)
            self._cfg.set("DETECTION", "ModelPath", path)

    def _run_detection(self):
        if not YOLO_AVAILABLE:
            QMessageBox.warning(self, "YOLO Missing",
                "ultralytics is not installed.\npip install ultralytics")
            return
        model_path = self._model_edit.text().strip()
        if not model_path or not Path(model_path).exists():
            QMessageBox.warning(self, "Model Required",
                "Please select a valid YOLO .pt model file.")
            return
        if not self._image_paths:
            QMessageBox.warning(self, "No Images", "Load images first.")
            return
        if not _model_mgr.load(model_path, self._device_combo.currentText()):
            QMessageBox.critical(self, "Model Error", "Failed to load model.")
            return

        conf_low = float(self._sl_low.value()) / 100
        iou      = float(self._sl_nms.value()) / 100
        img_size = int(self._imgsize_combo.currentText())

        self._worker = DetectionWorker(self._image_paths, _model_mgr,
                                       conf_low, iou, img_size)
        self._worker.signals.progress.connect(self._on_progress)
        self._worker.signals.image_done.connect(self._on_image_done)
        self._worker.signals.batch_done.connect(self._on_batch_done)
        self._worker.signals.error.connect(self._on_error)

        self._run_btn.setEnabled(False); self._cancel_btn.setEnabled(True)
        self._results.clear(); self._prog.setValue(0)
        self._pool.start(self._worker)

    def _cancel(self):
        if self._worker:
            self._worker.cancel()
        self._run_btn.setEnabled(True); self._cancel_btn.setEnabled(False)
        self._status_lbl.setText("Cancelled.")

    def _on_progress(self, done: int, total: int):
        self._prog.setValue(int(done / max(total, 1) * 100))
        self._status_lbl.setText(f"Processing {done}/{total}…")

    def _on_image_done(self, fp: str, results: list):
        self._results[fp] = results

    def _on_batch_done(self, total: int, total_dets: int):
        self._run_btn.setEnabled(True); self._cancel_btn.setEnabled(False)
        self._prog.setValue(100)
        self._status_lbl.setText(
            f"Done — {total_dets} detections across {total} images.")
        self.detection_finished.emit(self._results)

    def _on_error(self, fname: str, msg: str):
        log.warning(f"Detection error on {fname}: {msg}")
        self._status_lbl.setText(f"Error on {fname}")


# ==============================================================================
# TRAINING TAB  (Natalie Cross + Sarah Chen)
# ==============================================================================

class TrainingTab(QWidget):
    def __init__(self, project: Project, cfg: AppConfig, parent=None):
        super().__init__(parent)
        self._project = project
        self._cfg     = cfg
        self._worker: Optional[TrainingWorker] = None
        self._build_ui()

    def _build_ui(self):
        lay = QVBoxLayout(self)
        lay.setContentsMargins(12, 12, 12, 12)
        lay.setSpacing(10)

        grp_exp = QGroupBox("Dataset Export")
        ge = QFormLayout(grp_exp)
        exp_row = QHBoxLayout()
        self._exp_edit = QLineEdit(self._cfg.get("TRAINING", "ExportFolder"))
        self._exp_edit.setPlaceholderText("Folder to export YOLO dataset…")
        exp_browse = QPushButton("Browse…")
        exp_browse.setFixedWidth(80)
        exp_browse.clicked.connect(self._browse_export)
        exp_row.addWidget(self._exp_edit); exp_row.addWidget(exp_browse)
        ge.addRow("Export folder:", exp_row)
        lay.addWidget(grp_exp)

        grp_hp = QGroupBox("Hyperparameters")
        gh = QFormLayout(grp_hp)
        self._ep_spin  = QSpinBox(); self._ep_spin.setRange(1, 500)
        self._ep_spin.setValue(int(self._cfg.get("TRAINING", "Epochs", "50")))
        gh.addRow("Epochs:", self._ep_spin)
        self._bs_spin  = QSpinBox(); self._bs_spin.setRange(1, 256)
        self._bs_spin.setValue(int(self._cfg.get("TRAINING", "BatchSize", "16")))
        gh.addRow("Batch size:", self._bs_spin)
        self._imgsz_combo = QComboBox()
        self._imgsz_combo.addItems(["320", "416", "640", "1280"])
        self._imgsz_combo.setCurrentText(self._cfg.get("TRAINING", "ImgSize", "640"))
        gh.addRow("Image size:", self._imgsz_combo)
        self._dev_combo = QComboBox()
        self._dev_combo.addItems(["auto", "cuda", "cpu"])
        self._dev_combo.setCurrentText(self._cfg.get("TRAINING", "Device", "auto"))
        gh.addRow("Device:", self._dev_combo)
        lay.addWidget(grp_hp)

        grp_run = QGroupBox("Training")
        gr = QVBoxLayout(grp_run)
        self._train_btn = QPushButton("🏋  Start Training")
        self._train_btn.setStyleSheet(
            f"background:{UI_THEME['accent_orange']};color:#0d1117;"
            f"font-weight:bold;padding:9px;border-radius:6px;border:none;")
        self._train_btn.clicked.connect(self._start_training)
        self._stop_btn = QPushButton("⏹  Stop")
        self._stop_btn.setEnabled(False)
        self._stop_btn.clicked.connect(self._stop_training)
        gr.addWidget(self._train_btn); gr.addWidget(self._stop_btn)
        self._train_prog = QProgressBar()
        self._train_prog.setRange(0, 100); self._train_prog.setValue(0)
        gr.addWidget(self._train_prog)
        lay.addWidget(grp_run)

        self._console = QPlainTextEdit()
        self._console.setReadOnly(True)
        self._console.setStyleSheet(
            f"background:#0d1117;color:{UI_THEME['accent_green']};"
            f"font-family:Consolas,Monaco,monospace;font-size:8pt;"
            f"border:1px solid {UI_THEME['border']};border-radius:6px;")
        self._console.setMinimumHeight(160)
        lay.addWidget(self._console)

    def _browse_export(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Export Folder")
        if folder:
            self._exp_edit.setText(folder)
            self._cfg.set("TRAINING", "ExportFolder", folder)

    def _start_training(self):
        if not YOLO_AVAILABLE:
            QMessageBox.warning(self, "YOLO Missing",
                "pip install ultralytics")
            return
        export_dir = self._exp_edit.text().strip()
        if not export_dir:
            QMessageBox.warning(self, "Export Folder Required",
                "Please select a folder to export the dataset.")
            return
        self._cfg.set("TRAINING", "Epochs",    str(self._ep_spin.value()))
        self._cfg.set("TRAINING", "BatchSize", str(self._bs_spin.value()))
        self._cfg.set("TRAINING", "ImgSize",   self._imgsz_combo.currentText())
        self._cfg.set("TRAINING", "Device",    self._dev_combo.currentText())

        self._worker = TrainingWorker(self._project, export_dir, self._cfg)
        self._worker.signals.log_line.connect(self._console.appendPlainText)
        self._worker.signals.progress.connect(self._train_prog.setValue)
        self._worker.signals.finished.connect(self._on_finished)
        self._train_btn.setEnabled(False); self._stop_btn.setEnabled(True)
        self._console.clear()
        self._worker.start()

    def _stop_training(self):
        if self._worker:
            self._worker.stop()
        self._train_btn.setEnabled(True); self._stop_btn.setEnabled(False)

    def _on_finished(self, success: bool, msg: str):
        self._train_btn.setEnabled(True); self._stop_btn.setEnabled(False)
        self._train_prog.setValue(100 if success else 0)
        if success:
            self._console.appendPlainText(f"\n✅ Training complete.\nBest model: {msg}")
            QMessageBox.information(self, "Training Complete",
                f"Model saved to:\n{msg}")
        else:
            self._console.appendPlainText(f"\n❌ Training failed: {msg}")


# ==============================================================================
# ML DIALOG  (Sarah Chen + Sam Okafor)
# ==============================================================================

class MLDialog(QDialog):
    annotations_committed = pyqtSignal(int)

    def __init__(self, image_paths: List[str], project: Project,
                 cfg: AppConfig, parent=None):
        super().__init__(parent)
        self.setWindowTitle("🤖  ML — Detection & Training")
        self.setMinimumSize(780, 600)
        lay  = QVBoxLayout(self)
        tabs = QTabWidget()
        self._det_tab   = DetectionTab(image_paths, project, cfg, self)
        self._train_tab = TrainingTab(project, cfg, self)
        tabs.addTab(self._det_tab,   "🤖  Detection")
        tabs.addTab(self._train_tab, "🏋  Training")
        self._det_tab.detection_finished.connect(self.detection_finished)
        lay.addWidget(tabs)
        close = QPushButton("Close")
        close.setFixedWidth(90)
        close.clicked.connect(self.accept)
        lay.addWidget(close, alignment=Qt.AlignmentFlag.AlignRight)

    # Expose detection results to MainWindow
    def detection_finished(self, results: dict):
        self._last_results = results

    def last_results(self) -> dict:
        return getattr(self, "_last_results", {})

# ==============================================================================
# REPORT GENERATOR  (Tom K. + Dev Patel)
# PDF output matching Sample.pdf: cover page + per-annotation detail pages.
# ==============================================================================

class ReportGenerator:
    """
    Tom K. + Dev Patel: Generates a multi-page PDF inspection report.

    Structure:
      Page 1  — Cover: metadata table, severity overview band, legend table,
                        summary text
      Page 2+ — Per-annotation: metadata header, full image, cropped inset,
                                badge cards (severity, defect, remedy, notes)
    """

    _SEV_COLORS_HEX = _SEV_HEX   # v3.0.1: use central map

    def __init__(self, project: Project,
                 report_settings: Optional[Dict[str, str]] = None):
        self._project  = project
        self._settings = report_settings or {}   # Phase 8.1: company/logo/reviewer

    def generate(self, output_path: str, also_csv: bool = True) -> bool:
        if not REPORTLAB_AVAILABLE:
            log.error("ReportGenerator: pip install reportlab")
            return False
        try:
            doc = SimpleDocTemplate(
                output_path, pagesize=A4,
                rightMargin=15 * mm, leftMargin=15 * mm,
                topMargin=20 * mm,
                # FIX-12: increased from 15mm to 32mm to give the 3x-enlarged
                # company footer logo (max 84×24mm) clear clearance below content
                bottomMargin=32 * mm,
            )
            story  = []
            story += self._build_cover_page()          # Page 1: cover + narrative
            story += self._build_defect_summary_page() # Page 2: defect summary (after Results, before defects)
            story += self._build_annotation_pages()    # Page 3+: per-annotation pages
            doc.build(story,
                      onFirstPage=self._add_header_footer,
                      onLaterPages=self._add_header_footer)
            log.info(f"PDF saved → {output_path}")
            # Auto-export CSV alongside PDF
            if also_csv:
                csv_path = output_path.replace(".pdf", "_annotations.csv")
                if not csv_path.endswith(".csv"):
                    csv_path += "_annotations.csv"
                self.export_csv(csv_path)
            return True
        except Exception as exc:
            log.error(f"ReportGenerator.generate: {exc}")
            return False

    # ── Header / footer ────────────────────────────────────────────────────────

    def _add_header_footer(self, canvas, doc):
        from reportlab.lib.units import mm
        canvas.saveState()
        # Title: "[Site Name] Aerial Wind Tower Inspection"  (user request)
        _site          = self._project.site or self._project.name or "Wind Tower"
        _report_title  = f"{_site} Aerial Wind Tower Inspection"
        reviewer_name  = self._settings.get("reviewer_name", "")
        logo_path      = self._settings.get("logo_path", "")

        # FIX-16b: Header bar lightened from dark slate (#4A5568, RGB 0.29/0.33/0.41)
        # to a soft light slate (#B8C2D1, RGB 0.72/0.76/0.82).  White text still
        # passes WCAG AA on the lighter background, and logos are no longer lost
        # against a near-black bar.
        canvas.setFillColorRGB(0.72, 0.76, 0.82)
        canvas.rect(0, A4[1] - 18 * mm, A4[0], 18 * mm, fill=1, stroke=0)

        # FIX-13: Header logo moved to RIGHT side at HALF the original size.
        # Original: left side, max 40×14mm.  New: right side, max 20×7mm.
        # The logo is right-anchored at A4[0]-12mm, vertically centred in the bar.
        # logo_drawn_w is no longer needed (logo not on left) so report title
        # is always left-anchored at 12mm with no offset.
        _hdr_logo_right_x = A4[0] - 12 * mm   # right edge of header logo
        _hdr_logo_drawn_lx = _hdr_logo_right_x  # default: no logo → nothing shifts client logo
        if logo_path and os.path.exists(logo_path):
            try:
                from reportlab.lib.utils import ImageReader
                logo_img  = ImageReader(logo_path)
                lw, lh    = logo_img.getSize()
                # FIX-13: halved cap — 20×7mm (was 40×14mm)
                max_lw, max_lh = 20 * mm, 7 * mm
                scale    = min(max_lw / lw, max_lh / lh)
                draw_lw  = lw * scale
                draw_lh  = lh * scale
                lx = _hdr_logo_right_x - draw_lw          # right-anchored
                ly = A4[1] - 18 * mm + (18 * mm - draw_lh) / 2
                canvas.drawImage(logo_path, lx, ly,
                                 width=draw_lw, height=draw_lh,
                                 mask="auto", preserveAspectRatio=True)
                # Record the left edge of the header logo so client logo can
                # sit immediately to its left with a 3mm gap
                _hdr_logo_drawn_lx = lx - 3 * mm
            except Exception as exc:
                log.warning(f"Logo embed failed: {exc}")

        # Report title — dark text on light header bar (FIX-16b: bar is now light gray)
        text_x = 12 * mm
        canvas.setFont("Helvetica-Bold", 11)
        canvas.setFillColorRGB(0.10, 0.13, 0.20)   # near-black (#1A2133) — legible on light bg
        canvas.drawString(text_x, A4[1] - 11.5 * mm, _report_title)

        # Right side header: client logo sits left of the header logo (if any).
        # FIX-15b: Site name TEXT has been removed from the PDF header right
        # side per user request.  Only the client logo image is drawn here —
        # no drawRightString call remains.  _site_x is not needed.
        client_logo_path = self._settings.get("client_logo_path", "")
        if client_logo_path and os.path.exists(client_logo_path):
            try:
                from reportlab.lib.utils import ImageReader as _IR2
                _cli_img = _IR2(client_logo_path)
                _cw, _ch = _cli_img.getSize()
                _max_cw, _max_ch = 38 * mm, 14 * mm
                _sc = min(_max_cw / _cw, _max_ch / _ch)
                _dcw, _dch = _cw * _sc, _ch * _sc
                # Place client logo immediately left of the (possibly halved) header logo
                _cx = _hdr_logo_drawn_lx - _dcw
                _cy = A4[1] - 18 * mm + (18 * mm - _dch) / 2
                canvas.drawImage(client_logo_path, _cx, _cy,
                                 width=_dcw, height=_dch,
                                 mask="auto", preserveAspectRatio=True)
            except Exception as exc:
                log.warning(f"Client logo header embed failed: {exc}")

        # Cyan accent line below header
        canvas.setStrokeColorRGB(0.0, 0.83, 0.88)
        canvas.setLineWidth(2)
        canvas.line(0, A4[1] - 18 * mm, A4[0], A4[1] - 18 * mm)

        # ── Footer ──────────────────────────────────────────────────────────
        total_pg = self._count_pages()

        # Company logo — left side of footer (FIX-12: 3x larger, covers Turbine text)
        # Max size increased 28×8mm → 84×24mm (3×).  When the logo is drawn the
        # "Turbine: XXX" text is suppressed because the logo visually covers that area.
        co_logo_path   = self._settings.get("company_logo_path", "")
        _logo_drawn    = False          # flag: suppress Turbine text when True
        _FOOTER_BAND_H = 28 * mm       # total footer band height (y=2mm … y=30mm)
        _FOOTER_BASE_Y = 2  * mm       # bottom edge of the logo zone
        if co_logo_path and os.path.exists(co_logo_path):
            try:
                from reportlab.lib.utils import ImageReader as _IR3
                _co_img = _IR3(co_logo_path)
                _cow, _coh = _co_img.getSize()
                # FIX-14d: "a bit more" width increase — 110mm wide, 28mm tall
                # (was 84×24mm after FIX-12's 3× enlargement).  The footer band
                # is 28mm tall so the taller cap still fits with no overflow.
                _max_cow, _max_coh = 110 * mm, 28 * mm
                _cosc  = min(_max_cow / _cow, _max_coh / _coh)
                _dcow  = _cow * _cosc
                _dcoh  = _coh * _cosc
                # Centre the logo vertically in the footer band
                _co_y  = _FOOTER_BASE_Y + (_FOOTER_BAND_H - _dcoh) / 2
                canvas.drawImage(co_logo_path, 12 * mm, _co_y,
                                 width=_dcow, height=_dcoh,
                                 mask="auto", preserveAspectRatio=True)
                _logo_drawn = True
            except Exception as exc:
                log.warning(f"Company footer logo embed failed: {exc}")

        # "Turbine: XXX" text — only rendered when no logo is drawn.
        # When the logo is present it visually covers the same left-footer zone.
        canvas.setFont("Helvetica", 8)
        canvas.setFillColorRGB(0.45, 0.45, 0.5)
        if not _logo_drawn:
            footer_left = f"Turbine: {self._project.turbine_id or '—'}"
            canvas.drawString(15 * mm, 14 * mm, footer_left)

        # Right: page numbers — vertically centred in the new 28mm footer band
        canvas.setFont("Helvetica", 8)
        canvas.drawRightString(A4[0] - 15 * mm, 14 * mm,
                               f"Page {doc.page} of {total_pg}")
        canvas.restoreState()

    def _count_pages(self) -> int:
        # v2.2.0: Updated for 2-per-page layout (v2.1.3 TODO #4)
        # 1 fixed page (cover) + component headers + annotation pages (2 per page)
        components = {}
        for irec in self._project.images.values():
            comp = irec.blade or "Unknown"
            components.setdefault(comp, []).append(irec)
        
        total = 1  # cover page
        
        for comp, recs in components.items():
            # Count annotations for this component
            ann_count = sum(len(irec.annotations) for irec in recs if irec.annotations)
            if ann_count > 0:
                total += 1  # component header page
                # v2.1.3: Two defects per page, round up for odd numbers
                total += (ann_count + 1) // 2
        
        return max(total, 1)

    # ── Cover page ─────────────────────────────────────────────────────────────

    def _build_cover_page(self) -> list:
        """
        PDF cover page — revised per user spec:
          • "[Site] Aerial Wind Tower Inspection" heading
          • WTG placeholder image box (clearly labelled, user replaces)
          • 3-col meta row: Date | Type | Turbine  (Images/Annotations REMOVED)
          • Location row: site name + GPS coords (both — user request)
          • Blade serial numbers table
          • Defect index  (numbered, severity-sorted)
          • Severity overview band + legend table
          • Inspector info box
          • GL-16 sections: Objective, Scope, Data Collection, Turbine Specs
          • Results heading → PageBreak
          NOTE: Summary block REMOVED per user request.
        """
        story = []
        p     = self._project
        report_date = datetime.now().strftime("%B %d, %Y").upper()
        usable_w    = A4[0] - 30 * mm   # 180 mm

        # ── Cover heading ───────────────────────────────────────────────────────
        site_name    = p.site or p.name or "Wind Tower"
        cover_title  = f"{site_name} Aerial Wind Tower Inspection"
        story.append(Spacer(1, 3 * mm))
        # v4.2.0: Increased font size from 18→22pt and ensured bold for better visibility
        story.append(Paragraph(
            f"<b>{cover_title}</b>",
            ParagraphStyle("CvrTitle", fontName="Helvetica-Bold", fontSize=22,
                           textColor=rl_colors.HexColor("#1a202c"), spaceAfter=4,
                           leading=28)))
        # v4.2.0: Increased subtitle font size from 9→11pt and made bold
        story.append(Paragraph(
            "<b>Drone Visual Inspection — Executive Report</b>",
            ParagraphStyle("CvrSub", fontName="Helvetica-Bold", fontSize=11,
                           textColor=rl_colors.HexColor("#4a5568"), spaceAfter=2)))
        
        # FIX-16a: "Report generated on <date> and <HH:MM IST>" paragraph removed.
        # The subtitle already provides sufficient date context; the generated-on
        # line was noisy and cluttered the cover page header area.

        # ── WTG front-page image (user-selectable) or placeholder ──────────────
        wtg_img_path = self._settings.get("wtg_image_path", "")
        # T11 FIX: Auto-read GPS from any project image EXIF if settings GPS is empty.
        # Tries each image record until a non-empty gps_coords is found.
        if not self._settings.get("gps_coords", ""):
            for _irec in self._project.images.values():
                _exif = _read_exif_metadata(_irec.filepath)
                if _exif.get("gps_coords"):
                    # Temporarily set for this cover render; not saved to settings
                    self._settings = dict(self._settings)
                    self._settings["gps_coords"] = _exif["gps_coords"]
                    log.debug(f"Cover GPS auto-read from {_irec.filename}: "
                              f"{_exif['gps_coords']}")
                    break

        if wtg_img_path and os.path.exists(wtg_img_path):
            try:
                story.append(Spacer(1, 2 * mm))
                # T11 FIX: Increased cover image height from 52mm → 90mm
                story.append(RLImage(wtg_img_path, width=usable_w, height=90 * mm,
                                     kind="proportional"))
                # v4.3.1 MOD-6: Image caption removed per user request
                story.append(Spacer(1, 4 * mm))
            except Exception:
                story.append(Spacer(1, 4 * mm))
        else:
            ph_tbl = Table(
                [[Paragraph(
                    "<b>[ INSERT WTG FRONT IMAGE HERE ]</b><br/>"
                    "<font color='#718096' size='8'>"
                    "Set path in Report Settings → Identity → WTG Cover Photo, "
                    "or replace this box in your PDF editor."
                    "</font>",
                    ParagraphStyle("ph", fontName="Helvetica", fontSize=9,
                                   textColor=rl_colors.HexColor("#4a5568"),
                                   alignment=TA_CENTER, leading=14))]],
                colWidths=[usable_w], rowHeights=[90 * mm])
            ph_tbl.setStyle(TableStyle([
                ("BOX",          (0, 0), (-1, -1), 1.2, rl_colors.HexColor("#a0aec0")),
                ("BACKGROUND",   (0, 0), (-1, -1), rl_colors.HexColor("#edf2f7")),
                ("ALIGN",        (0, 0), (-1, -1), "CENTER"),
                ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING",   (0, 0), (-1, -1), 14),
                ("BOTTOMPADDING",(0, 0), (-1, -1), 14),
            ]))
            story.append(ph_tbl)
            story.append(Spacer(1, 4 * mm))

        # ── Single-column meta row: WTG No only (Component Type removed per spec) ──
        _raw_turb = p.turbine_id or "—"
        _wtg_no   = (f"WTG-{_raw_turb}"
                     if _raw_turb and _raw_turb != "—" and
                        not _raw_turb.upper().startswith("WTG")
                     else _raw_turb)
        usable_w_meta = A4[0] - 30 * mm
        meta_data = [
            ["WTG No:"],
            [_wtg_no],
        ]
        meta_tbl = Table(meta_data, colWidths=[usable_w_meta])
        meta_tbl.setStyle(TableStyle([
            # v4.2.1: Header label row made bold for visual prominence (user spec)
            ("FONTNAME",     (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTNAME",     (0, 1), (-1, 1),  "Helvetica-Bold"),
            ("FONTSIZE",     (0, 0), (-1, 0),  8),
            ("FONTSIZE",     (0, 1), (-1, 1),  11),
            ("TEXTCOLOR",    (0, 0), (-1, 0),  rl_colors.HexColor("#888888")),
            ("TEXTCOLOR",    (0, 1), (-1, 1),  rl_colors.HexColor("#111111")),
            # CHG-F: Centre both label and value rows for WTG No, matching reference
            ("ALIGN",        (0, 0), (-1, -1), "CENTER"),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
            ("TOPPADDING",   (0, 0), (-1, -1), 4),
            # v3.3.13 FIX: Add visible border around WTG No table like location table
            ("BOX",          (0, 0), (-1, -1), 0.5, rl_colors.HexColor("#cbd5e0")),
            ("BACKGROUND",   (0, 0), (-1, -1), rl_colors.HexColor("#f7f8fa")),
        ]))
        story.append(meta_tbl)

        # v4.2.0: Add Scan Date if provided by user in project metadata
        if p.scan_date:
            scan_date_data = [
                ["Scan Date:"],
                [p.scan_date],
            ]
            scan_date_tbl = Table(scan_date_data, colWidths=[usable_w_meta])
            scan_date_tbl.setStyle(TableStyle([
                # v4.2.1: Header label row made bold (newly added row, matches WTG No label style)
                ("FONTNAME",     (0, 0), (-1, 0),  "Helvetica-Bold"),
                ("FONTNAME",     (0, 1), (-1, 1),  "Helvetica-Bold"),
                ("FONTSIZE",     (0, 0), (-1, 0),  8),
                ("FONTSIZE",     (0, 1), (-1, 1),  11),
                ("TEXTCOLOR",    (0, 0), (-1, 0),  rl_colors.HexColor("#888888")),
                ("TEXTCOLOR",    (0, 1), (-1, 1),  rl_colors.HexColor("#111111")),
                ("ALIGN",        (0, 0), (-1, -1), "CENTER"),
                ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
                ("TOPPADDING",   (0, 0), (-1, -1), 4),
                ("BOX",          (0, 0), (-1, -1), 0.5, rl_colors.HexColor("#cbd5e0")),
                ("BACKGROUND",   (0, 0), (-1, -1), rl_colors.HexColor("#f7f8fa")),
            ]))
            story.append(Spacer(1, 3 * mm))   # FIX-16c: normalised from 2mm → 3mm
            story.append(scan_date_tbl)

        # ── Tower location: site name + GPS coords (both — user request) ────────
        gps_val = self._settings.get("gps_coords", "") or ""
        loc_data = [
            ["Tower Location (Site):", "GPS Coordinates:"],
            [p.site or "—", gps_val or "Not recorded"],
        ]
        loc_tbl = Table(loc_data, colWidths=[90 * mm, 90 * mm])
        loc_tbl.setStyle(TableStyle([
            # v4.2.1: Label row made bold to match other cover page header labels
            ("FONTNAME",     (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTNAME",     (0, 1), (-1, 1),  "Helvetica-Bold"),
            ("FONTSIZE",     (0, 0), (-1, 0),  8),
            ("FONTSIZE",     (0, 1), (-1, 1),  9),
            ("TEXTCOLOR",    (0, 0), (-1, 0),  rl_colors.HexColor("#888888")),
            ("TEXTCOLOR",    (0, 1), (-1, 1),  rl_colors.HexColor("#111111")),
            ("BACKGROUND",   (0, 0), (-1, -1), rl_colors.HexColor("#f7f8fa")),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
            ("TOPPADDING",   (0, 0), (-1, -1), 5),
            ("LEFTPADDING",  (0, 0), (-1, -1), 8),
            ("LINEBEFORE",   (1, 0), (-1, -1), 0.5, rl_colors.HexColor("#dddddd")),
            ("BOX",          (0, 0), (-1, -1), 0.5, rl_colors.HexColor("#cbd5e0")),
        ]))
        story.append(Spacer(1, 3 * mm))   # uniform 3 mm, matches all other cover gaps
        story.append(loc_tbl)
        # v4.4.9: Removed spurious HRFlowable that drew a visible grey separator line
        # between the Tower Location/GPS table and the Turbine Manufacturer/Power table.
        # Confirmed by cover-page screenshot review — the line was highlighted in yellow.
        # Inter-table spacing is now a uniform Spacer(1, 3*mm), identical to all other
        # gaps on the cover page (meta_tbl→scan_date, scan_date→loc_tbl, mfr_pwr→...).

        # ── v4.3.0: Two-column row — Turbine Manufacturer | Rated Power ───────────
        # Shown below Tower Location/GPS row; values come from project metadata or
        # fall back to report settings (turbine_model / rated_power keys).
        _mfr  = p.turbine_manufacturer or self._settings.get("turbine_manufacturer", "") or "—"
        _rpwr = (p.rated_power or
                 self._settings.get("rated_power", "") or
                 self._settings.get("turbine_model", "") or "—")
        mfr_pwr_data = [
            ["Turbine Manufacturer:", "Rated Power:"],
            [_mfr, _rpwr],
        ]
        mfr_pwr_tbl = Table(mfr_pwr_data, colWidths=[90 * mm, 90 * mm])
        mfr_pwr_tbl.setStyle(TableStyle([
            # v4.3.0: Both label and value rows bold; label row uses muted gray color
            ("FONTNAME",     (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTNAME",     (0, 1), (-1, 1),  "Helvetica-Bold"),
            ("FONTSIZE",     (0, 0), (-1, 0),  8),
            ("FONTSIZE",     (0, 1), (-1, 1),  9),
            ("TEXTCOLOR",    (0, 0), (-1, 0),  rl_colors.HexColor("#888888")),
            ("TEXTCOLOR",    (0, 1), (-1, 1),  rl_colors.HexColor("#111111")),
            ("BACKGROUND",   (0, 0), (-1, -1), rl_colors.HexColor("#f7f8fa")),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
            ("TOPPADDING",   (0, 0), (-1, -1), 5),
            ("LEFTPADDING",  (0, 0), (-1, -1), 8),
            ("LINEBEFORE",   (1, 0), (-1, -1), 0.5, rl_colors.HexColor("#dddddd")),
            ("BOX",          (0, 0), (-1, -1), 0.5, rl_colors.HexColor("#cbd5e0")),
        ]))
        story.append(Spacer(1, 3 * mm))   # FIX-16c: normalised (was no spacer before mfr_pwr_tbl)
        story.append(mfr_pwr_tbl)
        story.append(Spacer(1, 3 * mm))   # FIX-16c: normalised from 4mm → 3mm

        # BLADE SERIAL NUMBERS section removed per spec (image mark-up, orange X)
        # SEVERITY OVERVIEW moved to _build_defect_summary_page() — appears
        # after Results section, immediately before the defect summary table.

        story.append(Spacer(1, 4 * mm))
        story.append(PageBreak())

        # ══════════════════════════════════════════════════════════════════════
        # GL-16 STYLE NARRATIVE SECTIONS — placed before defect findings.
        # User request: Objective, Scope, Data Collection, Turbine Specs.
        # ══════════════════════════════════════════════════════════════════════
        _company = self._settings.get("company", "") or p.name or "the inspection company"
        _client  = self._settings.get("client", "") or "The Client"
        _model   = self._settings.get("turbine_model", "") or "2 MW capacity"

        # v4.2.0: Increased heading font size from 12→16pt for better prominence
        # MOD-10: spaceAfter on gl_h2 increased 3→8pt so there is visible gap between
        #         the heading text and the blue HR line beneath it.
        #         _gl_hr spaceAfter increased 8mm→12mm for more breathing room before body text.
        gl_h2 = ParagraphStyle(
            "GL16_H2", fontName="Helvetica-Bold", fontSize=16,
            textColor=rl_colors.HexColor("#2b6cb0"),
            spaceBefore=12, spaceAfter=8)
        gl_body = ParagraphStyle(
            "GL16_Body", fontName="Helvetica", fontSize=9,
            leading=14, spaceAfter=7)
        gl_bullet = ParagraphStyle(
            "GL16_Bullet", fontName="Helvetica", fontSize=9,
            leading=13, leftIndent=14, spaceAfter=2)

        def _gl_hr():
            # MOD-10: spaceAfter increased 8mm→12mm for more space between HR and body text
            return HRFlowable(width="100%", thickness=1,
                              color=rl_colors.HexColor("#2b6cb0"), spaceAfter=12)

        # Default narrative text — overridden by user-supplied settings if present
        _def_objective = (
            f"{_client} appointed {_company} to conduct drone-based visual inspection "
            f"of the {_model} Wind Turbine Generator at the site. "
            "Wind Turbine Generators (WTGs) are exposed to unexpected weather conditions "
            "which may affect component performance and efficiency. Drone inspection "
            "provides a visual record of WTG component condition — a cost-effective and "
            "efficient method compared to traditional manual inspection. The objective of "
            "this report is to analyse the visual images of the wind turbine and assess "
            "the condition of all WTG components. Drone-based visual surveys help ensure "
            "a safe working environment, support preventive maintenance planning, reduce "
            "machine breakdowns and erosion, and provide access to otherwise inaccessible "
            "components.")
        _def_scope = (
            "To conduct a visual drone survey of the Wind Turbine Generator commissioned "
            "at the site and identify any visual damage, failures, or erosion of wind "
            "turbine components.")
        _def_data_collection = (
            "During on-site inspection the site team checks weather conditions (wind "
            "speed, visibility, etc.) and flies the drone to the required position for "
            "data capture. The turbine is in stop condition with blades pitched out. "
            "The site team ensures that the drone collects data from 8–10 metres distance "
            "from the WTG.\n\n"
            "Drones are equipped with required sensors and digital cameras to collect "
            "high-quality visual images from the most optimal perspective. Collected data "
            "is scrutinised by the project team; identified anomalies are inspected before "
            "providing a summarised report. The drone operates stably in winds up to "
            "10 m/s.")
        _def_results_intro = (
            "The scan results are presented for each blade and body of the Wind "
            "Turbine in the following sections.")

        # Use user-supplied text if present, else fall back to defaults
        obj_txt   = self._settings.get("objective_text",       "").strip() or _def_objective
        scope_txt = self._settings.get("scope_text",           "").strip() or _def_scope
        data_txt  = self._settings.get("data_collection_text", "").strip() or _def_data_collection
        res_txt   = self._settings.get("results_text",         "").strip() or _def_results_intro

        # Objective
        story.append(Paragraph("Objective", gl_h2))
        story.append(_gl_hr())
        for para_txt in obj_txt.split("\n\n"):   # support multi-paragraph override
            story.append(Paragraph(para_txt.strip(), gl_body))

        # Scope of Work
        story.append(Paragraph("Scope of Work", gl_h2))
        story.append(_gl_hr())
        for para_txt in scope_txt.split("\n\n"):
            story.append(Paragraph(para_txt.strip(), gl_body))

        # Data Collection Methodology
        story.append(Paragraph("Data Collection Methodology & Definitions", gl_h2))
        story.append(_gl_hr())
        for para_txt in data_txt.split("\n\n"):
            story.append(Paragraph(para_txt.strip(), gl_body))

        # Turbine Specifications — all values read from _settings (per-site editable)
        story.append(Paragraph("Turbine Specifications", gl_h2))
        story.append(_gl_hr())
        story.append(Paragraph(
            "<b>Key Specifications:</b>", gl_body))
        _s = self._settings  # shorthand
        
        # v4.2.0: Add Turbine Manufacturer and Rated Power from project metadata at top
        spec_list = []
        if p.turbine_manufacturer:
            spec_list.append(("Turbine Manufacturer", "", p.turbine_manufacturer))
        if p.rated_power:
            spec_list.append(("Rated Power", "", p.rated_power))
        
        # Add standard specifications
        spec_list.extend([
            ("Rated Power",        "rated_power",        "2.0 MW (2,000 kW)"),
            ("Rotor Diameter",     "rotor_diameter",     "114 m"),
            ("Blade Length",       "blade_length",       "55.5 m (3 blades)"),
            ("Swept Area",         "swept_area",         "10,207 m²"),
            ("Wind Class",         "wind_class",         "IEC IIIa (Low Wind)"),
            ("Cut-in Wind Speed",  "cut_in_speed",       "2.5 m/s"),
            ("Rated Wind Speed",   "rated_speed",        "10.0 – 12.5 m/s"),
            ("Cut-out Wind Speed", "cut_out_speed",      "25.0 m/s"),
            ("Tower Height",       "tower_height",       "106 m"),
            ("Max Rotor Speed",    "max_rotor_speed",    "~16 RPM"),
            ("Generator Type",     "generator_type",     "Doubly Fed Induction Generator (DFIG)"),
        ])
        
        for item in spec_list:
            if len(item) == 3 and item[1] == "":  # Project metadata (no settings key)
                _lbl, _, _val = item
                story.append(Paragraph(f"• <b>{_lbl}:</b>  {_val}", gl_bullet))
            else:  # Settings-based spec
                _lbl, _key, _default = item
                # Skip Rated Power from settings if already added from project metadata
                if _lbl == "Rated Power" and p.rated_power:
                    continue
                _val = _s.get(_key, "").strip() or _default
                story.append(Paragraph(f"• <b>{_lbl}:</b>  {_val}", gl_bullet))


        story.append(Spacer(1, 5 * mm))
        story.append(Paragraph(
            "<b>Results</b>",
            ParagraphStyle("ResHdr", fontName="Helvetica-Bold", fontSize=13,
                           textColor=rl_colors.HexColor("#1a202c"), spaceAfter=3)))
        story.append(Paragraph(res_txt, gl_body))
        story.append(PageBreak())
        return story

    # ── Defect Summary Page (Page 2) ───────────────────────────────────────────

    def _build_defect_summary_page(self) -> list:
        """
        SENSEHAWK-STYLE Turbine Index page:
        - Blade diagram at top (SenseHawk blade map)
        - Severity legend row (Severity 1-5 with descriptions + actions)
        - Full defect summary table: Blade | Side | Issues | Severity | Size | MR | Page#
          Sorted: highest severity (5) first, then descending severity, then blade order
        Matches the SenseHawk PDF format exactly.
        """
        if not REPORTLAB_AVAILABLE:
            return []
        from reportlab.platypus import (
            Paragraph, Table, TableStyle, Spacer, HRFlowable,
            Image as RLImage, PageBreak,
        )
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib import colors as rl_colors
        from reportlab.lib.units import mm
        import io, collections

        story = []
        p = self._project

        # ── Turbine Summary header (CHG-E: was "Turbine index") ────────────────
        h1_ps = ParagraphStyle("TI_H1", fontName="Helvetica-Bold", fontSize=11,
                               textColor=rl_colors.HexColor("#111111"), spaceAfter=2)
        small_ps = ParagraphStyle("TI_sm", fontName="Helvetica", fontSize=8,
                                  textColor=rl_colors.HexColor("#555555"), spaceAfter=2)
        # v3.3.13 FIX: Format as "Turbine Summary - WTG-{number}" with dash
        _raw_turb = p.turbine_id or p.name or "Unknown"
        _wtg_id = (f"WTG-{_raw_turb}" if _raw_turb and _raw_turb != "Unknown" and
                   not _raw_turb.upper().startswith("WTG") else _raw_turb)
        story.append(Paragraph(
            f"Turbine Summary - {_wtg_id}", h1_ps))
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=rl_colors.HexColor("#cccccc"), spaceAfter=4))

        # ── Blade diagram (SenseHawk-style blade map image) ───────────────────
        diag_buf = self._render_blade_diagram_to_bytes()
        if diag_buf:
            usable_w = A4[0] - 30 * mm
            # MOD-12: height multiplier increased 0.42→0.60 to respect the figsize(16,9)
            #         natural aspect ratio (9/16=0.5625). Previously 0.42 caused vertical
            #         squashing — blades appeared too compressed in the PDF.
            diag_img = RLImage(diag_buf, width=usable_w, height=usable_w * 0.60)
            story.append(diag_img)
            story.append(Spacer(1, 4 * mm))

        # ── Severity overview + legend (moved from cover — appears after Results) ──
        sev_counts = self._count_by_severity()
        story.append(Paragraph(
            "<b>SEVERITY OVERVIEW</b>",
            ParagraphStyle("SH2", fontName="Helvetica-Bold", fontSize=9,
                           textColor=rl_colors.HexColor("#333333"), spaceAfter=4)
        ))
        sev_ov_keys = [
            ("POI",      "POI"),
            ("Minor",    "Minor"),
            ("Major",    "Major"),
            ("Critical", "Critical"),
        ]
        col_ov = (A4[0] - 30 * mm) / 4
        hdr_ov = [
            Paragraph(f"<b>{disp}</b>",
                      ParagraphStyle(f"ovh{j}", fontName="Helvetica-Bold",
                                     fontSize=8, textColor=rl_colors.white,
                                     alignment=TA_CENTER))
            for j, (_, disp) in enumerate(sev_ov_keys)
        ]
        cnt_ov = [
            Paragraph(f"<b>{sev_counts.get(key, 0)}</b>",
                      ParagraphStyle(f"ovc{j}", fontName="Helvetica-Bold",
                                     fontSize=22, textColor=rl_colors.black,
                                     alignment=TA_CENTER))
            for j, (key, _) in enumerate(sev_ov_keys)
        ]
        sev_ov_tbl = Table([hdr_ov, cnt_ov],
                           colWidths=[col_ov] * 4, rowHeights=[18, 38])
        ov_cmds = [
            ("ALIGN",        (0, 0), (-1, -1), "CENTER"),
            ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
            ("TOPPADDING",   (0, 0), (-1, -1), 4),
            ("GRID",         (0, 0), (-1, -1), 0.5, rl_colors.white),
        ]
        for j, (key, _) in enumerate(sev_ov_keys):
            bg_ov = rl_colors.HexColor(self._SEV_COLORS_HEX.get(key, "#cccccc"))
            ov_cmds.append(("BACKGROUND", (j, 0), (j, 0), bg_ov))
            ov_cmds.append(("LINEBELOW",  (j, 1), (j, 1), 3, bg_ov))
        sev_ov_tbl.setStyle(TableStyle(ov_cmds))
        story.append(sev_ov_tbl)
        story.append(Spacer(1, 4 * mm))

        # ── Severity legend (4-col: POI | Minor | Major | Critical) ──────────
        leg_data = [
            ("POI",      "Point of Interest", SEVERITY_REMEDY.get("POI", "")),
            ("Minor",    "Minor defect",      SEVERITY_REMEDY.get("Minor", "")),
            ("Major",    "Major defect",      SEVERITY_REMEDY.get("Major", "")),
            ("Critical", "Critical",          SEVERITY_REMEDY.get("Critical", "")),
        ]
        leg_col = (A4[0] - 30 * mm) / 4
        leg_hdrs = [
            Paragraph(f"<b>{k}</b>",
                      ParagraphStyle(f"lgh{j}", fontName="Helvetica-Bold",
                                     fontSize=8, textColor=rl_colors.white,
                                     alignment=TA_CENTER))  # v3.3.13 FIX: center header
            for j, (k, _, _a) in enumerate(leg_data)
        ]
        leg_tbl = Table(
            [leg_hdrs,
             [Paragraph(d, ParagraphStyle(f"lgd{j}", fontSize=8, alignment=TA_CENTER))
              for j, (_, d, _) in enumerate(leg_data)],
             [Paragraph(a, ParagraphStyle(f"lga{j}", fontSize=7,
                                          textColor=rl_colors.HexColor("#555555"),
                                          alignment=TA_CENTER))  # v3.3.13 FIX: center action
              for j, (_, _, a) in enumerate(leg_data)]],
            colWidths=[leg_col] * 4,
        )
        leg_cmds = [
            # v3.3.13 FIX: CENTER alignment instead of LEFT for legend table
            ("ALIGN",        (0, 0), (-1, -1), "CENTER"),
            ("VALIGN",       (0, 0), (-1, -1), "TOP"),
            ("GRID",         (0, 0), (-1, -1), 0.3, rl_colors.HexColor("#dddddd")),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
            ("TOPPADDING",   (0, 0), (-1, -1), 4),
        ]
        for j, (k, _, _) in enumerate(leg_data):
            leg_cmds.append(("BACKGROUND", (j, 0), (j, 0),
                             rl_colors.HexColor(_SEV_HEX.get(k, "#cccccc"))))
        leg_tbl.setStyle(TableStyle(leg_cmds))
        story.append(leg_tbl)
        story.append(Spacer(1, 5 * mm))

        # ── Defect summary table (SenseHawk format) ───────────────────────────
        # Collect all annotations with their ImageRecord
        all_pairs: List[Tuple] = []  # (irec, ann, page_num)
        # We assign page numbers in order: sorted by blade/face, 2 per page
        # Build page map first
        page_map: Dict[str, int] = {}  # ann_id → page
        page_counter = 3  # cover=1, turbine_index=2, then per-blade pages start at 3

        # Sort: blade A→C, face PS/LE/TE/SS, then by severity desc
        def _blade_order(b: str) -> int:
            return {"A": 0, "B": 1, "C": 2}.get(b, 9)
        # v4.5.0: Sort by surface then zone
        def _surface_order(s: str) -> int:
            return {"PS": 0, "SS": 1}.get(s, 9)
        def _zone_order(z: str) -> int:
            return {"LE": 0, "TE": 1, "MB": 2}.get(z, 9)
        def _sev_order_desc(s: str) -> int:
            # Higher severity = lower sort number = comes first
            return 10 - SEVERITY_RANK.get(s, 0)

        # v4.5.0: Build flat list of (blade, surface_abbr, zone_abbr, irec, ann) sorted
        all_blade_face_anns = []
        for irec in self._project.images.values():
            for ann in irec.annotations:
                # Extract abbreviations from surface and zone
                surf_abbr = (getattr(ann, 'surface', '').split("(")[-1].strip(")")
                           if "(" in getattr(ann, 'surface', '') else getattr(ann, 'surface', ''))
                zone_abbr = (getattr(ann, 'zone', '').split("(")[-1].strip(")")
                           if "(" in getattr(ann, 'zone', '') else getattr(ann, 'zone', ''))
                # Use dummy face for backward compat with page assignment logic
                all_blade_face_anns.append((ann.blade or "", f"{surf_abbr}-{zone_abbr}", irec, ann))

        all_blade_face_anns.sort(key=lambda x: (
            _blade_order(x[0]), x[1], _sev_order_desc(x[3].severity)))

        # Assign page numbers: each blade+surface-zone section starts a new page, 2 per page
        current_section = None
        ann_count_in_section = 0
        section_page = page_counter
        for blade, face, irec, ann in all_blade_face_anns:
            section = (blade, face)
            if section != current_section:
                current_section = section
                section_page = page_counter
                page_counter += 1
                ann_count_in_section = 0
            if ann_count_in_section > 0 and ann_count_in_section % 2 == 0:
                page_counter += 1
                section_page = page_counter
            page_map[ann.ann_id] = section_page
            ann_count_in_section += 1

        # Now build the table rows grouped by severity (highest first) like SenseHawk
        # SenseHawk shows: Severity 2 rows first, then Severity 1 rows
        # Bug #10 FIX: rename column from "MR" to "Root Dist. (m)" and switch from
        # auto-estimated distance_from_root_mm to user-entered root_distance_m.
        # v4.2.0: Changed "Blade" to "Component", added "Tip Dist. (m)", adjusted column widths,
        # center-aligned all columns for better readability
        # v4.5.0: Split "Side" into "Surface" and "Zone" columns (mutually exclusive selection)
        # MOD-11: Column widths rescaled from total 156mm → 180mm (= A4[0] - 30mm) so the
        #         defect summary table spans the exact same width as the POI/severity tables above.
        # FIX-16d: "Size" → "Size (cm)" so unit is explicit for readers.
        tbl_hdrs = ["Component", "Surface", "Zone", "Issues", "Severity", "Size (cm)", "Root (m)", "Tip (m)", "Page"]
        tbl_col_w = [
            16*mm, 11*mm, 11*mm, 48*mm, 20*mm, 30*mm, 15*mm, 15*mm, 14*mm
        ]  # sum = 180mm, matches usable_w = A4[0] - 30mm
        tbl_rows = [
            [Paragraph(f"<b>{h}</b>", ParagraphStyle(f"th{i}", fontName="Helvetica-Bold",
                        fontSize=7, textColor=rl_colors.white, alignment=TA_CENTER))
             for i, h in enumerate(tbl_hdrs)]
        ]

        sev_hex_map = {
            "1 - Cosmetic":     "#FFD700", "2 - Similar/Cosmetic": "#FFD700",  # v4.2.0: Gold
            "3 - Non-Serious":  "#FFA500", "4 - Serious": "#FFA500",            # v4.2.0: Amber
            "5 - Very Serious": "#FF0000",                                       # v4.2.0: Red
            "Minor":   "#FFD700",  # v4.2.0: Gold
            "Major":   "#FFA500",  # v4.2.0: Amber
            "Critical":"#FF0000",  # v4.2.0: Red
            "Low": "#FFD700", "Medium": "#FFA500",
        }

        row_bgs = []  # (row_index, bg_color) for table style
        for blade, face, irec, ann in all_blade_face_anns:
            # v4.5.0: Extract surface and zone abbreviations
            if blade in ["Hub", "Tower"]:
                component_name = blade
                surf_disp = ""  # Empty for Hub/Tower
                zone_disp = ""
            else:
                component_name = f"Blade {blade}"
                # Extract abbreviations: "Pressure Side (PS)" → "PS"
                surf_disp = (getattr(ann, 'surface', '').split("(")[-1].strip(")")
                           if "(" in getattr(ann, 'surface', '') else getattr(ann, 'surface', ''))
                zone_disp = (getattr(ann, 'zone', '').split("(")[-1].strip(")")
                           if "(" in getattr(ann, 'zone', '') else getattr(ann, 'zone', ''))
            
            # Build size string in cm (like SenseHawk)
            if ann.width_cm and ann.height_cm:
                size_str = f"{ann.width_cm:.1f} × {ann.height_cm:.1f}"
            else:
                size_str = "N/A"
            
            # Bug #10 FIX: use user-entered root_distance_m (meters) instead of
            # auto-estimated distance_from_root_mm (mm).
            root_m_sh = ann.root_distance_m
            root_s_sh = (f"{root_m_sh:.1f}" if root_m_sh else "—")
            
            # v4.2.0: Add tip distance display
            tip_m_sh = ann.tip_distance_m
            tip_s_sh = (f"{tip_m_sh:.1f}" if tip_m_sh else "—")
            
            # Severity short display — v3.0.1: use central map
            sev_num = SEVERITY_SHORT.get(ann.severity or "", "—")
            pg = page_map.get(ann.ann_id, "—")
            
            # v4.5.0: All columns center-aligned, reduced font size to 6pt for better fit
            row = [
                Paragraph(component_name,
                          ParagraphStyle("tr0", fontSize=6, fontName="Helvetica", alignment=TA_CENTER)),
                Paragraph(surf_disp,
                          ParagraphStyle("tr1", fontSize=6, fontName="Helvetica", alignment=TA_CENTER)),
                Paragraph(zone_disp,
                          ParagraphStyle("tr2", fontSize=6, fontName="Helvetica", alignment=TA_CENTER)),
                Paragraph(ann.defect or "—",
                          ParagraphStyle("tr3", fontSize=6, fontName="Helvetica", alignment=TA_CENTER)),
                Paragraph(sev_num,
                          ParagraphStyle("tr4", fontSize=6, fontName="Helvetica-Bold",
                                         alignment=TA_CENTER)),
                Paragraph(size_str,
                          ParagraphStyle("tr5", fontSize=6, fontName="Helvetica", alignment=TA_CENTER)),
                Paragraph(root_s_sh,
                          ParagraphStyle("tr6", fontSize=6, fontName="Helvetica", alignment=TA_CENTER)),
                Paragraph(tip_s_sh,
                          ParagraphStyle("tr7", fontSize=6, fontName="Helvetica", alignment=TA_CENTER)),
                Paragraph(str(pg),
                          ParagraphStyle("tr8", fontSize=6, fontName="Helvetica", alignment=TA_CENTER)),
            ]
            tbl_rows.append(row)
            row_bgs.append((len(tbl_rows) - 1, sev_hex_map.get(ann.severity or "", "#ffffff")))

        sum_tbl = Table(tbl_rows, colWidths=tbl_col_w, repeatRows=1)  # repeatRows: header on every page
        sum_tbl.hAlign = 'CENTER'   # FIX-16g: centre the full-width table on the page
        sum_cmds = [
            ("BACKGROUND",   (0, 0), (-1, 0), rl_colors.HexColor("#333333")),
            ("TEXTCOLOR",    (0, 0), (-1, 0), rl_colors.white),
            ("ALIGN",        (0, 0), (-1, -1), "CENTER"),  # v4.2.0: All columns center-aligned
            ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
            ("FONTSIZE",     (0, 0), (-1, -1), 6),  # v4.2.0: Reduced to fit all columns
            ("TOPPADDING",   (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 3),
            ("GRID",         (0, 0), (-1, -1), 0.3, rl_colors.HexColor("#dddddd")),
            ("ROWBACKGROUNDS",(0, 1), (-1, -1),
             [rl_colors.HexColor("#ffffff"), rl_colors.HexColor("#f9f9f9")]),
        ]
        # Colour the severity column (col 4 - after Surface and Zone split) for each row
        for row_idx, bg_hex in row_bgs:
            sum_cmds.append(("BACKGROUND", (4, row_idx), (4, row_idx),
                             rl_colors.HexColor(bg_hex)))
            sum_cmds.append(("TEXTCOLOR",  (4, row_idx), (4, row_idx),
                             rl_colors.white))
        sum_tbl.setStyle(TableStyle(sum_cmds))
        story.append(sum_tbl)
        story.append(PageBreak())

        # Store page map for use in annotation pages
        self._page_map = page_map
        return story

    def _render_blade_diagram_to_bytes(self) -> Optional[bytes]:
        """
        v1.7.0 REPLACEMENT: Generate the blade diagram PNG using matplotlib.

        Replaces the old QPainter cell-grid version which only coloured span zones and
        had no dot positioning.  This version renders tapered blade silhouettes (same
        as _dc_build_blade_diagram used by the DOCX generator) and places each
        annotation as a coloured scatter dot at the CORRECT Y-position derived from
        ann.distance_from_root_mm.

        Auto-estimate fallback: if distance_from_root_mm is None on an annotation,
        the dot is drawn at 50% (mid-blade) with reduced opacity to signal uncertainty.

        Returns PNG bytes suitable for a ReportLab RLImage, or None on failure.
        """
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as _plt
            import matplotlib.patches as _mpatches
            import numpy as _np
            import io as _io

            blade_length_mm = getattr(self._project, "blade_length_mm", 50_000.0) or 50_000.0
            FACES  = ["PS", "LE", "TE", "SS"]
            BLADES = ["A", "B", "C"]
            SEV_COL = _SEV_HEX   # v3.0.1: use central map

            # MOD-12: col_w increased 1.0→1.6 so face columns are wider and blade
            #         silhouettes are not cramped. gap increased 0.5→1.2 for clear
            #         separation between the three blade groups.
            #         figsize widened to (16, 9) to accommodate the extra space.
            col_w   = 1.6
            gap     = 1.2
            group_w = len(FACES) * col_w + gap
            fig_w   = len(BLADES) * group_w - gap

            fig, ax = _plt.subplots(figsize=(16, 9))
            fig.patch.set_facecolor("#f8f9fa")
            ax.set_facecolor("#f8f9fa")

            y_vals = _np.linspace(0, 1, 200)

            def _hw(y):
                """Half-width of blade cross-section — tapered root→tip."""
                return col_w * 0.44 * _np.sqrt(_np.maximum(0.0, 1.0 - y))

            for bi, bkey in enumerate(BLADES):
                gx = bi * group_w
                for si, face in enumerate(FACES):
                    cx = gx + si * col_w + col_w / 2
                    hw = _hw(y_vals)
                    ax.fill_betweenx(y_vals, cx - hw, cx + hw,
                                     color="#dde3ea", alpha=0.85, zorder=1)
                    ax.plot(cx - hw, y_vals, color="#aab4be", lw=0.6, zorder=2)
                    ax.plot(cx + hw, y_vals, color="#aab4be", lw=0.6, zorder=2)
                    ax.text(cx, -0.06, face, ha="center", va="bottom",
                            fontsize=11, fontweight="bold", color="#333")

                # Blade group label + serial number
                gcx    = gx + len(FACES) * col_w / 2
                serial = self._project.blade_numbers.get(bkey, "")
                lbl    = f"Blade {bkey}" + (f" ({serial})" if serial else "")
                ax.text(gcx, -0.14, lbl, ha="center", va="bottom",
                        fontsize=12, fontweight="bold", color="#1a1a2e")

                # ── Plot annotation dots ────────────────────────────────────────
                # ROOT-CAUSE FIX: was filtering by irec.blade (folder auto-detected,
                # defaults to "A") which caused ALL annotations from any image whose
                # irec.blade="A" to plot in the Blade A column — regardless of what
                # the user selected in the annotation panel (ann.blade).
                # Fix: use ann.blade as the authoritative blade assignment.
                # irec.blade is only used as a fallback when ann.blade is empty.
                for irec in self._project.images.values():
                    for ann in irec.annotations:
                        _ann_blade = (ann.blade or irec.blade or "").strip()
                        if _ann_blade != bkey:
                            continue
                        face_raw = ann.face
                        # Extract abbreviation from "Leading Edge (LE)" → "LE"
                        if "(" in face_raw:
                            face_abbr = face_raw.split("(")[-1].strip(")")
                        else:
                            face_abbr = face_raw
                        if face_abbr not in FACES:
                            continue
                        si  = FACES.index(face_abbr)
                        cx  = gx + si * col_w + col_w / 2
                        # Bug #12 FIX: use user-entered root_distance_m (meters →
                        # convert to mm for blade_length_mm comparison).  The old
                        # distance_from_root_mm was auto-estimated from Y-fraction
                        # and has been removed; only root_distance_m is reliable.
                        rm = ann.root_distance_m
                        if rm is not None and rm > 0:
                            dist = rm * 1000.0   # m → mm for normalisation
                            yn      = min(1.0, dist / blade_length_mm)
                            alpha   = 1.0
                            edge_c  = "white"
                        else:
                            # Estimate from span string ("Root"→15%, "Mid"→50%, "Tip"→80%)
                            span_s = (ann.span or "").split(" ")[0].lower()
                            yn     = {"root": 0.15, "mid": 0.50,
                                      "tip": 0.80}.get(span_s, 0.50)
                            alpha  = 0.45   # visual cue: uncertain position
                            edge_c = "#888888"
                        clr = SEV_COL.get(ann.severity, "#F97316")
                        ax.scatter(cx, yn, s=70, color=clr, edgecolors=edge_c,
                                   linewidths=0.8, alpha=alpha, zorder=5)

            # MOD-12: ylim top extended -0.14→-0.24 so face labels (y=-0.06) and
            #         blade group labels (y=-0.14) are fully visible without clipping.
            ax.set_ylim(1.05, -0.24)
            blade_len_km  = blade_length_mm / 1_000
            tick_step_m   = 10_000 if blade_length_mm >= 30_000 else 5_000
            ticks_mm      = list(range(0, int(blade_length_mm) + 1, tick_step_m))
            ax.set_yticks([t / blade_length_mm for t in ticks_mm])
            ax.set_yticklabels([f"{int(t/1000)}m" for t in ticks_mm],
                                fontsize=10, color="#555")
            ax.yaxis.set_tick_params(length=0)
            ax.set_xlim(-col_w * 0.6, fig_w - gap + col_w * 0.6)
            ax.set_xticks([])
            ax.spines[["top", "right", "bottom"]].set_visible(False)
            ax.spines["left"].set_color("#ccc")

            # v4.3.1 MOD-7: Legend removed per user request - blade diagram now shows
            # only the blade silhouettes with severity-colored dots, no legend box
            ax.set_title("Blade Defect Location Diagram", fontsize=12,
                         fontweight="bold", color="#1a1a2e", pad=18)
            fig.tight_layout(rect=[0, 0.02, 1, 1])  # Reduced bottom margin since no legend

            buf = _io.BytesIO()
            fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
            _plt.close(fig)
            buf.seek(0)
            return buf

        except Exception as exc:
            log.warning(f"_render_blade_diagram_to_bytes: {exc}")
            return None

    # ── Component-based annotation pages ──────────────────────────────────────
    # Phase 8.7 complete: dimension labels on images
    # NEW: group by component (blade/hub/tower), 6 images per grid page

    def _build_annotation_pages(self) -> list:
        """
        v3.3.2: Two defects per page layout.
        T04: Global defect numbering (1,2,3…) across A→B→C→Hub→Tower.
        T08: No empty pages — PageBreak inserted between components only when needed.
        FIX-10: Components sorted A→B→C→Hub→Tower using ann.blade (authoritative)
                as the grouping key; images within each component sorted by filename.
        """
        from io import BytesIO
        story = []

        # ── Build global sorted flat list ───────────────────────────────────────
        # T04: Pre-build the complete sorted list so each defect gets a
        # globally-sequential number (1, 2, 3 … N) ordered by component.
        def _comp_sort_key(b: str) -> int:
            return {"A": 0, "B": 1, "C": 2, "Hub": 3, "Tower": 4}.get(b, 10)

        # FIX-10 Root Cause 1: group by ann.blade (user-confirmed, authoritative)
        # with fallback to irec.blade then "Unknown".  The old irec.blade-only
        # grouping put images in the wrong section when irec.blade was stale
        # (e.g. images loaded via "Select Individual Files" keep irec.blade="A"
        # even after the annotation was saved with ann.blade="B").
        #
        # FIX-10 Root Cause 2: sort pairs within each component by filename so
        # the defect sequence is deterministic regardless of folder-load order.
        pair_comps: "Dict[str, list]" = {}   # comp_name → [(irec, ann), …]
        for irec in self._project.images.values():
            for ann in irec.annotations:
                comp = (ann.blade or irec.blade or "Unknown").strip()
                pair_comps.setdefault(comp, []).append((irec, ann))

        # Stable filename → ann_id sort within each component
        for _pairs in pair_comps.values():
            _pairs.sort(key=lambda t: (t[0].filename or "", t[1].ann_id or ""))

        sorted_comps = sorted(pair_comps.items(),
                              key=lambda kv: _comp_sort_key(kv[0]))

        # Flat global list of (comp_name, irec, ann) — used for T04 global numbering
        global_list: list = []
        for comp_name, pairs in sorted_comps:
            for irec, ann in pairs:
                global_list.append((comp_name, irec, ann))

        if not global_list:
            return story

        # ── Render component sections with correct page-break placement ─────────
        prev_comp = None

        for comp_idx, (comp_name, pairs) in enumerate(sorted_comps):
            if not pairs:
                continue

            # T08: add PageBreak BEFORE component section (except first component)
            if comp_idx > 0 and prev_comp is not None:
                story.append(PageBreak())

            # MOD-9 (revised): The company/WTG/date header row is REMOVED.
            # Only the chip breadcrumb row [WTG-X] [Blade X] is rendered above each
            # component section, giving context without the full inspector header.
            _raw_wtg_bc = (self._project.turbine_id or self._project.name or "WTG")
            _wtg_id_bc  = (_raw_wtg_bc if _raw_wtg_bc.upper().startswith("WTG")
                           else f"WTG-{_raw_wtg_bc}")
            _comp_labels_bc = {"A": "Blade A", "B": "Blade B", "C": "Blade C",
                               "Hub": "Hub", "Tower": "Tower"}
            _comp_disp_bc = _comp_labels_bc.get(comp_name, comp_name)
            _chips_bc = [_wtg_id_bc,
                         f"Blade {comp_name}" if comp_name in ("A", "B", "C")
                         else _comp_disp_bc]
            _chip_cells_bc = [
                Paragraph(
                    f"<b>{ch}</b>",
                    ParagraphStyle(
                        # FIX-16e: fontSize 8→11, backColor #f0f0f0→#DBEAFE (light blue),
                        # textColor #333333→#1E40AF (strong blue) for visual prominence.
                        f"bc_chip_{i}", fontName="Helvetica-Bold", fontSize=11,
                        textColor=rl_colors.HexColor("#1E40AF"),
                        backColor=rl_colors.HexColor("#DBEAFE"),
                        borderColor=rl_colors.HexColor("#93C5FD"),
                        borderWidth=0.5, borderRadius=4,
                        borderPadding=(4, 8, 4, 8)))
                for i, ch in enumerate(_chips_bc)
            ]
            _cw_bc = (A4[0] - 30 * mm) / len(_chip_cells_bc)
            _chip_tbl_bc = Table([_chip_cells_bc],
                                 colWidths=[_cw_bc] * len(_chip_cells_bc))
            _chip_tbl_bc.setStyle(TableStyle([
                ("ALIGN",  (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))
            story.append(_chip_tbl_bc)
            story.append(Spacer(1, 4 * mm))

            # Global start index for this component's defects (T04 sequential numbers)
            global_start = next(
                i for i, (cn, _, _) in enumerate(global_list) if cn == comp_name)

            # FIX-10 Root Cause 3: pass pre-sorted pairs directly so
            # _build_image_grid_pages does not re-build them from irecs (which
            # would discard the filename sort applied above).
            story += self._build_image_grid_pages(
                comp_name, irecs=[],
                global_list=global_list,
                global_start=global_start,
                sorted_pairs=pairs)

            prev_comp = comp_name

        return story


    def _build_component_header(self, comp_name: str, irecs: list) -> list:
        """
        SENSEHAWK-STYLE section header:
        White background, small header bar: "Company | WTG name | Date"
        T06: Chip row shows only [WTG-X] [Blade X] — faces removed.
        T08: NO PageBreak at start — caller (_build_annotation_pages) manages breaks.
        """
        from io import BytesIO
        # T08 FIX: PageBreak removed here; _build_annotation_pages inserts it before
        # each component header (except the first) to avoid empty pages between components.
        story = []
        p = self._project

        # Company / WTG / Date header row — light theme, like SenseHawk
        hdr_ps = ParagraphStyle(
            "SHHdr", fontName="Helvetica", fontSize=8,
            textColor=rl_colors.HexColor("#555555"),
        )
        date_str = datetime.now().strftime("%b %d, %Y")
        company  = self._settings.get("company", "") or p.name or "—"
        turb_id  = p.turbine_id or p.name or "—"
        # Bug #4 FIX: build prefixed WTG ID once; use everywhere in this header
        # so the centre cell and breadcrumb chip are consistent with grid pages.
        wtg_id   = turb_id if turb_id.upper().startswith("WTG") else f"WTG-{turb_id}"

        # BLADE-ABC: comp_labels
        comp_labels = {"A": "Blade A", "B": "Blade B", "C": "Blade C",
                       "Hub": "Hub", "Tower": "Tower"}
        comp_display = comp_labels.get(comp_name, comp_name)

        # Header table: left=company, centre=WTG name, right=date
        hdr_tbl = Table(
            [[Paragraph(company, hdr_ps),
              Paragraph(f"<b>{wtg_id}</b>", ParagraphStyle(
                  "SHHdrC", fontName="Helvetica-Bold", fontSize=8,
                  textColor=rl_colors.HexColor("#222222"),
                  alignment=TA_CENTER)),
              Paragraph(date_str, ParagraphStyle(
                  "SHHdrR", fontName="Helvetica", fontSize=8,
                  textColor=rl_colors.HexColor("#555555"),
                  alignment=2))]],   # 2=RIGHT
            colWidths=[60*mm, 80*mm, 40*mm]
        )
        hdr_tbl.setStyle(TableStyle([
            ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
            ("LINEBELOW",    (0, 0), (-1, 0),  0.5, rl_colors.HexColor("#cccccc")),
            ("TOPPADDING",   (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 3),
        ]))
        story.append(hdr_tbl)
        story.append(Spacer(1, 4*mm))

        # T06 FIX: chip row shows only [WTG-X] and [Blade X] — faces removed.
        # Faces were redundant visual clutter; per-defect chips (T04) use same two pills.
        chips = [wtg_id]
        if comp_name in ("A", "B", "C"):
            chips.append(f"Blade {comp_name}")
        else:
            chips.append(comp_display)

        chip_cells = []
        for chip in chips:
            chip_cells.append(Paragraph(
                f"<b>{chip}</b>",
                # FIX-16e: fontSize 8→11, backColor #f0f0f0→#DBEAFE, textColor →#1E40AF
                ParagraphStyle("chip", fontName="Helvetica-Bold", fontSize=11,
                               textColor=rl_colors.HexColor("#1E40AF"),
                               backColor=rl_colors.HexColor("#DBEAFE"),
                               borderColor=rl_colors.HexColor("#93C5FD"),
                               borderWidth=0.5, borderRadius=4,
                               borderPadding=(4, 8, 4, 8))
            ))
        cw = (A4[0] - 30*mm) / max(len(chip_cells), 1)
        chip_tbl = Table([chip_cells], colWidths=[cw]*len(chip_cells))
        chip_tbl.setStyle(TableStyle([
            ("ALIGN",  (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(chip_tbl)
        story.append(Spacer(1, 6*mm))
        return story

    def _build_image_grid_pages(self, comp_name: str, irecs: list,
                                global_list: list = None,
                                global_start: int = 0,
                                sorted_pairs: list = None) -> list:
        """
        v3.3.2: Two defects per page layout.
        T04: Per-defect chip [WTG-X | Blade X] with global sequential number.
        T05: Sub-header serial line removed (header already contains serial).
        T07: Pill table and cr_tbl widths match the image row width exactly.
        T08: Page breaks only between defect pairs; no trailing break on last defect.
        FIX-10: sorted_pairs kwarg — when provided by _build_annotation_pages the
                pre-sorted [(irec, ann), …] list is used directly so the filename
                sort and ann.blade-based grouping are preserved.  Falls back to
                rebuilding from irecs when sorted_pairs is None (legacy callers).
        """
        from io import BytesIO
        story = []

        # FIX-10 Root Cause 3: use pre-sorted pairs when supplied; fall back to
        # rebuilding from irecs for any legacy call-site that omits sorted_pairs.
        if sorted_pairs is not None:
            all_pairs = sorted_pairs
        else:
            all_pairs = [(irec, ann)
                         for irec in irecs for ann in irec.annotations]
        if not all_pairs:
            return story

        usable_w = A4[0] - 30 * mm   # 180 mm
        # IMG_H reduced 85mm → 65mm so that 2 KeepTogether defect blocks fit on
        # one A4 page without ReportLab splitting them.
        # Budget: A4 usable height 262mm − 2×57mm fixed overhead − 9mm separator
        # = 139mm for 2 image rows → 69mm each; 65mm gives 9mm safety margin.
        WIDE_W   = 98 * mm            # wide-shot width
        ZOOM_W   = 55 * mm            # zoom-crop width
        MINI_W   = 23 * mm            # mini blade diagram width
        IMG_H    = 65 * mm            # reduced from 85mm to allow 2 defects per page

        # T07 FIX: pill table / cr_tbl width must match the image row above
        BLADE_ROW_W = WIDE_W + ZOOM_W + MINI_W   # 176 mm for blade pages
        HUB_ROW_W   = usable_w                    # 180 mm for Hub/Tower pages

        _raw_wtg = (self._project.turbine_id or self._project.name or "WTG")
        _wtg_id  = (_raw_wtg if _raw_wtg.upper().startswith("WTG")
                    else f"WTG-{_raw_wtg}")
        comp_labels = {"A": "Blade A", "B": "Blade B", "C": "Blade C",
                       "Hub": "Hub", "Tower": "Tower"}
        comp_label = comp_labels.get(comp_name, comp_name)

        for pair_idx, (irec, ann) in enumerate(all_pairs):
            face_abbr = (ann.face.split("(")[-1].strip(")")
                         if "(" in ann.face else ann.face)
            _is_blade_c = comp_name in ("A", "B", "C")
            # Bug D fix: check ann.blade not comp_name — Hub images may load with blade=""
            is_hub_tower = (ann.blade or "").strip() not in ("A", "B", "C")

            # T04: global defect number (1-based, across all components)
            global_num = global_start + pair_idx + 1 if global_list else pair_idx + 1

            # Distances and size
            root_m = ann.root_distance_m
            tip_m  = ann.tip_distance_m
            size_s = (f"{ann.width_cm:.1f} × {ann.height_cm:.1f} cm"
                      if ann.width_cm else "—")
            sev_c_hex = _SEV_HEX.get(ann.severity, "#cccccc")
            sev_num   = SEVERITY_SHORT.get(ann.severity or "", "—")

            serial_str   = ann.serial_number or _generate_defect_serial(
                               self._project, ann.blade, ann.surface, ann.zone)
            # v3.3.5 DOC-7: heading_str replaced by 2-col [Defect Name | Serial] table below
            # MODIFICATION: Defect Name table removed per user request

            _blk = []

            # FIX-16g: All tables on this page use a single canonical width = usable_w
            # (A4[0]-30mm = 180mm) so blade pages (previously 176mm) and Hub/Tower
            # pages (previously 180mm) are now identical width.  hAlign='CENTER' on
            # each table guarantees left-margin alignment is centred on the page.
            _tbl_w_pre = usable_w   # was BLADE_ROW_W (176mm) for blades, HUB_ROW_W (180mm) for hub

            # FIX-BUG4: Defect name table — background now matches Issue Type header
            # row (#f5f5f5 = same as pill_tbl header row).  Font / colour for BOTH
            # columns matches Comments/Remedy body text: Helvetica 8pt #222222.
            _dn_ps_lbl = ParagraphStyle("dnl", fontName="Helvetica-Bold", fontSize=8,
                                         textColor=rl_colors.HexColor("#222222"),
                                         alignment=TA_LEFT)
            _dn_ps_val = ParagraphStyle("dnv", fontName="Helvetica", fontSize=8,
                                         textColor=rl_colors.HexColor("#222222"),
                                         alignment=TA_LEFT)
            _dn_tbl_pdf = Table(
                [[Paragraph("Defect Name", _dn_ps_lbl),
                  Paragraph(serial_str,    _dn_ps_val)]],
                colWidths=[_tbl_w_pre * 0.25, _tbl_w_pre * 0.75],
                rowHeights=[16],
            )
            _dn_tbl_pdf.hAlign = 'CENTER'
            _dn_tbl_pdf.setStyle(TableStyle([
                ("BACKGROUND",   (0, 0), (-1, -1), rl_colors.HexColor("#f5f5f5")),  # matches pill header
                ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN",        (0, 0), (-1, -1), "LEFT"),
                ("TOPPADDING",   (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
                ("LEFTPADDING",  (0, 0), (-1, -1), 6),
                ("GRID",         (0, 0), (-1, -1), 0.3, rl_colors.HexColor("#dddddd")),
            ]))
            _blk.append(_dn_tbl_pdf)
            _blk.append(Spacer(1, 5 * mm))

            # ── 1. Wide annotated image ─────────────────────────────────────────
            wide_img = self._render_wide_image(irec, ann, WIDE_W, IMG_H)

            # ── 2. Zoomed defect crop ───────────────────────────────────────────
            zoom_img = self._render_zoom_crop(irec, ann, ZOOM_W, IMG_H)

            # ── 3. Mini blade pinpoint diagram — skip for Hub/Tower ─────────────
            mini_diag = None
            if not is_hub_tower:
                blade_pos = getattr(ann, "pinpoint_blade_pos", None)
                if blade_pos is None:
                    bl_mm = getattr(self._project, "blade_length_mm", 50_000.0) or 50_000.0
                    blade_pos = (min(1.0, ann.root_distance_m * 1000 / bl_mm)
                                 if ann.root_distance_m else 0.5)
                dist_for_diag = blade_pos * (
                    getattr(self._project, "blade_length_mm", 50_000.0) or 50_000.0)
                # v4.3.0: Pass edge_side so the mini diagram shows the edge highlight
                _ann_edge = getattr(ann, "edge_side", None)
                mini_diag = self._render_mini_blade_diagram(
                    ann.blade or irec.blade or "A", face_abbr, dist_for_diag,
                    ann.severity, MINI_W, IMG_H, edge_side=_ann_edge)

            # ── 4. Image row assembly ───────────────────────────────────────────
            # FIX-BUG5: Blade wide-image column widened so img_row total = usable_w
            # (180mm), matching pill_tbl and cr_tbl exactly.  Previously WIDE_W=98mm
            # + ZOOM_W=55mm + MINI_W=23mm = 176mm while other tables were 180mm;
            # 4mm mismatch caused the wide image's left edge to be offset 2mm right
            # when all tables use hAlign='CENTER'.
            # FIX-16g: _tbl_w = usable_w always so all tables below (pill, cr_tbl)
            # are the same width regardless of blade vs hub/tower pages.
            _WIDE_W_BLD = usable_w - ZOOM_W - MINI_W  # 180-55-23 = 102mm (was 98mm)
            if is_hub_tower or mini_diag is None:
                WIDE_W2 = usable_w * 0.60
                ZOOM_W2 = usable_w * 0.40
                img_row = Table([[wide_img, zoom_img]],
                                colWidths=[WIDE_W2, ZOOM_W2],
                                rowHeights=[IMG_H])
            else:
                img_row = Table([[wide_img, zoom_img, mini_diag]],
                                colWidths=[_WIDE_W_BLD, ZOOM_W, MINI_W],  # total = 180mm
                                rowHeights=[IMG_H])
            _tbl_w = usable_w  # FIX-16g: canonical width for pill + cr_tbl (was BLADE_ROW_W / HUB_ROW_W)
            img_row.hAlign = 'CENTER'  # FIX-16g
            img_row.setStyle(TableStyle([
                # v4.3.0: MIDDLE so the (usually shorter) zoom crop is vertically centred
                ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN",        (0, 0), (-1, -1), "CENTER"),
                ("LEFTPADDING",  (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING",   (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING",(0, 0), (-1, -1), 0),
            ]))
            _blk.append(img_row)
            # CHG-B: 5 mm spacer between image and pill table (one visible line gap)
            _blk.append(Spacer(1, 5 * mm))

            # ── 5. Severity / Issue / Root Dist / Tip Dist / Size pill row ───────
            # FIX-16h revised: Issue-Type 28%→24%, Severity 16%→12%, Size 24%→32%.
            # Root/Tip stay at 16% each.  Total = 24+12+16+16+32 = 100%.
            # Size at 32% (57.6mm) gives ~51mm content — holds "205.59 × 246.71 cm"
            # (≈32mm at 8pt Helvetica) on one line without wrapping.
            # FIX-16g: _tbl_w = usable_w (set above) so pill_tbl = same 180mm as all others.
            _p24i = _tbl_w * 0.24  # Issue Type  (was 0.28)
            _p12  = _tbl_w * 0.12  # Severity    (was 0.16)
            _p16  = _tbl_w * 0.16  # Root / Tip  (unchanged)
            _p32  = _tbl_w * 0.32  # Size        (was 0.24)
            sev_bg   = rl_colors.HexColor(sev_c_hex)
            # v3.3.13 FIX: Add alignment=TA_CENTER to both label and value styles
            ps_val   = lambda n: ParagraphStyle(
                n, fontName="Helvetica", fontSize=8,
                textColor=rl_colors.HexColor("#222222"), alignment=TA_CENTER)
            # v4.2.0: Made label style bold to match Comments heading
            ps_lbl   = lambda n: ParagraphStyle(
                n, fontName="Helvetica-Bold", fontSize=7,
                textColor=rl_colors.HexColor("#888888"), alignment=TA_CENTER)

            _root_s = (f"{root_m:.1f} m" if (_is_blade_c and root_m) else "—")
            _tip_s  = (f"{tip_m:.1f} m"  if (_is_blade_c and tip_m)  else "—")

            pill_tbl = Table(
                [
                    [Paragraph("Issue Type",                           ps_lbl("pl0")),
                     Paragraph("Severity",                             ps_lbl("pl1")),
                     Paragraph("Root Dist." if _is_blade_c else "—",  ps_lbl("pl2")),
                     Paragraph("Tip Dist."  if _is_blade_c else "—",  ps_lbl("pl3")),
                     Paragraph("Size",                                 ps_lbl("pl4"))],
                    [Paragraph(ann.defect or "—",     ps_val("pv0")),
                     Paragraph(f"  {sev_num}  ",
                               ParagraphStyle("psev", fontName="Helvetica-Bold",
                                              fontSize=10, textColor=rl_colors.white,
                                              alignment=TA_CENTER)),
                     Paragraph(_root_s,               ps_val("pv2")),
                     Paragraph(_tip_s,                ps_val("pv3")),
                     Paragraph(size_s,                ps_val("pv4"))],
                ],
                colWidths=[_p24i, _p12, _p16, _p16, _p32],  # 24/12/16/16/32 % = 100%
                rowHeights=[12, 18],
            )
            pill_tbl.hAlign = 'CENTER'  # FIX-16g
            pill_tbl.setStyle(TableStyle([
                ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
                # v3.3.13 FIX: CENTER all cells (both label and value rows)
                # Paragraph styles now have alignment=TA_CENTER for all cells
                ("ALIGN",        (0, 0), (-1, -1), "CENTER"),
                ("TOPPADDING",   (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING",(0, 0), (-1, -1), 3),
                ("LEFTPADDING",  (0, 0), (-1, -1), 6),
                ("BACKGROUND",   (1, 1), (1, 1),   sev_bg),
                ("GRID",         (0, 0), (-1, -1), 0.3, rl_colors.HexColor("#dddddd")),
                ("BACKGROUND",   (0, 0), (-1, 0),  rl_colors.HexColor("#f5f5f5")),
            ]))
            _blk.append(pill_tbl)
            # CHG-H: NO spacer between pill table and comments/remedy table
            # (tables are directly adjacent per user spec)

            # ── 6. Comments + Remedy Action side-by-side ─────────────────────────
            # T07 FIX: cr_tbl width = same _tbl_w as pill table above
            # CHG-D: Remedy Action = severity action from reference document legend
            # (SEVERITY_REMEDY updated to match: Monitoring/Intervention/Immediate)
            # v3.3.13 FIX: Always prefer SEVERITY_REMEDY based on severity.
            # Ignore stored remedy if it matches ANY auto-generated pattern.
            remedy = SEVERITY_REMEDY.get(ann.severity or "",
                     _auto_remedy(ann.defect or ""))
            if ann.remedy_action and ann.remedy_action.strip():
                stored = ann.remedy_action.strip()
                # Check if stored text matches auto-remedy pattern (standard suffix)
                is_auto_pattern = stored.endswith("repair recommended during the next planned inspection.")
                # Check if it matches any SEVERITY_REMEDY value
                is_severity_text = stored in SEVERITY_REMEDY.values()
                # Only use stored text if it's truly custom (not auto-generated)
                if not is_auto_pattern and not is_severity_text:
                    remedy = stored
            notes  = ann.notes or "—"
            cr_ps  = ParagraphStyle("crp", fontName="Helvetica", fontSize=8, leading=11)
            cr_hdr = ParagraphStyle("crh", fontName="Helvetica-Bold", fontSize=8)
            _half  = _tbl_w / 2
            cr_tbl = Table(
                [[Paragraph("■  Comments",      cr_hdr),
                  Paragraph("■  Remedy Action", cr_hdr)],
                 [Paragraph(notes,  cr_ps),
                  Paragraph(remedy, cr_ps)]],
                colWidths=[_half, _half],
            )
            cr_tbl.hAlign = 'CENTER'  # FIX-16g
            cr_tbl.setStyle(TableStyle([
                ("GRID",         (0, 0), (-1, -1), 0.3, rl_colors.HexColor("#dddddd")),
                ("TOPPADDING",   (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
                ("LEFTPADDING",  (0, 0), (-1, -1), 6),
                ("BACKGROUND",   (0, 0), (-1, 0),  rl_colors.HexColor("#f5f5f5")),
                ("VALIGN",       (0, 0), (-1, -1), "TOP"),
            ]))
            _blk.append(cr_tbl)

            # v4.2.0: REMOVED calibration metadata display per user request

            # KeepTogether ensures each defect block (defect name table + images +
            # pill table + comments) is never split across pages. IMG_H was reduced
            # from 85mm to 65mm so that exactly 2 KeepTogether blocks fit per A4 page.
            story.append(KeepTogether(_blk))

            # ── T08 FIX: Two defects per page — clean page-break logic ──────────
            # Rules:
            #   • After defect #2, 4, 6… (even position) AND more defects follow in
            #     this component → insert PageBreak
            #   • After defect #1, 3, 5… (odd position, not last) → thin HR separator
            #   • After last defect of component → nothing (caller handles inter-comp break)
            is_last_in_comp = (pair_idx + 1 == len(all_pairs))
            is_even_pos     = ((pair_idx + 1) % 2 == 0)

            if not is_last_in_comp:
                if is_even_pos:
                    # T08: End of 2-defect block → hard page break within component
                    story.append(PageBreak())
                else:
                    # Between two defects on the same page → thin separator
                    story.append(Spacer(1, 3 * mm))
                    story.append(HRFlowable(
                        width="100%", thickness=0.5,
                        color=rl_colors.HexColor("#cccccc"), spaceAfter=3 * mm))
            # is_last_in_comp: NO trailing PageBreak — _build_annotation_pages
            # inserts the inter-component PageBreak before the NEXT component header.

        return story

    # ── Image rendering helpers ────────────────────────────────────────────────

    def _resolve_fp(self, irec: "ImageRecord") -> "Optional[str]":
        """Resolve filepath with fallback to project_folder subdirectory scan."""
        fp = irec.filepath
        if fp and os.path.exists(fp):
            return fp
        pf = getattr(self._project, "project_folder", "")
        fname = irec.filename or (os.path.basename(fp) if fp else "")
        if pf and fname:
            candidate = os.path.join(pf, fname)
            if os.path.exists(candidate):
                return candidate
            try:
                for entry in os.scandir(pf):
                    if entry.is_dir():
                        c2 = os.path.join(entry.path, fname)
                        if os.path.exists(c2):
                            return c2
            except Exception:
                pass
        return None

    def _render_wide_image(self, irec: "ImageRecord", ann: "Annotation",
                           max_w: float, max_h: float):
        """Full annotated image with all annotation boxes burned in."""
        fp = self._resolve_fp(irec)
        if not fp:
            return Spacer(max_w, max_h * 0.5)
        try:
            from io import BytesIO as _BIO
            with Image.open(fp) as src:
                img = src.convert("RGB")
            iw, ih = img.size
            draw = ImageDraw.Draw(img, "RGBA")
            try:
                font_sm = ImageFont.truetype("arial.ttf", max(10, iw // 120))
            except Exception:
                font_sm = ImageFont.load_default()

            for a in irec.annotations:
                sev_h = self._SEV_COLORS_HEX.get(a.severity, "#FFC107")
                r_, g_, b_ = int(sev_h[1:3], 16), int(sev_h[3:5], 16), int(sev_h[5:7], 16)
                outline = (r_, g_, b_, 255); fill_t = (r_, g_, b_, 55)
                if a.mode == "box":
                    x1 = min(a.x1_px, a.x2_px); y1 = min(a.y1_px, a.y2_px)
                    x2 = max(a.x1_px, a.x2_px); y2 = max(a.y1_px, a.y2_px)
                    rot = getattr(a, "rotation_deg", 0.0) or 0.0
                    if abs(rot) > 0.5:
                        # Rotated box: draw as filled polygon matching canvas appearance
                        corners = _rotated_box_corners(x1, y1, x2, y2, rot)
                        draw.polygon(corners, outline=outline, fill=fill_t)
                        # Draw border explicitly (polygon fill doesn't include width)
                        draw.line(corners + [corners[0]], fill=outline, width=3)
                    else:
                        draw.rectangle([x1, y1, x2, y2], outline=outline, fill=fill_t, width=3)
                    lbl = (f"{a.width_cm:.1f}×{a.height_cm:.1f}cm" if a.width_cm
                           else a.defect or "")
                    if lbl:
                        try:
                            bb = font_sm.getbbox(lbl); lw, lh = bb[2]-bb[0]+6, bb[3]-bb[1]+4
                        except Exception:
                            lw, lh = len(lbl)*7, 14
                        ly = max(0, y1 - lh - 2)
                        draw.rectangle([x1, ly, x1+lw, ly+lh], fill=(r_, g_, b_, 200))
                        draw.text((x1+3, ly+2), lbl, fill=(255, 255, 255, 255), font=font_sm)
                elif a.mode == "pin":
                    cx, cy, rp = int(a.x1_px), int(a.y1_px), 14
                    draw.ellipse([cx-rp, cy-rp, cx+rp, cy+rp],
                                 outline=outline, fill=(r_, g_, b_, 160), width=3)
                elif a.mode == "polygon" and len(a.poly_pts) >= 6:
                    pts = [(a.poly_pts[i], a.poly_pts[i+1])
                           for i in range(0, len(a.poly_pts) - 1, 2)]
                    draw.polygon(pts, fill=fill_t, outline=outline)

            scale = min(max_w / iw, max_h / ih)
            buf   = _BIO()
            img.save(buf, "JPEG", quality=84)
            buf.seek(0)
            return RLImage(buf, width=int(iw*scale), height=int(ih*scale))
        except Exception as exc:
            log.warning(f"_render_wide_image {fp}: {exc}")
            return Spacer(max_w, max_h * 0.5)

    def _render_zoom_crop(self, irec: "ImageRecord", ann: "Annotation",
                          max_w: float, max_h: float):
        """Zoom crop of just the defect bbox with 20% padding, for box/polygon anns."""
        fp = self._resolve_fp(irec)
        if not fp:
            return Spacer(max_w, max_h * 0.5)
        try:
            from io import BytesIO as _BIO
            with Image.open(fp) as src:
                img = src.convert("RGB")
            iw, ih = img.size

            # Determine crop region from annotation mode
            if ann.mode == "box":
                x1 = min(ann.x1_px, ann.x2_px); y1 = min(ann.y1_px, ann.y2_px)
                x2 = max(ann.x1_px, ann.x2_px); y2 = max(ann.y1_px, ann.y2_px)
                rot = getattr(ann, "rotation_deg", 0.0) or 0.0
                if abs(rot) > 0.5:
                    # Use the rotated bounding envelope for the crop region
                    x1, y1, x2, y2 = _rotated_box_bounds(x1, y1, x2, y2, rot)
            elif ann.mode == "polygon" and len(ann.poly_pts) >= 6:
                xs = ann.poly_pts[0::2]; ys = ann.poly_pts[1::2]
                x1, y1, x2, y2 = min(xs), min(ys), max(xs), max(ys)
            elif ann.mode == "pin":
                r = 60
                x1 = max(0, ann.x1_px - r); y1 = max(0, ann.y1_px - r)
                x2 = min(iw, ann.x1_px + r); y2 = min(ih, ann.y1_px + r)
            else:
                return Spacer(max_w, max_h * 0.5)

            # Fixed 30px padding on all sides, keeping annotation centered
            pad_px = 30  # Fixed padding in pixels
            cx1 = max(0, x1 - pad_px); cy1 = max(0, y1 - pad_px)
            cx2 = min(iw, x2 + pad_px); cy2 = min(ih, y2 + pad_px)
            crop = img.crop((cx1, cy1, cx2, cy2))
            cw, ch = crop.size

            # Draw box on the crop
            draw = ImageDraw.Draw(crop, "RGBA")
            sev_h = self._SEV_COLORS_HEX.get(ann.severity, "#FFC107")
            r_, g_, b_ = int(sev_h[1:3], 16), int(sev_h[3:5], 16), int(sev_h[5:7], 16)
            bx1, by1 = x1 - cx1, y1 - cy1
            bx2, by2 = x2 - cx1, y2 - cy1
            if ann.mode != "pin":
                _rot_z = getattr(ann, "rotation_deg", 0.0) or 0.0
                if abs(_rot_z) > 0.5:
                    # Re-derive un-cropped corners from original stored coords and
                    # translate into crop-local space
                    _ox1 = min(ann.x1_px, ann.x2_px); _oy1 = min(ann.y1_px, ann.y2_px)
                    _ox2 = max(ann.x1_px, ann.x2_px); _oy2 = max(ann.y1_px, ann.y2_px)
                    _corners = _rotated_box_corners(_ox1, _oy1, _ox2, _oy2, _rot_z)
                    _crop_corners = [(px - cx1, py - cy1) for px, py in _corners]
                    draw.polygon(_crop_corners, outline=(r_, g_, b_, 255))
                    draw.line(_crop_corners + [_crop_corners[0]],
                              fill=(r_, g_, b_, 255), width=2)
                else:
                    draw.rectangle([bx1, by1, bx2, by2],
                                   outline=(r_, g_, b_, 255), width=2)

            # v4.3.0: Measurement overlay REMOVED per user request —
            # dimension label (cm) is shown in the pill table below, not burned onto the crop.

            scale = min(max_w / cw, max_h / ch)
            # Vertically centre the cropped image: pad top so it sits in the middle of IMG_H
            rendered_h = int(ch * scale)
            top_pad = max(0, int((max_h - rendered_h) / 2))
            buf   = _BIO()
            crop.save(buf, "JPEG", quality=88)
            buf.seek(0)
            # Return image — vertical centering handled via Table VALIGN MIDDLE on img_row
            return RLImage(buf, width=int(cw*scale), height=rendered_h)
        except Exception as exc:
            log.warning(f"_render_zoom_crop {fp}: {exc}")
            return Spacer(max_w, max_h * 0.5)

    def _find_pinpoint_image(self, annotation: "Annotation", image_record: "ImageRecord") -> "Optional[str]":
        """
        Phase 3 (TODO #7): Search for Scopito pinpoint images.
        Returns the path to a pinpoint image if found, otherwise None.
        
        Search patterns:
        - {project_dir}/pinpoints/{image_name}_{annotation_id}_pinpoint.png
        - {project_dir}/pinpoints/{image_name}_pinpoint.png
        - {project_dir}/scopito_pinpoints/{image_name}_{annotation_id}.png
        
        Returns None if:
        - No pinpoint directory exists
        - No matching pinpoint image found
        - Component is Hub or Tower (skip pinpoint logic)
        """
        # Skip pinpoint logic for Hub and Tower
        component = image_record.blade or ""
        if component.lower() in ("hub", "tower"):
            log.debug(f"Skipping pinpoint search for {component} (Hub/Tower rule)")
            return None
        
        # Get project directory
        project_dir = self._project.project_folder if self._project.project_folder else None
        if not project_dir:
            return None
        
        # Get base image name (without extension)
        image_name = os.path.splitext(os.path.basename(image_record.filepath))[0]
        annotation_id = annotation.id if hasattr(annotation, 'id') else ""
        
        # Search in common pinpoint directories
        pinpoint_dirs = ["pinpoints", "scopito_pinpoints", "pinpoint_images"]
        
        for pinpoint_dir in pinpoint_dirs:
            search_dir = os.path.join(project_dir, pinpoint_dir)
            if not os.path.isdir(search_dir):
                continue
            
            # Try different filename patterns
            patterns = [
                f"{image_name}_{annotation_id}_pinpoint.png",
                f"{image_name}_{annotation_id}.png",
                f"{image_name}_pinpoint.png",
            ]
            
            for pattern in patterns:
                pinpoint_path = os.path.join(search_dir, pattern)
                if os.path.isfile(pinpoint_path):
                    log.info(f"✓ Found pinpoint image: {os.path.basename(pinpoint_path)}")
                    return pinpoint_path
        
        log.debug(f"No pinpoint image found for {image_name}, using burned image")
        return None

    def _render_mini_blade_diagram(self, blade_key: str, face_abbr: str,
                                   dist_mm: float, severity: str,
                                   max_w: float, max_h: float,
                                   edge_side: Optional[str] = None):
        """
        Mini single-blade silhouette with ONE coloured dot at the defect position.
        Vertical orientation: root at top, tip at bottom.
        """
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as _plt
            import numpy as _np
            from io import BytesIO as _BIO

            blade_length_mm = getattr(self._project, "blade_length_mm", 50_000.0) or 50_000.0
            SEV_COL = _SEV_HEX   # v3.0.1: use central map
            dot_col = SEV_COL.get(severity, "#F97316")

            fig, ax = _plt.subplots(figsize=(1.2, 4.5))
            fig.patch.set_facecolor("#f8f9fa")
            ax.set_facecolor("#f8f9fa")

            y_vals = _np.linspace(0, 1, 200)
            hw = 0.44 * _np.sqrt(_np.maximum(0.0, 1.0 - y_vals))

            # Draw single tapered blade
            ax.fill_betweenx(y_vals, -hw, hw, color="#dde3ea", alpha=0.85, zorder=1)
            ax.plot(-hw, y_vals, color="#aab4be", lw=0.8, zorder=2)
            ax.plot( hw, y_vals, color="#aab4be", lw=0.8, zorder=2)

            # FIX: guard is division-by-zero only (blade_length_mm==0), NOT
            # dist_mm==0.  The previous `if dist_mm > 0` mapped any root-end
            # annotation (dist_mm=0.0) to the midpoint (yn=0.5), silently
            # ignoring the user's explicit pinpoint placement.
            yn = min(1.0, max(0.0, dist_mm / blade_length_mm)) if blade_length_mm > 0 else 0.5
            ax.scatter(0.0, yn, s=120, color=dot_col,
                       edgecolors="white", linewidths=1.2, zorder=5)

            # v4.3.0: Draw edge highlight stripe if edge_side is set
            if edge_side in ("LE", "TE"):
                # LE = left edge (-hw), TE = right edge (+hw)
                edge_x = -hw if edge_side == "LE" else hw
                ax.plot(edge_x, y_vals, color="#00d4e0", lw=2.5, zorder=4,
                        solid_capstyle="round")
                # Label the edge
                label_y = 0.25   # middle of blade
                label_x = -0.62 if edge_side == "LE" else 0.58
                ax.text(label_x, label_y, edge_side, fontsize=6,
                        fontweight="bold", color="#00d4e0",
                        ha="left" if edge_side == "TE" else "right", va="center")

            # Face label
            ax.text(0, -0.06, face_abbr or "?", ha="center", va="bottom",
                    fontsize=7, fontweight="bold", color="#333")
            ax.text(0, -0.13, f"Blade {blade_key}", ha="center", va="bottom",
                    fontsize=6, color="#666")

            # Y-axis distance labels
            ticks_mm = [0, blade_length_mm * 0.5, blade_length_mm]
            ax.set_yticks([t / blade_length_mm for t in ticks_mm])
            ax.set_yticklabels([f"{int(t/1000)}m" for t in ticks_mm],
                               fontsize=5, color="#777")
            ax.yaxis.set_tick_params(length=0)
            ax.set_ylim(1.08, -0.18)
            ax.set_xlim(-0.65, 0.65)
            ax.set_xticks([])
            ax.spines[["top", "right", "bottom"]].set_visible(False)
            ax.spines["left"].set_color("#ccc")
            fig.tight_layout(pad=0.3)

            buf = _BIO()
            fig.savefig(buf, format="png", dpi=130, bbox_inches="tight")
            _plt.close(fig)
            buf.seek(0)
            scale = min(max_w / (1.2 * 25.4), max_h / (4.5 * 25.4))  # inch→mm
            return RLImage(buf, width=max_w * 0.95, height=max_h * 0.95)
        except Exception as exc:
            log.warning(f"_render_mini_blade_diagram: {exc}")
            return Spacer(max_w, max_h)

    def _build_component_detail_table(self, irecs: list) -> list:
        """
        Compact per-annotation detail table appended at the end of a
        component's grid pages — one row per annotation, all detail fields.
        """
        story = [Spacer(1, 3*mm)]
        all_pairs = [(irec, ann)
                     for irec in irecs for ann in irec.annotations]
        if not all_pairs:
            return story

        hdr_ps = ParagraphStyle("dh", fontName="Helvetica-Bold", fontSize=7,
                                textColor=rl_colors.white)
        val_ps = ParagraphStyle("dv", fontName="Helvetica", fontSize=7,
                                leading=9)

        headers = ["#", "Image", "Span / Face", "Defect", "Severity", "Size"]
        col_ws  = [6*mm, 42*mm, 34*mm, 34*mm, 26*mm, 38*mm]
        tbl_data = [[Paragraph(h, hdr_ps) for h in headers]]

        for idx, (irec, ann) in enumerate(all_pairs):
            fname = os.path.basename(irec.filepath)[:24]
            size_s = (f"{ann.width_cm:.1f}×{ann.height_cm:.1f}cm"
                      if ann.width_cm else "—")
            span_face = f"{ann.span or '—'} / {ann.face or '—'}"
            sev_c = self._SEV_COLORS_HEX.get(ann.severity, "#999999")
            sev_s = SEVERITY_SHORT.get(ann.severity, ann.severity)

            tbl_data.append([
                Paragraph(str(idx + 1), val_ps),
                Paragraph(fname, val_ps),
                Paragraph(span_face, val_ps),
                Paragraph(ann.defect or "—", val_ps),
                Paragraph(f"<font color='{sev_c}'><b>{sev_s}</b></font>", val_ps),
                Paragraph(size_s, val_ps),
            ])

        det_tbl = Table(tbl_data, colWidths=col_ws, repeatRows=1)
        det_style = TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0),  rl_colors.HexColor("#0d1117")),
            ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0), (-1, -1), 7),
            ("LINEBELOW",     (0, 0), (-1, 0),  0.8, rl_colors.HexColor("#00d4e0")),
            ("TOPPADDING",    (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("GRID",          (0, 0), (-1, -1), 0.2, rl_colors.HexColor("#cccccc")),
        ])
        for i in range(1, len(tbl_data)):
            bg = rl_colors.HexColor("#f8f8f8") if i % 2 == 1 else rl_colors.white
            det_style.add("BACKGROUND", (0, i), (-1, i), bg)
        det_tbl.setStyle(det_style)
        story.append(det_tbl)
        return story

    @staticmethod
    def _worst_severity(anns: list) -> str:
        order = {
            "Critical": 0, "5 - Very Serious": 0,
            "Major": 1, "Serious": 1, "4 - Serious": 1, "3 - Non-Serious": 1,
            "POI": 2, "Minor": 3, "1 - Cosmetic": 3,
            "2 - Similar/Cosmetic": 3,
        }
        sevs = [a.severity for a in anns if a.severity]
        if not sevs:
            return "—"
        return min(sevs, key=lambda s: order.get(s, 99))

    # ── CSV export ─────────────────────────────────────────────────────────────

    def export_csv(self, csv_path: str) -> bool:
        """
        Export all annotation data to a flat CSV file.
        Columns: annotation_id, image_file, component, blade, span, face,
                 defect_type, severity, size_w_cm, size_h_cm, mode,
                 gsd_cm_per_px, gsd_source, notes, status, created_by,
                 reviewed_by, reviewed_at, reviewer_note,
                 gps_coords, altitude_m, date_taken, heading,
                 turbine_id, site, project_name, report_date
        """
        import csv
        FIELDS = [
            "annotation_id", "image_file", "component",
            "blade", "span", "face",
            "defect_type", "severity",
            "size_w_cm", "size_h_cm", "mode",
            "gsd_cm_per_px", "gsd_source",
            "notes", "status",
            "created_by", "reviewed_by", "reviewed_at", "reviewer_note",
            "gps_coords", "altitude_m", "date_taken", "heading",
            "turbine_id", "site", "project_name", "report_date",
        ]
        p = self._project
        report_date = datetime.now().strftime("%Y-%m-%d %H:%M")
        try:
            with open(csv_path, "w", newline="", encoding="utf-8") as fh:
                writer = csv.DictWriter(fh, fieldnames=FIELDS)
                writer.writeheader()
                for irec in p.images.values():
                    for ann in irec.annotations:
                        writer.writerow({
                            "annotation_id":  ann.ann_id,
                            "image_file":     irec.filename,
                            "component":      irec.blade or "—",
                            "blade":          ann.blade or irec.blade or "—",
                            "span":           ann.span  or "—",
                            "face":           ann.face  or "—",
                            "defect_type":    ann.defect or "—",
                            "severity":       ann.severity or "—",
                            "size_w_cm":      f"{ann.width_cm:.3f}"
                                              if ann.width_cm is not None else "",
                            "size_h_cm":      f"{ann.height_cm:.3f}"
                                              if ann.height_cm is not None else "",
                            "mode":           ann.mode or "—",
                            "gsd_cm_per_px":  f"{ann.gsd_value:.5f}"
                                              if ann.gsd_value is not None else "",
                            "gsd_source":     ann.gsd_source or "—",
                            "notes":          ann.notes or "",
                            "status":         ann.status or "pending",
                            "created_by":     ann.created_by or "",
                            "reviewed_by":    ann.reviewed_by or "",
                            "reviewed_at":    ann.reviewed_at or "",
                            "reviewer_note":  ann.reviewer_note or "",
                            "gps_coords":     irec.gps_coords or "",
                            "altitude_m":     irec.altitude_m or "",
                            "date_taken":     irec.date_taken or "",
                            "heading":        irec.heading or "",
                            "turbine_id":     irec.turbine_id or p.turbine_id or "",
                            "site":           p.site or "",
                            "project_name":   p.name or "",
                            "report_date":    report_date,
                        })
            log.info(f"CSV exported → {csv_path}")
            return True
        except Exception as exc:
            log.error(f"CSV export failed: {exc}")
            return False

    def _count_by_severity(self) -> Dict[str, int]:
        counts: Dict[str, int] = {}
        for irec in self._project.images.values():
            for ann in irec.annotations:
                counts[ann.severity] = counts.get(ann.severity, 0) + 1
        return counts

# ==============================================================================
# TOAST NOTIFICATION  (Dev Patel — UX)
# Slide-in non-blocking message shown on top of the main window.
# Auto-dismisses after `duration_ms` milliseconds.
# ==============================================================================

class ToastNotification(QLabel):
    """Dev Patel: Transient, non-blocking toast notification overlay."""
    def __init__(self, parent: QWidget, message: str, colour: str = "",
                 duration_ms: int = 2800):
        super().__init__(message, parent)
        c = colour or UI_THEME["accent_cyan"]
        self.setStyleSheet(
            f"background:{UI_THEME['bg_card']};color:{c};"
            f"border:1px solid {c};border-radius:8px;"
            f"padding:8px 16px;font-size:10pt;font-weight:600;")
        self.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.adjustSize()
        self._reposition(parent)
        self.show()
        self.raise_()
        QTimer.singleShot(duration_ms, self.deleteLater)

    def _reposition(self, parent: QWidget):
        pw, ph = parent.width(), parent.height()
        self.adjustSize()
        x = (pw - self.width()) // 2
        y = ph - self.height() - 48
        self.move(x, y)

# ==============================================================================
# PHASE 8.1 — REPORT SETTINGS DIALOG  (Sarah Chen)
# ==============================================================================

# ==============================================================================
# DOCX REPORT GENERATOR  (inlined — no companion file needed)
# Originally docx_report_generator.py — merged here for monolithic distribution.
# Updated: uses ann.distance_from_root_mm, ann.remedy_action, project.blade_numbers
# ==============================================================================

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches as _DInches, Pt as _DPt, RGBColor as _DRGBColor, Mm as _DMm
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _DAP
    from docx.enum.table import WD_TABLE_ALIGNMENT as _DTA
    from docx.oxml.ns import qn as _dqn
    from docx.oxml import OxmlElement as _DOEL
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    # Stub out docx names so module-level code below doesn't NameError on import
    def _DRGBColor(*a): return None   # type: ignore[misc]
    def _DPt(*a):       return None   # type: ignore[misc]
    def _DOEL(*a):      return None   # type: ignore[misc]
    def _dqn(*a):       return ""     # type: ignore[misc]
    class _DAP:         pass          # type: ignore[misc]
    class _DTA:         pass          # type: ignore[misc]
    _DInches = _DMm = lambda x: x    # type: ignore[misc]


def _dc_set_cell_bg(cell, rgb):
    """Set cell background. rgb can be RGBColor (int subclass), tuple, or int."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = _DOEL("w:shd")
    shd.set(_dqn("w:val"),   "clear")
    shd.set(_dqn("w:color"), "auto")
    # RGBColor from python-docx inherits int; extract r/g/b via bit-ops.
    # Fallback to subscript for plain tuples.
    try:
        _iv = int(rgb)
        hex_s = f"{(_iv >> 16) & 0xFF:02X}{(_iv >> 8) & 0xFF:02X}{_iv & 0xFF:02X}"
    except (TypeError, ValueError):
        hex_s = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(_dqn("w:fill"), hex_s)
    tcPr.append(shd)


def _dc_set_margins(cell, top=60, bottom=60, left=100, right=100):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = _DOEL("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = _DOEL(f"w:{side}")
        el.set(_dqn("w:w"),    str(val))
        el.set(_dqn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _dc_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spc = _DOEL("w:spacing")
    spc.set(_dqn("w:before"), str(before))
    spc.set(_dqn("w:after"),  str(after))
    pPr.append(spc)


def _dc_run(para, text, bold=False, italic=False, size_pt=10,
            color=None, font="Arial"):
    run = para.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.name = font; run.font.size = _DPt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def _dc_page_break(doc):
    para = doc.add_paragraph()
    from docx.enum.text import WD_BREAK
    para.add_run().add_break(WD_BREAK.PAGE)


# Colour palette
_DC_DARK    = _DRGBColor(0x0D, 0x11, 0x17)
_DC_CYAN    = _DRGBColor(0x00, 0xD4, 0xE0)
_DC_WHITE   = _DRGBColor(0xFF, 0xFF, 0xFF)
_DC_GREY    = _DRGBColor(0x88, 0x88, 0x88)
_DC_BODY    = _DRGBColor(0x11, 0x11, 0x11)
_DC_ROW_ALT = _DRGBColor(0xF4, 0xF6, 0xF8)
_DC_AMBER   = _DRGBColor(0xE8, 0x77, 0x0A)

_DC_SEV_RGB = {
    "Minor":    _DRGBColor(0x3F, 0xB9, 0x50),
    "Major":    _DRGBColor(0xD2, 0x99, 0x22),
    "Critical": _DRGBColor(0xF8, 0x51, 0x49),
    "POI":      _DRGBColor(0x38, 0x8B, 0xFD),
    # Legacy
    "1 - Cosmetic":         _DRGBColor(0x3F, 0xB9, 0x50),
    "2 - Similar/Cosmetic": _DRGBColor(0x7E, 0xE7, 0x87),
    "3 - Non-Serious":      _DRGBColor(0xD2, 0x99, 0x22),
    "4 - Serious":          _DRGBColor(0xD2, 0x99, 0x22),
    "5 - Very Serious":     _DRGBColor(0xF8, 0x51, 0x49),
    "Serious":  _DRGBColor(0xD2, 0x99, 0x22),
    "Low":      _DRGBColor(0x3F, 0xB9, 0x50),
    "Medium":   _DRGBColor(0xD2, 0x99, 0x22),
    "High":     _DRGBColor(0xD2, 0x99, 0x22),
}

# Blade diagram: matplotlib-based, Image-2 reference style
def _dc_build_blade_diagram(project: "Project") -> Optional[bytes]:
    """
    Generate a blade diagram PNG (Image-2 Scopito style):
      3 blade columns (B1/B2/B3) × 4 face lanes (LE/TE/PS/SS)
      Tapered blade silhouettes, orange dots at proportional distance-from-root,
      coloured by severity.  Returns PNG bytes or None.
    """
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as _plt
        import matplotlib.patches as _mpatches
        import numpy as _np

        # Read blade_length from project (matches PDF generator behaviour)
        blade_length_mm = getattr(project, "blade_length_mm", 50_000.0) or 50_000.0
        FACES  = ["PS", "LE", "TE", "SS"]
        BLADES = ["A", "B", "C"]
        SEV_COL = _SEV_HEX   # v3.0.1: use central map

        col_w   = 1.0
        gap     = 0.5
        group_w = len(FACES) * col_w + gap
        fig_w   = len(BLADES) * group_w - gap

        fig, ax = _plt.subplots(figsize=(12, 9))
        fig.patch.set_facecolor('#f8f9fa')
        ax.set_facecolor('#f8f9fa')

        y_vals  = _np.linspace(0, 1, 200)
        def _hw(y): return col_w * 0.44 * _np.sqrt(_np.maximum(0.0, 1.0 - y))

        for bi, bkey in enumerate(BLADES):
            gx = bi * group_w
            for si, face in enumerate(FACES):
                cx  = gx + si * col_w + col_w / 2
                hw  = _hw(y_vals)
                ax.fill_betweenx(y_vals, cx - hw, cx + hw, color='#dde3ea', alpha=0.85, zorder=1)
                ax.plot(cx - hw, y_vals, color='#aab4be', lw=0.6, zorder=2)
                ax.plot(cx + hw, y_vals, color='#aab4be', lw=0.6, zorder=2)
                ax.text(cx, -0.04, face, ha='center', va='bottom', fontsize=9, fontweight='bold', color='#333')

            gcx  = gx + len(FACES) * col_w / 2
            serial = project.blade_numbers.get(bkey, "")
            lbl    = f"Blade {bkey[-1]}" + (f" ({serial})" if serial else "")
            ax.text(gcx, -0.10, lbl, ha='center', va='bottom', fontsize=10, fontweight='bold', color='#1a1a2e')

            # ROOT-CAUSE FIX: use ann.blade (user-chosen in panel) not irec.blade
            # (folder auto-detected, defaults to "A") to route dots correctly.
            for irec in project.images.values():
                for ann in irec.annotations:
                    _ann_blade = (ann.blade or irec.blade or "").strip()
                    if _ann_blade != bkey:
                        continue
                    dist = ann.distance_from_root_mm or 0
                    parts = ann.face.split("(")
                    face_abbr = parts[-1].strip(")") if len(parts) > 1 else ann.face
                    if face_abbr not in FACES:
                        continue
                    si = FACES.index(face_abbr)
                    cx = gx + si * col_w + col_w / 2
                    yn = min(1.0, dist / blade_length_mm)
                    clr = SEV_COL.get(ann.severity, "#F97316")
                    ax.scatter(cx, yn, s=70, color=clr, edgecolors='white', linewidths=0.8, zorder=5)

        ax.set_ylim(1.05, -0.14)
        ticks   = [0, 10_000, 20_000, 30_000, 40_000, 50_000]
        ax.set_yticks([t / blade_length_mm for t in ticks])
        ax.set_yticklabels([f"{int(t/1000)}m" for t in ticks], fontsize=8, color='#555')
        ax.yaxis.set_tick_params(length=0)
        ax.set_xlim(-col_w * 0.6, fig_w - gap + col_w * 0.6)
        ax.set_xticks([])
        ax.spines[['top','right','bottom']].set_visible(False)
        ax.spines['left'].set_color('#ccc')

        # v4.3.0: Updated to Yellow/Amber/Red palette
        handles = [_mpatches.Patch(color=c, label=l) for l, c in [
            ("Minor (Sev 1)",    "#FFD700"),
            ("Major (Sev 3/4)", "#FFA500"),
            ("Critical (Sev 5)", "#FF0000"),
            ("POI",              "#388bfd"),
        ]]
        ax.legend(handles=handles, loc='lower center', bbox_to_anchor=(0.5, -0.06),
                  ncol=4, fontsize=7, frameon=False)
        ax.set_title("Blade Defect Location Diagram", fontsize=12,
                     fontweight='bold', color='#1a1a2e', pad=18)
        fig.tight_layout(rect=[0, 0.03, 1, 1])

        import io as _io
        buf = _io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        _plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception as exc:
        log.warning(f"_dc_build_blade_diagram: {exc}")
        return None


class DocxReportGenerator:
    """
    MONOLITHIC: Generates a Word (.docx) inspection report with full Scopito
    field parity — distance_from_root_mm, remedy_action, blade serial numbers,
    blade diagram, turbine index table (SenseHawk PDF pages 3-4 layout).
    """

    def __init__(self, project: "Project",
                 report_settings: Optional[Dict[str, Any]] = None):
        self._project  = project
        self._settings = report_settings or {}

    def generate(self, output_path: str) -> bool:
        if not PYTHON_DOCX_AVAILABLE:
            log.error("DocxReportGenerator: pip install python-docx")
            return False
        try:
            doc = DocxDocument()
            self._setup_doc(doc)
            self._build_cover(doc)
            self._build_docx_defect_summary(doc)   # own page-broken page after cover
            # Per-annotation pages directly follow the defect summary.
            self._build_annotation_pages(doc)
            doc.save(output_path)
            log.info(f"DOCX saved → {output_path}")
            return True
        except Exception as exc:
            log.exception(f"DocxReportGenerator.generate: {exc}")
            return False

    def _setup_doc(self, doc):
        sec = doc.sections[0]
        sec.page_width    = _DInches(8.5)
        sec.page_height   = _DInches(11)
        sec.left_margin   = _DInches(1)
        sec.right_margin  = _DInches(1)
        sec.top_margin    = _DInches(1)
        sec.bottom_margin = _DInches(0.75)
        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = _DPt(10)
        self._build_header_footer(doc)

    def _build_header_footer(self, doc):
        p       = self._project
        site    = p.site or p.name or "Wind Tower"
        # FIX-15c: Report title mirrors PDF — "{site} Aerial Wind Tower Inspection"
        report_title = f"{site} Aerial Wind Tower Inspection"

        # ── Header: 3-col table mirrors PDF bar layout ─────────────────────────
        # PDF layout (left → right):
        #   [Report title, bold white, 11pt]  ···  [header logo | client logo]
        # DOCX matches:
        #   Left cell  — report title (bold white 11pt, same text as PDF)
        #   Middle cell — empty  (site name removed, mirrors FIX-15b in PDF)
        #   Right cell  — header logo (logo_path) then client logo, right-aligned
        #                 mirroring the PDF right-to-left stacking order
        header = doc.sections[0].header
        for para in header.paragraphs: para.clear()
        ht = header.add_table(rows=1, cols=3, width=_DInches(6.5))
        ht.alignment = _DTA.LEFT
        lc, mc, rc = ht.rows[0].cells
        # Slate-gray background — same RGB as PDF header bar (#4A5568 → approx 0x71,0x82,0x96)
        _HEADER_BG = _DRGBColor(0x4A, 0x55, 0x68)
        for c_ in (lc, mc, rc):
            _dc_set_cell_bg(c_, _HEADER_BG); _dc_set_margins(c_, 60, 60, 140, 140)
        lc.width = _DInches(3.0); mc.width = _DInches(1.5); rc.width = _DInches(2.0)

        # Left cell: report title — mirrors PDF left-anchored title text
        _dc_run(lc.paragraphs[0], report_title, bold=True, size_pt=11, color=_DC_WHITE)

        # Middle cell: intentionally empty — mirrors FIX-15b removal of site text
        # mc is already cleared above; nothing added.

        # Right cell: header logo (logo_path) RIGHT-MOST + client logo to its left,
        # mirroring PDF right-to-left stacking.  DOCX can only stack inline so we
        # place both runs in right-aligned paragraph: [client_logo_run][header_logo_run]
        # This produces: client logo ← header logo (right edge) — matches PDF visual.
        rp = rc.paragraphs[0]; rp.alignment = _DAP.RIGHT
        # Client logo drawn first (left of header logo in the right cell)
        cl_logo = self._settings.get("client_logo_path", "")
        if cl_logo and os.path.exists(cl_logo):
            try:
                rp.add_run().add_picture(cl_logo, height=_DInches(0.55))  # ≈14mm
            except Exception:
                _dc_run(rp, "Client ", size_pt=8, color=_DC_WHITE)
        # Header logo drawn second (rightmost position — mirrors PDF far-right anchor)
        hdr_logo = self._settings.get("logo_path", "")
        if hdr_logo and os.path.exists(hdr_logo):
            try:
                rp.add_run().add_picture(hdr_logo, height=_DInches(0.27))  # ≈7mm
            except Exception:
                pass  # Logo missing — cell still has client logo or is blank

        # ── Footer: 3-col table [company_logo + turbine | timestamp | page] ──
        footer = doc.sections[0].footer
        for para in footer.paragraphs: para.clear()
        ft = footer.add_table(rows=1, cols=3, width=_DInches(6.5))
        ft.alignment = _DTA.LEFT
        flc, fmc, frc = ft.rows[0].cells
        flc.width = _DInches(2.5); fmc.width = _DInches(2.5); frc.width = _DInches(1.5)

        # Footer left: company logo + turbine ID
        co_logo = self._settings.get("company_logo_path", "")
        if co_logo and os.path.exists(co_logo):
            try: flc.paragraphs[0].add_run().add_picture(co_logo, height=_DInches(0.25))
            except Exception: pass
        _dc_run(flc.paragraphs[0], f"  Turbine: {p.turbine_id or '—'}",
                size_pt=8, color=_DC_GREY)

        # Footer centre: (timestamp removed - now on cover page)
        # Leave empty

        # Footer right: page number field
        frp = frc.paragraphs[0]; frp.alignment = _DAP.RIGHT
        _dc_run(frp, "Page ", size_pt=8, color=_DC_GREY)
        # Add auto page-number field via OOXML
        try:
            from docx.oxml.ns import qn as _qn2
            from docx.oxml import OxmlElement as _OE2
            _fld = _OE2("w:fldChar"); _fld.set(_qn2("w:fldCharType"), "begin")
            frp.runs[-1]._r.append(_fld)
            _instrR = _OE2("w:r"); _instr = _OE2("w:instrText")
            _instr.text = "PAGE "; _instrR.append(_instr); frp._p.append(_instrR)
            _fld2 = _OE2("w:fldChar"); _fld2.set(_qn2("w:fldCharType"), "end")
            frp.runs[-1]._r.append(_fld2)
        except Exception:
            pass

    def _build_cover(self, doc):
        p       = self._project
        company = self._settings.get("company", "") or p.name or "Inspection Report"
        date_s  = datetime.now().strftime("%B %d, %Y").upper()
        site    = p.site or "—"

        # Cover client logo
        cl = self._settings.get("client_logo_path", "")
        if cl and os.path.exists(cl):
            try:
                lp = doc.add_paragraph(); lp.alignment = _DAP.CENTER
                lp.add_run().add_picture(cl, width=_DInches(2.5))
            except Exception: pass

        # ── Title matching PDF: "{site} Aerial Wind Tower Inspection" ─────────
        tp = doc.add_paragraph()
        _dc_run(tp, f"{site} Aerial Wind Tower Inspection",
                bold=True, size_pt=24, color=_DC_DARK)
        tp.alignment = _DAP.CENTER
        _dc_spacing(tp, before=0, after=80)
        # v3.3.5 DOC-1: subtitle now includes generation date + IST time
        # matching reference DOCX: "generated on '{date}' and {IST}"
        from datetime import timezone, timedelta
        _ist_offset = timedelta(hours=5, minutes=30)
        _ist_now    = datetime.now(timezone.utc).astimezone(timezone(_ist_offset))
        _ist_str    = _ist_now.strftime("%B %d, %Y") + " and " + _ist_now.strftime("%H:%M IST")
        sub = doc.add_paragraph()
        _dc_run(sub, f"Drone Visual Inspection \u2014 Executive Report generated on {_ist_str}",
                size_pt=11, color=_DC_GREY)
        sub.alignment = _DAP.CENTER
        _dc_spacing(sub, before=0, after=160)

        # ── WTG front page photo ──────────────────────────────────────────────
        wtg_path = self._settings.get("wtg_image_path", "")
        # T11 FIX: auto-read EXIF GPS from any project image if not set manually
        if not self._settings.get("gps_coords", "").strip():
            try:
                for _irec in self._project.images.values():
                    _exif = _read_exif_metadata(_irec.filepath)
                    _gps  = _exif.get("gps_coords", "")
                    if _gps:
                        self._settings["gps_coords"] = _gps
                        break
            except Exception:
                pass
        if wtg_path and os.path.exists(wtg_path):
            try:
                wp = doc.add_paragraph(); wp.alignment = _DAP.CENTER
                wp.add_run().add_picture(wtg_path, width=_DInches(5.5))
                _dc_spacing(wp, before=0, after=40)
                # T11 FIX: filename caption below cover photo
                cap_p = doc.add_paragraph()
                _dc_run(cap_p, os.path.basename(wtg_path),
                        italic=True, size_pt=8, color=_DC_GREY)
                cap_p.alignment = _DAP.CENTER
                _dc_spacing(cap_p, before=0, after=80)
            except Exception: pass
        else:
            # Placeholder box
            ph_t = doc.add_table(rows=1, cols=1)
            ph_c = ph_t.rows[0].cells[0]
            _dc_set_cell_bg(ph_c, _DRGBColor(0xED, 0xF2, 0xF7))
            _dc_set_margins(ph_c, 560, 560, 200, 200)
            ph_p = ph_c.paragraphs[0]; ph_p.alignment = _DAP.CENTER
            _dc_run(ph_p, "[ INSERT WTG FRONT IMAGE HERE ]\n"
                    "Set path in Report Settings → Identity → WTG Cover Photo",
                    italic=True, size_pt=9, color=_DC_GREY)
            doc.add_paragraph()

        # ── v3.3.5 DOC-2: Single 4-row × 2-col meta table (reference structure)
        # Row 0: label row "WTG No" | "WTG No"             bg #F4F6F8, text #888888
        # Row 1: value row {wtg_no} | {wtg_no}             bg white, bold
        # Row 2: label row "Tower Location (Site)" | "GPS Coordinates"  same label style
        # Row 3: value row {site} | {gps}                  bg white, bold
        # Previously this was two separate tables (rows 2×2 + rows 2×2).
        _raw_t  = p.turbine_id or "—"
        _wtg_no = (f"WTG-{_raw_t}" if _raw_t and _raw_t != "—"
                   and not _raw_t.upper().startswith("WTG") else _raw_t)
        gps_val = self._settings.get("gps_coords", "") or "Not recorded"
        _DC_META_LBL_BG = _DRGBColor(0xF4, 0xF6, 0xF8)  # #F4F6F8 from reference
        _DC_META_LBL    = _DRGBColor(0x88, 0x88, 0x88)   # #888888 from reference
        meta4 = doc.add_table(rows=4, cols=2)
        meta4.style = "Table Grid"; meta4.alignment = _DTA.LEFT
        _meta_rows = [
            # (label0, label1, is_label_row)
            ("WTG No",                "WTG No",            True),
            (_wtg_no,                 _wtg_no,             False),
            ("Tower Location (Site)", "GPS Coordinates",   True),
            (site,                    gps_val,             False),
        ]
        for ri, (v0, v1, is_lbl) in enumerate(_meta_rows):
            c0 = meta4.rows[ri].cells[0]; c1 = meta4.rows[ri].cells[1]
            c0.width = c1.width = _DInches(3.25)
            _dc_set_margins(c0); _dc_set_margins(c1)
            if is_lbl:
                _dc_set_cell_bg(c0, _DC_META_LBL_BG)
                _dc_set_cell_bg(c1, _DC_META_LBL_BG)
                _dc_run(c0.paragraphs[0], v0, bold=True, size_pt=8, color=_DC_META_LBL)
                _dc_run(c1.paragraphs[0], v1, bold=True, size_pt=8, color=_DC_META_LBL)
                c0.paragraphs[0].alignment = _DAP.CENTER
                c1.paragraphs[0].alignment = _DAP.CENTER
            else:
                _dc_run(c0.paragraphs[0], v0, bold=True, size_pt=10)
                _dc_run(c1.paragraphs[0], v1, bold=True, size_pt=10)
                # CHG-F: Centre WTG No value to match reference report
                c0.paragraphs[0].alignment = _DAP.CENTER
                c1.paragraphs[0].alignment = _DAP.CENTER
        doc.add_paragraph()

        # BLADE SERIAL NUMBERS section removed per spec (image mark-up, orange X)
        # SEVERITY OVERVIEW moved to _build_docx_defect_summary() — after Results.

        # ── Blade diagram ─────────────────────────────────────────────────────
        bd_bytes = _dc_build_blade_diagram(p)
        if bd_bytes:
            bdh = doc.add_paragraph()
            _dc_run(bdh, "BLADE DEFECT LOCATION DIAGRAM", bold=True, size_pt=10, color=_DC_DARK)
            _dc_spacing(bdh, before=200, after=60)
            import io as _io
            bd_para = doc.add_paragraph(); bd_para.alignment = _DAP.CENTER
            bd_para.add_run().add_picture(_io.BytesIO(bd_bytes), width=_DInches(6.2))
            doc.add_paragraph()

        # ── v3.3.5 DOC-11/13: GL-16 narrative sections with blue divider lines ─
        # These sections mirror the PDF generator's GL-16 narrative and match the
        # reference DOCX layout: Heading 1 + blue underline rule + body text.
        # Blue divider helper: 1-row, full-width, 3pt height, #2A6BAF fill.
        _DC_BLUE_RULE = _DRGBColor(0x2A, 0x6B, 0xAF)

        def _dc_add_blue_rule(doc_):
            """v3.3.5 DOC-11: thin full-width blue rule under Heading 1 sections."""
            rule_t = doc_.add_table(rows=1, cols=1)
            rule_t.alignment = _DTA.LEFT
            c_ = rule_t.rows[0].cells[0]
            c_.width = _DInches(6.5)
            _dc_set_cell_bg(c_, _DC_BLUE_RULE)
            # Collapse to ~3pt height via cell margins and fixed-height paragraph
            _dc_set_margins(c_, 0, 0, 0, 0)
            # Remove border so it looks like a plain rule
            from docx.oxml.ns import qn as _qnr
            from docx.oxml import OxmlElement as _OEr
            # Access or create table properties (python-docx compatible)
            tblPr = rule_t._tbl.tblPr
            if tblPr is None:
                tblPr = _OEr('w:tblPr')
                rule_t._tbl.insert(0, tblPr)
            tblBrd = _OEr('w:tblBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                brd = _OEr(f'w:{side}')
                brd.set(_qnr('w:val'), 'none')
                tblBrd.append(brd)
            tblPr.append(tblBrd)
            # Set exact row height to 3pt
            tr = rule_t.rows[0]._tr
            trPr = tr.get_or_add_trPr()
            trH = _OEr('w:trHeight')
            trH.set(_qnr('w:val'), '45')   # 45 twips ≈ 0.8mm ≈ 3pt
            trH.set(_qnr('w:hRule'), 'exact')
            trPr.append(trH)

        _company = self._settings.get("company", "") or p.name or "the inspection company"
        _client  = self._settings.get("client", "") or "The Client"
        _model   = self._settings.get("turbine_model", "") or "2 MW capacity"

        _def_objective = (
            f"{_client} appointed {_company} to conduct drone-based visual inspection "
            f"of the {_model} Wind Turbine Generator at the site. "
            "Wind Turbine Generators (WTGs) are exposed to unexpected weather conditions "
            "which may affect component performance and efficiency. Drone inspection "
            "provides a visual record of WTG component condition — a cost-effective and "
            "efficient method compared to traditional manual inspection. The objective of "
            "this report is to analyse the visual images of the wind turbine and assess "
            "the condition of all WTG components. Drone-based visual surveys help ensure "
            "a safe working environment, support preventive maintenance planning, reduce "
            "machine breakdowns and erosion, and provide access to otherwise inaccessible "
            "components.")
        _def_scope = (
            "To conduct a visual drone survey of the Wind Turbine Generator commissioned "
            "at the site and identify any visual damage, failures, or erosion of wind "
            "turbine components.")
        _def_data_col = (
            "During on-site inspection the site team checks weather conditions (wind "
            "speed, visibility, etc.) and flies the drone to the required position for "
            "data capture. The turbine is in stop condition with blades pitched out. "
            "The site team ensures that the drone collects data from 8\u201310 metres "
            "distance from the WTG.\n\nDrones are equipped with required sensors and "
            "digital cameras to collect high-quality visual images from the most optimal "
            "perspective. Collected data is scrutinised by the project team; identified "
            "anomalies are inspected before providing a summarised report. The drone "
            "operates stably in winds up to 10 m/s.")

        # Turbine spec list — use project blade_length_mm if available
        _bl_m  = (getattr(p, "blade_length_mm", 55_500) or 55_500) / 1000.0
        _turbine_specs = self._settings.get("turbine_specs") or [
            ("Rated Power",         "2.0 MW (2,000 kW)"),
            ("Rotor Diameter",      "114 m"),
            (f"Blade Length",       f"{_bl_m:.1f} m (3 blades)"),
            ("Swept Area",          "10,207 m\u00b2"),
            ("Wind Class",          "IEC IIIa (Low Wind)"),
            ("Cut-in Wind Speed",   "2.5 m/s"),
            ("Rated Wind Speed",    "10.0 \u2013 12.5 m/s"),
            ("Cut-out Wind Speed",  "25.0 m/s"),
            ("Tower Height",        f"{self._settings.get('tower_height_m', '106')} m"),
            ("Max Rotor Speed",     "~16 RPM"),
            ("Generator Type",      "Doubly Fed Induction Generator (DFIG)"),
        ]

        # Helper: add a Heading 1 paragraph styled blue (#2A6BAF) + blue rule below
        def _dc_h1(doc_, text_):
            h1 = doc_.add_paragraph()
            _dc_run(h1, text_, bold=True, size_pt=12,
                    color=_DRGBColor(0x2A, 0x6B, 0xAF))
            _dc_spacing(h1, before=160, after=40)
            _dc_add_blue_rule(doc_)

        _dc_h1(doc, "Objective")
        obj_p = doc.add_paragraph()
        _dc_run(obj_p, self._settings.get("objective", _def_objective), size_pt=9)
        _dc_spacing(obj_p, before=80, after=80)

        _dc_h1(doc, "Scope of Work")
        scope_p = doc.add_paragraph()
        _dc_run(scope_p, self._settings.get("scope", _def_scope), size_pt=9)
        _dc_spacing(scope_p, before=80, after=80)

        _dc_h1(doc, "Data Collection Methodology & Definitions")
        dc_p = doc.add_paragraph()
        _dc_run(dc_p, self._settings.get("data_collection", _def_data_col), size_pt=9)
        _dc_spacing(dc_p, before=80, after=80)

        _dc_h1(doc, "Turbine Specifications")
        spec_intro = doc.add_paragraph()
        _dc_run(spec_intro, f"{_model} \u2014 Key Specifications:", bold=True, size_pt=9)
        _dc_spacing(spec_intro, before=60, after=40)
        for spec_label, spec_val in (_turbine_specs if isinstance(_turbine_specs, list)
                                      else []):
            sp = doc.add_paragraph(style="List Bullet")
            _dc_run(sp, f"{spec_label}: ", bold=True, size_pt=9)
            sp.add_run(spec_val).font.size = _DPt(9)

        # ── Results heading ───────────────────────────────────────────────────
        _dc_h1(doc, "Results")
        res_p = doc.add_paragraph()
        _dc_run(res_p,
                "The scan results are presented for each blade and body of the "
                "Wind Turbine in the following sections.",
                size_pt=9)
        _dc_spacing(res_p, before=60, after=120)

    def _build_docx_defect_summary(self, doc):
        """
        Defect summary table on its own page-broken page after the cover.
        Mirrors PDF _build_defect_summary_page() — called from generate() after _build_cover().
        """
        _dc_page_break(doc)
        p = self._project

        ds_hdr = doc.add_paragraph()
        # CHG-E: "Turbine Summary" (was "DEFECT SUMMARY" header) — mirrors PDF heading
        # v3.3.13 FIX: Format as "Turbine Summary - WTG-{number}" to match PDF
        _raw_turb = self._project.turbine_id or self._project.name or "Unknown"
        _wtg_id = (f"WTG-{_raw_turb}" if _raw_turb and _raw_turb != "Unknown" and
                   not _raw_turb.upper().startswith("WTG") else _raw_turb)
        _dc_run(ds_hdr, f"Turbine Summary - {_wtg_id}", bold=True, size_pt=11,
                color=_DRGBColor(0x11, 0x11, 0x11))
        _dc_spacing(ds_hdr, before=80, after=80)

        # ── Severity overview band (POI | Minor | Major | Critical) ──────────
        sev_counts = self._count_by_severity()
        _sov_keys = [
            ("POI",      "POI"),
            ("Minor",    "Minor"),
            ("Major",    "Major"),
            ("Critical", "Critical"),
        ]
        so_hdr = doc.add_paragraph()
        _dc_run(so_hdr, "SEVERITY OVERVIEW", bold=True, size_pt=9, color=_DC_DARK)
        _dc_spacing(so_hdr, before=80, after=60)
        st_ds = doc.add_table(rows=2, cols=4); st_ds.alignment = _DTA.LEFT
        cw_ds = _DInches(6.5 / 4)
        for _ci, (_key, _disp) in enumerate(_sov_keys):
            _bg_sov = _DC_SEV_RGB.get(_key, _DRGBColor(0x71, 0x82, 0x96))
            _lc = st_ds.rows[0].cells[_ci]; _cc2 = st_ds.rows[1].cells[_ci]
            for _c in (_lc, _cc2):
                _c.width = cw_ds
                _dc_set_cell_bg(_c, _bg_sov)
                _dc_set_margins(_c, 60, 60, 80, 80)
            _lp = _lc.paragraphs[0]; _lp.alignment = _DAP.CENTER
            _dc_run(_lp, _disp, bold=True, size_pt=8, color=_DC_WHITE)
            _cp = _cc2.paragraphs[0]; _cp.alignment = _DAP.CENTER
            _dc_run(_cp, str(sev_counts.get(_key, 0)), bold=True, size_pt=20, color=_DC_WHITE)
        doc.add_paragraph()

        # ── Severity legend (3 rows: key | description | action) ─────────────
        _sov_legend = [
            ("POI",      "Point of interest",  "Needs further investigation"),
            ("Minor",    "Minor defect",        "Monitoring — no intervention"),
            ("Major",    "Major defect",        "Intervention at planned inspection"),
            ("Critical", "Critical",            "Immediate intervention required"),
        ]
        lt_ds = doc.add_table(rows=3, cols=4); lt_ds.alignment = _DTA.LEFT
        for _ci, (_skey, _desc, _act) in enumerate(_sov_legend):
            _bg_l = _DC_SEV_RGB.get(_skey, _DRGBColor(0x71, 0x82, 0x96))
            _r0c = lt_ds.rows[0].cells[_ci]
            _r1c = lt_ds.rows[1].cells[_ci]
            _r2c = lt_ds.rows[2].cells[_ci]
            _r0c.width = _DInches(6.5 / 4)
            _dc_set_cell_bg(_r0c, _bg_l)
            _dc_set_margins(_r0c); _dc_set_margins(_r1c); _dc_set_margins(_r2c)
            _dc_run(_r0c.paragraphs[0], _skey, bold=True, size_pt=8, color=_DC_WHITE)
            _dc_run(_r1c.paragraphs[0], _desc, size_pt=8)
            _dc_run(_r2c.paragraphs[0], _act, size_pt=7, italic=True, color=_DC_GREY)
        doc.add_paragraph()

        # Build sorted flat list of all annotations
        _all_ds: list = []
        for _irec in p.images.values():
            for _ann in _irec.annotations:
                _all_ds.append((_irec, _ann))

        def _sev_rank_ds(s): return -SEVERITY_RANK.get(s, 0)
        def _b_ord_ds(b):    return {"A": 0, "B": 1, "C": 2}.get(b, 9)
        def _f_ord_ds(f):    return {"PS": 0, "LE": 1, "TE": 2, "SS": 3}.get(f, 9)
        _all_ds.sort(key=lambda x: (
            _sev_rank_ds(x[1].severity), _b_ord_ds(x[1].blade),
            _f_ord_ds(x[1].face.split("(")[-1].strip(")") if "(" in x[1].face else x[1].face)))

        if _all_ds:
            _ds_cols  = ["#", "Blade", "Issues", "Severity", "Size", "Root Dist.", "Tip Dist."]
            _ds_cws   = [0.35, 0.9, 1.3, 0.75, 1.1, 0.85, 0.85]
            _total_ds = sum(_ds_cws)
            _ds_cws   = [w * 6.5 / _total_ds for w in _ds_cws]
            _ds_tbl   = doc.add_table(rows=1 + len(_all_ds), cols=7)
            _ds_tbl.style = "Table Grid"; _ds_tbl.alignment = _DTA.LEFT
            _DC_DS_HDR = _DRGBColor(0x49, 0x54, 0x67)
            for _ci, (_h, _cw) in enumerate(zip(_ds_cols, _ds_cws)):
                _hc = _ds_tbl.rows[0].cells[_ci]; _hc.width = _DInches(_cw)
                _dc_set_cell_bg(_hc, _DC_DS_HDR); _dc_set_margins(_hc, 50, 50, 70, 70)
                _dc_run(_hc.paragraphs[0], _h, bold=True, size_pt=7, color=_DC_WHITE)
                _hc.paragraphs[0].alignment = _DAP.CENTER
            for _seq, (_irec_ds, _ann_ds) in enumerate(_all_ds, 1):
                _row_bg    = _DC_ROW_ALT if _seq % 2 == 0 else _DRGBColor(0xFF, 0xFF, 0xFF)
                _sev_bg_ds = _DC_SEV_RGB.get(_ann_ds.severity)
                _root_ds   = (f"{_ann_ds.root_distance_m:.1f} m"
                              if _ann_ds.root_distance_m else "\u2014")
                _tip_ds    = (f"{_ann_ds.tip_distance_m:.1f} m"
                              if _ann_ds.tip_distance_m  else "\u2014")
                _size_ds   = (f"{_ann_ds.width_cm:.1f} \u00d7 {_ann_ds.height_cm:.1f} cm"
                              if _ann_ds.width_cm else "N/A")
                _blade_lbl_ds = (f"Blade {_ann_ds.blade}" if _ann_ds.blade in ("A","B","C")
                                 else (_ann_ds.blade or "?"))
                _sn_ds = SEVERITY_SHORT.get(_ann_ds.severity or "", "\u2014")
                _vals_ds = [str(_seq), _blade_lbl_ds, _ann_ds.defect or "\u2014",
                            _sn_ds, _size_ds, _root_ds, _tip_ds]
                # Short cols (0=#, 1=Blade, 3=Severity, 5=Root, 6=Tip): centre; long col (2=Issues): left
                _short_cols = {0, 1, 3, 5, 6}
                for _ci, (_v, _cw) in enumerate(zip(_vals_ds, _ds_cws)):
                    _vc = _ds_tbl.rows[_seq].cells[_ci]; _vc.width = _DInches(_cw)
                    if _ci == 4 and _sev_bg_ds:
                        _dc_set_cell_bg(_vc, _sev_bg_ds)
                        _dc_set_margins(_vc, 50, 50, 70, 70)
                        _dc_run(_vc.paragraphs[0], _v, size_pt=7, color=_DC_WHITE)
                        _vc.paragraphs[0].alignment = _DAP.CENTER
                    else:
                        _dc_set_cell_bg(_vc, _row_bg)
                        _dc_set_margins(_vc, 50, 50, 70, 70)
                        _dc_run(_vc.paragraphs[0], _v, size_pt=7, bold=(_ci == 3))
                        _vc.paragraphs[0].alignment = (
                            _DAP.CENTER if _ci in _short_cols else _DAP.LEFT)
        else:
            no_def = doc.add_paragraph()
            _dc_run(no_def, "No defects recorded.", italic=True, size_pt=9)

    def _build_annotation_pages(self, doc):
        """T04/T08: sort A→B→C→Hub→Tower; 2 defects per page; cantSplit per block.
        FIX-10: group by ann.blade (authoritative, user-confirmed) instead of
        irec.blade (stale dataclass default). Sort within each component by filename
        so the defect sequence is deterministic regardless of folder-load order.
        """
        def _comp_sort_key(b: str) -> int:
            return {"A": 0, "B": 1, "C": 2, "Hub": 3, "Tower": 4}.get(b, 10)

        # FIX-10: build (irec, ann) pairs keyed by ann.blade (authoritative).
        # irec.blade defaults to "A" and may be stale; ann.blade is set by the
        # user in the annotation panel and persisted to project.json on every save.
        pair_comps: dict = {}
        for irec in self._project.images.values():
            for ann in irec.annotations:
                comp = (ann.blade or irec.blade or "Unknown").strip()
                pair_comps.setdefault(comp, []).append((irec, ann))

        # Stable filename-alphabetical order within each component (FIX-10)
        for _pairs in pair_comps.values():
            _pairs.sort(key=lambda t: (t[0].filename or "", t[1].ann_id or ""))

        sorted_comps = sorted(pair_comps.items(), key=lambda kv: _comp_sort_key(kv[0]))

        global_num = 0
        for comp_name, pairs in sorted_comps:
            for irec, ann in pairs:
                global_num += 1
                # page break before 1st defect and every odd-numbered one (0-based pair start)
                page_break = (global_num == 1) or ((global_num % 2) == 1)
                self._build_single_page(doc, irec, ann, global_num, page_break)

    @staticmethod
    def _cant_split_table(table):
        """Prevent any table row from splitting across pages (cantSplit)."""
        from docx.oxml.ns import qn as _qncx
        from docx.oxml import OxmlElement as _OEcx
        for row in table.rows:
            trPr = row._tr.get_or_add_trPr()
            cs = _OEcx('w:cantSplit')
            cs.set(_qncx('w:val'), '1')
            trPr.append(cs)

    def _build_single_page(self, doc, ir: "ImageRecord", ann: "Annotation",
                           global_num: int = 0, do_page_break: bool = True):
        """
        Scopito-style per-annotation block (2 per page).
        do_page_break=True  → hard page break before this block (1st of pair).
        do_page_break=False → thin separator rule only (2nd of pair on same page).
        """
        if do_page_break:
            _dc_page_break(doc)
        else:
            # Thin horizontal rule between the two defects sharing a page
            sep = doc.add_paragraph()
            from docx.oxml.ns import qn as _qns
            from docx.oxml import OxmlElement as _OEs
            pPr = sep._p.get_or_add_pPr()
            pb = _OEs('w:pBdr')
            top = _OEs('w:top')
            top.set(_qns('w:val'), 'single')
            top.set(_qns('w:sz'), '6')
            top.set(_qns('w:space'), '1')
            top.set(_qns('w:color'), 'BBBBBB')
            pb.append(top); pPr.append(pb)
            _dc_spacing(sep, before=80, after=80)

        fname      = ir.filename
        blade_key  = ir.blade or "?"
        serial     = self._project.blade_numbers.get(ir.blade, "")
        parts      = ann.face.split("("); face_abbr = parts[-1].strip(")") if len(parts) > 1 else ann.face
        blade_lbl  = f"Blade {blade_key}" + (f" ({serial})" if serial else "")
        dist       = ann.distance_from_root_mm or 0
        size_s     = (f"{ann.width_cm:.1f} × {ann.height_cm:.1f} cm" if ann.width_cm else "—")

        # ── Compact header row (v3.2.0: WTG No | Component) ─────────────────────
        _is_blade_p = blade_key in ("A", "B", "C")
        _wtg_no_p   = self._project.turbine_id or "—"
        if _wtg_no_p and _wtg_no_p != "—" and not _wtg_no_p.upper().startswith("WTG"):
            _wtg_no_p = f"WTG-{_wtg_no_p}"
        _comp_hdr = blade_lbl if _is_blade_p else (blade_key or "?")
        mt = doc.add_table(rows=2, cols=2)
        mt.style = "Table Grid"; mt.alignment = _DTA.LEFT
        for ci, (h, v) in enumerate(zip(
            ["WTG No",  "Component"],
            [_wtg_no_p, _comp_hdr]
        )):
            hc = mt.rows[0].cells[ci]; vc = mt.rows[1].cells[ci]
            _dc_set_cell_bg(hc, _DC_ROW_ALT); _dc_set_margins(hc); _dc_set_margins(vc)
            hc.width = vc.width = _DInches(3.25)
            _dc_run(hc.paragraphs[0], h, bold=True, size_pt=7, color=_DC_GREY)
            _dc_run(vc.paragraphs[0], v, bold=True, size_pt=9)
        self._cant_split_table(mt)
        # NOTE: python-docx Table objects do NOT expose .paragraphs at the table
        # level — only individual cells do.  The previous line
        #   mt.paragraphs[0].paragraph_format.keep_with_next = True
        # raised AttributeError on every call, which was silently swallowed by
        # generate()'s outer try/except, causing generate() to always return
        # False and producing zero DOCX output.  Removed; _cant_split_table()
        # above already prevents rows from splitting across pages.

        doc.add_paragraph()

        # ── Generate wide image + zoom crop bytes ──────────────────────────────
        wide_bytes = self._make_wide_bytes(ir, ann)
        zoom_bytes = self._make_zoom_bytes(ir, ann)

        # v3.2.0: Mini blade diagram — only for A/B/C blades, never Hub/Tower.
        # Position comes from ann.pinpoint_blade_pos (user-selected in Location tab).
        _is_blade_dp = blade_key in ("A", "B", "C")
        import io as _io
        if _is_blade_dp:
            _blade_pos = getattr(ann, "pinpoint_blade_pos", None)
            if _blade_pos is None:
                _bl_mm = getattr(self._project, "blade_length_mm", 50_000.0) or 50_000.0
                _blade_pos = (min(1.0, ann.root_distance_m * 1000 / _bl_mm)
                              if ann.root_distance_m else 0.5)
            _dist_for_mini = _blade_pos * (getattr(self._project, "blade_length_mm", 50_000.0) or 50_000.0)
            mini_bytes = self._make_mini_blade_bytes(blade_key, face_abbr, _dist_for_mini, ann.severity)
        else:
            mini_bytes = None

        # ── Image table: height-capped so 2 defects fit per page ─────────────
        # CHG-G: Max image height increased 2.0"→3.0" for better visibility.
        # Increased further from 2.5" to 3.0" for larger vertical display per user request.
        MAX_IMG_H_IN = 3.0

        def _add_img_sized(cell, img_bytes, max_w_in, max_h_in=MAX_IMG_H_IN):
            """Insert image capped to max_w×max_h; use height constraint when image is tall."""
            if not img_bytes:
                _dc_run(cell.paragraphs[0], "[ no image ]", italic=True, size_pt=8, color=_DC_GREY)
                return
            try:
                from PIL import Image as _PILI
                _sz = _PILI.open(_io.BytesIO(img_bytes)).size
                _ar = _sz[0] / _sz[1]   # width / height aspect ratio
                # Determine which dimension is the binding constraint
                if max_w_in / _ar > max_h_in:
                    # Tall image — constrain height
                    cell.paragraphs[0].alignment = _DAP.CENTER
                    cell.paragraphs[0].add_run().add_picture(
                        _io.BytesIO(img_bytes), height=_DInches(max_h_in))
                else:
                    cell.paragraphs[0].alignment = _DAP.CENTER
                    cell.paragraphs[0].add_run().add_picture(
                        _io.BytesIO(img_bytes), width=_DInches(max_w_in))
            except Exception:
                try:
                    cell.paragraphs[0].alignment = _DAP.CENTER
                    cell.paragraphs[0].add_run().add_picture(
                        _io.BytesIO(img_bytes), height=_DInches(max_h_in))
                except Exception:
                    _dc_run(cell.paragraphs[0], "[ no image ]", italic=True, size_pt=8, color=_DC_GREY)

        if _is_blade_dp and mini_bytes:
            # CHG-G: Increased WIDE_IN 3.5→3.8, ZOOM_IN 1.8→2.0 for better visibility
            WIDE_IN = 3.8; ZOOM_IN = 2.0; MINI_IN = 0.45
            img_t = doc.add_table(rows=1, cols=3)
            img_t.alignment = _DTA.LEFT
            cells = img_t.rows[0].cells
            cells[0].width = _DInches(WIDE_IN)
            cells[1].width = _DInches(ZOOM_IN)
            cells[2].width = _DInches(MINI_IN)
            _add_img_sized(cells[0], wide_bytes, WIDE_IN)
            _add_img_sized(cells[1], zoom_bytes, ZOOM_IN)
            _add_img_sized(cells[2], mini_bytes, MINI_IN)
        else:
            # Hub/Tower — wide + zoom only; CHG-G: 3.8→4.2, 2.2→2.3
            WIDE_IN = 4.2; ZOOM_IN = 2.3
            img_t = doc.add_table(rows=1, cols=2)
            img_t.alignment = _DTA.LEFT
            cells = img_t.rows[0].cells
            cells[0].width = _DInches(WIDE_IN)
            cells[1].width = _DInches(ZOOM_IN)
            _add_img_sized(cells[0], wide_bytes, WIDE_IN)
            _add_img_sized(cells[1], zoom_bytes, ZOOM_IN)

        self._cant_split_table(img_t)
        # CHG-B: One empty paragraph between image table and pill table (one line space)
        doc.add_paragraph()

        # ── v3.3.5 DOC-7: Defect name block variables ────────────────────────
        # _hdr_text removed — the bold heading line was replaced by the 2-col
        # [Defect Name | Serial] table below.  Only _serial_d, _comp_d, _wtg_d,
        # _face_d are needed.
        comp_labels_d = {"A": "Blade A", "B": "Blade B", "C": "Blade C",
                         "Hub": "Hub", "Tower": "Tower"}
        _comp_d    = comp_labels_d.get(ir.blade or "", ir.blade or "?")
        _raw_wtg_d = self._project.turbine_id or self._project.name or "WTG"
        _wtg_d     = (_raw_wtg_d if _raw_wtg_d.upper().startswith("WTG")
                      else f"WTG-{_raw_wtg_d}")
        # v4.5.0: Extract both surface and zone abbreviations
        _surf_abbr_d = (ann.surface.split("(")[-1].strip(")") if "(" in ann.surface else ann.surface)
        _zone_abbr_d = (ann.zone.split("(")[-1].strip(")") if "(" in ann.zone else ann.zone)
        _serial_d  = ann.serial_number or _generate_defect_serial(
            self._project, ann.blade, ann.surface, ann.zone)
        # chip row [#N][WTG][Component] removed per spec (CHG-1c)

        # MODIFICATION: Defect Name table removed per user request
        # The "Defect Name | Serial" table has been removed from DOCX defect pages

        # ── Severity badge + pill strip ────────────────────────────────────────
        # Bug #8 FIX: expand from 4→5 cols; add Tip Dist alongside Root Dist.
        # v3.3.5 DOC-9: label row uses bg #F4F4F4 and text #878787, NOT bold —
        # matches reference DOCX Table 4 row 0 exactly.
        sev_bg = _DC_SEV_RGB.get(ann.severity, _DRGBColor(0x71, 0x82, 0x96))
        sev_num = SEVERITY_SHORT.get(ann.severity or "", "—")
        _DC_PILL_LBL_BG  = _DRGBColor(0xF4, 0xF4, 0xF4)  # #F4F4F4 from reference
        _DC_PILL_LBL_TXT = _DRGBColor(0x87, 0x87, 0x87)  # #878787 from reference

        bt = doc.add_table(rows=2, cols=5); bt.alignment = _DTA.LEFT
        _is_blade_dd = (ir.blade or "") in ("A", "B", "C")
        _root_m_dd   = ann.root_distance_m
        _tip_m_dd    = ann.tip_distance_m
        _dist_lbl_r  = "Root Dist." if _is_blade_dd else "—"
        _dist_lbl_t  = "Tip Dist."  if _is_blade_dd else "—"
        _root_val    = (f"{_root_m_dd:.1f} m" if (_is_blade_dd and _root_m_dd) else "—")
        _tip_val     = (f"{_tip_m_dd:.1f} m"  if (_is_blade_dd and _tip_m_dd)  else "—")
        bhdrs = ["Issue Type", "Severity", _dist_lbl_r, _dist_lbl_t, "Size"]
        bvals = [ann.defect or "—", f"  {sev_num}  ", _root_val, _tip_val, size_s]
        bws   = [1.8, 1.0, 1.2, 1.2, 1.3]
        for ci, (bh, bv, bw) in enumerate(zip(bhdrs, bvals, bws)):
            hc = bt.rows[0].cells[ci]; vc = bt.rows[1].cells[ci]
            hc.width = vc.width = _DInches(bw)
            # v3.3.5 DOC-9: label row bg #F4F4F4, text #878787, not bold
            _dc_set_cell_bg(hc, _DC_PILL_LBL_BG); _dc_set_margins(hc); _dc_set_margins(vc)
            _dc_run(hc.paragraphs[0], bh, bold=False, size_pt=7, color=_DC_PILL_LBL_TXT)
            hc.paragraphs[0].alignment = _DAP.CENTER  # all header labels → centred
            if ci == 1:  # severity — coloured bg, white bold text, larger font
                _dc_set_cell_bg(vc, sev_bg)
                _dc_run(vc.paragraphs[0], bv, bold=True, size_pt=10, color=_DC_WHITE)
                vc.paragraphs[0].alignment = _DAP.CENTER
            elif ci == 0:
                # CHG-C: Issue Type (col 0) value LEFT-aligned — text can be long
                _dc_run(vc.paragraphs[0], bv, size_pt=8,
                        color=_DRGBColor(0x21, 0x21, 0x21))
                vc.paragraphs[0].alignment = _DAP.LEFT
            else:
                # CHG-C: Short value cells (Root/Tip/Size) → CENTRE-aligned
                _dc_run(vc.paragraphs[0], bv, size_pt=8,
                        color=_DRGBColor(0x21, 0x21, 0x21))
                vc.paragraphs[0].alignment = _DAP.CENTER
        self._cant_split_table(bt)
        # CHG-H: NO paragraph between pill table and comments table
        # (tables are directly adjacent — spacer removed per user spec)

        # ── Comments + Remedy Action ───────────────────────────────────────────
        # v3.3.5 DOC-10: both header cells bg #F4F4F4 (was left=lightblue),
        # plain bold text "Comments" / "Remedy Action" — no emoji (matches reference)
        # v3.3.13 FIX: Always prefer SEVERITY_REMEDY based on severity.
        # Ignore stored remedy if it matches ANY auto-generated pattern.
        remedy = SEVERITY_REMEDY.get(ann.severity or "", _auto_remedy(ann.defect or ""))
        if ann.remedy_action and ann.remedy_action.strip():
            stored = ann.remedy_action.strip()
            # Check if stored text matches auto-remedy pattern or severity text
            is_auto_pattern = stored.endswith("repair recommended during the next planned inspection.")
            is_severity_text = stored in SEVERITY_REMEDY.values()
            # Only use stored text if it's truly custom
            if not is_auto_pattern and not is_severity_text:
                remedy = stored
        notes  = ann.notes or "—"
        inset  = doc.add_table(rows=1, cols=2)
        inset.style = "Table Grid"; inset.alignment = _DTA.LEFT
        lc = inset.rows[0].cells[0]; rc = inset.rows[0].cells[1]
        lc.width = _DInches(3.0); rc.width = _DInches(3.5)
        _DC_COMMENTS_HDR_BG = _DRGBColor(0xF4, 0xF4, 0xF4)  # #F4F4F4 from reference
        _dc_set_cell_bg(lc, _DC_COMMENTS_HDR_BG)
        _dc_set_cell_bg(rc, _DC_COMMENTS_HDR_BG)
        _dc_set_margins(lc, 80, 80, 80, 80); _dc_set_margins(rc, 80, 80, 120, 120)
        _dc_run(lc.paragraphs[0], "Comments",      bold=True, size_pt=8)
        lc.paragraphs[0].alignment = _DAP.CENTER  # header label small → centred
        lc.add_paragraph().add_run(notes).font.size = _DPt(8)   # body text left (default)
        _dc_run(rc.paragraphs[0], "Remedy Action", bold=True, size_pt=8)
        rc.paragraphs[0].alignment = _DAP.CENTER  # header label small → centred
        rc.add_paragraph().add_run(remedy).font.size = _DPt(8)  # body text left (default)
        self._cant_split_table(inset)

        # v4.2.0: REMOVED calibration metadata display per user request
        # Calibration info no longer shown in DOCX reports for cleaner presentation

    # ── DOCX image helpers ─────────────────────────────────────────────────────

    def _resolve_filepath(self, ir: "ImageRecord") -> "Optional[str]":
        """Resolve ir.filepath with fallback to project_folder subdirectory scan."""
        fp = ir.filepath
        if fp and os.path.exists(fp):
            return fp
        pf = getattr(self._project, "project_folder", "")
        fname = ir.filename or (os.path.basename(fp) if fp else "")
        if pf and fname:
            candidate = os.path.join(pf, fname)
            if os.path.exists(candidate):
                return candidate
            try:
                for entry in os.scandir(pf):
                    if entry.is_dir():
                        c2 = os.path.join(entry.path, fname)
                        if os.path.exists(c2):
                            return c2
            except Exception:
                pass
        return None

    def _sev_rgb_pil(self, severity: str) -> tuple:
        """Return (r, g, b) for a severity string, safe for PIL.
        Uses _SEV_HEX (hex strings) — avoids docx RGBColor type mismatch."""
        h = _SEV_HEX.get(severity, "#FFC107").lstrip("#")
        if len(h) == 6:
            return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        return 255, 193, 7

    def _make_wide_bytes(self, ir: "ImageRecord", ann: "Annotation") -> Optional[bytes]:
        """Annotated full image → JPEG bytes."""
        fp = self._resolve_filepath(ir)
        if not fp:
            return None
        try:
            import io as _io
            with Image.open(fp) as src:
                img = src.convert("RGB")
            iw, ih = img.size
            draw = ImageDraw.Draw(img, "RGBA")
            try:
                font_sm = ImageFont.truetype("arial.ttf", max(10, iw // 120))
            except Exception:
                font_sm = ImageFont.load_default()
            for a in ir.annotations:
                r_, g_, b_ = self._sev_rgb_pil(a.severity)
                outline = (r_, g_, b_, 255); fill_t = (r_, g_, b_, 55)
                if a.mode == "box":
                    x1 = min(a.x1_px, a.x2_px); y1 = min(a.y1_px, a.y2_px)
                    x2 = max(a.x1_px, a.x2_px); y2 = max(a.y1_px, a.y2_px)
                    rot = getattr(a, "rotation_deg", 0.0) or 0.0
                    if abs(rot) > 0.5:
                        corners = _rotated_box_corners(x1, y1, x2, y2, rot)
                        draw.polygon(corners, outline=outline, fill=fill_t)
                        draw.line(corners + [corners[0]], fill=outline, width=3)
                    else:
                        draw.rectangle([x1, y1, x2, y2], outline=outline, fill=fill_t, width=3)
                    lbl = (f"{a.width_cm:.1f}×{a.height_cm:.1f}cm"
                           if a.width_cm else a.defect or "")
                    if lbl:
                        try:
                            bb = font_sm.getbbox(lbl); lw, lh = bb[2]-bb[0]+6, bb[3]-bb[1]+4
                        except Exception:
                            lw, lh = len(lbl)*7, 14
                        ly = max(0, y1 - lh - 2)
                        draw.rectangle([x1, ly, x1+lw, ly+lh], fill=(r_, g_, b_, 200))
                        draw.text((x1+3, ly+2), lbl, fill=(255, 255, 255, 255), font=font_sm)
                elif a.mode == "pin":
                    cx, cy, rp = int(a.x1_px), int(a.y1_px), 14
                    draw.ellipse([cx-rp, cy-rp, cx+rp, cy+rp],
                                 outline=outline, fill=(r_, g_, b_, 160), width=3)
                elif a.mode == "polygon" and len(a.poly_pts) >= 6:
                    pts = [(a.poly_pts[i], a.poly_pts[i+1])
                           for i in range(0, len(a.poly_pts) - 1, 2)]
                    draw.polygon(pts, fill=fill_t, outline=outline)
            buf = _io.BytesIO(); img.save(buf, "JPEG", quality=85)
            buf.seek(0); return buf.read()
        except Exception as exc:
            log.warning(f"_make_wide_bytes {fp}: {exc}"); return None

    def _make_zoom_bytes(self, ir: "ImageRecord", ann: "Annotation") -> Optional[bytes]:
        """Zoomed crop of defect bbox with 20% padding → JPEG bytes.

        T03 FIX: cm dimension overlay at bottom (DOCX parity with PDF _render_zoom_crop).
        SEV_HEX used instead of _DC_SEV_RGB to avoid RGBColor/int conversion errors.

        FIX-9 — ROTATION: two bugs corrected to match _render_zoom_crop (PDF path):

          Bug A — CROP REGION WRONG FOR ROTATED BOXES:
            x1/y1/x2/y2 stored by _commit_geometry are the *unrotated* half-extents
            centred on the rotation centre (centre ± w/2, ± h/2).  For a rotated box
            the corners extend beyond that axis-aligned envelope, so the previous
            min/max crop missed the actual visual shape corners.
            Fix: call _rotated_box_bounds() to get the true pixel envelope before
            computing cx1/cy1/cx2/cy2.  The unrotated originals are preserved in
            _ox1/_oy1/_ox2/_oy2 for the corner-drawing step (Bug B).

          Bug B — DRAWN OUTLINE WAS ALWAYS AXIS-ALIGNED:
            draw.rectangle([x1-cx1, y1-cy1, ...]) ignores rotation_deg entirely —
            it paints a straight box even for a 45° annotation.
            Fix: for abs(rot) > 0.5, compute the four rotated corners via
            _rotated_box_corners() using the unrotated stored coords, translate each
            corner into crop-local space (subtract cx1, cy1), then paint as
            draw.polygon() + draw.line() — identical to _render_zoom_crop.
        """
        fp = self._resolve_filepath(ir)
        if not fp:
            return None
        try:
            import io as _io
            with Image.open(fp) as src:
                img = src.convert("RGB")
            iw, ih = img.size

            # ── Determine raw annotation bounding coords ───────────────────────
            if ann.mode == "box":
                # Preserve unrotated originals for corner drawing (FIX-9 Bug B)
                _ox1 = min(ann.x1_px, ann.x2_px); _oy1 = min(ann.y1_px, ann.y2_px)
                _ox2 = max(ann.x1_px, ann.x2_px); _oy2 = max(ann.y1_px, ann.y2_px)
                x1, y1, x2, y2 = _ox1, _oy1, _ox2, _oy2
                # FIX-9 Bug A: replace with rotated bounding envelope so the crop
                # covers the full visual extent of the tilted shape.
                _rot = getattr(ann, "rotation_deg", 0.0) or 0.0
                if abs(_rot) > 0.5:
                    x1, y1, x2, y2 = _rotated_box_bounds(_ox1, _oy1, _ox2, _oy2, _rot)
            elif ann.mode == "polygon" and len(ann.poly_pts) >= 6:
                xs = ann.poly_pts[0::2]; ys = ann.poly_pts[1::2]
                x1, y1, x2, y2 = min(xs), min(ys), max(xs), max(ys)
                _rot = 0.0
            elif ann.mode == "pin":
                r = 60
                x1 = max(0, ann.x1_px - r); y1 = max(0, ann.y1_px - r)
                x2 = min(iw, ann.x1_px + r); y2 = min(ih, ann.y1_px + r)
                _rot = 0.0
            else:
                return None

            # ── Crop region: 20% padding around the (possibly rotated) envelope ─
            pad_x = max(int((x2 - x1) * 0.2), 10)
            pad_y = max(int((y2 - y1) * 0.2), 10)
            cx1 = max(0, int(x1) - pad_x); cy1 = max(0, int(y1) - pad_y)
            cx2 = min(iw, int(x2) + pad_x); cy2 = min(ih, int(y2) + pad_y)
            crop = img.crop((cx1, cy1, cx2, cy2))
            cw, ch = crop.size
            draw = ImageDraw.Draw(crop, "RGBA")

            # ── Draw annotation outline on the crop ───────────────────────────
            if ann.mode != "pin":
                r_, g_, b_ = self._sev_rgb_pil(ann.severity)
                outline_c = (r_, g_, b_, 255)
                if ann.mode == "box" and abs(_rot) > 0.5:
                    # FIX-9 Bug B: rotated box → polygon drawn with rotated corners.
                    # Use the original *unrotated* stored coords (_ox1/_oy1/_ox2/_oy2)
                    # so _rotated_box_corners computes the correct rotation from centre;
                    # then translate each corner into crop-local pixel space.
                    _corners_abs = _rotated_box_corners(_ox1, _oy1, _ox2, _oy2, _rot)
                    _corners_crop = [(px - cx1, py - cy1) for px, py in _corners_abs]
                    draw.polygon(_corners_crop, outline=outline_c)
                    # draw.line provides explicit stroke width (polygon outline is 1px)
                    draw.line(_corners_crop + [_corners_crop[0]],
                              fill=outline_c, width=4)
                elif ann.mode == "polygon" and len(ann.poly_pts) >= 6:
                    pts_crop = [(ann.poly_pts[i] - cx1, ann.poly_pts[i+1] - cy1)
                                for i in range(0, len(ann.poly_pts) - 1, 2)]
                    draw.polygon(pts_crop, outline=outline_c)
                    draw.line(pts_crop + [pts_crop[0]], fill=outline_c, width=2)
                else:
                    # Non-rotated box or polygon fallback: simple axis-aligned rect
                    bx1 = x1 - cx1; by1 = y1 - cy1
                    bx2 = x2 - cx1; by2 = y2 - cy1
                    draw.rectangle([bx1, by1, bx2, by2], outline=outline_c, width=4)

            # T03 FIX: cm dimension overlay at bottom
            if ann.width_cm and ann.height_cm:
                dim_txt = f"{ann.width_cm:.1f} × {ann.height_cm:.1f} cm"
                banner_h = max(18, int(ch * 0.08))
                draw.rectangle([(0, ch - banner_h), (cw, ch)], fill=(0, 0, 0, 160))
                try:
                    from PIL import ImageFont as _IFont
                    _fnt = _IFont.load_default()
                except Exception:
                    _fnt = None
                draw.text((6, ch - banner_h + 3), dim_txt,
                          fill=(255, 255, 255, 240), font=_fnt)

            buf = _io.BytesIO(); crop.save(buf, "JPEG", quality=90)
            buf.seek(0); return buf.read()
        except Exception as exc:
            log.warning(f"_make_zoom_bytes {fp}: {exc}"); return None


    def _make_mini_blade_bytes(self, blade_key: str, face_abbr: str,
                               dist_mm: float, severity: str) -> Optional[bytes]:
        """Mini single-blade silhouette with pinpoint → PNG bytes."""
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as _plt
            import numpy as _np
            import io as _io

            blade_length_mm = getattr(self._project, "blade_length_mm", 50_000.0) or 50_000.0
            SEV_COL = {
                **_SEV_HEX,
            }
            dot_col = SEV_COL.get(severity, "#F97316")

            fig, ax = _plt.subplots(figsize=(1.1, 4.0))
            fig.patch.set_facecolor("#f8f9fa"); ax.set_facecolor("#f8f9fa")
            y_vals = _np.linspace(0, 1, 200)
            hw = 0.44 * _np.sqrt(_np.maximum(0.0, 1.0 - y_vals))
            ax.fill_betweenx(y_vals, -hw, hw, color="#dde3ea", alpha=0.85, zorder=1)
            ax.plot(-hw, y_vals, color="#aab4be", lw=0.8, zorder=2)
            ax.plot( hw, y_vals, color="#aab4be", lw=0.8, zorder=2)
            # FIX: guard is division-by-zero only (blade_length_mm==0), NOT
            # dist_mm==0.  The previous `if dist_mm > 0` mapped any root-end
            # annotation (dist_mm=0.0) to the midpoint (yn=0.5), silently
            # ignoring the user's explicit pinpoint placement in the Location tab.
            yn = min(1.0, max(0.0, dist_mm / blade_length_mm)) if blade_length_mm > 0 else 0.5
            ax.scatter(0, yn, s=100, color=dot_col, edgecolors="white", linewidths=1.0, zorder=5)
            ax.text(0, -0.06, face_abbr or "?", ha="center", va="bottom",
                    fontsize=6, fontweight="bold", color="#333")
            ticks_mm = [0, blade_length_mm * 0.5, blade_length_mm]
            ax.set_yticks([t / blade_length_mm for t in ticks_mm])
            ax.set_yticklabels([f"{int(t/1000)}m" for t in ticks_mm], fontsize=5, color="#777")
            ax.yaxis.set_tick_params(length=0)
            ax.set_ylim(1.06, -0.14); ax.set_xlim(-0.65, 0.65); ax.set_xticks([])
            ax.spines[["top", "right", "bottom"]].set_visible(False)
            ax.spines["left"].set_color("#ccc")
            fig.tight_layout(pad=0.3)
            buf = _io.BytesIO()
            fig.savefig(buf, format="png", dpi=130, bbox_inches="tight")
            _plt.close(fig); buf.seek(0); return buf.read()
        except Exception as exc:
            log.warning(f"_make_mini_blade_bytes: {exc}"); return None

    def _count_by_severity(self) -> Dict[str, int]:
        counts: Dict[str, int] = {}
        for ir in self._project.images.values():
            for ann in ir.annotations:
                k = ann.severity or "Unknown"
                counts[k] = counts.get(k, 0) + 1
        return counts


class ReportSettingsDialog(QDialog):
    """
    v3.3.2: Full per-site report settings — 3 tabs covering identity/branding,
    turbine specifications, and GL-16 narrative text.

    Persistence strategy:
      • Project-level: stored in project.summary_notes as a __rpt_settings__: JSON line.
      • Global defaults: stored in settings.ini [REPORT_DEFAULTS] when the user
        enables "Save as global defaults" toggle. Global defaults are loaded as
        fallback whenever a project has no project-level settings saved yet.
      • The toggle state itself is persisted in settings.ini [UI] GlobalReportSettings.
    """

    _GLOBAL_SECTION = "REPORT_DEFAULTS"
    _GLOBAL_KEY     = "json"
    _TOGGLE_SECTION = "UI"
    _TOGGLE_KEY     = "GlobalReportSettings"

    def __init__(self, project: "Project", parent=None):
        super().__init__(parent)
        self.setWindowTitle("Report Settings")
        self.setMinimumSize(600, 660)
        self._project  = project
        self._settings = self._load()
        self._build_ui()

    # ── Persistence ────────────────────────────────────────────────────────────

    @staticmethod
    def _read_global() -> Dict[str, str]:
        """Read global defaults from settings.ini [REPORT_DEFAULTS]."""
        import json as _json
        try:
            raw = CFG.get(ReportSettingsDialog._GLOBAL_SECTION,
                          ReportSettingsDialog._GLOBAL_KEY, "")
            if raw:
                return _json.loads(raw)
        except Exception:
            pass
        return {}

    @staticmethod
    def _write_global(s: Dict[str, str]):
        """Write global defaults to settings.ini [REPORT_DEFAULTS]."""
        import json as _json
        try:
            if not CFG._cfg.has_section(ReportSettingsDialog._GLOBAL_SECTION):
                CFG._cfg.add_section(ReportSettingsDialog._GLOBAL_SECTION)
            CFG._cfg.set(ReportSettingsDialog._GLOBAL_SECTION,
                         ReportSettingsDialog._GLOBAL_KEY, _json.dumps(s))
            CFG.save()
        except Exception as exc:
            log.warning(f"ReportSettingsDialog._write_global: {exc}")

    @staticmethod
    def _global_toggle_on() -> bool:
        """Return True if the 'save as global' toggle is enabled."""
        return CFG.get(ReportSettingsDialog._TOGGLE_SECTION,
                       ReportSettingsDialog._TOGGLE_KEY, "false").lower() == "true"

    @staticmethod
    def _set_global_toggle(value: bool):
        """Persist the toggle state."""
        try:
            if not CFG._cfg.has_section(ReportSettingsDialog._TOGGLE_SECTION):
                CFG._cfg.add_section(ReportSettingsDialog._TOGGLE_SECTION)
            CFG._cfg.set(ReportSettingsDialog._TOGGLE_SECTION,
                         ReportSettingsDialog._TOGGLE_KEY,
                         "true" if value else "false")
            CFG.save()
        except Exception as exc:
            log.warning(f"ReportSettingsDialog._set_global_toggle: {exc}")

    def _load(self) -> Dict[str, str]:
        """Load settings: project-level first, fall back to global defaults."""
        import json as _json
        for line in (self._project.summary_notes or "").splitlines():
            if line.startswith("__rpt_settings__:"):
                try:
                    return _json.loads(line[len("__rpt_settings__:"):])
                except Exception:
                    pass
        # No project-level settings → try global defaults
        return self._read_global()

    def _save(self, s: Dict[str, str], also_global: bool):
        """Save settings to project and optionally to global defaults."""
        import json as _json
        # Always update project-level settings
        lines = [l for l in (self._project.summary_notes or "").splitlines()
                 if not l.startswith("__rpt_settings__:")]
        self._project.summary_notes = (
            "\n".join(lines).rstrip("\n") +
            f"\n__rpt_settings__:{_json.dumps(s)}"
        ).lstrip("\n")
        self._settings = s
        # Optionally persist as global defaults
        if also_global:
            self._write_global(s)
        self._set_global_toggle(also_global)

    # ── UI ─────────────────────────────────────────────────────────────────────

    def _build_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(16, 12, 16, 12); root.setSpacing(10)

        hdr = QLabel(
            "Settings are saved with the project and applied to every PDF/DOCX report.<br/>"
            "<small style='color:#888'>Changes take effect on next report generation.</small>")
        hdr.setWordWrap(True)
        root.addWidget(hdr)

        tabs = QTabWidget()
        tabs.addTab(self._tab_identity(),   "🏢  Identity & Branding")
        tabs.addTab(self._tab_turbine(),    "⚙  Turbine Specifications")
        tabs.addTab(self._tab_narrative(),  "📝  Report Narrative")
        root.addWidget(tabs)

        # T02: Global-defaults toggle — persists settings across project folders
        self._global_chk = QCheckBox(
            "💾  Save as global defaults (pre-fill these settings for every new project)")
        self._global_chk.setChecked(self._global_toggle_on())
        self._global_chk.setStyleSheet(
            f"color:{UI_THEME['text_primary']};font-size:9pt;")
        self._global_chk.setToolTip(
            "When checked, your filled-in settings are saved globally in settings.ini.\n"
            "Any project that hasn't overridden them will inherit these values.\n"
            "Uncheck to keep settings project-only.")
        root.addWidget(self._global_chk)

        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok |
            QDialogButtonBox.StandardButton.Cancel)
        btns.accepted.connect(self._on_ok)
        btns.rejected.connect(self.reject)
        root.addWidget(btns)

    # ── Helper: labelled line edit ─────────────────────────────────────────────

    @staticmethod
    def _le(settings: Dict, key: str, placeholder: str = "") -> QLineEdit:
        w = QLineEdit(settings.get(key, ""))
        if placeholder:
            w.setPlaceholderText(placeholder)
        return w

    @staticmethod
    def _logo_row(le: QLineEdit, parent: "ReportSettingsDialog") -> QHBoxLayout:
        row = QHBoxLayout()
        btn = QPushButton("Browse…")
        btn.setFixedWidth(80)
        btn.clicked.connect(lambda: (
            fn := QFileDialog.getOpenFileName(parent, "Select Image", "", "Images (*.png *.jpg)")[0],
            le.setText(fn) if fn else None
        ))
        row.addWidget(le); row.addWidget(btn)
        return row

    # ── Tab 1: Identity & Branding ─────────────────────────────────────────────

    def _tab_identity(self) -> QWidget:
        s = self._settings
        w = QWidget(); form = QFormLayout(w)
        form.setSpacing(10); form.setContentsMargins(12, 12, 12, 12)

        self._f_company   = self._le(s, "company",         "e.g. SenseHawk Wind Services")
        self._f_client    = self._le(s, "client",          "e.g. Sekura Energy Ltd.")
        self._f_reviewer  = self._le(s, "reviewer_name",   "e.g. Jane Smith")
        self._f_gps       = self._le(s, "gps_coords",      "e.g. 15.123456, 74.567890")

        # v4.3.0: Added scan_date, turbine_manufacturer, rated_power — editable from settings
        # so users can update these fields without creating a new project.
        self._f_scan_date = self._le(s, "scan_date",           "e.g. 15 Feb 2026")
        self._f_turb_mfr  = self._le(s, "turbine_manufacturer","e.g. Vestas, Siemens Gamesa")
        self._f_rated_pwr = self._le(s, "rated_power",         "e.g. 2.5 MW, 3.6 MW")

        # FIX-17d: Tower base GPS + MSL altitude for 3D distance model.
        # Pre-populate from project fields if already set; fall back to
        # settings blob so global defaults also work.
        _t_lat = (str(self._project.tower_lat)
                  if self._project.tower_lat is not None
                  else s.get("tower_lat", ""))
        _t_lon = (str(self._project.tower_lon)
                  if self._project.tower_lon is not None
                  else s.get("tower_lon", ""))
        _t_alt = (str(self._project.tower_base_alt_msl)
                  if self._project.tower_base_alt_msl is not None
                  else s.get("tower_base_alt_msl", ""))
        self._f_tower_lat = QLineEdit(_t_lat)
        self._f_tower_lat.setPlaceholderText("e.g. 15.1238  (decimal degrees, N positive)")
        self._f_tower_lon = QLineEdit(_t_lon)
        self._f_tower_lon.setPlaceholderText("e.g. 74.5680  (decimal degrees, E positive)")
        self._f_tower_alt = QLineEdit(_t_alt)
        self._f_tower_alt.setPlaceholderText("e.g. 100.0  (metres above mean sea level)")

        self._f_logo      = self._le(s, "logo_path",        "Path to company header logo (.png)")
        self._f_cli_logo  = self._le(s, "client_logo_path", "Path to client cover logo (.png)")
        self._f_co_logo   = self._le(s, "company_logo_path","Path to company footer logo (.png)")
        # WTG front-page image — shown on cover page instead of placeholder box
        self._f_wtg_img   = self._le(s, "wtg_image_path",   "Path to WTG front photo (.jpg/.png)")

        form.addRow("Inspection Company:", self._f_company)
        form.addRow("Client / Site Owner:", self._f_client)
        form.addRow("Lead Reviewer:",       self._f_reviewer)
        form.addRow("GPS Coordinates:",     self._f_gps)
        # v4.3.0: New identity fields
        form.addRow("Scan Date:",              self._f_scan_date)
        form.addRow("Turbine Manufacturer:",   self._f_turb_mfr)
        form.addRow("Rated Power:",            self._f_rated_pwr)
        # FIX-17d: Tower base location — used by 3D GSD distance model
        form.addRow(QLabel())  # visual spacer
        _tower_hdr = QLabel(
            "<b>Tower Base Location</b> &nbsp;"
            "<small style='color:#888'>— used for 3D GPS distance calibration</small>")
        _tower_hdr.setWordWrap(True)
        form.addRow(_tower_hdr)
        form.addRow("Tower Lat (DD):",          self._f_tower_lat)
        form.addRow("Tower Lon (DD):",          self._f_tower_lon)
        form.addRow("Tower Base Alt MSL (m):",  self._f_tower_alt)
        form.addRow(QLabel())  # spacer row
        form.addRow("WTG Cover Photo:",  self._logo_row_w(self._f_wtg_img))
        form.addRow("Header Logo:",      self._logo_row_w(self._f_logo))
        form.addRow("Cover Logo:",       self._logo_row_w(self._f_cli_logo))
        form.addRow("Footer Logo:",      self._logo_row_w(self._f_co_logo))
        return w

    def _logo_row_w(self, le: QLineEdit) -> QWidget:
        """Return a QWidget containing the line-edit + Browse button."""
        row_w = QWidget()
        row_l = QHBoxLayout(row_w)
        row_l.setContentsMargins(0, 0, 0, 0); row_l.setSpacing(4)
        btn = QPushButton("Browse…"); btn.setFixedWidth(80)
        btn.clicked.connect(lambda _=False, _le=le: self._browse_image(_le))
        row_l.addWidget(le); row_l.addWidget(btn)
        return row_w

    def _browse_image(self, target: QLineEdit):
        p, _ = QFileDialog.getOpenFileName(self, "Select Image", "", "Images (*.png *.jpg *.jpeg)")
        if p:
            target.setText(p)

    # ── Tab 2: Turbine Specifications ──────────────────────────────────────────

    def _tab_turbine(self) -> QWidget:
        s   = self._settings
        w   = QWidget()
        sc  = QScrollArea(); sc.setWidgetResizable(True)
        inner = QWidget(); form = QFormLayout(inner)
        form.setSpacing(10); form.setContentsMargins(12, 12, 12, 12)

        # All 11 spec fields — user overrides; PDF falls back to defaults if blank
        specs = [
            ("turbine_model",   "Turbine Model / Name",   "e.g. Siemens Gamesa G114-2.0MW"),
            ("rated_power",     "Rated Power",            "e.g. 2.0 MW (2,000 kW)"),
            ("rotor_diameter",  "Rotor Diameter",         "e.g. 114 m"),
            ("blade_length",    "Blade Length",           "e.g. 55.5 m (3 blades)"),
            ("swept_area",      "Swept Area",             "e.g. 10,207 m²"),
            ("wind_class",      "Wind Class",             "e.g. IEC IIIa (Low Wind)"),
            ("cut_in_speed",    "Cut-in Wind Speed",      "e.g. 2.5 m/s"),
            ("rated_speed",     "Rated Wind Speed",       "e.g. 10.0 – 12.5 m/s"),
            ("cut_out_speed",   "Cut-out Wind Speed",     "e.g. 25.0 m/s"),
            ("tower_height",    "Tower Height",           "e.g. 106 m"),
            ("max_rotor_speed", "Max Rotor Speed",        "e.g. ~16 RPM"),
            ("generator_type",  "Generator Type",         "e.g. DFIG"),
        ]
        self._spec_fields: Dict[str, QLineEdit] = {}
        for key, label, ph in specs:
            le = self._le(s, key, ph)
            self._spec_fields[key] = le
            form.addRow(label + ":", le)

        note = QLabel(
            "<small style='color:#888'>Leave blank to use the built-in default value.<br/>"
            "Applies to both PDF and DOCX reports.</small>")
        note.setWordWrap(True)
        form.addRow(note)

        sc.setWidget(inner)
        outer = QVBoxLayout(w)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.addWidget(sc)
        return w

    # ── Tab 3: Report Narrative ────────────────────────────────────────────────

    def _tab_narrative(self) -> QWidget:
        s  = self._settings
        w  = QWidget()
        sc = QScrollArea(); sc.setWidgetResizable(True)
        inner = QWidget(); lay = QVBoxLayout(inner)
        lay.setSpacing(10); lay.setContentsMargins(12, 12, 12, 12)

        note = QLabel(
            "<small style='color:#888'>Leave any field blank to use the default GL-16 "
            "template text.  Use blank lines to separate paragraphs.</small>")
        note.setWordWrap(True)
        lay.addWidget(note)

        narratives = [
            ("objective_text",       "Objective",
             "Override the Objective section (blank = use default GL-16 text)"),
            ("scope_text",           "Scope of Work",
             "Override the Scope of Work section"),
            ("data_collection_text", "Data Collection Methodology",
             "Override the Data Collection section (use blank line for paragraph break)"),
            ("results_text",         "Results Introduction",
             "Override the Results intro paragraph"),
            ("summary_text",         "Executive Summary (DOCX cover)",
             "Appears on the DOCX cover page summary block"),
        ]
        self._narrative_fields: Dict[str, QTextEdit] = {}
        for key, label, tooltip in narratives:
            lbl = QLabel(f"<b>{label}</b>")
            lbl.setStyleSheet("color:#ccc;font-size:9pt;")
            lay.addWidget(lbl)
            te = QTextEdit()
            te.setPlainText(s.get(key, ""))
            te.setPlaceholderText(tooltip)
            te.setFixedHeight(90)
            lay.addWidget(te)
            self._narrative_fields[key] = te

        sc.setWidget(inner)
        outer = QVBoxLayout(w)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.addWidget(sc)
        return w

    # ── Save ───────────────────────────────────────────────────────────────────

    def _on_ok(self):
        s: Dict[str, str] = {
            # Tab 1 — Identity
            "company":            self._f_company.text().strip(),
            "client":             self._f_client.text().strip(),
            "reviewer_name":      self._f_reviewer.text().strip(),
            "gps_coords":         self._f_gps.text().strip(),
            # v4.3.0: Scan date + turbine metadata editable from settings
            "scan_date":              self._f_scan_date.text().strip(),
            "turbine_manufacturer":   self._f_turb_mfr.text().strip(),
            "rated_power":            self._f_rated_pwr.text().strip(),
            # FIX-17d: Tower base GPS + MSL altitude
            "tower_lat":          self._f_tower_lat.text().strip(),
            "tower_lon":          self._f_tower_lon.text().strip(),
            "tower_base_alt_msl": self._f_tower_alt.text().strip(),
            "logo_path":          self._f_logo.text().strip(),
            "client_logo_path":   self._f_cli_logo.text().strip(),
            "company_logo_path":  self._f_co_logo.text().strip(),
            "wtg_image_path":     self._f_wtg_img.text().strip(),
        }
        # Tab 2 — Turbine specs
        for key, le in self._spec_fields.items():
            s[key] = le.text().strip()
        # Tab 3 — Narrative text
        for key, te in self._narrative_fields.items():
            s[key] = te.toPlainText().strip()

        # FIX-17d: Write tower GPS directly onto the project object so that
        # _batch_auto_calibrate() can read them without re-loading settings.
        # Invalid or blank values are stored as None (not as an error).
        def _parse_float_or_none(raw: str) -> Optional[float]:
            try:
                return float(raw) if raw else None
            except ValueError:
                return None

        self._project.tower_lat          = _parse_float_or_none(s["tower_lat"])
        self._project.tower_lon          = _parse_float_or_none(s["tower_lon"])
        self._project.tower_base_alt_msl = _parse_float_or_none(s["tower_base_alt_msl"])

        # T02: pass global flag so _save decides whether to write settings.ini
        also_global = self._global_chk.isChecked()
        self._save(s, also_global)
        self.accept()

    def get_settings(self) -> Dict[str, str]:
        return self._settings


# ==============================================================================
# PHASE 9.1 — KEYBOARD SHORTCUTS HELP DIALOG  (Jamie Liu)
# ==============================================================================

class ShortcutsDialog(QDialog):
    """Phase 9.1: Non-modal reference of all keyboard shortcuts."""
    _SHORTCUTS = [
        ("Project",     [
            ("Ctrl+N",       "New Project"),
            ("Ctrl+O",       "Open Project"),
            ("Ctrl+S",       "Save Project"),
            ("Ctrl+L",       "Load Images"),
            ("Ctrl+R",       "Generate PDF Report"),
            ("Ctrl+Q",       "Quit"),
        ]),
        ("Draw Modes",  [
            ("S",            "Select / Pan"),
            ("B",            "Box annotation"),
            ("P",            "Pin annotation"),
            ("G",            "Polygon annotation"),
            ("C",            "Calibrate (draw GSD reference line)"),
            ("Esc",          "Cancel polygon in progress"),
        ]),
        ("Annotation",  [
            ("Ctrl+Z",       "Undo last annotation"),
            ("Ctrl+Y",       "Redo annotation"),
            ("Delete",       "Delete selected annotation"),
        ]),
        ("Navigation",  [
            ("← / ↑",        "Previous image"),
            ("→ / ↓",        "Next image"),
        ]),
        ("ML & Review", [
            ("Ctrl+M",       "Open ML dialog (Detection & Training)"),
            ("Ctrl+K",       "Open QC Viewer"),
        ]),
        ("Blade Diagram",[
            ("Left-click",   "Drill-down: list annotations on cell"),
            ("Right-click",  "Copy diagram to clipboard"),
        ]),
        ("QC Viewer",   [
            ("Ctrl+Z",       "Undo QC checkbox change"),
            ("Ctrl+Y",       "Redo QC checkbox change"),
        ]),
    ]

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Keyboard Shortcuts")
        self.setMinimumSize(520, 580)
        self.setAttribute(Qt.WidgetAttribute.WA_DeleteOnClose, False)
        self._build_ui()

    def _build_ui(self):
        lay = QVBoxLayout(self)
        lay.setContentsMargins(20, 16, 20, 16); lay.setSpacing(8)
        lay.addWidget(QLabel(
            "<b>All Keyboard Shortcuts</b>  —  Wind Tower Inspection App"))

        scroll = QScrollArea(); scroll.setWidgetResizable(True)
        container = QWidget()
        c_lay = QVBoxLayout(container)
        c_lay.setContentsMargins(4, 4, 4, 4); c_lay.setSpacing(10)

        for group_name, shortcuts in self._SHORTCUTS:
            # Group header
            hdr = QLabel(group_name.upper())
            hdr.setStyleSheet(
                f"color:{UI_THEME['accent_cyan']};font-size:8pt;"
                f"font-weight:bold;letter-spacing:1px;background:transparent;")
            c_lay.addWidget(hdr)

            # Shortcut rows
            for keys, desc in shortcuts:
                row_w  = QWidget()
                row_l  = QHBoxLayout(row_w)
                row_l.setContentsMargins(4, 2, 4, 2)
                key_l  = QLabel(keys)
                key_l.setFixedWidth(130)
                key_l.setStyleSheet(
                    f"color:{UI_THEME['text_primary']};font-family:monospace;"
                    f"font-size:9pt;background:{UI_THEME['bg_elevated']};"
                    f"border:1px solid {UI_THEME['border']};"
                    f"border-radius:4px;padding:2px 6px;")
                desc_l = QLabel(desc)
                desc_l.setStyleSheet(
                    f"color:{UI_THEME['text_secondary']};font-size:9pt;"
                    f"background:transparent;")
                row_l.addWidget(key_l)
                row_l.addSpacing(8)
                row_l.addWidget(desc_l)
                row_l.addStretch()
                c_lay.addWidget(row_w)

            c_lay.addSpacing(4)

        scroll.setWidget(container)
        lay.addWidget(scroll)
        btns = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        btns.rejected.connect(self.reject)
        lay.addWidget(btns)


# ==============================================================================
# MAIN WINDOW  (Sarah Chen + Sam Okafor)
# ==============================================================================

class MainWindow(QMainWindow):
    """
    Sarah Chen + Sam Okafor:
    Orchestrates all widgets.  v1.2.1 additions:
      • Project header info bar (always visible)
      • 3-zone toolbar: Project | Drawing | ML+Export
      • QStackedWidget: annotation mode ↔ QC viewer
      • Toast notifications for non-critical feedback
      • Thumbnail borders coloured by worst annotation severity
      • Global annotation Undo (Ctrl+Z) / Redo (Ctrl+Y) wired to viewer
    """

    def __init__(self):
        super().__init__()
        self._session = SESSION   # global singleton
        self.setWindowTitle(
            f"{APP_TITLE}  v{APP_VERSION}")
        # SCREEN-FIT FIX: Check available screen height at runtime.
        # The accumulated minimum height of all panels (~1372px) exceeds a
        # typical laptop screen (1366×768, ~697px usable after taskbar).
        # Always starting maximized prevents the QWindowsWindow::setGeometry
        # warning and ensures the splitter / thumbnail dock are visible.
        # Users can restore/resize normally after the app is open.
        _screen = QApplication.primaryScreen()
        _avail_h = _screen.availableGeometry().height() if _screen else 768
        if _avail_h < 900:
            # Small screen — start maximized so Qt never has to clip the window.
            # We still call resize() as a hint for when the user later un-maximizes.
            self.resize(1366, _avail_h)
            self._start_maximized = True
        else:
            self.resize(1440, 920)
            self._start_maximized = False

        self._project          : Optional[Project]     = None
        self._current_rec      : Optional[ImageRecord] = None
        self._current_filepath : str                   = ""
        self._image_paths      : List[str]             = []
        self._thumb_pool       = QThreadPool()
        self._thumb_pool.setMaxThreadCount(4)

        self._build_header_bar()
        self._build_menu()
        self._build_toolbar()
        self._build_central()
        self._build_thumb_dock()
        self._build_statusbar()

        # Wire signals
        self._viewer.signals.annotation_ready.connect(self._on_annotation_ready)
        self._viewer.signals.annotation_selected.connect(self._on_annotation_selected)
        self._viewer.signals.annotation_deleted.connect(self._on_delete_annotation)
        self._viewer.signals.annotation_modified.connect(self._on_annotation_modified)
        self._viewer.signals.gsd_updated.connect(self._on_gsd_updated)
        self._viewer.signals.calibration_metadata.connect(self._store_exif_metadata)  # v3.4.0: EXIF metadata
        # v4.1.1: viewer Key_S → sync toolbar button without re-calling set_mode
        self._viewer.signals.mode_change_requested.connect(
            lambda m: self._mode_btns[m].setChecked(True))
        # v4.1.1: ann_list click in Annotate tab → also feed QC Review
        self._ann_panel.ann_selected_for_qc.connect(
            self._qc_review_panel.load_annotation)
        self._ann_panel.save_requested.connect(self._on_save_annotation)
        self._ann_panel.delete_requested.connect(self._on_delete_annotation)
        # FIX-UX: Discard — removes unsaved annotation from scene without persisting
        self._ann_panel.discard_requested.connect(self._on_discard_annotation)
        self._ann_panel.approve_requested.connect(self._on_approve_annotation)
        self._ann_panel.reject_requested.connect(self._on_reject_annotation)
        # v1.7.0: inline renamer
        self._ann_panel.rename_requested.connect(self._on_rename_file_from_panel)
        self._blade_diag.cell_clicked.connect(self._on_diagram_cell_click)
        self._blade_diag.ann_clicked.connect(self._on_diagram_ann_click)   # v3.2.0
        self._qc_widget.annotations_committed.connect(self._on_qc_committed)
        self._qc_widget.back_requested.connect(self._switch_to_annotation_mode)

        log.info(f"{APP_TITLE} v{APP_VERSION} started")

    # ── Header bar ─────────────────────────────────────────────────────────────

    def _build_header_bar(self):
        """Dev Patel: Always-visible slim info bar showing project context."""
        bar = QWidget()
        bar.setFixedHeight(32)
        bar.setStyleSheet(
            f"background:{UI_THEME['bg_toolbar']};"
            f"border-bottom:1px solid {UI_THEME['border']};")
        lay = QHBoxLayout(bar)
        lay.setContentsMargins(12, 0, 12, 0)
        lay.setSpacing(0)

        def _info_lbl(text: str, is_value: bool = False) -> QLabel:
            lbl = QLabel(text)
            lbl.setStyleSheet(
                f"color:{UI_THEME['accent_cyan'] if is_value else UI_THEME['text_tertiary']};"
                f"font-size:9pt;{'font-weight:bold;' if is_value else ''}"
                f"background:transparent;padding:0 4px;")
            return lbl

        lay.addWidget(_info_lbl("PROJECT "))
        self._hb_project  = _info_lbl("—", True)
        lay.addWidget(self._hb_project)

        lay.addWidget(_info_lbl("  ·  "))

        lay.addWidget(_info_lbl("SITE "))
        self._hb_site = _info_lbl("—", True)
        lay.addWidget(self._hb_site)

        lay.addWidget(_info_lbl("  ·  "))
        lay.addWidget(_info_lbl("INSPECTOR "))
        self._hb_inspector = _info_lbl("—", True)
        lay.addWidget(self._hb_inspector)

        lay.addStretch()

        # FIX-15a: Reinstated — was incorrectly removed by FIX-14c (v4.4.0).
        # The user's request was to remove text from the PDF report header, not
        # here.  GSD feedback in the top bar gives at-a-glance calibration
        # status without the user needing to look at the status bar.
        lay.addWidget(_info_lbl("GSD "))
        self._hb_gsd = _info_lbl("not calibrated", True)
        lay.addWidget(self._hb_gsd)

        self._header_dock = QDockWidget()
        self._header_dock.setTitleBarWidget(QWidget())  # hide title bar
        self._header_dock.setWidget(bar)
        self._header_dock.setFeatures(QDockWidget.DockWidgetFeature.NoDockWidgetFeatures)
        self.addDockWidget(Qt.DockWidgetArea.TopDockWidgetArea, self._header_dock)

    def _update_header_bar(self):
        if self._project:
            self._hb_project.setText(self._project.name or "—")
            self._hb_site.setText(self._project.site or "—")
            self._hb_inspector.setText(self._project.inspector or "—")
        else:
            self._hb_project.setText("—")
            self._hb_site.setText("—")
            self._hb_inspector.setText("—")

    # ── Menu ───────────────────────────────────────────────────────────────────

    def _build_menu(self):
        mb = self.menuBar()

        file_m = mb.addMenu("&File")
        acts = {
            "new":  QAction("&New Project…",          self, shortcut="Ctrl+N"),
            "open": QAction("&Open Project…",         self, shortcut="Ctrl+O"),
            "save": QAction("&Save Project",          self, shortcut="Ctrl+S"),
            "imgs": QAction("&Load Images…",          self, shortcut="Ctrl+L"),
            "rpt":  QAction("📄 &Generate PDF Report…", self, shortcut="Ctrl+R"),
            "jpeg": QAction("💾 Save Annotated JPEG",  self),
            "quit": QAction("&Quit",                  self, shortcut="Ctrl+Q"),
        }
        acts["new"].triggered.connect(self._new_project)
        acts["open"].triggered.connect(self._open_project)
        acts["save"].triggered.connect(self._save_project_action)
        acts["imgs"].triggered.connect(self._load_images)
        acts["rpt"].triggered.connect(self._generate_report)
        acts["jpeg"].triggered.connect(self._save_annotated_jpeg)
        acts["quit"].triggered.connect(self.close)
        acts["rpt"].setEnabled(True)
        acts["rpt"].setToolTip("Generate PDF report")
        rpt_settings_act = QAction("⚙ Report Settings…", self)
        rpt_settings_act.triggered.connect(self._open_report_settings)
        rpt_settings_act.setEnabled(True)
        for k in ["new", "open", "save", "imgs", None, "rpt", None, "jpeg", None, "quit"]:
            if k is None:
                file_m.addSeparator()
            else:
                file_m.addAction(acts[k])
        file_m.insertAction(acts["jpeg"], rpt_settings_act)

        edit_m = mb.addMenu("&Edit")
        undo_a = QAction("↩ &Undo Annotation", self, shortcut="Ctrl+Z")
        redo_a = QAction("↪ &Redo Annotation", self, shortcut="Ctrl+Y")
        undo_a.triggered.connect(self._undo_annotation)
        redo_a.triggered.connect(self._redo_annotation)
        edit_m.addAction(undo_a); edit_m.addAction(redo_a)

        tools_m = mb.addMenu("&Tools")
        gsd_a   = QAction("Set Session &GSD…",     self)
        types_a = QAction("Manage &Defect Types…", self)
        gsd_a.triggered.connect(self._set_session_gsd)
        types_a.triggered.connect(self._manage_defect_types)
        tools_m.addAction(gsd_a)
        tools_m.addAction(types_a)

        # v2.1.1: ML Safe Mode - hide menu if ultralytics not available
        if YOLO_AVAILABLE:
            ml_m  = mb.addMenu("🤖 &ML")
            ml_a  = QAction("&Detection && Training…", self, shortcut="Ctrl+M")
            qc_a  = QAction("🔍 &QC Viewer…",          self, shortcut="Ctrl+K")
            ml_a.triggered.connect(self._open_ml_dialog)
            qc_a.triggered.connect(self._launch_qc_guard)
            ml_m.addAction(ml_a); ml_m.addAction(qc_a)

        help_m = mb.addMenu("&Help")
        # v4.1.0: Comprehensive user guide
        userguide_a = QAction("📖  &User Guide…", self, shortcut="F1")
        userguide_a.triggered.connect(self._show_user_guide)
        shortcuts_a = QAction("⌨  &Keyboard Shortcuts…", self,
                              shortcut="Ctrl+/")
        shortcuts_a.triggered.connect(self._show_shortcuts)
        about_a = QAction("&About", self)
        about_a.triggered.connect(self._show_about)
        help_m.addAction(userguide_a)
        help_m.addSeparator()
        help_m.addAction(shortcuts_a)
        help_m.addSeparator()
        help_m.addAction(about_a)

        # Phase 9.3: Recent Projects submenu (populated on open)
        self._recent_menu = mb.addMenu("🕐 &Recent")
        self._refresh_recent_menu()

    # ── 3-Zone Toolbar ─────────────────────────────────────────────────────────

    def _build_toolbar(self):
        """Dev Patel: 3-zone dark toolbar — Project | Drawing | ML+Export."""
        tb = QToolBar("Main", self)
        tb.setMovable(False)
        tb.setFloatable(False)
        tb.setStyleSheet(
            f"QToolBar{{background:{UI_THEME['bg_toolbar']};"
            f"border-bottom:1px solid {UI_THEME['border']};"
            f"padding:4px 8px;spacing:4px;}}"
            f"QToolBar::separator{{"
            f"background:{UI_THEME['border']};width:1px;margin:4px 6px;}}"
        )
        self.addToolBar(Qt.ToolBarArea.TopToolBarArea, tb)

        def _btn(text: str, colour: str = "", tooltip: str = "",
                 checkable: bool = False, fw: int = 0) -> QPushButton:
            b = QPushButton(text)
            b.setToolTip(tooltip)
            b.setCheckable(checkable)
            bg = colour or UI_THEME["bg_card"]
            txt_col = "#0d1117" if colour else UI_THEME["text_primary"]
            b.setStyleSheet(f"""
                QPushButton {{
                    background: {bg};
                    color: {txt_col};
                    border: 1px solid {UI_THEME['border']};
                    border-radius: 6px;
                    padding: 5px {'10px' if not fw else '0'};
                    font-weight: 600;
                    font-size: 9pt;
                }}
                QPushButton:hover {{
                    background: {UI_THEME['bg_elevated']};
                    border-color: {UI_THEME['accent_cyan']};
                    color: {UI_THEME['accent_cyan']};
                }}
                QPushButton:checked {{
                    background: {UI_THEME['accent_cyan']};
                    color: #0d1117;
                    border-color: {UI_THEME['accent_cyan']};
                    font-weight: bold;
                }}
                QPushButton:disabled {{
                    background: {UI_THEME['bg_secondary']};
                    color: {UI_THEME['text_tertiary']};
                }}
            """)
            if fw:
                b.setFixedWidth(fw)
            return b

        # ── Zone 1: Project ─────────────────────────────────────────────────────
        z1_lbl = QLabel("  PROJECT  ")
        z1_lbl.setStyleSheet(
            f"color:{UI_THEME['text_tertiary']};font-size:7pt;font-weight:bold;"
            f"letter-spacing:1px;background:transparent;")
        tb.addWidget(z1_lbl)

        new_btn  = _btn("＋ New",   tooltip="New project  (Ctrl+N)", fw=72)
        open_btn = _btn("📂 Open",  tooltip="Open project  (Ctrl+O)", fw=80)
        save_btn = _btn("💾 Save",  tooltip="Save project  (Ctrl+S)", fw=72)
        load_btn = _btn("🖼 Images", tooltip="Load images  (Ctrl+L)", fw=90)
        new_btn.clicked.connect(self._new_project)
        open_btn.clicked.connect(self._open_project)
        save_btn.clicked.connect(self._save_project_action)
        load_btn.clicked.connect(self._load_images)
        for b in [new_btn, open_btn, save_btn, load_btn]:
            tb.addWidget(b)

        tb.addSeparator()

        # ── Zone 2: Drawing tools ───────────────────────────────────────────────
        z2_lbl = QLabel("  ANNOTATE  ")
        z2_lbl.setStyleSheet(z1_lbl.styleSheet())
        tb.addWidget(z2_lbl)

        self._mode_group = QButtonGroup(self)
        draw_modes = [
            (ImageViewer.MODE_SEL,  "🖱 Select",   "Select / pan (S)",           82),
            (ImageViewer.MODE_BOX,  "▭  Box",      "Draw bounding box (B)",      72),
            (ImageViewer.MODE_PIN,  "📌 Pin",      "Drop point pin (P)",          68),
            (ImageViewer.MODE_POLY, "⬡  Polygon",  "Draw polygon (double-click to close) (G)", 90),
        ]
        self._mode_btns: Dict[str, QPushButton] = {}
        for mode, label, tip, fw in draw_modes:
            btn = _btn(label, tooltip=tip, checkable=True, fw=fw)
            self._mode_group.addButton(btn)
            tb.addWidget(btn)
            self._mode_btns[mode] = btn
            btn.clicked.connect(lambda _, m=mode: self._set_viewer_mode(m))
        self._mode_btns[ImageViewer.MODE_SEL].setChecked(True)
        
        # v4.1.1: Both auto (EXIF) and manual (per-image) calibration buttons restored
        tb.addSeparator()
        self._auto_cal_btn = _btn("🤖 Auto-Calibrate All",
                                   UI_THEME["accent_green"],
                                   "Batch auto-calibrate all images using EXIF data (per-image)",
                                   fw=155)
        self._manual_cal_btn = _btn("📏 Calibrate Image",
                                     tooltip="Draw calibration line on current image (C) - per-image only",
                                     checkable=True, fw=140)
        self._mode_group.addButton(self._manual_cal_btn)
        self._mode_btns[ImageViewer.MODE_CAL] = self._manual_cal_btn
        self._auto_cal_btn.clicked.connect(self._batch_auto_calibrate)
        self._manual_cal_btn.clicked.connect(lambda: self._set_viewer_mode(ImageViewer.MODE_CAL))
        tb.addWidget(self._auto_cal_btn)
        tb.addWidget(self._manual_cal_btn)

        # Undo / Redo
        tb.addSeparator()
        undo_tb = _btn("↩", tooltip="Undo last annotation  (Ctrl+Z)", fw=36)
        redo_tb = _btn("↪", tooltip="Redo annotation  (Ctrl+Y)", fw=36)
        undo_tb.clicked.connect(self._undo_annotation)
        redo_tb.clicked.connect(self._redo_annotation)
        tb.addWidget(undo_tb); tb.addWidget(redo_tb)

        # GSD label
        tb.addSeparator()
        self._gsd_tb_lbl = QLabel("GSD: not set")
        self._gsd_tb_lbl.setStyleSheet(
            f"color:{UI_THEME['text_tertiary']};font-size:8pt;background:transparent;padding:0 6px;")
        tb.addWidget(self._gsd_tb_lbl)

        # v1.7.0: per-component calibration status badge
        self._cal_badge = QLabel("📏 No components calibrated  (press C to calibrate)")
        self._cal_badge.setStyleSheet(
            f"color:{UI_THEME['accent_amber']};font-size:8pt;background:transparent;padding:0 6px;")
        tb.addWidget(self._cal_badge)

        tb.addSeparator()

        # ── Zone 3: ML + Export ─────────────────────────────────────────────────
        z3_lbl = QLabel("  ML & EXPORT  ")
        z3_lbl.setStyleSheet(z1_lbl.styleSheet())
        tb.addWidget(z3_lbl)

        self._ml_btn  = _btn("🤖 ML",      UI_THEME["accent_amber"],
                             "Detection & Training  (Ctrl+M)", fw=70)
        self._qc_btn  = _btn("🔍 QC",      UI_THEME["accent_blue"],
                             "QC Viewer  (Ctrl+K)", fw=65)
        self._rpt_btn = _btn("📄 Report",  UI_THEME["accent_purple"],
                             "Generate PDF report  (Ctrl+R)", fw=90)
        self._jpg_btn = _btn("💾 JPEG",    UI_THEME["accent_green"],
                             "Save annotated JPEG for current image", fw=80)
        # RENAME/REPORT: New buttons for image renaming and selected-image report
        self._ren_btn = _btn("✏️ Rename",  UI_THEME["accent_cyan"],
                             "Batch rename images (add blade/face prefix)", fw=90)
        self._sel_rpt_btn = _btn("📋 Selection Report", UI_THEME["accent_amber"],
                             "Generate report from selected/filtered images", fw=140)
        self._ml_btn.clicked.connect(self._open_ml_dialog)
        self._qc_btn.clicked.connect(self._launch_qc_guard)
        self._rpt_btn.clicked.connect(self._generate_report)
        self._jpg_btn.clicked.connect(self._save_annotated_jpeg)
        self._ren_btn.clicked.connect(self._rename_images_dialog)
        self._sel_rpt_btn.clicked.connect(self._generate_selection_report)
        for b in [self._ml_btn, self._qc_btn, self._rpt_btn,
                  self._jpg_btn, self._ren_btn, self._sel_rpt_btn]:
            tb.addWidget(b)

        # Phase 9.1: Hide ML/QC buttons when YOLO not available
        if not YOLO_AVAILABLE:
            self._ml_btn.setVisible(False)
            self._qc_btn.setVisible(False)
            self._ml_btn.setToolTip("ML disabled — ultralytics not installed")
            self._qc_btn.setToolTip("QC disabled — ultralytics not installed")

    # ── Central layout ─────────────────────────────────────────────────────────

    def _build_central(self):
        """Sarah Chen: QStackedWidget — index 0 = annotate, index 1 = QC."""
        self._stacked = QStackedWidget()

        # ── Index 0: annotation splitter ───────────────────────────────────────
        splitter = QSplitter(Qt.Orientation.Horizontal)

        # Left: tabbed panel — "Images" tab + "Blade Diagram" tab
        left = QWidget()
        left.setStyleSheet(
            f"background:{UI_THEME['bg_secondary']};"
            f"border-right:1px solid {UI_THEME['border']};")
        left.setMaximumWidth(340)
        left.setMinimumWidth(180)
        left_main = QVBoxLayout(left)
        left_main.setContentsMargins(4, 4, 4, 4)
        left_main.setSpacing(2)

        self._proj_info = QLabel("No project loaded")
        self._proj_info.setWordWrap(True)
        self._proj_info.setStyleSheet(
            f"font-size:8pt;color:{UI_THEME['text_secondary']};"
            f"background:transparent;padding:2px;")
        left_main.addWidget(self._proj_info)

        # ── Tab widget ──────────────────────────────────────────────────────────
        left_tabs = QTabWidget()
        left_tabs.setDocumentMode(True)
        left_tabs.setTabPosition(QTabWidget.TabPosition.North)
        left_tabs.setStyleSheet(
            f"QTabWidget::pane{{border:none;background:{UI_THEME['bg_secondary']};}}"
            f"QTabBar::tab{{background:{UI_THEME['bg_primary']};color:{UI_THEME['text_secondary']};"
            f"padding:4px 10px;border:none;font-size:8pt;}}"
            f"QTabBar::tab:selected{{color:{UI_THEME['accent_cyan']};"
            f"border-bottom:2px solid {UI_THEME['accent_cyan']};background:{UI_THEME['bg_secondary']};}}"
        )

        # Tab 1: Images (thumbnail strip)
        tab_images = QWidget()
        tab_images.setStyleSheet("background:transparent;")
        tab_img_layout = QVBoxLayout(tab_images)
        tab_img_layout.setContentsMargins(0, 4, 0, 0)
        tab_img_layout.setSpacing(2)

        self._thumb_strip = QListWidget()
        self._thumb_strip.setFlow(QListWidget.Flow.TopToBottom)
        self._thumb_strip.setResizeMode(QListWidget.ResizeMode.Adjust)
        self._thumb_strip.setViewMode(QListWidget.ViewMode.IconMode)
        self._thumb_strip.setIconSize(QSize(120, 80))
        self._thumb_strip.setSpacing(2)
        self._thumb_strip.setStyleSheet(
            f"QListWidget{{background:{UI_THEME['bg_secondary']};border:none;outline:none;}}"
            f"QListWidget::item{{border:2px solid {UI_THEME['border']};"
            f"border-radius:4px;margin:2px;padding:0;}}"
            f"QListWidget::item:selected{{border-color:{UI_THEME['accent_cyan']};}}"
        )
        self._thumb_strip.currentRowChanged.connect(self._on_thumb_selected)
        tab_img_layout.addWidget(self._thumb_strip)
        left_tabs.addTab(tab_images, "Images")

        # Tab 2: Blade Diagram
        tab_diag = QWidget()
        tab_diag.setStyleSheet("background:transparent;")
        tab_diag_layout = QVBoxLayout(tab_diag)
        tab_diag_layout.setContentsMargins(0, 4, 0, 0)
        tab_diag_layout.setSpacing(2)

        diag_sub = QLabel("Click a dot to jump to defect")
        diag_sub.setStyleSheet(
            f"color:{UI_THEME['text_tertiary']};font-size:7pt;"
            f"background:transparent;padding:0 0 2px 0;")
        self._blade_diag = BladeDiagram()
        self._blade_diag.setMinimumHeight(180)
        tab_diag_layout.addWidget(diag_sub)
        tab_diag_layout.addWidget(self._blade_diag)
        left_tabs.addTab(tab_diag, "Blade Diagram")

        left_main.addWidget(left_tabs)
        splitter.addWidget(left)

        # Centre: image viewer
        self._viewer = ImageViewer()
        splitter.addWidget(self._viewer)

        # Right: Tabbed panel with Annotation and QC Review tabs
        right_tabs = QTabWidget()
        self._right_tabs = right_tabs  # v4.1.1: stored for programmatic tab switch
        right_tabs.setDocumentMode(True)
        right_tabs.setStyleSheet(
            f"QTabWidget::pane{{border:none;background:{UI_THEME['bg_secondary']};}}"
            f"QTabBar::tab{{background:{UI_THEME['bg_primary']};color:{UI_THEME['text_secondary']};"
            f"padding:6px 12px;border:none;font-size:9pt;}}"
            f"QTabBar::tab:selected{{color:{UI_THEME['accent_cyan']};"
            f"border-bottom:2px solid {UI_THEME['accent_cyan']};background:{UI_THEME['bg_secondary']};}}"
        )
        
        # Tab 1: Annotation Panel (existing)
        self._ann_panel = AnnotationPanel()
        right_tabs.addTab(self._ann_panel, "📝 Annotate")
        
        # Tab 2: QC Review Panel (NEW in v4.1.0)
        self._qc_review_panel = QCReviewPanel()
        self._qc_review_panel.approve_requested.connect(self._on_approve_annotation)
        self._qc_review_panel.reject_requested.connect(self._on_reject_annotation)
        right_tabs.addTab(self._qc_review_panel, "🔍 QC Review")
        
        # Show QC Review tab only if user has approve permission
        if not SESSION.can_do("approve"):
            right_tabs.setTabVisible(1, False)
        
        splitter.addWidget(right_tabs)

        splitter.setSizes([260, 840, 320])
        self._stacked.addWidget(splitter)   # index 0

        # ── Index 1: inline QC viewer ───────────────────────────────────────────
        self._qc_widget = QCViewerWidget()
        self._stacked.addWidget(self._qc_widget)   # index 1

        self.setCentralWidget(self._stacked)

    # ── Thumbnail dock ─────────────────────────────────────────────────────────

    def _build_thumb_dock(self):
        """v3.3.7: Thumbnail strip moved to left panel. This method is now a no-op."""
        pass  # strip already built in _build_main_ui inside the left splitter panel

    # ── Status bar ─────────────────────────────────────────────────────────────

    def _build_statusbar(self):
        sb = QStatusBar(self)
        self._status_main = QLabel("Ready")
        self._status_main.setStyleSheet(f"color:{UI_THEME['text_secondary']};background:transparent;")
        self._status_gsd  = QLabel("GSD: not calibrated")
        self._status_gsd.setStyleSheet(f"color:{UI_THEME['text_tertiary']};background:transparent;")
        sb.addWidget(self._status_main, 1)
        sb.addPermanentWidget(self._status_gsd)
        self.setStatusBar(sb)

    # ── Toast helper ───────────────────────────────────────────────────────────

    def _toast(self, message: str, colour: str = "", duration: int = 2800):
        """Dev Patel: Show a non-blocking toast overlay notification."""
        ToastNotification(self, message, colour, duration)

    # ── EXIF Calibration Metadata ──────────────────────────────────────────────────

    def _store_exif_metadata(self, metadata_dict: dict):
        """
        v3.4.0: Store EXIF calibration metadata in current image record
        
        Called via calibration_metadata signal from ImageViewer when calibration completes.
        
        Args:
            metadata_dict: Dictionary with keys:
                - 'exif_cal': ExifCalibrationData object (None for manual)
                - 'gsd': Calculated GSD value
                - 'method': "exif-auto", "manual", "component", or "session"
        """
        if not self._current_rec:
            return
        
        exif_cal = metadata_dict.get('exif_cal')
        gsd = metadata_dict.get('gsd')
        method = metadata_dict.get('method', 'manual')
        
        self._current_rec.calibration_method = method
        
        if exif_cal:
            self._current_rec.camera_model = exif_cal.camera_model
            self._current_rec.focal_length_mm = exif_cal.focal_length_mm
            self._current_rec.confidence_level = exif_cal.confidence.value
            self._current_rec.exif_distance_m = exif_cal.estimate_distance_from_gps()
            
            log.info(
                f"[EXIF] Stored metadata: {exif_cal.camera_model}, "
                f"{exif_cal.focal_length_mm:.1f}mm, "
                f"conf={exif_cal.confidence.value}"
            )
        else:
            self._current_rec.camera_model = "Manual Entry"
            self._current_rec.confidence_level = "MANUAL"
            log.info("[EXIF] Stored manual calibration metadata")
        
        save_project(self._project)

    def _batch_auto_calibrate(self):
        """
        v4.0.1: Batch auto-calibrate all loaded images using EXIF data.
        No reference line required - purely EXIF-based calibration.
        
        CRITICAL: Each image is calibrated INDIVIDUALLY with its own EXIF data.
        No calibration is shared between images - every image gets its own:
        - Camera parameters from its EXIF/database
        - Distance estimate from its GPS altitude + gimbal pitch
        - GSD calculation based on its specific parameters
        
        Process:
        1. Extract EXIF from each loaded image INDIVIDUALLY
        2. Calculate GSD from GPS + gimbal + camera specs FOR THAT IMAGE
        3. Apply GSD to image if confidence is acceptable
        4. Show summary dialog with results
        """
        if not self._project or not self._image_paths:
            QMessageBox.warning(
                self, "No Images",
                "Please load images first before auto-calibrating.")
            return
        
        # Progress dialog
        progress = QProgressDialog(
            "Auto-calibrating images using EXIF data...",
            "Cancel", 0, len(self._image_paths), self)
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.setMinimumDuration(500)  # Show if takes >0.5s
        
        calibrated_count    = 0
        failed_count        = 0
        low_confidence_count = 0
        locked_count        = 0  # v4.4.8: annotations skipped because size_locked=True
        migrated_count      = 0  # v4.4.8: annotations recalculated after phantom-GSD migration
        results_by_component = {}  # Track by blade/component
        tier_counts         = {}   # FIX-17e/18b: {"3d-gps": N, "pitch-altitude": N, "assumed-45deg-pitch": N}
        
        for idx, filepath in enumerate(self._image_paths):
            if progress.wasCanceled():
                break
            
            progress.setValue(idx)
            progress.setLabelText(
                f"Processing {idx+1}/{len(self._image_paths)}: {os.path.basename(filepath)}")
            QApplication.processEvents()  # Keep UI responsive
            
            # Find the ImageRecord for this filepath
            irec = None
            for img_rec in self._project.images.values():
                if img_rec.filepath == filepath:
                    irec = img_rec
                    break
            
            if not irec:
                failed_count += 1
                continue
            
            # Try EXIF extraction - INDIVIDUAL per image, not shared
            try:
                # Create a NEW calibrator for THIS specific image
                log.info(f"\n[EXIF] ═══ Processing: {os.path.basename(filepath)} ═══")
                calibrator = EXIFCalibrator(filepath)
                exif_cal = calibrator.calibrate()
                
                # VALIDATION GATE: Must have valid calibration data
                # This enforces the requirement that each image has complete EXIF
                if not exif_cal:
                    failed_count += 1
                    log.warning(f"[EXIF] ✗ Calibration FAILED for {os.path.basename(filepath)}")
                    log.warning(f"[EXIF]   Reason: Could not extract sufficient EXIF/XMP data")
                    continue
                
                # Check confidence level
                if exif_cal.confidence == ConfidenceLevel.FAILED:
                    failed_count += 1
                    log.warning(f"[EXIF] Calibration failed for {os.path.basename(filepath)}")
                    continue
                
                if exif_cal.confidence == ConfidenceLevel.LOW:
                    low_confidence_count += 1
                    # Still try to calibrate with LOW confidence, but track separately
                
                # ── FIX-17e: 3-tier distance estimation ───────────────────────
                # Tier 1 (BEST) — full 3D GPS: uses drone lat/lon/AbsAlt AND
                #   tower base lat/lon/MSL stored in the project.  Handles
                #   varying pilot altitude naturally because every image carries
                #   its own AbsoluteAltitude from XMP.
                # Tier 2 — legacy pitch/altitude: RelativeAltitude / tan(pitch).
                #   Works without tower GPS but is inaccurate on uneven terrain.
                # Tier 3 — 30 m hard assumption: logged prominently; flags image
                #   as LOW confidence so the inspector knows to verify manually.
                dist_method   = "unknown"
                estimated_dist = None

                # Tier 1: 3D GPS — requires tower coords on project AND drone GPS
                _t_lat = getattr(self._project, 'tower_lat',          None)
                _t_lon = getattr(self._project, 'tower_lon',          None)
                _t_alt = getattr(self._project, 'tower_base_alt_msl', None)
                if all(v is not None for v in [_t_lat, _t_lon, _t_alt,
                                               exif_cal.drone_lat,
                                               exif_cal.drone_lon,
                                               exif_cal.absolute_altitude]):
                    estimated_dist = exif_cal.estimate_distance_3d(
                        tower_lat=_t_lat,
                        tower_lon=_t_lon,
                        tower_base_msl=_t_alt,
                    )
                    if estimated_dist:
                        dist_method = "3d-gps"
                        log.info(
                            f"[EXIF-3D] ✓ Tier-1 (3D GPS) dist={estimated_dist:.2f}m "
                            f"for {os.path.basename(filepath)}")

                # Tier 2: legacy pitch / relative altitude
                if not estimated_dist:
                    estimated_dist = exif_cal.estimate_distance_from_gps()
                    if estimated_dist:
                        dist_method = "pitch-altitude"
                        log.info(
                            f"[EXIF] ✓ Tier-2 (pitch/alt) dist={estimated_dist:.2f}m "
                            f"for {os.path.basename(filepath)}")

                # Tier 3: physically-grounded fallback (no pitch data)
                # 30 m was wrong by 3–9× on real DJI data (FIX-18b).
                # If relative_altitude is known, assume 45° camera pitch
                # (industry midpoint for wind-tower inspection: 35–60°).
                # distance = rel_alt / tan(45°) = rel_alt × 1.0
                # If rel_alt is also absent, skip — no GSD can be estimated.
                if not estimated_dist:
                    if exif_cal.relative_altitude:
                        estimated_dist = exif_cal.relative_altitude / math.tan(math.radians(45))
                        dist_method    = "assumed-45deg-pitch"
                        log.warning(
                            f"[EXIF] ⚠ Tier-3 (assumed 45° pitch) dist={estimated_dist:.2f}m "
                            f"for {os.path.basename(filepath)} — gimbal pitch absent. "
                            "Set tower GPS in Report Settings for Tier-1 accuracy.")
                    else:
                        log.warning(
                            f"[EXIF] ✗ Tier-3 SKIP — no altitude or pitch data for "
                            f"{os.path.basename(filepath)}. Cannot estimate distance.")
                        failed_count += 1
                        continue
                # ── end 3-tier ─────────────────────────────────────────────────
                
                auto_gsd = exif_cal.calculate_gsd_cm_per_px(estimated_dist)
                
                if not auto_gsd or auto_gsd <= 0:
                    failed_count += 1
                    log.warning(f"[EXIF] Invalid GSD calculated for {os.path.basename(filepath)}")
                    continue
                
                # Apply GSD to this image
                irec.gsd_cm_per_px = auto_gsd

                # Store metadata — FIX-17e: calibration_method reflects which tier was used
                irec.calibration_method = f"exif-auto-{dist_method}"
                irec.camera_model       = exif_cal.camera_model
                irec.focal_length_mm    = exif_cal.focal_length_mm
                irec.confidence_level   = exif_cal.confidence.value
                irec.exif_distance_m    = estimated_dist
                
                # Recompute annotation sizes for this image
                # v4.4.8: Respect size_locked — never overwrite a size the user
                # has explicitly verified or manually set.  Stamp size_calibrated_at
                # on every annotation we DO update so future migration code can
                # identify which calibration generation wrote the value.
                _now_iso   = datetime.now().isoformat()
                _recalced  = 0
                _locked    = 0
                for ann in irec.annotations:
                    if ann.size_locked:
                        # User explicitly locked this size — respect it unconditionally
                        _locked      += 1
                        locked_count += 1
                        continue
                    x1, y1, w_px, h_px = ann.bounding_rect()
                    if w_px > 0 or h_px > 0:
                        _was_phantom = (ann.gsd_value is not None
                                        and ann.gsd_value > 20.0
                                        and ann.gsd_source == "image")
                        ann.width_cm           = round(w_px * auto_gsd, 2)
                        ann.height_cm          = round(h_px * auto_gsd, 2)
                        ann.gsd_source         = "image"
                        ann.gsd_value          = auto_gsd
                        ann.size_calibrated_at = _now_iso
                        _recalced += 1
                        if _was_phantom:
                            migrated_count += 1
                if _locked:
                    log.info(
                        f"[EXIF] {os.path.basename(filepath)}: "
                        f"{_recalced} sizes recalculated, "
                        f"{_locked} locked annotation(s) preserved")
                
                # Track by component
                comp = irec.blade or "Unknown"
                if comp not in results_by_component:
                    results_by_component[comp] = {
                        'count': 0, 'gsd_avg': 0.0, 'gsd_values': []
                    }
                results_by_component[comp]['count'] += 1
                results_by_component[comp]['gsd_values'].append(auto_gsd)

                # FIX-17e: Count images by tier so the summary dialog can show
                # how many used each distance method.
                tier_counts[dist_method] = tier_counts.get(dist_method, 0) + 1
                
                calibrated_count += 1
                log.info(
                    f"[EXIF] Auto-calibrated {os.path.basename(filepath)}: "
                    f"{auto_gsd:.4f} cm/px, {exif_cal.camera_model}, "
                    f"conf={exif_cal.confidence.value}")
                
            except Exception as e:
                failed_count += 1
                log.error(f"[EXIF] Batch calibration error for {os.path.basename(filepath)}: {e}")
        
        progress.setValue(len(self._image_paths))
        
        # Save project with all calibrations
        if calibrated_count > 0:
            save_project(self._project)
            
            # Update display for current image if visible
            if self._current_rec and self._current_rec.gsd_cm_per_px:
                self._viewer.set_gsd(self._current_rec.gsd_cm_per_px)
                self._update_gsd_labels(self._current_rec.gsd_cm_per_px, "image")
                self._ann_panel.update_gsd_display(self._current_rec.gsd_cm_per_px, "image")
        
        # Show results dialog
        self._show_batch_calibration_results(
            calibrated_count, failed_count, low_confidence_count,
            results_by_component, len(self._image_paths), tier_counts,
            locked_count, migrated_count)
    
    def _show_batch_calibration_results(self, calibrated: int, failed: int,
                                        low_conf: int, by_component: dict,
                                        total: int,
                                        tier_counts: dict = None,
                                        locked: int = 0,
                                        migrated: int = 0):
        """
        v3.4.0: Show batch calibration results dialog with summary.
        FIX-17e: tier_counts shows how many images used each distance method.
        v4.4.8: locked/migrated show backward-compat annotation handling.
        """
        tier_counts = tier_counts or {}
        summary_parts = []

        if calibrated > 0:
            summary_parts.append(
                f"✅ <b>{calibrated} images</b> calibrated successfully")

        if failed > 0:
            summary_parts.append(
                f"❌ <b>{failed} images</b> failed (no EXIF or invalid data)")

        if low_conf > 0:
            summary_parts.append(
                f"⚠️ <b>{low_conf} images</b> calibrated with LOW confidence")

        # v4.4.8: Backward-compat annotation status lines
        if migrated > 0:
            summary_parts.append(
                f"🔄 <b>{migrated} annotation(s)</b> recalculated "
                f"(corrected from broken calibration in v4.4.0–v4.4.6)")
        if locked > 0:
            summary_parts.append(
                f"🔒 <b>{locked} annotation(s)</b> preserved "
                f"(manually verified — size_locked=True)")

        summary = "<br>".join(summary_parts)

        # FIX-17e / FIX-18b: Distance method (tier) breakdown
        tier_text = ""
        if tier_counts:
            tier_labels = {
                "3d-gps":              "🛰 Tier 1 — 3D GPS (most accurate)",
                "pitch-altitude":      "📐 Tier 2 — Pitch / Relative Altitude",
                "assumed-45deg-pitch": "⚠ Tier 3 — Assumed 45° pitch (no gimbal data)",
            }
            tier_text = "<br><br><b>Distance Method Used:</b><br>"
            for method, count in sorted(tier_counts.items()):
                label = tier_labels.get(method, method)
                tier_text += f"&nbsp;&nbsp;• {label}: <b>{count} images</b><br>"

            if tier_counts.get("assumed-45deg-pitch", 0) > 0:
                tier_text += (
                    "<br><font color='orange'><b>💡 Tip:</b> "
                    f"{tier_counts['assumed-45deg-pitch']} image(s) used the 45° assumed-pitch "
                    "fallback (gimbal pitch absent). Ensure original DJI RAW files are loaded "
                    "so GimbalPitchDegree XMP is available, enabling Tier-2 accuracy or better."
                    "</font>")

        # Component breakdown
        comp_text = ""
        if by_component:
            comp_text = "<br><br><b>By Component:</b><br>"
            for comp, data in sorted(by_component.items()):
                if data['gsd_values']:
                    avg_gsd = sum(data['gsd_values']) / len(data['gsd_values'])
                    min_gsd = min(data['gsd_values'])
                    max_gsd = max(data['gsd_values'])
                    comp_text += (
                        f"&nbsp;&nbsp;• <b>{comp}</b>: {data['count']} images, "
                        f"GSD avg: {avg_gsd:.4f} cm/px "
                        f"(range: {min_gsd:.4f} – {max_gsd:.4f})<br>")

        # Full message
        if calibrated == 0:
            title   = "Auto-Calibration Failed"
            icon    = QMessageBox.Icon.Warning
            message = (
                f"<b>No images could be auto-calibrated.</b><br><br>"
                f"{summary}<br><br>"
                f"<i>Possible reasons:</i><br>"
                f"• Images don't have EXIF metadata<br>"
                f"• Camera model not in database<br>"
                f"• Missing GPS or gimbal data<br><br>"
                f"Please use Manual Calibrate instead.")
        else:
            title   = "Auto-Calibration Complete"
            icon    = QMessageBox.Icon.Information
            message = (
                f"<b>Batch auto-calibration complete!</b><br><br>"
                f"{summary}{tier_text}{comp_text}<br><br>"
                f"<i>Total processed: {total} images</i>")

            if low_conf > 0:
                message += (
                    f"<br><br><font color='orange'><b>Note:</b> "
                    f"{low_conf} images were calibrated with LOW confidence. "
                    f"Consider verifying these manually.</font>")

        msg_box = QMessageBox(self)
        msg_box.setWindowTitle(title)
        msg_box.setIcon(icon)
        msg_box.setText(message)
        msg_box.setStandardButtons(QMessageBox.StandardButton.Ok)
        msg_box.exec()

        # Toast with quick summary
        if calibrated > 0:
            self._toast(
                f"✓ {calibrated} images auto-calibrated successfully",
                UI_THEME["accent_green"], 3000)
        else:
            self._toast(
                "✗ Auto-calibration failed — use Manual Calibrate",
                UI_THEME["accent_orange"], 3000)

    # ── Project actions ────────────────────────────────────────────────────────

    def _new_project(self):
        log.debug("[PROJECT] _new_project: dialog opened")
        name, ok = QInputDialog.getText(self, "New Project", "Project name:")
        if not ok or not name.strip():
            log.debug("[PROJECT] _new_project: cancelled at name prompt")
            return
        site, ok2     = QInputDialog.getText(self, "New Project", "Site / Company:")
        turbine, ok3  = QInputDialog.getText(self, "New Project", "Turbine ID:")
        inspector, ok4= QInputDialog.getText(self, "New Project", "Inspector name:")

        # v4.2.0: Add new report metadata fields
        scan_date, ok5 = QInputDialog.getText(self, "New Project",
            "Scan Date (e.g., 15 Feb 2026):",
            text=datetime.now().strftime("%d %b %Y"))
        turbine_mfr, ok6 = QInputDialog.getText(self, "New Project",
            "Turbine Manufacturer (e.g., Vestas, Siemens Gamesa):")
        rated_pwr, ok7 = QInputDialog.getText(self, "New Project",
            "Rated Power (e.g., 2.5 MW, 3.6 MW):")

        # Blade serial numbers removed — set as empty by default
        blade_numbers = {"A": "", "B": "", "C": ""}

        folder = QFileDialog.getExistingDirectory(
            self, "Select project folder (or create new)")
        if not folder:
            log.debug("[PROJECT] _new_project: cancelled at folder selection")
            return
        self._project = Project(
            name=name.strip(), site=site.strip() if ok2 else "",
            turbine_id=turbine.strip() if ok3 else "",
            inspector=inspector.strip() if ok4 else "",
            project_folder=folder,
            blade_numbers=blade_numbers,
            # v4.2.0: Set new metadata fields
            scan_date=scan_date.strip() if ok5 else "",
            turbine_manufacturer=turbine_mfr.strip() if ok6 else "",
            rated_power=rated_pwr.strip() if ok7 else "",
        )
        save_project(self._project)
        log.info(
            f"[PROJECT] Created: '{self._project.name}'  "
            f"site='{self._project.site}'  turbine='{self._project.turbine_id}'  "
            f"inspector='{self._project.inspector}'  folder='{folder}'")
        # Phase 9.3: track in recent projects list
        CFG.add_recent_project(
            str(Path(folder) / "project.json"))
        self._refresh_recent_menu()
        self._update_project_ui()
        self._toast(f"Project '{name}' created", UI_THEME["accent_green"])

    def _open_project(self):
        folder = QFileDialog.getExistingDirectory(self, "Open Project Folder")
        if not folder:
            log.debug("[PROJECT] _open_project: cancelled")
            return
        log.info(f"[PROJECT] Opening project from: {folder}")
        p = load_project(folder)
        if not p:
            log.error(f"[PROJECT] Failed to load project.json from: {folder}")
            QMessageBox.warning(self, "Load Error",
                "Could not load project.json from that folder.")
            return
        self._project = p
        n_images = len(p.images)
        n_anns   = sum(len(ir.annotations) for ir in p.images.values())
        log.info(
            f"[PROJECT] Loaded: '{p.name}'  site='{p.site}'  "
            f"turbine='{p.turbine_id}'  images={n_images}  annotations={n_anns}")
        CFG.set("GENERAL", "LastProjectFolder", folder)
        CFG.save()
        # Phase 9.3: track in recent projects list
        project_file = str(Path(folder) / "project.json")
        CFG.add_recent_project(project_file)
        self._refresh_recent_menu()
        self._update_project_ui()
        self._update_header_bar()
        self._toast(f"Loaded: {p.name}", UI_THEME["accent_cyan"])

    def _save_project_action(self):
        if not self._project:
            QMessageBox.information(self, "No Project", "Create or open a project first.")
            return
        log.debug(f"[PROJECT] Manual save requested: '{self._project.name}'")
        if save_project(self._project):
            log.info(f"[PROJECT] Saved: '{self._project.name}' → {self._project.project_folder}")
            self._toast("Project saved ✓", UI_THEME["accent_green"])
        else:
            log.error(f"[PROJECT] Save FAILED for: '{self._project.name}'")
            self._toast("Save failed!", UI_THEME["accent_red"])

    def _update_project_ui(self):
        if not self._project:
            return
        self._update_header_bar()
        # FIX-BUG: Push current project reference into AnnotationPanel so that
        # load_pending() can build auto-suggested filenames (turbine ID, blade).
        # Must happen before any annotation UI interaction — doing it here covers
        # both new-project and open-project flows since both paths call this method.
        if hasattr(self, "_ann_panel"):
            self._ann_panel.set_project(self._project)
        gsd = self._project.session_gsd
        self._update_gsd_labels(gsd, "session")
        p = self._project
        self._proj_info.setText(
            f"<b>{p.name}</b><br/>"
            f"<span style='color:{UI_THEME['text_tertiary']};'>"
            f"Site: {p.site or '—'}<br/>"
            f"Turbine: {p.turbine_id or '—'}<br/>"
            f"Inspector: {p.inspector or '—'}</span>"
        )
        self._blade_diag.update_project(p)
        # Restore thumb strip from project.images so images are visible after open
        self._restore_strip_from_project()
        # Phase 6: show review progress in status bar
        all_anns   = [a for ir in p.images.values() for a in ir.annotations]
        n_total    = len(all_anns)
        n_approved = sum(1 for a in all_anns if a.status == "approved")
        n_rejected = sum(1 for a in all_anns if a.status == "rejected")
        n_pending  = n_total - n_approved - n_rejected
        review_str = (f"  ·  Review: ✔{n_approved} ○{n_pending} ✕{n_rejected}"
                      if SESSION.can_do("approve") else "")
        self._status_main.setText(
            f"Project: {p.name}  |  {len(p.images)} images  |  "
            f"{n_total} annotations{review_str}"
        )

    def _restore_strip_from_project(self):
        """
        Repopulate the thumbnail strip and _image_paths from project.images on
        project open.  Only adds files that are not already in the strip so it
        is safe to call redundantly (e.g. on every _update_project_ui call).
        Files whose irec.filepath no longer exists on disk are silently skipped
        so a moved/deleted folder does not crash the restore.
        """
        if not self._project:
            return
        existing_set = set(getattr(self, "_image_paths", []))
        if not hasattr(self, "_image_paths"):
            self._image_paths = []

        cache_dir = Path(self._project.project_folder) / ".thumbcache"
        base_idx  = self._thumb_strip.count()
        added_cnt = 0

        for irec in self._project.images.values():
            fp = irec.filepath
            if not fp or fp in existing_set:
                continue  # already in strip or no path stored
            if not os.path.exists(fp):
                log.debug(f"_restore_strip: skipping missing file {fp}")
                continue  # file moved/deleted — skip silently
            self._image_paths.append(fp)
            existing_set.add(fp)
            fname = os.path.basename(fp)
            item  = QListWidgetItem(fname)
            item.setData(Qt.ItemDataRole.UserRole, fp)
            self._thumb_strip.addItem(item)
            worker = ThumbnailWorker(base_idx + added_cnt, fp, cache_dir)
            worker.signals.done.connect(self._on_thumb_done)
            self._thumb_pool.start(worker)
            added_cnt += 1

        if added_cnt:
            self._update_thumbnail_borders()
            log.info(f"_restore_strip_from_project: restored {added_cnt} image(s)")
            # FIX-18a: auto-fill tower GPS from restored images if not yet set
            self._autofill_tower_gps_from_images(self._image_paths)

    # ── Image loading ──────────────────────────────────────────────────────────

    def _load_images(self):
        """
        FIX-FOLDER: Let user pick either individual files OR an entire folder.
        When a folder is chosen, all supported image files in it are auto-populated —
        no need to Ctrl+click dozens of files individually.
        """
        if not self._project:
            QMessageBox.information(self, "No Project", "Create or open a project first.")
            return

        # Ask: files or folder?
        choice_dlg = QDialog(self)
        choice_dlg.setWindowTitle("Load Images - Folder Selection Guide")
        choice_dlg.setFixedWidth(480)
        c_lay = QVBoxLayout(choice_dlg)
        c_lay.setContentsMargins(20, 18, 20, 18)
        c_lay.setSpacing(12)
        c_lbl = QLabel(
            "<b>How would you like to load images?</b><br/>"
            "<span style='color:#7d8590;font-size:9pt;'>"
            "Choose a folder to auto-import all images inside it, "
            "or pick specific files.</span><br/><br/>"
            "<b style='color:#76e3ea;'>📁 FOLDER SELECTION GUIDE:</b><br/>"
            "<span style='color:#adbac7;font-size:9pt;'>"
            "<b>• Single WTG Folder</b> (for individual turbine inspection):<br/>"
            "&nbsp;&nbsp;Select folder containing <b>one turbine's</b> images<br/>"
            "&nbsp;&nbsp;Example: <code>WTGS001/</code> with subfolders Blade_1, Blade_2, Tower, Hub<br/><br/>"
            "<b>• Full Site Parent Folder</b> (for complete site report):<br/>"
            "&nbsp;&nbsp;Select parent folder containing <b>multiple WTGs</b><br/>"
            "&nbsp;&nbsp;Example: <code>WindFarm_Project/</code> with WTGS001/, WTGS002/, WTGS003/<br/><br/>"
            "<b>Tower and Hub folders are automatically included!</b><br/>"
            "All subdirectories (Blade_1, Blade_2, Blade_3, Tower, Hub) will be scanned.</span>")
        c_lbl.setWordWrap(True)
        c_lbl.setStyleSheet("background:transparent;")
        c_lay.addWidget(c_lbl)

        folder_btn = QPushButton("📁  Load from Folder  (auto-import all images)")
        folder_btn.setFixedHeight(38)
        folder_btn.setStyleSheet(
            f"background:{UI_THEME['accent_cyan']};color:#0d1117;"
            f"font-weight:bold;border-radius:6px;border:none;")
        folder_btn.clicked.connect(lambda: (choice_dlg.done(1)))

        files_btn = QPushButton("🖼  Select Individual Files…")
        files_btn.setFixedHeight(38)
        files_btn.setStyleSheet(
            f"background:{UI_THEME['bg_card']};color:{UI_THEME['text_primary']};"
            f"border:1px solid {UI_THEME['border']};border-radius:6px;")
        files_btn.clicked.connect(lambda: (choice_dlg.done(2)))

        cancel_btn = QPushButton("Cancel")
        cancel_btn.setStyleSheet(
            f"color:{UI_THEME['text_secondary']};background:transparent;"
            f"border:none;font-size:9pt;")
        cancel_btn.clicked.connect(choice_dlg.reject)

        c_lay.addWidget(folder_btn)
        c_lay.addWidget(files_btn)
        c_lay.addWidget(cancel_btn)
        result = choice_dlg.exec()

        paths: List[Tuple[str, str, str]] = []  # (filepath, blade, face)
        IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"}

        if result == 1:
            # Folder mode — scan for sub-folders named by blade/face (Scopito-style)
            # e.g. project_root/Blade A PS/*.jpg, Blade A LE/*.jpg etc.
            folder = QFileDialog.getExistingDirectory(
                self, "Select Image Folder", "")
            if not folder:
                return

            IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"}
            folder_p   = Path(folder)

            # Directories the app itself creates inside the project/image folder.
            # Excluded from scanning so burned copies, thumb cache, and pinpoint
            # images are never loaded as real inspection images.
            _SCAN_EXCLUDE_DIRS = frozenset({
                "annotated",          # _burn_in_jpeg_annotations output
                ".thumbcache",        # ThumbnailWorker JPEG cache (visible on Windows)
                "pinpoints",          # Scopito / app-generated pinpoint overlays
                "scopito_pinpoints",  # alternate Scopito pinpoint dir name
                "pinpoint_images",    # third variant pinpoint dir name
                "crops",              # report-crop output folder
                "reports",            # exported report assets
            })

            # FIX v4.1.0: Enhanced folder scanning - ALWAYS include Tower and Hub
            # Try Scopito-style: look for sub-folders that have blade/face in name
            sub_dirs = [
                d for d in sorted(folder_p.iterdir())
                if d.is_dir()
                and d.name.lower() not in _SCAN_EXCLUDE_DIRS  # skip app-internal dirs
                and not d.name.startswith(".")                 # skip all hidden dirs
            ]
            auto_map: Dict[str, Tuple[str, str]] = {}  # subfolder → (blade, face)
            scopito_style = False

            # NEW v4.1.0: Explicit Tower and Hub detection with logging
            for sd in sub_dirs:
                b, f = _parse_blade_face_from_folder(sd.name)
                if b or f:
                    auto_map[str(sd)] = (b, f)
                    scopito_style = True
                    if b in ["Tower", "Hub"]:
                        log.info(f"[v4.1.0 FOLDER SCAN] ✓ Detected {b} folder: {sd.name}")
                        print(f"[LOADING] Found {b} images in: {sd.name}")

            if scopito_style and auto_map:
                # FOLDER-AUTO: Scopito mode — images inherit blade/face from sub-folder.
                # BUG-3 FIX: also recurse into unrecognised sub-dirs so parent-folder
                # scans (A/, B/, C/ etc.) pick up ALL images, not just named ones.
                for sd in sub_dirs:
                    if str(sd) in auto_map:
                        blade_auto, face_auto = auto_map[str(sd)]
                    else:
                        # unrecognised subfolder: try name-parse; empty = unknown
                        blade_auto, face_auto = _parse_blade_face_from_folder(sd.name)
                    for entry in sorted(sd.iterdir()):
                        if entry.is_file() and entry.suffix.lower() in IMAGE_EXTS:
                            paths.append((str(entry), blade_auto, face_auto))
                # Also collect any images sitting directly in the parent folder
                for entry in sorted(folder_p.iterdir()):
                    if entry.is_file() and entry.suffix.lower() in IMAGE_EXTS:
                        b_fn, f_fn = _parse_blade_face_from_folder(entry.stem)
                        paths.append((str(entry), b_fn, f_fn))
                n_auto = sum(1 for _, b, _ in paths if b)
                self._toast(
                    f"Auto-detected {n_auto}/{len(paths)} images → blade/face from folders",
                    UI_THEME["accent_cyan"])
            else:
                # Flat / parent folder: parse blade/face from filenames at root level,
                # BUG-3 FIX: then recursively scan ALL sub-directories so a parent folder
                # containing A/, B/, C/ sub-folders collects their images too.
                for entry in sorted(folder_p.iterdir()):
                    if entry.is_file() and entry.suffix.lower() in IMAGE_EXTS:
                        b_fn, f_fn = _parse_blade_face_from_folder(entry.stem)
                        paths.append((str(entry), b_fn, f_fn))
                for sd in sub_dirs:
                    b_sd, f_sd = _parse_blade_face_from_folder(sd.name)
                    for entry in sorted(sd.iterdir()):
                        if entry.is_file() and entry.suffix.lower() in IMAGE_EXTS:
                            b_fn, f_fn = _parse_blade_face_from_folder(entry.stem)
                            # sub-dir name blade overrides filename parse when available
                            paths.append((str(entry), b_sd or b_fn, f_sd or f_fn))
        elif result == 2:
            # File selection mode — classic multi-file picker
            raw_paths, _ = QFileDialog.getOpenFileNames(
                self, "Select Images", "",
                "Images (*.jpg *.jpeg *.png *.tif *.tiff *.bmp)")
            if not raw_paths:
                return
            paths = [(p, "", "") for p in sorted(raw_paths)]
        else:
            return  # cancelled

        if not paths:
            QMessageBox.information(
                self, "No Images Found",
                "No supported image files found.\n"
                "Supported: .jpg .jpeg .png .tif .tiff .bmp")
            return

        # BUG-2 FIX: accumulate into existing _image_paths instead of replacing.
        # Replacing wiped all previously loaded images from the thumbnail strip,
        # making blade-diagram click-navigation show "not in session" for older folders.
        existing_set = set(getattr(self, "_image_paths", []))
        if not hasattr(self, "_image_paths"):
            self._image_paths = []

        # Only clear + rebuild the strip when the very first batch is loaded so
        # that previously annotated images (already in the strip) are preserved.
        if not existing_set:
            self._thumb_strip.clear()
            self._image_paths = []

        # Starting index for ThumbnailWorker (continues from current strip length)
        base_idx  = self._thumb_strip.count()
        added_cnt = 0  # tracks only items actually appended to the strip

        cache_dir = Path(self._project.project_folder) / ".thumbcache"
        for fp, blade_auto, face_auto in paths:
            if fp in existing_set:
                continue  # already in strip — skip duplicate
            fname = os.path.basename(fp)
            if fname not in self._project.images:
                irec = ImageRecord(filename=fname, filepath=fp)
                # FOLDER-AUTO: pre-assign blade and face from folder name detection
                if blade_auto:
                    irec.blade = blade_auto
                if face_auto:
                    irec.default_face = face_auto  # type: ignore[attr-defined]
                # Phase 3.6: auto-populate metadata from EXIF
                exif_data = _read_exif_metadata(fp)
                if exif_data.get("gps_coords"):
                    irec.gps_coords  = exif_data["gps_coords"]
                if exif_data.get("altitude_m"):
                    irec.altitude_m  = exif_data["altitude_m"]
                if exif_data.get("date_taken"):
                    irec.date_taken  = exif_data["date_taken"]
                if exif_data.get("heading"):
                    irec.heading     = exif_data["heading"]
                self._project.images[fname] = irec
            else:
                # BUG-4 FIX: update blade/filepath on existing record when the new
                # scan provides a better blade assignment (e.g. parent-folder re-scan
                # after images were first loaded via "Select Individual Files" which
                # defaults irec.blade to "A"). Only overwrite if no user annotation has
                # already set a confirmed blade (ann.blade is authoritative post-save).
                irec = self._project.images[fname]
                irec.filepath = fp  # refresh path in case folder was moved
                has_user_blade = any(a.blade for a in irec.annotations)
                if blade_auto and not has_user_blade:
                    irec.blade = blade_auto  # apply corrected component from folder name
                if face_auto and not irec.annotations:
                    irec.default_face = face_auto  # type: ignore[attr-defined]

            self._image_paths.append(fp)
            item = QListWidgetItem(fname)
            item.setData(Qt.ItemDataRole.UserRole, fp)
            self._thumb_strip.addItem(item)

            worker = ThumbnailWorker(base_idx + added_cnt, fp, cache_dir)
            worker.signals.done.connect(self._on_thumb_done)
            self._thumb_pool.start(worker)
            added_cnt += 1  # only count items actually added

        save_project(self._project)
        self._update_thumbnail_borders()
        # FIX-5: force blade diagram repaint so subfolder annotations appear immediately.
        # paintEvent reads project.images live but only fires on Qt-triggered repaints;
        # calling update_project() guarantees a fresh repaint with all loaded annotations.
        self._blade_diag.update_project(self._project)
        # Refresh status bar annotation / image count after accumulating new images.
        self._update_project_ui()
        self._toast(f"{len(paths)} images loaded", UI_THEME["accent_cyan"])
        # FIX-18a: Auto-fill tower GPS from EXIF with zero user input.
        self._autofill_tower_gps_from_images(self._image_paths)

    def _autofill_tower_gps_from_images(self, image_paths: list) -> None:
        """
        FIX-18a: Automatically populate project.tower_lat/lon/tower_base_alt_msl
        from the first image with valid EXIF GPS.  Zero user input required.
        Only fires when all three tower fields are still None.

        Algorithm:
          tower_lat/lon          = drone GPS (drone hovers near tower — good proxy)
          tower_base_alt_msl     = absolute_altitude − relative_altitude
                                   (drone MSL minus height above takeoff pad)
        Validated: two DJI images of the same tower both give 633.18 m MSL. ✅
        """
        if not self._project:
            return
        # Skip if tower fields already set — don't overwrite user edits
        if all(v is not None for v in [
                self._project.tower_lat,
                self._project.tower_lon,
                self._project.tower_base_alt_msl]):
            return
        for fp in image_paths:
            try:
                cal = EXIFCalibrator(fp).calibrate()
                if not cal:
                    continue
                if cal.drone_lat is None or cal.drone_lon is None or cal.absolute_altitude is None:
                    continue
                base_msl = (float(cal.absolute_altitude) - float(cal.relative_altitude)
                            if cal.relative_altitude is not None
                            else float(cal.absolute_altitude))
                self._project.tower_lat          = cal.drone_lat
                self._project.tower_lon          = cal.drone_lon
                self._project.tower_base_alt_msl = round(base_msl, 2)
                save_project(self._project)
                self._toast(
                    f"📍 Tower GPS auto-filled: {cal.drone_lat:.5f}, {cal.drone_lon:.5f}"
                    f"  Base MSL: {base_msl:.1f} m",
                    UI_THEME["accent_cyan"], 5000)
                log.info(
                    f"[FIX-18a] Tower GPS auto-filled from {os.path.basename(fp)}: "
                    f"lat={cal.drone_lat:.6f} lon={cal.drone_lon:.6f} "
                    f"base_msl={base_msl:.2f}m")
                return  # one valid image is enough
            except Exception as exc:
                log.debug(f"[FIX-18a] GPS extract failed for {fp}: {exc}")

    @pyqtSlot(int, QPixmap)
    def _on_thumb_done(self, index: int, pixmap: QPixmap):
        if index < self._thumb_strip.count():
            item = self._thumb_strip.item(index)
            item.setIcon(QIcon(pixmap))
            self._update_one_thumb_border(index)

    def _update_thumbnail_borders(self):
        for i in range(self._thumb_strip.count()):
            self._update_one_thumb_border(i)

    def _update_one_thumb_border(self, index: int):
        """Dev Patel: Colour thumbnail border by worst annotation severity."""
        if not self._project or index >= self._thumb_strip.count():
            return
        item  = self._thumb_strip.item(index)
        fp    = item.data(Qt.ItemDataRole.UserRole)
        fname = os.path.basename(fp) if fp else ""
        irec  = self._project.images.get(fname)
        if not irec or not irec.annotations:
            return
        worst = max(
            irec.annotations,
            key=lambda a: SEVERITY_RANK.get(a.severity, 0)).severity
        col   = SEVERITY_COLORS.get(worst, QColor(UI_THEME["border"]))
        # Embed border via stylesheet; set foreground as proxy for severity hint
        item.setForeground(col)

    @pyqtSlot(int)
    def _on_thumb_selected(self, row: int):
        if row < 0 or row >= len(self._image_paths):
            log.debug(f"[VIEWER] _on_thumb_selected: row {row} out of range "
                      f"(total {len(self._image_paths)})")
            return
        fp    = self._image_paths[row]
        fname = os.path.basename(fp)
        if not self._project or fname not in self._project.images:
            log.warning(f"[VIEWER] _on_thumb_selected: '{fname}' not in project images dict")
            return
        self._current_filepath = fp
        self._current_rec      = self._project.images[fname]
        gsd         = self._current_rec.gsd_cm_per_px
        # v4.1.1: per-image only — no component GSD fallback; image GSD or session GSD only
        session_gsd = self._project.session_gsd
        eff_gsd     = gsd or session_gsd
        log.debug(
            f"[VIEWER] Selected [{row+1}/{len(self._image_paths)}] '{fname}'  "
            f"gsd={gsd:.4f if gsd else 'None'}  "
            f"session_gsd={session_gsd:.4f if session_gsd else 'None'}  "
            f"annotations={len(self._current_rec.annotations)}")
        self._viewer.load_image(fp, self._current_rec, eff_gsd, None)
        self._ann_panel.refresh_ann_list(self._current_rec)
        # v1.7.0: tell the annotation panel which file is active (for inline renamer)
        self._ann_panel.set_current_filepath(fp)
        src = "image" if gsd else ("session" if session_gsd else None)
        self._update_gsd_labels(eff_gsd, src or "none")
        self._status_main.setText(f"Viewing: {fname}")
        # Highlight the active blade column in the diagram so it's clear which
        # blade the currently-open image belongs to (cyan outline + white label)
        self._blade_diag.set_active_blade(self._current_rec.blade or None)
        # v4.1.1: refresh QC Review with all annotations on the new image
        if hasattr(self, '_qc_review_panel'):
            self._qc_review_panel.load_image_annotations(self._current_rec)

    # ── Annotation signals ─────────────────────────────────────────────────────

    def _on_annotation_ready(self, ann: Annotation):
        """Sam Okafor: New annotation drawn — load into panel for editing.

        v4.1.1: Compulsory calibration gate — per-image only.
        Before accepting any annotation the current image must have its own GSD
        (from EXIF auto-calibration or manual calibration) or a session GSD.
        Component GSD is no longer accepted — each image must be individually
        calibrated if required.

        Calibration priority (highest first):
          1. irec.gsd_cm_per_px  — calibrated on this specific image (required)
          2. project.session_gsd — global session GSD (also accepted as fallback)
        """
        if not self._current_rec:
            log.debug("[ANNOTATION] _on_annotation_ready: no current image selected")
            self._toast("Select an image first", UI_THEME["accent_orange"])
            return

        # --- Compulsory per-image calibration check ---
        img_gsd  = self._current_rec.gsd_cm_per_px
        sess_gsd = self._project.session_gsd if self._project else None
        if not (img_gsd or sess_gsd):
            log.warning(
                f"[ANNOTATION] Calibration gate blocked annotation on "
                f"'{self._current_rec.filename}' — "
                f"img_gsd=None  session_gsd=None")
            # Remove the just-drawn item so the canvas stays clean
            self._viewer.remove_annotation_item(ann)
            if ann in self._viewer._undo_stack:
                self._viewer._undo_stack.remove(ann)
            comp     = self._current_rec.blade or "Unknown"
            comp_label = {"A": "Blade A", "B": "Blade B", "C": "Blade C"}.get(comp, comp)
            QMessageBox.warning(
                self, "Calibration Required",
                f"<b>Calibration is required before annotating this image.</b><br/><br/>"
                f"Component: <b>{comp_label}</b><br/><br/>"
                f"Use <b>🤖 Auto-Calibrate All</b> to calibrate via EXIF data, or "
                f"press <b>C</b> (📏 Calibrate Image) to manually draw a calibration "
                f"line on a feature of known size in <i>this specific image</i>.<br/><br/>"
                f"Each image must be individually calibrated — calibrating one image "
                f"does not apply to others."
            )
            # Auto-switch to calibration mode so the user can act immediately
            self._mode_btns[ImageViewer.MODE_CAL].setChecked(True)
            self._set_viewer_mode(ImageViewer.MODE_CAL)
            return

        # FIX-5: Pre-fill ann.blade from irec so Hub/Tower images don't default to "A".
        # Without this, load_pending → _validate_and_update_save_button reads the
        # blade combo's stale value from the previous image, showing face/dist/pinpoint
        # for Hub and Tower annotations incorrectly.
        if self._current_rec.blade:
            ann.blade = self._current_rec.blade
        # v4.5.0: Carry over default surface and zone if record has them
        _default_surf = getattr(self._current_rec, "default_surface", "")
        _default_zone = getattr(self._current_rec, "default_zone", "")
        if _default_surf and not getattr(ann, 'surface', ''):
            ann.surface = _default_surf
        if _default_zone and not getattr(ann, 'zone', ''):
            ann.zone = _default_zone
        log.debug(
            f"[ANNOTATION] Ready: mode={ann.mode}  blade={ann.blade}  "
            f"surface={getattr(ann, 'surface', '')}  zone={getattr(ann, 'zone', '')}  "
            f"pos=({ann.x1_px:.0f},{ann.y1_px:.0f})  "
            f"size=({abs(ann.x2_px-ann.x1_px):.0f}×{abs(ann.y2_px-ann.y1_px):.0f})px  "
            f"gsd={'img' if img_gsd else 'session'}={img_gsd or sess_gsd:.4f if (img_gsd or sess_gsd) else 'None'}")
        self._ann_panel.load_pending(ann)
        self._ann_panel.setFocus()

    def _on_discard_annotation(self, ann: Annotation):
        """
        FIX-UX: Remove a just-drawn (unsaved) annotation from the scene.
        The annotation has not been saved to the data model yet, so only the
        visual item needs to be removed — no project save required.
        """
        self._viewer.remove_annotation_item(ann)
        # Also pop it from the viewer's undo stack if it got pushed there
        if ann in self._viewer._undo_stack:
            self._viewer._undo_stack.remove(ann)
        self._toast("Annotation discarded", UI_THEME["accent_orange"], 1500)

    def _on_annotation_selected(self, ann: Annotation):
        """Sam Okafor (BUG-01 fix): Clicking an existing annotation populates panel."""
        self._ann_panel.load_existing(ann)
        # v4.1.1: Load annotation into QC Review and auto-switch to QC tab
        if hasattr(self, '_qc_review_panel'):
            self._qc_review_panel.load_annotation(ann)
        if hasattr(self, '_right_tabs'):
            self._right_tabs.setCurrentIndex(1)  # switch to QC Review tab

    def _on_annotation_modified(self, ann: Annotation):
        """
        EditableBoxItem committed a geometry change (resize / rotate / move).
        Update the Annotation in the data model in-place and auto-save.
        The item already mutated ann directly (shared reference), so we only
        need to make sure it is listed in current_rec.annotations and persist.
        """
        if not self._current_rec or not self._project:
            return
        # Ensure ann is in the annotations list (it always should be, but guard)
        ids = [a.ann_id for a in self._current_rec.annotations]
        if ann.ann_id not in ids:
            log.debug(f"[ANNOTATION] _on_annotation_modified: ann {ann.ann_id[:8]} "
                      f"not in list — appending (unexpected; check signal wiring)")
            self._current_rec.annotations.append(ann)
        log.debug(
            f"[ANNOTATION] Modified: {ann.ann_id[:8]}  "
            f"pos=({ann.x1_px:.0f},{ann.y1_px:.0f})→({ann.x2_px:.0f},{ann.y2_px:.0f})  "
            f"rot={ann.rotation_deg:.1f}°  "
            f"size={'%.2f×%.2f cm' % (ann.width_cm, ann.height_cm) if ann.width_cm else '? cm'}")
        # Refresh panel fields if this annotation is currently displayed
        try:
            if (hasattr(self._ann_panel, "_current_ann")
                    and self._ann_panel._current_ann is not None
                    and self._ann_panel._current_ann.ann_id == ann.ann_id):
                # Reload panel so cm dimensions reflect the new geometry
                self._ann_panel.load_existing(ann)
        except Exception:
            pass
        # Auto-save — geometry edits are data-model mutations
        save_project(self._project)
        self._toast(
            f"Box updated — {ann.width_cm:.1f}×{ann.height_cm:.1f} cm"
            if ann.width_cm else "Box geometry updated",
            UI_THEME["accent_cyan"], 1200)

    def _on_save_annotation(self, ann: Annotation):
        """Tom K. + Sam Okafor: Persist annotation to data model + draw.

        v1.7.0 additions:
          • Auto-estimate distance_from_root_mm from annotation Y-position if the
            inspector left it at 0. Only applies to blade components (A/B/C); skipped
            for Hub/Tower where the concept doesn't apply.
          • Uses project.blade_length_mm (default 50 000 mm) for the estimate.
          • The estimate is stored as-is so the inspector can correct it in the panel.
        """
        if not self._current_rec or not self._project:
            log.debug("[ANNOTATION] _on_save_annotation: no current_rec or project — aborting")
            return
        if not SESSION.can_do("save_annotation"):
            log.warning(f"[PERMISSION] {SESSION.username} attempted save_annotation — denied")
            QMessageBox.warning(self, "Access Denied",
                "Your role does not permit saving annotations.")
            return

        log.info(
            f"[ANNOTATION] Saving: id={ann.ann_id[:8]}  mode={ann.mode}  "
            f"defect='{ann.defect}'  severity={ann.severity}  "
            f"blade={ann.blade}  surface={getattr(ann, 'surface', '')}  zone={getattr(ann, 'zone', '')}  "
            f"size={'%.2f×%.2f cm' % (ann.width_cm, ann.height_cm) if ann.width_cm else '? cm'}  "
            f"gsd_source={ann.gsd_source}  "
            f"image='{self._current_rec.filename}'")

        # Stamp creator on first save only
        if not ann.created_by:
            ann.created_by = SESSION.username

        # v4.5.0: Generate initial serial number using surface+zone split
        ann.serial_number = _generate_defect_serial(self._project, ann.blade, ann.surface, ann.zone)
        log.debug(f"[ANNOTATION] Initial serial assigned: {ann.serial_number}")

        # v3.3.1: Auto-estimate block removed (Bug #6 FIX).
        # root_distance_m and tip_distance_m are always entered by the inspector
        # via the panel spinboxes; the old Y-fraction estimate was a legacy
        # workaround that wrote incorrect distance_from_root_mm values and
        # conflicted with the user-entered workflow introduced in v2.1.1.

        # Update or append
        existing = next(
            (a for a in self._current_rec.annotations if a.ann_id == ann.ann_id),
            None)
        if existing:
            log.debug(f"[ANNOTATION] Replacing existing annotation {ann.ann_id[:8]}")
            self._current_rec.annotations.remove(existing)
            self._viewer.remove_annotation_item(existing)
        else:
            log.debug(f"[ANNOTATION] Appending new annotation {ann.ann_id[:8]}")
        self._current_rec.annotations.append(ann)

        # ROOT-CAUSE FIX: sync irec.blade ← ann.blade on every save.
        # irec.blade is set by folder-name auto-detection and defaults to "A".
        # When the user picks a different blade in the annotation panel, ann.blade
        # is correct but irec.blade stays wrong, causing _build_annotation_pages
        # to group the image under the wrong component section in the report.
        # Edge case: if one image has annotations on multiple blades (unusual),
        # irec.blade takes the most-recent ann.blade — the diagram uses ann.blade
        # directly (see blade diagram ROOT-CAUSE FIX) so it is unaffected.
        if ann.blade:
            if self._current_rec.blade != ann.blade:
                log.debug(
                    f"[ANNOTATION] irec.blade updated: "
                    f"'{self._current_rec.blade}' → '{ann.blade}' "
                    f"(image: {self._current_rec.filename})")
            self._current_rec.blade = ann.blade

        # v4.4.8: Lock the physical size of this annotation when the user explicitly
        # saves it with valid dimensions AND a real calibration source.  This marks
        # the annotation as "inspector-verified" so _batch_auto_calibrate() will
        # never silently overwrite it.  The lock applies to:
        #   • manual calibration (gsd_source="manual") — inspector drew a reference
        #     line and entered a known distance: the most accurate possible source.
        #   • image calibration (gsd_source="image") — the annotation was saved
        #     after the image was individually calibrated (EXIF or manual line).
        # Annotations saved with session or component GSD are NOT locked because
        # those are shared estimates; per-image calibration should take precedence.
        _lockable_sources = ("manual", "image")
        if (ann.width_cm is not None and ann.height_cm is not None
                and ann.gsd_source in _lockable_sources):
            ann.size_locked        = True
            ann.size_calibrated_at = ann.size_calibrated_at or datetime.now().isoformat()
            log.debug(
                f"[CAL] Locked size for {ann.ann_id[:8]}: "
                f"{ann.width_cm:.2f}×{ann.height_cm:.2f} cm "
                f"(source={ann.gsd_source})")
        elif ann.width_cm is not None:
            log.debug(
                f"[CAL] Size NOT locked for {ann.ann_id[:8]}: "
                f"gsd_source='{ann.gsd_source}' is not in lockable sources "
                f"{_lockable_sources} — auto-calibrate may overwrite later")

        self._viewer.draw_annotation(ann)
        save_project(self._project)

        # FIX-11: Reassign ALL serial endings in canonical report order
        # (A→B→C→Hub→Tower, filename-sorted) so the _001/_002/… suffix
        # always matches the top-to-bottom position in the generated PDF/DOCX.
        # Called after save_project so the new annotation is in the data model;
        # save_project is called again (implicitly via session close or next
        # save) to persist the corrected serials.  The panel refresh below
        # picks up corrected values directly from the in-memory annotation.
        _repair_serial_numbers(self._project)
        save_project(self._project)   # persist corrected serials immediately
        log.debug(f"[ANNOTATION] Serial after repair: {ann.serial_number}  "
                  f"total_anns_on_image={len(self._current_rec.annotations)}")

        self._ann_panel.refresh_ann_list(self._current_rec)
        self._blade_diag.update_project(self._project)
        self._update_one_thumb_border(self._thumb_strip.currentRow())
        # FIX-BUG3: Rename BEFORE burn-in so annotated copy uses the defect name
        self._auto_rename_after_annotation(ann)
        # Auto burn-in (uses self._current_filepath which is now the renamed path)
        if self._project.project_folder:
            log.debug(f"[ANNOTATION] Starting burn-in for: {self._current_filepath}")
            _burn_in_jpeg_annotations(
                self._current_filepath, self._current_rec,
                self._project.project_folder)
        self._update_project_ui()
        # v4.1.1: refresh QC list to include newly saved annotation
        if hasattr(self, '_qc_review_panel'):
            self._qc_review_panel.load_image_annotations(self._current_rec)
        self._toast("Annotation saved ✓", UI_THEME["accent_green"])

    def _auto_rename_after_annotation(self, ann: "Annotation") -> None:
        """
        FIX-19: Automatically rename the image file to a report-friendly name
        the first time an annotation is saved on a raw-camera-named image.

        Triggers only when:
          • Current filename begins with a raw camera prefix (DJI_, IMG_, DSC,
            DCIM, P1_, P_) — never renames already-meaningful filenames.
          • The image has exactly ONE annotation (the one just saved) — prevents
            re-renames on subsequent edits to the same image.

        Rename format (slug-safe, unique, report-consistent):
          {serial_number}_{DefectTypeSlug}{ext}
          e.g. DJI_0736_W.JPG → WTG234_BladeA_LE_001_LeadingEdgeErosion.JPG

        Conflict-safe: appends _2, _3 … if the target path already exists.
        Updates project.images dict, self._image_paths, thumbnail strip, and
        self._current_filepath atomically so all downstream code sees new name.
        """
        if not self._current_rec or not self._project or not self._current_filepath:
            return
        # Only rename on the very first annotation for this image
        if len(self._current_rec.annotations) != 1:
            return
        # Only rename raw camera filenames
        stem = os.path.splitext(os.path.basename(self._current_filepath))[0]
        raw_prefixes = ("DJI_", "IMG_", "DSC", "DCIM", "P1_", "P_")
        if not any(stem.upper().startswith(p.upper()) for p in raw_prefixes):
            return
        try:
            # Build slug: serial_number already encodes WTG+component+seq
            serial  = (ann.serial_number or "").strip()
            defect  = (ann.defect or "Defect").strip()
            # Slug: remove non-alphanumeric chars; use module-level re (already imported)
            defect_slug = re.sub(r"[^A-Za-z0-9]", "", defect)
            new_stem    = f"{serial}_{defect_slug}" if serial else defect_slug
            old_dir     = os.path.dirname(self._current_filepath)
            old_ext     = os.path.splitext(self._current_filepath)[1]
            # Conflict-safe: append _2, _3 … until filename is free
            candidate = new_stem
            counter   = 2
            while os.path.exists(os.path.join(old_dir, candidate + old_ext)):
                if os.path.join(old_dir, candidate + old_ext) == self._current_filepath:
                    break  # same file — no conflict
                candidate = f"{new_stem}_{counter}"
                counter  += 1
            new_fname = candidate + old_ext
            new_path  = os.path.join(old_dir, new_fname)
            if new_path == self._current_filepath:
                return  # nothing to do
            os.rename(self._current_filepath, new_path)
            old_fname = os.path.basename(self._current_filepath)
            # Update project images dict
            if old_fname in self._project.images:
                irec = self._project.images.pop(old_fname)
                irec.filename = new_fname
                irec.filepath = new_path
                self._project.images[new_fname] = irec
                self._current_rec = irec
            # Update image_paths list
            for i, fp in enumerate(self._image_paths):
                if fp == self._current_filepath:
                    self._image_paths[i] = new_path
                    break
            # Update thumbnail strip item text + data
            for i in range(self._thumb_strip.count()):
                item = self._thumb_strip.item(i)
                if item and item.data(Qt.ItemDataRole.UserRole) == self._current_filepath:
                    item.setText(new_fname)
                    item.setData(Qt.ItemDataRole.UserRole, new_path)
                    break
            self._current_filepath = new_path
            self._ann_panel.set_current_filepath(new_path)
            save_project(self._project)
            self._toast(
                f"✏️ Renamed → {new_fname}", UI_THEME["accent_amber"], 4000)
            log.info(f"[FIX-19] Auto-renamed '{old_fname}' → '{new_fname}'")
        except Exception as exc:
            log.warning(f"[FIX-19] Auto-rename failed: {exc}")

    def _on_delete_annotation(self, ann: Annotation):
        """Sam Okafor (BUG-02 fix): Delete removes data model entry.
        CTO-AUDIT: Logic-layer permission check (not only UI-layer).
        """
        if not self._current_rec or not self._project:
            log.debug("[ANNOTATION] _on_delete_annotation: no current_rec or project — aborting")
            return
        log.info(
            f"[ANNOTATION] Delete requested: id={ann.ann_id[:8]}  "
            f"defect='{ann.defect}'  severity={ann.severity}  "
            f"blade={ann.blade}  created_by={ann.created_by}  "
            f"by_user={SESSION.username}  image='{self._current_rec.filename}'")
        # Logic-layer enforcement: check delete permissions before touching data
        can_delete = (SESSION.can_do("delete_any") or
                      (SESSION.can_do("delete_own") and ann.created_by == SESSION.username))
        if not can_delete:
            log.warning(f"[PERMISSION] {SESSION.username} attempted to delete annotation "
                        f"{ann.ann_id} owned by {ann.created_by}")
            QMessageBox.warning(self, "Access Denied",
                "You can only delete annotations you created.\n"
                "You do not have permission to remove this annotation.")
            return
        before_count = len(self._current_rec.annotations)
        self._current_rec.annotations = [
            a for a in self._current_rec.annotations if a.ann_id != ann.ann_id]
        after_count = len(self._current_rec.annotations)
        log.debug(f"[ANNOTATION] Deleted from data model: "
                  f"{before_count} → {after_count} annotations on '{self._current_rec.filename}'")
        self._viewer.remove_annotation_item(ann)
        save_project(self._project)
        self._ann_panel.refresh_ann_list(self._current_rec)
        self._blade_diag.update_project(self._project)
        self._update_one_thumb_border(self._thumb_strip.currentRow())
        self._update_project_ui()
        self._toast("Annotation deleted", UI_THEME["accent_orange"])

    # ── Undo / Redo ────────────────────────────────────────────────────────────

    def _undo_annotation(self):
        ann = self._viewer.undo_last()
        if ann and self._current_rec and self._project:
            self._current_rec.annotations = [
                a for a in self._current_rec.annotations if a.ann_id != ann.ann_id]
            save_project(self._project)
            self._ann_panel.refresh_ann_list(self._current_rec)
            self._blade_diag.update_project(self._project)
            self._update_project_ui()
            self._toast("Undo ✓", UI_THEME["text_secondary"], 1500)
        else:
            self._toast("Nothing to undo", UI_THEME["text_tertiary"], 1500)

    def _redo_annotation(self):
        ann = self._viewer.redo_last()
        if ann and self._current_rec and self._project:
            if not any(a.ann_id == ann.ann_id
                       for a in self._current_rec.annotations):
                self._current_rec.annotations.append(ann)
            save_project(self._project)
            self._ann_panel.refresh_ann_list(self._current_rec)
            self._blade_diag.update_project(self._project)
            self._update_project_ui()
            self._toast("Redo ✓", UI_THEME["text_secondary"], 1500)
        else:
            self._toast("Nothing to redo", UI_THEME["text_tertiary"], 1500)

    # ── GSD calibration ────────────────────────────────────────────────────────

    def _on_gsd_updated(self, gsd: float):
        if not self._project:
            return
        if self._current_rec:
            self._current_rec.gsd_cm_per_px = gsd
            self._viewer.set_gsd(gsd)
            # Phase 5.7: retroactively recompute sizes for annotations on this image only
            n = self._recompute_annotation_sizes(self._current_rec, gsd, "image")
            # v4.1.1: manual cal is per-image only — no component_gsd write, no propagation
            save_project(self._project)
            self._update_gsd_labels(gsd, "image")
            self._ann_panel.update_gsd_display(gsd, "image")
            self._ann_panel.refresh_ann_list(self._current_rec)
            self._toast(
                f"GSD: {gsd:.4f} cm/px calibrated ✓  ({n} sizes recomputed)",
                UI_THEME["accent_cyan"])

    def _set_session_gsd(self):
        if not self._project:
            QMessageBox.information(self, "No Project", "Load a project first.")
            return
        val, ok = QInputDialog.getDouble(
            self, "Set Session GSD",
            "Enter session ground sampling distance (cm/px):",
            value=self._project.session_gsd or 0.0,
            decimals=4, min=0.0001, max=100.0)
        if ok and val > 0:
            self._project.session_gsd = val
            self._viewer.set_session_gsd(val)
            # Phase 5.7: retroactively recompute all annotations that used session GSD
            n_total = 0
            for irec in self._project.images.values():
                if irec.gsd_cm_per_px:
                    continue   # image has own calibration — skip
                n_total += self._recompute_annotation_sizes(irec, val, "session")
            save_project(self._project)
            self._update_gsd_labels(val, "session")
            self._ann_panel.update_gsd_display(val, "session")
            self._ann_panel.refresh_ann_list(self._current_rec)
            self._toast(
                f"Session GSD set: {val:.4f} cm/px  ({n_total} sizes recomputed)",
                UI_THEME["accent_green"])

    @staticmethod
    def _recompute_annotation_sizes(
            irec: "ImageRecord", gsd: float, source: str) -> int:
        """
        Phase 5.7: Recompute width_cm / height_cm for all annotations on irec
        using the given GSD.  Returns count of annotations updated.

        v4.4.8: Respects size_locked — locked annotations are skipped.
        Stamps size_calibrated_at on every updated annotation.
        """
        count    = 0
        _now_iso = datetime.now().isoformat()
        for ann in irec.annotations:
            # v4.4.8: never overwrite a user-locked size
            if ann.size_locked:
                continue
            x1, y1, w_px, h_px = ann.bounding_rect()
            if w_px > 0 or h_px > 0:
                ann.width_cm           = round(w_px * gsd, 2)
                ann.height_cm          = round(h_px * gsd, 2)
                ann.gsd_source         = source
                ann.gsd_value          = gsd
                ann.size_calibrated_at = _now_iso
                count += 1
        return count

    def _update_gsd_labels(self, gsd: Optional[float], source: str):
        # v4.1.1: badge derived from per-image gsd_cm_per_px records (component_gsd no longer written by manual cal)
        if self._project and hasattr(self, "_cal_badge"):
            components  = ["A", "B", "C", "Hub", "Tower"]
            comp_labels = {"A": "Blade A", "B": "Blade B", "C": "Blade C",
                           "Hub": "Hub", "Tower": "Tower"}
            # Collect components that have at least one image with its own GSD
            calibrated = {irec.blade for irec in self._project.images.values()
                          if irec.gsd_cm_per_px and irec.blade}
            if calibrated:
                cal_str = "  ".join(
                    f"✅ {comp_labels.get(c, c)}" for c in components if c in calibrated)
                uncal = [comp_labels.get(c, c) for c in components if c not in calibrated]
                uncal_str = ("  " + "  ".join(f"⬜ {c}" for c in uncal)) if uncal else ""
                self._cal_badge.setText(f"📏 Calibrated: {cal_str}{uncal_str}")
                self._cal_badge.setStyleSheet(
                    f"color:{UI_THEME['accent_cyan']};font-size:8pt;background:transparent;")
            else:
                self._cal_badge.setText("📏 No images calibrated yet  (🤖 Auto or C to calibrate)")
                self._cal_badge.setStyleSheet(
                    f"color:{UI_THEME['accent_amber']};font-size:8pt;background:transparent;")

        # CTO-AUDIT: Calibration confidence indicator
        # Source hierarchy: image > component > session > none
        # Confidence score: image=High, component=Medium, session=Low
        # v3.4.0: Enhanced with EXIF auto-calibration metadata
        _conf_map = {
            "image":     ("HIGH",   UI_THEME["accent_green"]),
            "component": ("MEDIUM", UI_THEME["accent_amber"]),
            "session":   ("LOW",    UI_THEME["accent_orange"]),
        }
        conf_label, conf_colour = _conf_map.get(source, ("—", UI_THEME["text_tertiary"]))

        # v3.4.0: Check for EXIF calibration metadata
        conf_emoji = {
            "HIGH": "🟢",
            "MEDIUM": "🟡",
            "LOW": "🟠",
            "MANUAL": "🔵",
        }
        
        # Build the GSD text with confidence indicator
        if gsd:
            # v3.4.0: Enhanced display with EXIF metadata when available
            if self._current_rec and self._current_rec.calibration_method:
                method_label = {
                    "exif-auto": "EXIF Auto",
                    "manual": "Manual",
                    "component": "Component",
                    "session": "Session"
                }.get(self._current_rec.calibration_method, source)
                
                confidence = self._current_rec.confidence_level or conf_label
                emoji = conf_emoji.get(confidence, "⚪")
                
                txt = (
                    f"GSD: {gsd:.4f} cm/px  |  "
                    f"{emoji} {confidence}  |  "
                    f"Method: {method_label}"
                )
                
                # Add camera model if available
                if self._current_rec.camera_model and self._current_rec.camera_model != "Manual Entry":
                    txt += f"  |  📷 {self._current_rec.camera_model}"
                
                # Override confidence colour based on EXIF confidence
                _exif_conf_map = {
                    "HIGH": UI_THEME["accent_green"],
                    "MEDIUM": UI_THEME["accent_amber"],
                    "LOW": UI_THEME["accent_orange"],
                    "MANUAL": UI_THEME["accent_cyan"],
                }
                conf_colour = _exif_conf_map.get(confidence, conf_colour)
            else:
                # Original format when no EXIF metadata
                txt = f"GSD: {gsd:.4f} cm/px  [{source}]  confidence: {conf_label}"
        else:
            txt = "GSD: not calibrated"
            conf_colour = UI_THEME["text_tertiary"]

        for attr in ("_gsd_tb_lbl", "_status_gsd"):
            w = getattr(self, attr, None)
            if w:
                w.setText(txt)
                w.setStyleSheet(f"color:{conf_colour};background:transparent;")
        hb = getattr(self, "_hb_gsd", None)
        if hb:
            if gsd and self._current_rec and self._current_rec.confidence_level:
                confidence = self._current_rec.confidence_level
                emoji = conf_emoji.get(confidence, "⚪")
                hb.setText(f"{gsd:.4f} cm/px [{emoji} {confidence}]")
            else:
                hb.setText(f"{gsd:.4f} cm/px [{conf_label}]" if gsd else "not calibrated")

    # ── Blade diagram ──────────────────────────────────────────────────────────

    def _on_diagram_ann_click(self, ann):
        """
        v3.2.0: User clicked a defect dot on the Scopito-style blade position panel.
        Navigate to the source image and load the annotation into the panel.
        """
        if not self._project or not self._image_paths:
            return
        # Find which ImageRecord owns this annotation
        for irec in self._project.images.values():
            if any(a.ann_id == ann.ann_id for a in irec.annotations):
                fp = irec.filepath
                if fp in self._image_paths:
                    idx = self._image_paths.index(fp)
                    self._thumb_strip.setCurrentRow(idx)
                    # load_existing after viewer has loaded the image
                    QTimer.singleShot(
                        120, lambda a=ann: self._ann_panel.load_existing(a))
                    self._toast(
                        f"Jumped to  {os.path.basename(fp)}",
                        UI_THEME['accent_cyan'], 1400)
                else:
                    self._toast(
                        f"{irec.filename} not in session — load it first",
                        UI_THEME['accent_orange'], 2000)
                return

    def _on_diagram_cell_click(self, blade: str, span: str, face: str):
        """
        Phase 7.7: Cell click → dialog listing all annotations on that blade+span.
        Double-click an annotation row to jump to its source image.
        """
        if not self._project:
            return

        # Gather matching annotations with their source ImageRecord
        matches: List[tuple] = []   # (irec, ann)
        for irec in self._project.images.values():
            for ann in irec.annotations:
                if ann.blade == blade and ann.span == span:
                    matches.append((irec, ann))

        # Build dialog
        dlg = QDialog(self)
        dlg.setWindowTitle(f"Annotations — {blade}  /  {span}")
        dlg.setMinimumSize(600, 380)
        lay = QVBoxLayout(dlg)
        lay.setContentsMargins(16, 14, 16, 14); lay.setSpacing(10)

        hdr = QLabel(
            f"<b>{blade}</b>  ·  {span}  —  "
            f"<span style='color:{UI_THEME['accent_cyan']}'>"
            f"{len(matches)} annotation(s)</span>")
        hdr.setStyleSheet(f"color:{UI_THEME['text_primary']};font-size:10pt;")
        lay.addWidget(hdr)

        table = QListWidget()
        table.setAlternatingRowColors(True)
        status_icons = {"approved": "✔", "rejected": "✕", "pending": "○"}
        for irec, ann in matches:
            short   = SEVERITY_SHORT.get(ann.severity, ann.severity)
            size    = (f"{ann.width_cm:.1f}×{ann.height_cm:.1f}cm"
                       if ann.width_cm else "—")
            status  = ann.status or "pending"
            icon    = status_icons.get(status, "○")
            label   = (f"{icon} [{short}]  {ann.defect}  {size}"
                       f"  —  {irec.filename}")
            item = QListWidgetItem(label)
            item.setData(Qt.ItemDataRole.UserRole, (irec, ann))
            col = SEVERITY_COLORS.get(ann.severity, QColor("#7d8590"))
            item.setForeground(col)
            table.addItem(item)
        table.setToolTip("Double-click to jump to image")
        lay.addWidget(table)

        hint = QLabel("Double-click a row to navigate to the source image.")
        hint.setStyleSheet(
            f"color:{UI_THEME['text_tertiary']};font-size:8pt;")
        lay.addWidget(hint)

        btns = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        btns.rejected.connect(dlg.reject)
        lay.addWidget(btns)

        def _jump(item: QListWidgetItem):
            """Jump to source image and select annotation."""
            data = item.data(Qt.ItemDataRole.UserRole)
            if not data:
                return
            irec_target, ann_target = data
            dlg.accept()
            # Find filepath in image_paths list and load it
            target_fp = irec_target.filepath
            if target_fp in self._image_paths:
                idx = self._image_paths.index(target_fp)
                self._thumb_strip.setCurrentRow(idx)
                self._load_image_idx(idx)
                # Select the annotation on the viewer
                QTimer.singleShot(
                    150, lambda a=ann_target: self._ann_panel.load_existing(a))
            else:
                QMessageBox.warning(self, "Image Not Loaded",
                    f"Image {irec_target.filename} is not in the current session.\n"
                    "Load it via File → Load Images.")

        table.itemDoubleClicked.connect(_jump)
        dlg.exec()

    # ── Mode ───────────────────────────────────────────────────────────────────

    def _set_viewer_mode(self, mode: str):
        self._viewer.set_mode(mode)
        labels = {
            ImageViewer.MODE_SEL:  "Select / Pan mode — click annotation to edit",
            ImageViewer.MODE_BOX:  "Box mode — drag to draw bounding box",
            ImageViewer.MODE_PIN:  "Pin mode — click to drop a point",
            ImageViewer.MODE_POLY: "Polygon mode — click vertices, double-click to close (Esc to cancel)",
            ImageViewer.MODE_CAL:  "Calibrate — click two points of known distance",
        }
        self._status_main.setText(labels.get(mode, ""))

    # ── Defect types ────────────────────────────────────────────────────────────

    # ── Image Rename Dialog ─────────────────────────────────────────────────────

    def _rename_images_dialog(self):
        """
        RENAME: Batch image renaming dialog.
        Shows all loaded images with detected blade/face prefix.
        User can set blade (A/B/C) and face (PS/LE/TE/SS) then rename on disk.
        Format: {Blade}_{Face}_{original_filename}
        e.g. "A_PS_DJI_0042.JPG"
        """
        if not self._project or not self._image_paths:
            QMessageBox.information(self, "No Images", "Load images first.")
            return

        dlg = QDialog(self)
        dlg.setWindowTitle("Batch Rename Images")
        dlg.setMinimumSize(720, 500)
        dlg.setStyleSheet(f"background:{UI_THEME['bg_primary']};color:{UI_THEME['text_primary']};")
        lay = QVBoxLayout(dlg)

        info = QLabel(
            "Assign a Blade (A/B/C) and Face (PS/LE/TE/SS) to each image, "
            "then click Rename to add that prefix to the filename on disk.\n"
            "Folders named 'Blade A PS' etc. will already be auto-detected."
        )
        info.setWordWrap(True)
        info.setStyleSheet(f"color:{UI_THEME['text_secondary']};font-size:9pt;")
        lay.addWidget(info)

        # Table of images
        tbl = QTableWidget(len(self._image_paths), 4)
        tbl.setHorizontalHeaderLabels(["Original filename", "Blade", "Face", "Preview new name"])
        tbl.horizontalHeader().setStretchLastSection(True)
        tbl.setStyleSheet(
            f"QTableWidget{{background:{UI_THEME['bg_card']};color:{UI_THEME['text_primary']};"
            f"gridline-color:{UI_THEME['border']};}}")
        tbl.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        tbl.setAlternatingRowColors(True)

        def _update_preview(row):
            blade_cb = tbl.cellWidget(row, 1)
            face_cb  = tbl.cellWidget(row, 2)
            if blade_cb and face_cb:
                b = blade_cb.currentText()
                f = face_cb.currentText()
                orig = tbl.item(row, 0).text() if tbl.item(row, 0) else ""
                prefix = ""
                if b and b != "—":
                    prefix += b + "_"
                if f and f != "—":
                    prefix += f + "_"
                preview = prefix + orig if prefix else orig
                tbl.setItem(row, 3, QTableWidgetItem(preview))

        for i, fp in enumerate(self._image_paths):
            fname = os.path.basename(fp)
            tbl.setItem(i, 0, QTableWidgetItem(fname))

            # Auto-detect from filename or parent folder
            b_auto, f_auto = _parse_blade_face_from_folder(
                Path(fp).parent.name or "")
            if not b_auto:
                b_auto, f_auto = _parse_blade_face_from_folder(fname)

            # Blade combo
            blade_cb = QComboBox()
            blade_cb.addItems(["—", "A", "B", "C"])
            if b_auto:
                blade_cb.setCurrentText(b_auto)
            tbl.setCellWidget(i, 1, blade_cb)

            # Face combo
            face_cb = QComboBox()
            face_cb.addItems(["—", "PS", "LE", "TE", "SS"])
            if f_auto:
                face_cb.setCurrentText(f_auto)
            tbl.setCellWidget(i, 2, face_cb)

            # Wire signals (capture row with default arg trick)
            blade_cb.currentTextChanged.connect(lambda _, r=i: _update_preview(r))
            face_cb.currentTextChanged.connect(lambda _, r=i: _update_preview(r))

            _update_preview(i)

        tbl.resizeColumnsToContents()
        lay.addWidget(tbl)

        # Select-all blade/face row
        bulk_row = QHBoxLayout()
        bulk_row.addWidget(QLabel("Set ALL to:"))
        bulk_blade = QComboBox(); bulk_blade.addItems(["—", "A", "B", "C"])
        bulk_face  = QComboBox(); bulk_face.addItems(["—", "PS", "LE", "TE", "SS"])
        apply_btn  = QPushButton("Apply to all selected rows")
        apply_btn.setStyleSheet(
            f"background:{UI_THEME['accent_blue']};color:white;"
            f"border-radius:4px;padding:4px 12px;border:none;")

        def _apply_bulk():
            sel_rows = {idx.row() for idx in tbl.selectedIndexes()}
            rows = sel_rows if sel_rows else range(tbl.rowCount())
            for r in rows:
                b = bulk_blade.currentText()
                f = bulk_face.currentText()
                if b != "—":
                    tbl.cellWidget(r, 1).setCurrentText(b)
                if f != "—":
                    tbl.cellWidget(r, 2).setCurrentText(f)
                _update_preview(r)
        apply_btn.clicked.connect(_apply_bulk)
        bulk_row.addWidget(bulk_blade); bulk_row.addWidget(bulk_face)
        bulk_row.addWidget(apply_btn); bulk_row.addStretch()
        lay.addLayout(bulk_row)

        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        btns.button(QDialogButtonBox.StandardButton.Ok).setText("✏️  Rename Files")
        btns.accepted.connect(dlg.accept)
        btns.rejected.connect(dlg.reject)
        lay.addWidget(btns)

        if dlg.exec() != QDialog.DialogCode.Accepted:
            return

        # Perform renames
        renamed = 0
        errors  = []
        for i in range(tbl.rowCount()):
            preview_item = tbl.item(i, 3)
            orig_item    = tbl.item(i, 0)
            if not preview_item or not orig_item:
                continue
            new_name = preview_item.text()
            old_path = self._image_paths[i]
            if new_name == orig_item.text():
                continue  # no change
            new_path = str(Path(old_path).parent / new_name)
            try:
                os.rename(old_path, new_path)
                # Update internal paths and project
                self._image_paths[i] = new_path
                old_fname = orig_item.text()
                if old_fname in self._project.images:
                    irec = self._project.images.pop(old_fname)
                    irec.filename = new_name
                    irec.filepath = new_path
                    # Auto-assign blade/face from prefix
                    b, f = _parse_blade_face_from_folder(new_name)
                    if b and not irec.blade:
                        irec.blade = b
                    self._project.images[new_name] = irec
                # Update thumbnail strip item
                item = self._thumb_strip.item(i)
                if item:
                    item.setText(new_name)
                    item.setData(Qt.ItemDataRole.UserRole, new_path)
                renamed += 1
            except Exception as e:
                errors.append(f"{orig_item.text()}: {e}")

        save_project(self._project)
        msg = f"Renamed {renamed} files."
        if errors:
            msg += f"\nErrors ({len(errors)}):\n" + "\n".join(errors[:5])
        self._toast(msg, UI_THEME["accent_green"] if not errors else UI_THEME["accent_orange"])

    # ── Report filename helper ──────────────────────────────────────────────────

    def _report_stem(self) -> str:
        """
        T12 FIX: Build the default report filename stem as "{Sitename}-{WTG_ID}".
        Falls back gracefully when either field is empty.
        """
        if not self._project:
            return "inspection_report"
        _site = (self._project.site or self._project.name or "Inspection").strip()
        _wtg  = (self._project.turbine_id or "").strip()
        if _wtg and not _wtg.upper().startswith("WTG"):
            _wtg = f"WTG-{_wtg}"
        # Sanitise for filesystem (remove/replace chars not valid in filenames)
        def _safe(s: str) -> str:
            import re as _re
            return _re.sub(r'[\\/:*?"<>|]', "_", s).strip("_. ") or "Unknown"
        if _wtg:
            return f"{_safe(_site)}-{_safe(_wtg)}"
        return _safe(_site)

    # ── Selection-based Report ──────────────────────────────────────────────────

    def _generate_selection_report(self):
        """
        SELECTION-REPORT: Let user pick which blades/faces/images to include,
        then generate a focused PDF for just that selection.
        Mimics Scopito's per-blade PDF output.
        """
        if not self._project:
            QMessageBox.information(self, "No Project", "Open a project first.")
            return
        if not REPORTLAB_AVAILABLE:
            QMessageBox.warning(self, "ReportLab Missing",
                "pip install reportlab --break-system-packages")
            return

        dlg = QDialog(self)
        dlg.setWindowTitle("Select Images for Report")
        dlg.setMinimumSize(560, 440)
        dlg.setStyleSheet(f"background:{UI_THEME['bg_primary']};color:{UI_THEME['text_primary']};")
        lay = QVBoxLayout(dlg)

        info = QLabel(
            "Choose which blades, faces, and images to include in the report.\n"
            "All images with annotations in the selected filter will be exported."
        )
        info.setWordWrap(True)
        info.setStyleSheet(f"color:{UI_THEME['text_secondary']};font-size:9pt;")
        lay.addWidget(info)

        # Blade filter checkboxes
        b_grp = QGroupBox("Blades to include")
        b_grp.setStyleSheet(
            f"QGroupBox{{color:{UI_THEME['text_primary']};border:1px solid {UI_THEME['border']};"
            f"border-radius:4px;margin-top:6px;}}"
            f"QGroupBox::title{{subcontrol-origin:margin;padding:0 4px;}}")
        b_lay = QHBoxLayout(b_grp)
        blade_cbs = {}
        for bname in BLADE_NAMES:
            cb = QCheckBox(bname)
            cb.setChecked(True)
            blade_cbs[bname] = cb
            b_lay.addWidget(cb)
        lay.addWidget(b_grp)

        # Face filter
        f_grp = QGroupBox("Faces to include")
        f_grp.setStyleSheet(b_grp.styleSheet())
        f_lay = QHBoxLayout(f_grp)
        face_cbs = {}
        for face in ["PS", "LE", "TE", "SS"]:
            cb = QCheckBox(face)
            cb.setChecked(True)
            face_cbs[face] = cb
            f_lay.addWidget(cb)
        lay.addWidget(f_grp)

        # List of images that have annotations — user can deselect individually
        img_lbl = QLabel("Images with annotations (uncheck to exclude):")
        img_lbl.setStyleSheet(f"color:{UI_THEME['text_secondary']};font-size:8pt;")
        lay.addWidget(img_lbl)

        img_list = QListWidget()
        img_list.setStyleSheet(
            f"background:{UI_THEME['bg_card']};color:{UI_THEME['text_primary']};"
            f"border:1px solid {UI_THEME['border']};")
        img_list.setSelectionMode(QListWidget.SelectionMode.MultiSelection)
        annotated_records = {
            fname: irec for fname, irec in self._project.images.items()
            if irec.annotations
        }
        for fname, irec in sorted(annotated_records.items()):
            item = QListWidgetItem(
                f"{fname}  [{irec.blade or '?'}]  {len(irec.annotations)} annotation(s)")
            item.setData(Qt.ItemDataRole.UserRole, fname)
            item.setCheckState(Qt.CheckState.Checked)
            img_list.addItem(item)
        lay.addWidget(img_list)

        sel_all_btn = QPushButton("Select All")
        sel_all_btn.clicked.connect(lambda: [
            img_list.item(i).setCheckState(Qt.CheckState.Checked)
            for i in range(img_list.count())])
        sel_none_btn = QPushButton("Select None")
        sel_none_btn.clicked.connect(lambda: [
            img_list.item(i).setCheckState(Qt.CheckState.Unchecked)
            for i in range(img_list.count())])
        sb = QHBoxLayout()
        sb.addWidget(sel_all_btn); sb.addWidget(sel_none_btn); sb.addStretch()
        lay.addLayout(sb)

        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        btns.button(QDialogButtonBox.StandardButton.Ok).setText("📄  Generate Report")
        btns.accepted.connect(dlg.accept)
        btns.rejected.connect(dlg.reject)
        lay.addWidget(btns)

        if dlg.exec() != QDialog.DialogCode.Accepted:
            return

        # Build selected blade/face sets and image set
        sel_blades = {b for b, cb in blade_cbs.items() if cb.isChecked()}
        sel_faces  = {f for f, cb in face_cbs.items() if cb.isChecked()}
        sel_fnames = set()
        for i in range(img_list.count()):
            item = img_list.item(i)
            if item.checkState() == Qt.CheckState.Checked:
                sel_fnames.add(item.data(Qt.ItemDataRole.UserRole))

        if not sel_fnames:
            QMessageBox.information(self, "Nothing Selected",
                "No images selected — report cancelled.")
            return

        # Build a temporary filtered Project for the generator
        import copy
        filtered_proj = copy.deepcopy(self._project)
        filtered_proj.images = {}
        for fname, irec in self._project.images.items():
            if fname not in sel_fnames:
                continue
            # Filter annotations by blade/face
            filtered_anns = []
            for ann in irec.annotations:
                blade_ok = not sel_blades or (ann.blade in sel_blades)
                face_s = ann.face.split("(")[-1].strip(")") if "(" in ann.face else ann.face
                face_ok  = not sel_faces  or (face_s in sel_faces)
                if blade_ok and face_ok:
                    filtered_anns.append(ann)
            if filtered_anns:
                new_rec = copy.deepcopy(irec)
                new_rec.annotations = filtered_anns
                filtered_proj.images[fname] = new_rec

        if not filtered_proj.images:
            QMessageBox.information(self, "Nothing to Report",
                "No annotations match the selected blade/face filter.")
            return

        out_path, _ = QFileDialog.getSaveFileName(
            self, "Save Selection Report",
            str(Path(self._project.project_folder) /
                f"{self._report_stem()}_selection_report.pdf"),
            "PDF Files (*.pdf)")
        if not out_path:
            return

        dlg_prog = QProgressDialog("Generating selection report…", None, 0, 0, self)
        dlg_prog.setWindowModality(Qt.WindowModality.WindowModal)
        dlg_prog.show()
        QApplication.processEvents()

        rpt_settings_dlg = ReportSettingsDialog(self._project, self)
        settings = rpt_settings_dlg.get_settings()
        ok = ReportGenerator(filtered_proj, settings).generate(out_path)
        dlg_prog.close()

        if ok:
            n = sum(len(ir.annotations) for ir in filtered_proj.images.values())
            self._toast(
                f"Selection report saved  ·  {len(filtered_proj.images)} images  ·  {n} annotations",
                UI_THEME["accent_purple"])
            QMessageBox.information(self, "Report Generated",
                f"Saved to:\n{out_path}\n\n"
                f"{len(filtered_proj.images)} images  ·  {n} annotations")
        else:
            QMessageBox.warning(self, "Report Failed",
                "Could not generate the selection report.")

    def _manage_defect_types(self):
        if not self._project:
            QMessageBox.information(self, "No Project", "Load a project first.")
            return
        dlg = QDialog(self)
        dlg.setWindowTitle("Manage Defect Types")
        dlg.setMinimumSize(360, 420)
        lay = QVBoxLayout(dlg)
        lay.addWidget(QLabel("Defect types (one per line):"))
        txt = QPlainTextEdit("\n".join(self._project.defect_types))
        lay.addWidget(txt)
        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        btns.accepted.connect(dlg.accept)
        btns.rejected.connect(dlg.reject)
        lay.addWidget(btns)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            new_types = [t.strip() for t in txt.toPlainText().split("\n") if t.strip()]
            if new_types:
                self._project.defect_types = new_types
                self._ann_panel.update_defect_types(new_types)
                save_project(self._project)

    # ── ML ─────────────────────────────────────────────────────────────────────

    def _open_ml_dialog(self):
        if not self._project:
            QMessageBox.information(self, "No Project", "Load a project first.")
            return
        if not self._image_paths:
            QMessageBox.information(self, "No Images", "Load images first.")
            return
        dlg = MLDialog(self._image_paths, self._project, CFG, self)
        dlg.exec()
        # If detection was run, offer to open QC
        results = dlg.last_results()
        if results:
            ret = QMessageBox.question(
                self, "Detection Complete",
                f"Detection finished with {sum(len(v) for v in results.values())} "
                "detections.\nOpen QC Viewer now?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            )
            if ret == QMessageBox.StandardButton.Yes:
                self._launch_qc_with(results)

    def _launch_qc_guard(self):
        self._toast("Run ML detection first, then use QC Viewer.",
                    UI_THEME["accent_orange"])

    def _launch_qc_with(self, results: Dict):
        conf_high = float(CFG.get("DETECTION", "ConfHigh", "0.45"))
        self._qc_widget.load_results(results, self._project, conf_high)
        self._switch_to_qc_mode()

    def _switch_to_qc_mode(self):
        self._stacked.setCurrentIndex(1)
        self._status_main.setText("QC Viewer — review detections, then Commit Approved")

    def _switch_to_annotation_mode(self):
        self._stacked.setCurrentIndex(0)
        self._status_main.setText("Annotation mode")

    def _on_qc_committed(self, count: int):
        self._switch_to_annotation_mode()
        self._blade_diag.update_project(self._project)
        self._update_project_ui()
        self._update_thumbnail_borders()
        self._toast(f"{count} annotations committed from QC",
                    UI_THEME["accent_green"])

    # ── Burn-in JPEG ───────────────────────────────────────────────────────────

    def _save_annotated_jpeg(self):
        if not self._current_rec:
            QMessageBox.information(self, "No Image", "Select an image first.")
            return
        if not self._project or not self._project.project_folder:
            QMessageBox.information(self, "No Project", "Open a project first.")
            return
        if not self._current_rec.annotations:
            self._toast("No annotations on this image", UI_THEME["accent_orange"])
            return
        out = _burn_in_jpeg_annotations(
            self._current_filepath, self._current_rec,
            self._project.project_folder)
        if out:
            self._toast(f"Annotated JPEG saved → {Path(out).name}",
                        UI_THEME["accent_green"])
        else:
            QMessageBox.warning(self, "Burn-in Failed",
                "Could not save the annotated JPEG.\nEnsure Pillow is installed.")

    # ── v1.7.0: Inline file renamer ────────────────────────────────────────────

    def _on_rename_file_from_panel(self, old_filepath: str, new_stem: str):
        """v1.7.0: Rename the image file on disk and update the project + thumbnail.

        Called when the inspector clicks "✏️ Rename on Disk" in the annotation panel.

        Steps:
          1. Build new path: same directory, same extension, new stem.
          2. Warn + abort if target already exists (to prevent overwriting).
          3. os.rename on disk.
          4. Update self._project.images dict key (filename changed).
          5. Update self._image_paths list and thumbnail strip item.
          6. Update self._current_filepath so future saves use the new path.
          7. Save project JSON.
          8. Toast confirmation.
        """
        import os as _os
        if not old_filepath or not new_stem or not self._project:
            return
        old_dir  = _os.path.dirname(old_filepath)
        old_ext  = _os.path.splitext(old_filepath)[1]
        new_path = _os.path.join(old_dir, new_stem + old_ext)

        if _os.path.exists(new_path) and new_path != old_filepath:
            QMessageBox.warning(
                self, "File Already Exists",
                f"A file named '{new_stem}{old_ext}' already exists in that folder. "
                "Choose a different name.")
            return

        try:
            _os.rename(old_filepath, new_path)
        except OSError as exc:
            QMessageBox.critical(self, "Rename Failed", str(exc))
            return

        old_fname = _os.path.basename(old_filepath)
        new_fname = _os.path.basename(new_path)

        # Update project images dict
        if old_fname in self._project.images:
            irec = self._project.images.pop(old_fname)
            irec.filename = new_fname
            irec.filepath = new_path
            self._project.images[new_fname] = irec

        # Update image_paths list
        for i, fp in enumerate(self._image_paths):
            if fp == old_filepath:
                self._image_paths[i] = new_path
                break

        # Update thumbnail strip item text + data
        for i in range(self._thumb_strip.count()):
            item = self._thumb_strip.item(i)
            if item.data(Qt.ItemDataRole.UserRole) == old_filepath:
                item.setText(new_fname)
                item.setData(Qt.ItemDataRole.UserRole, new_path)
                break

        # Update current state
        self._current_filepath = new_path
        self._current_rec = self._project.images.get(new_fname)

        # Update the renamer field to show the committed new name
        self._ann_panel.set_current_filepath(new_path)

        save_project(self._project)
        self._toast(
            f"✏️ Renamed → {new_fname}", UI_THEME["accent_amber"], 3000)

    def _on_approve_annotation(self, ann: Annotation):
        """Phase 6: Reviewer approved — persist status to project JSON."""
        if not self._current_rec or not self._project:
            return
        for a in self._current_rec.annotations:
            if a.ann_id == ann.ann_id:
                a.status       = ann.status
                a.reviewed_by  = ann.reviewed_by
                a.reviewed_at  = ann.reviewed_at
                a.reviewer_note= ann.reviewer_note
                # v4.4.8: Reviewer approval implicitly certifies physical size.
                # Lock so auto-calibrate cannot silently change the certified value.
                if a.width_cm is not None:
                    a.size_locked = True
                break
        save_project(self._project)
        self._ann_panel.refresh_ann_list(self._current_rec)
        # v4.1.1: refresh QC panel status + button states immediately
        if hasattr(self, '_qc_review_panel'):
            self._qc_review_panel.refresh_current()
        self._toast("Annotation approved ✔", UI_THEME["accent_green"])
        self._update_project_ui()

    def _on_reject_annotation(self, ann: Annotation):
        """Phase 6: Reviewer rejected — persist status to project JSON."""
        if not self._current_rec or not self._project:
            return
        for a in self._current_rec.annotations:
            if a.ann_id == ann.ann_id:
                a.status       = ann.status
                a.reviewed_by  = ann.reviewed_by
                a.reviewed_at  = ann.reviewed_at
                a.reviewer_note= ann.reviewer_note
                break
        save_project(self._project)
        self._ann_panel.refresh_ann_list(self._current_rec)
        # v4.1.1: refresh QC panel status + button states immediately
        if hasattr(self, '_qc_review_panel'):
            self._qc_review_panel.refresh_current()
        self._toast("Annotation rejected ✕", UI_THEME["accent_red"])
        self._update_project_ui()

    # ── PDF report ─────────────────────────────────────────────────────────────

    def _generate_report(self):
        if not SESSION.can_do("generate_report"):
            log.warning(f"[REPORT] {SESSION.username} attempted generate_report — denied")
            QMessageBox.warning(self, "Access Denied",
                "PDF report generation is not available.")
            return
        if not self._project:
            QMessageBox.information(self, "No Project", "Open a project first.")
            return
        if not REPORTLAB_AVAILABLE:
            log.error("[REPORT] ReportLab not installed — PDF generation impossible")
            QMessageBox.warning(self, "ReportLab Missing",
                "pip install reportlab --break-system-packages")
            return
        total_ann = sum(len(ir.annotations)
                        for ir in self._project.images.values())
        if total_ann == 0:
            log.info("[REPORT] Aborted — project has no annotations")
            QMessageBox.information(self, "No Annotations",
                "No annotations in this project — nothing to report.")
            return

        # v4.5.0: Validate surface and zone for blade annotations
        blade_anns = [a for ir in self._project.images.values()
                     for a in ir.annotations if a.blade in ("A", "B", "C")]
        missing_surf_zone = []
        for ann in blade_anns:
            surf = getattr(ann, 'surface', '')
            zone = getattr(ann, 'zone', '')
            if not surf or not zone:
                missing_surf_zone.append(ann.ann_id[:8])
        
        if missing_surf_zone:
            log.warning(f"[REPORT] {len(missing_surf_zone)} blade annotations missing surface/zone")
            reply = QMessageBox.warning(self, "Missing Surface/Zone",
                f"⚠️  {len(missing_surf_zone)} blade annotation(s) are missing Surface or Zone selection.\n\n"
                f"Blade annotations require both:\n"
                f"• Surface: PS (Pressure Side) OR SS (Suction Side)\n"
                f"• Zone: LE (Leading Edge) OR TE (Trailing Edge) OR MB (Midbody)\n\n"
                f"Hub and Tower annotations do not require these fields.\n\n"
                f"Please edit these annotations before generating the report.",
                QMessageBox.StandardButton.Ok | QMessageBox.StandardButton.Cancel)
            if reply == QMessageBox.StandardButton.Cancel:
                return

        # ── Phase 6 Report Gate ────────────────────────────────────────────────
        all_anns  = [a for ir in self._project.images.values()
                     for a in ir.annotations]
        n_pending  = sum(1 for a in all_anns if a.status == "pending")
        n_rejected = sum(1 for a in all_anns if a.status == "rejected")
        n_approved = sum(1 for a in all_anns if a.status == "approved")
        log.info(
            f"[REPORT] Gate check — total={total_ann}  "
            f"approved={n_approved}  pending={n_pending}  rejected={n_rejected}")
        if n_pending > 0 or n_rejected > 0:
            gate_msg = (
                f"Review status: {n_approved} approved  ·  "
                f"{n_pending} pending  ·  {n_rejected} rejected\n\n"
                "Not all annotations have been approved.\n"
                "Rejected annotations will appear in the report with a ✕ badge.\n\n"
                "Generate the report anyway?"
            )
            ret = QMessageBox.warning(
                self, "Incomplete Review", gate_msg,
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.Cancel)
            if ret != QMessageBox.StandardButton.Yes:
                log.info("[REPORT] User cancelled at review gate")
                return

        # Pre-report dialog: DOCX option + shortcut to edit report settings
        dlg = QDialog(self)
        dlg.setWindowTitle("Generate Report")
        dlg.setMinimumWidth(420)
        lay = QVBoxLayout(dlg)
        lay.setSpacing(12); lay.setContentsMargins(20, 16, 20, 16)
        lay.addWidget(QLabel(
            f"<b>{total_ann}</b> annotation(s) ready for report generation.<br/>"
            "<small style='color:#888'>Use <i>Report Settings</i> to set company, "
            "client, turbine specs, and narrative text before generating.</small>"))
        docx_chk = QCheckBox("Also generate Word (.docx) report")
        docx_chk.setChecked(True)
        lay.addWidget(docx_chk)

        # Quick-access settings button so user doesn't need to close this dialog
        def _open_settings_inline():
            s_dlg = ReportSettingsDialog(self._project, dlg)
            if s_dlg.exec() == QDialog.DialogCode.Accepted:
                save_project(self._project)
        settings_btn = QPushButton("⚙  Edit Report Settings…")
        settings_btn.clicked.connect(_open_settings_inline)
        lay.addWidget(settings_btn)

        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        btns.button(QDialogButtonBox.StandardButton.Ok).setText("📄  Generate PDF")
        btns.accepted.connect(dlg.accept); btns.rejected.connect(dlg.reject)
        lay.addWidget(btns)
        also_docx = False
        if dlg.exec() != QDialog.DialogCode.Accepted:
            log.debug("[REPORT] User cancelled report dialog")
            return
        also_docx = docx_chk.isChecked()

        out_path, _ = QFileDialog.getSaveFileName(
            self, "Save PDF Report",
            str(Path(self._project.project_folder) /
                f"{self._report_stem()}_inspection_report.pdf"),
            "PDF Files (*.pdf)")
        if not out_path:
            log.debug("[REPORT] User cancelled save-file dialog")
            return

        log.info(
            f"[REPORT] Starting generation — "
            f"project='{self._project.name}'  annotations={total_ann}  "
            f"also_docx={also_docx}  output='{out_path}'")

        dlg_prog = QProgressDialog("Generating reports…", None, 0, 0, self)
        dlg_prog.setWindowModality(Qt.WindowModality.WindowModal)
        dlg_prog.show()
        QApplication.processEvents()

        # Phase 8.1: load saved report settings (includes narrative, specs, identity)
        rpt_settings_dlg = ReportSettingsDialog(self._project, self)
        settings = rpt_settings_dlg.get_settings()
        # summary_text from Tab 3 narrative feeds DOCX cover; no extra dialog needed

        _t0  = datetime.now()
        ok   = ReportGenerator(self._project, settings).generate(out_path)
        _t1  = datetime.now()
        _pdf_ms = int((_t1 - _t0).total_seconds() * 1000)
        if ok:
            log.info(f"[REPORT] PDF generated in {_pdf_ms} ms → {out_path}")
        else:
            log.error(f"[REPORT] PDF generation FAILED after {_pdf_ms} ms — see ReportGenerator logs")

        # Phase 8X: also generate DOCX if the user checked the box
        docx_ok   = False
        docx_path = ""
        if also_docx:
            if PYTHON_DOCX_AVAILABLE:
                docx_path = out_path.replace(".pdf", ".docx")
                if not docx_path.endswith(".docx"):
                    docx_path += ".docx"
                _t2   = datetime.now()
                docx_ok = DocxReportGenerator(self._project, settings).generate(docx_path)
                _docx_ms = int((datetime.now() - _t2).total_seconds() * 1000)
                if docx_ok:
                    log.info(f"[REPORT] DOCX generated in {_docx_ms} ms → {docx_path}")
                else:
                    log.error(f"[REPORT] DOCX generation FAILED after {_docx_ms} ms")
            else:
                log.warning("[REPORT] python-docx not available — pip install python-docx")

        dlg_prog.close()

        if ok:
            self._toast("Reports saved ✓", UI_THEME["accent_purple"])
            csv_path_shown = out_path.replace(".pdf", "_annotations.csv")
            lines = [f"PDF  → {out_path}",
                     f"CSV  → {csv_path_shown}"]
            if docx_ok:
                lines.append(f"DOCX → {docx_path}")
            elif also_docx:
                lines.append("DOCX: failed — install python-docx (pip install python-docx)")
            msg = QMessageBox(self)
            msg.setWindowTitle("Reports Generated")
            msg.setText("\n".join(lines) + f"\n\n{total_ann} annotation(s) included.")
            open_pdf_btn  = msg.addButton("📄  Open PDF",  QMessageBox.ButtonRole.ActionRole)
            open_csv_btn  = msg.addButton("📊  Open CSV",  QMessageBox.ButtonRole.ActionRole)
            if docx_ok:
                open_docx_btn = msg.addButton("📝  Open DOCX", QMessageBox.ButtonRole.ActionRole)
            else:
                open_docx_btn = None
            msg.addButton(QMessageBox.StandardButton.Ok)
            msg.exec()
            clicked = msg.clickedButton()
            target = None
            if clicked == open_pdf_btn:
                target = out_path
            elif clicked == open_csv_btn:
                target = csv_path_shown
            elif open_docx_btn and clicked == open_docx_btn:
                target = docx_path
            if target:
                try:
                    import subprocess, platform
                    if platform.system() == "Windows":
                        import os as _os; _os.startfile(target)  # type: ignore
                    elif platform.system() == "Darwin":
                        subprocess.Popen(["open", target])
                    else:
                        subprocess.Popen(["xdg-open", target])
                except Exception as exc:
                    log.warning(f"Could not open file: {exc}")
        else:
            QMessageBox.critical(self, "Report Failed",
                "Could not generate the PDF report.\n"
                "Check the log for details.")

    def _open_report_settings(self):
        """Phase 8.1: Open report settings dialog (Reviewer/Admin)."""
        if not self._project:
            QMessageBox.information(self, "No Project", "Open a project first.")
            return
        if not SESSION.can_do("generate_report"):
            QMessageBox.warning(self, "Access Denied",
                "Report settings are not available.")
            return
        dlg = ReportSettingsDialog(self._project, self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            save_project(self._project)
            self._toast("Report settings saved ✓", UI_THEME["accent_cyan"])

    # ── Role / User management ─────────────────────────────────────────────────
    # Sign-in and role management removed — not applicable without authentication.

    def _switch_role(self):
        """Sign-in removed — this function is no longer used."""
        pass

    def _manage_users(self):
        """Sign-in removed — user management is no longer applicable."""
        pass

    def _show_user_guide(self):
        """v4.1.0: Comprehensive user guide dialog with all features explained."""
        dlg = QDialog(self)
        dlg.setWindowTitle("Wind Tower Inspection Tool - User Guide")
        dlg.resize(900, 700)
        
        lay = QVBoxLayout(dlg)
        lay.setContentsMargins(0, 0, 0, 0)
        
        # Scrollable content area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        
        content = QWidget()
        content_lay = QVBoxLayout(content)
        content_lay.setContentsMargins(30, 20, 30, 20)
        content_lay.setSpacing(15)
        
        # HTML content with comprehensive guide
        guide_html = """
        <html>
        <body style='font-family: Segoe UI, Arial, sans-serif; color: #adbac7; background: #22272e;'>
        
        <h1 style='color: #539bf5; border-bottom: 2px solid #444c56; padding-bottom: 10px;'>
        🌬️ Wind Tower Inspection Tool - Complete User Guide
        </h1>
        
        <h2 style='color: #76e3ea; margin-top: 25px;'>📂 Getting Started - Folder Structure</h2>
        
        <div style='background: #2d333b; border-left: 4px solid #57ab5a; padding: 15px; margin: 15px 0;'>
        <h3 style='color: #57ab5a; margin-top: 0;'>Understanding Project Folders</h3>
        <p><strong>Option 1: Single WTG Folder</strong> (for individual turbine inspection)</p>
        <pre style='background: #1c2128; padding: 10px; border-radius: 4px;'>
WTGS001/
├── Blade_1/ (or Blade_A/)
│   ├── IMG_001.jpg
│   ├── IMG_002.jpg
│   └── ...
├── Blade_2/ (or Blade_B/)
├── Blade_3/ (or Blade_C/)
├── Tower/
│   ├── TOWER_001.jpg
│   └── ...
└── Hub/
    ├── HUB_001.jpg
    └── ...</pre>
        <p><strong>Select the WTGS001 folder</strong> when creating/opening the project.</p>
        
        <p style='margin-top: 20px;'><strong>Option 2: Full Site Parent Folder</strong> (for complete site report with multiple turbines)</p>
        <pre style='background: #1c2128; padding: 10px; border-radius: 4px;'>
WindFarm_Site/
├── WTGS001/
│   ├── Blade_1/
│   ├── Blade_2/
│   ├── Blade_3/
│   ├── Tower/
│   └── Hub/
├── WTGS002/
│   └── ...
└── WTGS003/
    └── ...</pre>
        <p><strong>Select the WindFarm_Site folder</strong> for site-wide reports.</p>
        </div>
        
        <div style='background: #373e47; border-left: 4px solid #d29922; padding: 15px; margin: 15px 0;'>
        <strong>⚠️ Important:</strong> Tower and Hub folders are <strong>automatically included</strong> when loading images. 
        You don't need to do anything special - just make sure they're named "Tower" or "Hub" and they'll be detected!
        </div>
        
        <h2 style='color: #76e3ea; margin-top: 30px;'>🎨 Annotation Workflow</h2>
        
        <h3 style='color: #b083f0;'>Step 1: Draw Annotations</h3>
        <ol style='line-height: 1.8;'>
            <li><strong>Select Draw Mode:</strong> Click Box (▭), Pin (📌), or Polygon (⬡) in toolbar</li>
            <li><strong>Draw on Image:</strong>
                <ul>
                    <li><strong>Box:</strong> Click and drag to create rectangle - it appears immediately on canvas</li>
                    <li><strong>Pin:</strong> Single click to place marker</li>
                    <li><strong>Polygon:</strong> Click to add vertices, double-click to close shape</li>
                </ul>
            </li>
            <li><strong>Fill Details:</strong> Select severity (Critical/Major/Minor/Info), defect type, location</li>
            <li><strong>Save:</strong> Click 💾 Save button to store annotation permanently</li>
        </ol>
        
        <h3 style='color: #b083f0;'>Step 2: Edit Existing Annotations</h3>
        <div style='background: #373e47; border-left: 4px solid #57ab5a; padding: 15px; margin: 15px 0;'>
        <strong>📝 How to Resize & Rotate:</strong>
        <ol style='line-height: 1.8;'>
            <li>Click the <strong>Select tool (🖱️)</strong> in toolbar</li>
            <li>Click on any saved annotation - white resize handles and cyan rotation handle will appear</li>
            <li><strong>Resize:</strong> Drag corner handles or edge midpoint handles</li>
            <li><strong>Rotate:</strong> Drag the cyan circle above the box</li>
            <li><strong>Move:</strong> Click and drag the center of the annotation</li>
            <li>Changes are saved automatically when you release the mouse</li>
        </ol>
        <p><strong>Troubleshooting:</strong> If handles don't appear after drawing, click Select mode then click the annotation again.</p>
        </div>
        
        <h2 style='color: #76e3ea; margin-top: 30px;'>📏 Calibration Guide</h2>
        
        <div style='background: #2d333b; border-left: 4px solid #3fb950; padding: 15px; margin: 15px 0;'>
        <strong>ℹ️ Calibration Policy (v4.1.1):</strong> Two calibration modes are available:<br/>
        • <strong>🤖 Auto-Calibrate All</strong> — Uses EXIF/GPS data to calibrate every loaded image individually in batch. Each image gets its own GSD calculated from its own drone metadata.<br/>
        • <strong>📏 Calibrate Image (C)</strong> — Manually draw a calibration line on the <em>current image only</em>. This does <strong>not</strong> apply to any other image. Use this when auto-calibration fails for specific images.
        </div>
        
        <h3 style='color: #b083f0;'>Manual Calibration Steps:</h3>
        <ol style='line-height: 1.8;'>
            <li>Click "📏 Calibrate Image" button in toolbar (or press C)</li>
            <li>Find a feature of <strong>known size</strong> in the image (bolt, panel edge, marking, etc.)</li>
            <li>Click once at the start point, then click again at the end point to draw calibration line</li>
            <li>System will try auto-calibration from EXIF data first (if available)</li>
            <li>If EXIF unavailable or you decline, enter the <strong>real-world distance in centimeters</strong></li>
            <li>GSD (Ground Sample Distance) is calculated and saved for this image</li>
            <li>All measurements on this image will now show accurate cm dimensions</li>
        </ol>
        
        <h2 style='color: #76e3ea; margin-top: 30px;'>🔍 QC Review Workflow (NEW in v4.1.0)</h2>
        
        <div style='background: #2d333b; border-left: 4px solid #539bf5; padding: 15px; margin: 15px 0;'>
        <p>The <strong>QC Review tab</strong> is a dedicated panel for quality control and final approval.</p>
        <h3 style='color: #76e3ea; margin-top: 10px;'>Features:</h3>
        <ul style='line-height: 1.8;'>
            <li>Automatically shows <strong>only annotated images</strong></li>
            <li>All annotation vectors are <strong>fully editable</strong></li>
            <li>You can resize, rotate, reclassify, and correct any annotation</li>
            <li>All annotation tools available (same as annotation mode)</li>
            <li>Final <strong>"Approve"</strong> button replaces "Save"</li>
            <li>Track status: pending → approved or rejected</li>
        </ul>
        <p><strong>To access:</strong> Click "🔍 QC" button in toolbar or press Ctrl+K</p>
        </div>
        
        <h2 style='color: #76e3ea; margin-top: 30px;'>⌨️ Keyboard Shortcuts</h2>
        
        <table style='width: 100%; border-collapse: collapse; margin-top: 10px;'>
        <tr style='background: #2d333b;'>
            <th style='padding: 8px; text-align: left; border: 1px solid #444c56;'>Action</th>
            <th style='padding: 8px; text-align: left; border: 1px solid #444c56;'>Shortcut</th>
        </tr>
        <tr><td style='padding: 6px; border: 1px solid #444c56;'>Select Mode</td><td style='padding: 6px; border: 1px solid #444c56;'><code>S</code></td></tr>
        <tr style='background: #2d333b;'><td style='padding: 6px; border: 1px solid #444c56;'>Box Mode</td><td style='padding: 6px; border: 1px solid #444c56;'><code>B</code></td></tr>
        <tr><td style='padding: 6px; border: 1px solid #444c56;'>Pin Mode</td><td style='padding: 6px; border: 1px solid #444c56;'><code>P</code></td></tr>
        <tr style='background: #2d333b;'><td style='padding: 6px; border: 1px solid #444c56;'>Polygon Mode</td><td style='padding: 6px; border: 1px solid #444c56;'><code>G</code></td></tr>
        <tr><td style='padding: 6px; border: 1px solid #444c56;'>Calibrate Mode</td><td style='padding: 6px; border: 1px solid #444c56;'><code>C</code></td></tr>
        <tr style='background: #2d333b;'><td style='padding: 6px; border: 1px solid #444c56;'>Save Annotation</td><td style='padding: 6px; border: 1px solid #444c56;'><code>Space</code></td></tr>
        <tr><td style='padding: 6px; border: 1px solid #444c56;'>Discard Annotation</td><td style='padding: 6px; border: 1px solid #444c56;'><code>Esc</code></td></tr>
        <tr style='background: #2d333b;'><td style='padding: 6px; border: 1px solid #444c56;'>Previous Image</td><td style='padding: 6px; border: 1px solid #444c56;'><code>←</code></td></tr>
        <tr><td style='padding: 6px; border: 1px solid #444c56;'>Next Image</td><td style='padding: 6px; border: 1px solid #444c56;'><code>→</code></td></tr>
        <tr style='background: #2d333b;'><td style='padding: 6px; border: 1px solid #444c56;'>Undo</td><td style='padding: 6px; border: 1px solid #444c56;'><code>Ctrl+Z</code></td></tr>
        <tr><td style='padding: 6px; border: 1px solid #444c56;'>Redo</td><td style='padding: 6px; border: 1px solid #444c56;'><code>Ctrl+Y</code></td></tr>
        <tr style='background: #2d333b;'><td style='padding: 6px; border: 1px solid #444c56;'>Zoom In</td><td style='padding: 6px; border: 1px solid #444c56;'><code>Ctrl + Scroll Up</code></td></tr>
        <tr><td style='padding: 6px; border: 1px solid #444c56;'>Zoom Out</td><td style='padding: 6px; border: 1px solid #444c56;'><code>Ctrl + Scroll Down</code></td></tr>
        </table>
        
        <h2 style='color: #76e3ea; margin-top: 30px;'>📄 Report Generation</h2>
        
        <p>When ready to generate the final inspection report:</p>
        <ol style='line-height: 1.8;'>
            <li>Click <strong>"📄 Report"</strong> button in toolbar or press Ctrl+R</li>
            <li>System shows review status (approved/pending/rejected annotations)</li>
            <li>Configure report settings (company name, client, turbine specs)</li>
            <li>Choose PDF only or include Word (.docx) version</li>
            <li>Report is generated with all annotations, statistics, and images</li>
        </ol>
        
        <div style='background: #373e47; border-left: 4px solid #57ab5a; padding: 15px; margin: 15px 0;'>
        <strong>✅ Report Structure Unchanged:</strong> All v4.1.0 changes maintain the existing report format. 
        Adding new defect types does not break report generation.
        </div>
        
        <h2 style='color: #76e3ea; margin-top: 30px;'>🆘 Troubleshooting</h2>
        
        <div style='background: #2d333b; border: 1px solid #444c56; padding: 15px; margin: 15px 0; border-radius: 6px;'>
        <h4 style='color: #f85149; margin-top: 0;'>Problem: Box doesn't show after drawing</h4>
        <p><strong>Solution:</strong> Fixed in v4.1.0! Box annotations now appear immediately on canvas when you release the mouse. 
        No need to wait for save button.</p>
        </div>
        
        <div style='background: #2d333b; border: 1px solid #444c56; padding: 15px; margin: 15px 0; border-radius: 6px;'>
        <h4 style='color: #f85149; margin-top: 0;'>Problem: Resize/rotate handles don't appear</h4>
        <p><strong>Solution:</strong></p>
        <ol>
            <li>Click the <strong>Select tool (🖱️)</strong> in toolbar</li>
            <li>Click on the annotation you want to edit</li>
            <li>White handles (corners/edges) and cyan rotation handle should appear</li>
            <li>If still not working, save the annotation first, then select it</li>
        </ol>
        </div>
        
        <div style='background: #2d333b; border: 1px solid #444c56; padding: 15px; margin: 15px 0; border-radius: 6px;'>
        <h4 style='color: #f85149; margin-top: 0;'>Problem: Tower/Hub images not loading</h4>
        <p><strong>Solution:</strong> Fixed in v4.1.0! Folder scanning now explicitly detects and includes Tower and Hub folders. 
        Make sure your folders are named "Tower" or "Hub" (case-insensitive) and they'll be automatically found.</p>
        </div>
        
        <h2 style='color: #76e3ea; margin-top: 30px;'>📝 Input Field Reference</h2>
        
        <p>All input fields now have tooltips. Hover over any field to see its description:</p>
        <ul style='line-height: 1.8;'>
            <li><strong>Defect Type:</strong> Select the type of defect observed</li>
            <li><strong>Blade:</strong> Which blade (A/B/C) or component (Tower/Hub)</li>
            <li><strong>Surface:</strong> PS (Pressure Side) OR SS (Suction Side) — mutually exclusive</li>
            <li><strong>Zone:</strong> LE (Leading Edge) OR TE (Trailing Edge) OR MB (Midbody) — mutually exclusive</li>
            <li><strong>Span:</strong> Root (0-33%), Mid (33-66%), or Tip (66-100%)</li>
            <li><strong>Root Distance:</strong> Distance from blade root in meters (auto-estimated if 0)</li>
            <li><strong>Tip Distance:</strong> Distance from blade tip in meters (auto-calculated)</li>
            <li><strong>Notes:</strong> Additional observations or details about the defect</li>
            <li><strong>Remedy Action:</strong> Recommended repair action (auto-filled, editable)</li>
        </ul>
        
        <p style='margin-top: 30px; padding: 20px; background: #2d333b; border-radius: 6px; text-align: center;'>
        <strong>Need more help?</strong> Check the keyboard shortcuts (Ctrl+/) or contact your system administrator.
        </p>
        
        </body>
        </html>
        """
        
        guide_label = QLabel(guide_html)
        guide_label.setWordWrap(True)
        guide_label.setTextFormat(Qt.TextFormat.RichText)
        guide_label.setOpenExternalLinks(True)
        guide_label.setStyleSheet("background: transparent;")
        content_lay.addWidget(guide_label)
        content_lay.addStretch()
        
        scroll.setWidget(content)
        lay.addWidget(scroll)
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.setFixedHeight(36)
        close_btn.clicked.connect(dlg.accept)
        close_btn.setStyleSheet(f"""
            QPushButton {{
                background: {UI_THEME['accent_blue']};
                color: white;
                border: none;
                border-radius: 6px;
                font-weight: bold;
                padding: 8px 24px;
            }}
            QPushButton:hover {{
                background: {UI_THEME['accent_cyan']};
            }}
        """)
        btn_lay = QHBoxLayout()
        btn_lay.addStretch()
        btn_lay.addWidget(close_btn)
        btn_lay.addStretch()
        lay.addLayout(btn_lay)
        
        dlg.exec()

    def _show_shortcuts(self):
        """Phase 9.1: Show keyboard shortcuts reference dialog."""
        dlg = ShortcutsDialog(self)
        dlg.exec()

    def _refresh_recent_menu(self):
        """Phase 9.3: Rebuild the Recent Projects menu from settings.ini."""
        self._recent_menu.clear()
        recents = CFG.get_recent_projects()
        if not recents:
            no_act = self._recent_menu.addAction("(no recent projects)")
            no_act.setEnabled(False)
            return
        for path in recents:
            name = os.path.basename(path)
            act  = self._recent_menu.addAction(f"  {name}")
            act.setToolTip(path)
            act.triggered.connect(lambda checked, p=path: self._open_recent(p))
        self._recent_menu.addSeparator()
        clear_act = self._recent_menu.addAction("🗑  Clear Recent List")
        clear_act.triggered.connect(self._clear_recent)

    def _open_recent(self, path: str):
        """Phase 9.3: Open a project from the recent list."""
        if not os.path.exists(path):
            QMessageBox.warning(self, "File Not Found",
                f"Project file no longer exists:\n{path}")
            self._refresh_recent_menu()
            return
        project = load_project(Path(path))
        if project:
            self._project = project
            self._update_project_ui()
            self._update_header_bar()
            CFG.add_recent_project(path)
            self._refresh_recent_menu()
            self._toast(f"Opened: {project.name}", UI_THEME["accent_cyan"])
        else:
            QMessageBox.critical(self, "Open Failed",
                f"Could not load project:\n{path}")

    def _clear_recent(self):
        for i in range(10):
            CFG.set("RECENT", f"path{i}", "")
        CFG.save()
        self._refresh_recent_menu()

    # ── About ──────────────────────────────────────────────────────────────────

    def _show_about(self):
        QMessageBox.about(
            self, f"About {APP_TITLE}",
            f"<b>{APP_TITLE}</b> — v{APP_VERSION}<br/><br/>"
            f"Wind turbine drone image annotation and inspection tool.<br/>"
            f"Annotation modes: Box · Pin · Polygon · Calibrate<br/><br/>"
            f"<b>v1.7.0 Changes:</b><br/>"
            f"• 🐛 FIXED: distance_from_root_mm always showed 0mm in reports<br/>"
            f"• ⌨️  Space → save annotation; Esc → discard; ←/→ → prev/next image<br/>"
            f"• 🔒 Compulsory calibration per component (Blade/Hub/Tower) before annotating<br/>"
            f"• 📏 Each image must be individually calibrated (EXIF auto or manual per-image)<br/>"
            f"• 🟠 Blade diagram dots now placed at correct distance-from-root Y position<br/>"
            f"• 📍 Auto-estimates distance_from_root_mm from annotation Y-centre on save<br/>"
            f"• ✏️  Inline image renamer in the annotation panel (Details tab)<br/>"
            f"• 💡 Design notes for FLIR-style full-screen viewer (future v2.0 UI)<br/><br/>"
            f"<b>Features:</b><br/>"
            f"• Scopito 1–5 + POI severity scale<br/>"
            f"• FLIR-style dark UI with 3-zone toolbar<br/>"
            f"• Polygon annotation with double-click close<br/>"
            f"• Global Undo / Redo (Ctrl+Z / Ctrl+Y)<br/>"
            f"• QC Viewer: Undo / Redo / Clear All<br/>"
            f"• Thumbnail borders coloured by worst severity<br/>"
            f"• Auto burn-in JPEG after every save<br/>"
            f"• YOLO ML detection &amp; training pipeline<br/>"
            f"• PDF + DOCX inspection report (ReportLab + python-docx)<br/><br/>"
            f"<b>Keyboard shortcuts:</b>  Space=Save  Esc=Discard  ←/→=Navigate<br/>"
            f"B=Box  P=Pin  G=Polygon  C=Calibrate  S=Select<br/><br/>"
            f"<b>Team:</b> Sarah Chen · Marcus Webb · Priya Nair · Tom K.<br/>"
            f"Elena Vasquez · Dev Patel · Alex Stone · Jamie Liu<br/>"
            f"Chris Murphy · Sam Okafor · Natalie Cross"
        )

    # ── Keyboard shortcuts ─────────────────────────────────────────────────────

    def keyPressEvent(self, event):
        """
        v1.7.0 keyboard shortcut additions (FLIR-app inspired):
          Space   — Save the currently-loaded annotation (same as clicking 💾 Save)
                    If no annotation is pending, advances to the next image instead.
          ←/→     — Navigate to previous/next image in the thumbnail strip.
                    Was previously ←↑/→↓; arrow keys now consistently mean "prev/next".
          Escape  — Discard the pending (unsaved) annotation if one exists.
          C       — Switch to Calibrate mode (also auto-triggered if no GSD is set).
        """
        key  = event.key()
        mods = event.modifiers()

        if mods & Qt.KeyboardModifier.ControlModifier:
            if key == Qt.Key.Key_Z:
                self._undo_annotation(); return
            if key == Qt.Key.Key_Y:
                self._redo_annotation(); return

        # v1.7.0: Space → save pending annotation, or advance to next image
        if key == Qt.Key.Key_Space and not mods:
            if self._ann_panel._pending_ann and self._ann_panel._save_btn.isEnabled():
                self._ann_panel._on_save()
                self._toast("Annotation saved ✓  [Space]", UI_THEME["accent_green"], 1800)
            else:
                # No pending annotation — advance to next image (FLIR QCViewer behaviour)
                r = self._thumb_strip.currentRow()
                self._thumb_strip.setCurrentRow(
                    min(r + 1, self._thumb_strip.count() - 1))
            return

        # v1.7.0: Escape → discard pending (unsaved) annotation
        if key == Qt.Key.Key_Escape and not mods:
            if (self._ann_panel._pending_ann and
                    self._ann_panel._is_new_annotation and
                    self._ann_panel._discard_btn.isVisible()):
                self._ann_panel._on_discard()
                self._toast("Annotation discarded  [Esc]", UI_THEME["accent_orange"], 1800)
            return

        shortcut_modes = {
            Qt.Key.Key_S: ImageViewer.MODE_SEL,
            Qt.Key.Key_B: ImageViewer.MODE_BOX,
            Qt.Key.Key_P: ImageViewer.MODE_PIN,
            Qt.Key.Key_G: ImageViewer.MODE_POLY,
            Qt.Key.Key_C: ImageViewer.MODE_CAL,
        }
        if key in shortcut_modes and not mods:
            mode = shortcut_modes[key]
            self._mode_btns[mode].setChecked(True)
            self._set_viewer_mode(mode)
            return

        # v1.7.0: ←/→ navigate images.  ↑/↓ also kept for consistency.
        if key in (Qt.Key.Key_Right, Qt.Key.Key_Down):
            r = self._thumb_strip.currentRow()
            self._thumb_strip.setCurrentRow(
                min(r + 1, self._thumb_strip.count() - 1))
            return
        if key in (Qt.Key.Key_Left, Qt.Key.Key_Up):
            r = self._thumb_strip.currentRow()
            self._thumb_strip.setCurrentRow(max(r - 1, 0))
            return

        super().keyPressEvent(event)

    def closeEvent(self, event):
        if self._project:
            save_project(self._project)
        CFG.save()
        self._thumb_pool.waitForDone(2000)
        event.accept()

# ==============================================================================
# ENTRY POINT
# ==============================================================================

if __name__ == "__main__":
    # CTO-AUDIT: Install crash handler before anything else runs
    _install_crash_handler()

    # v2.1.1 / Phase 9.1: Log ML availability on startup
    if YOLO_AVAILABLE:
        log.info("✅ ML Features Enabled (ultralytics + torch detected)")
    elif TORCH_AVAILABLE:
        log.info("⚠️  ML Features Disabled — ultralytics not installed (pip install ultralytics)")
    else:
        log.info("⚠️  ML Features Disabled — ultralytics/torch not installed or DLL error detected")

    QApplication.setHighDpiScaleFactorRoundingPolicy(
        Qt.HighDpiScaleFactorRoundingPolicy.PassThrough)
    app = QApplication(sys.argv)
    app.setApplicationName(APP_TITLE)
    app.setApplicationVersion(APP_VERSION)
    app.setStyle("Fusion")
    app.setStyleSheet(DARK_STYLESHEET)

    # Sign-in removed — auto-start session with full access
    SESSION.login("user", "Admin")
    log.info("Session started (no sign-in required)")

    # Phase 9.8: Wrap MainWindow() in try/except — prevent silent crash on init failure
    try:
        window = MainWindow()
    except Exception as _mw_exc:
        import traceback as _tb
        _crash_msg = _tb.format_exc()
        log.critical(f"MainWindow init failed:\n{_crash_msg}")
        # Write crash log
        _crash_path = SCRIPT_DIR / f"crash_startup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
        try:
            _crash_path.write_text(
                f"Wind Tower Inspection — Startup Crash\n"
                f"Version: {APP_VERSION}\n"
                f"Time: {datetime.now().isoformat()}\n\n{_crash_msg}",
                encoding="utf-8")
        except Exception:
            pass
        # Show user-friendly dialog before exiting
        _err_dlg = QMessageBox()
        _err_dlg.setWindowTitle("Startup Error")
        _err_dlg.setIcon(QMessageBox.Icon.Critical)
        _err_dlg.setText(
            f"<b>Wind Tower Inspection failed to start.</b><br/><br/>"
            f"A crash log has been saved to:<br/><code>{_crash_path}</code><br/><br/>"
            f"<b>Error:</b> {type(_mw_exc).__name__}: {_mw_exc}"
        )
        _err_dlg.setDetailedText(_crash_msg)
        _err_dlg.exec()
        sys.exit(1)
    window.show()
    if getattr(window, "_start_maximized", False):
        window.showMaximized()
    sys.exit(app.exec())
